"""
Append Lecture 9 (Autoimmune Connective Tissue Disease) WITH IMAGES
to SCE_ConnectAC_Study_Guide_Lec1-8.docx

Saves as: SCE_ConnectAC_Study_Guide_Lec1-9.docx

Uses Pillow re-save technique to fix JPEG header issues.
"""

import re, os, io
from docx import Document
from docx.shared import Inches
from PIL import Image

# ── Paths ────────────────────────────────────────────────────────────────────
BASE_PATH = r"C:\Users\ahaww\Downloads\mypdf\SCE_ConnectAC_Study_Guide_Lec1-8.docx"
OUT_PATH  = r"C:\Users\ahaww\Downloads\mypdf\SCE_ConnectAC_Study_Guide_Lec1-9.docx"
LEC9_IMG_DIR = r"C:\Users\ahaww\Downloads\mypdf\ConnectAC_Slides\lec9_images"

# ── Load base document ───────────────────────────────────────────────────────
doc = Document(BASE_PATH)

# ── Image helper (Pillow re-save technique) ──────────────────────────────────

def add_image(page_num, width_inches=5.0):
    """Add an image from extracted slides, re-saving via Pillow to fix JPEG headers."""
    fname = f"lec9_p{page_num}_0.jpeg"
    path = os.path.join(LEC9_IMG_DIR, fname)
    if os.path.exists(path):
        img = Image.open(path)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        stream = io.BytesIO()
        img.save(stream, format='JPEG', quality=90)
        stream.seek(0)
        try:
            doc.add_picture(stream, width=Inches(width_inches))
            print(f"  ✓ Inserted image: {fname}")
        except Exception as e:
            print(f"  ✗ Failed to insert {fname}: {e}")
    else:
        print(f"  ✗ Image not found: {path}")

# ── inline-markdown parser  **bold** *italic* ────────────────────────────────

def add_inline_runs(para, text: str):
    tokens = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = para.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            run = para.add_run(tok[1:-1])
            run.italic = True
        else:
            para.add_run(tok)

# ── paragraph adders ─────────────────────────────────────────────────────────

def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.add_run(text)

def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.add_run(text)

def add_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.add_run(text)

def add_h4(text):
    p = doc.add_paragraph(style='Heading 4')
    p.add_run(text)

def add_normal(text: str):
    p = doc.add_paragraph(style='Normal')
    speaker_pat = re.match(
        r'^(\*\*)?((Dr\.|Student|Professor)[^:]*:)(\*\*)?\s*(.*)',
        text, re.DOTALL
    )
    if speaker_pat:
        label = speaker_pat.group(2).strip()
        rest  = speaker_pat.group(5).strip()
        run_l = p.add_run(label + '  ')
        run_l.bold = True
        add_inline_runs(p, rest)
    else:
        add_inline_runs(p, text.strip())

def add_bullet(text: str, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text.strip())

def add_list_number(text: str):
    p = doc.add_paragraph(style='List Number')
    add_inline_runs(p, text.strip())



# ══════════════════════════════════════════════════════════════════════════════
# ██  LECTURE 9: AUTOIMMUNE CONNECTIVE TISSUE DISEASE (WITH IMAGES)
# ══════════════════════════════════════════════════════════════════════════════
print("\n" + "="*60)
print("  APPENDING LECTURE 9: AUTOIMMUNE CONNECTIVE TISSUE DISEASE")
print("="*60)

add_h1("Lecture 9: Autoimmune Connective Tissue Disease")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 1 – Introduction to Autoimmune CTD & Autoantibodies
# ═══════════════════════════════════════════════════════════════════════
add_h2("1. Introduction to Autoimmune CTD & Autoantibodies")
add_h3("Summary")

add_h4("Definition")
add_bullet("**Autoimmune Connective Tissue Diseases (CTD):** A group of autoimmune disorders characterized by the production of circulating autoantibodies against multiple body systems.")
add_bullet("These diseases affect both the skin and internal organs (kidney, neurological, lung, heart, etc.).")

add_h4("Difference from Other Immunological Diseases")
add_bullet("**Psoriasis / Lichen Planus** = T-cell mediated immunological diseases.")
add_bullet("**Autoimmune CTD** = B-cell / antibody mediated — autoantibodies attack the body's own tissues.")

add_h4("Types of Autoimmune CTD")
add_bullet("**Lupus Erythematosus** (LE) — the focus of this lecture.")
add_bullet("**Scleroderma** — Diffuse disease / Limited (CREST syndrome) / Acrosclerosis.")
add_bullet("**Morphea**")
add_bullet("**Dermatomyositis** — affects skin + muscle.")
add_bullet("**Antiphospholipid Syndrome**")
add_bullet("**Mixed Connective Tissue Disease (MCTD)**")
add_bullet("**Overlap Syndrome**")
add_bullet("**Sjögren's Syndrome**")
add_bullet("**ANCA-associated Vasculitis** — Churg-Strauss, Granulomatosis with Polyangiitis (Wegener's).")

# >>> IMAGE: Types of CTD diagram (page 2)
add_image(2)

add_h4("Autoantibodies — Targets in the Cell")
add_bullet("Autoantibodies are directed against all cellular components:")
add_bullet("**Nuclear molecules** → ANA (Antinuclear Antibody) — the main screening test.", level=1)
add_bullet("**Cytoplasmic molecules** → ANCA (Anti-neutrophil cytoplasmic antibody).", level=1)
add_bullet("**Cell membrane molecules** → Antiphospholipid antibodies (Antiphospholipid syndrome).", level=1)

add_h4("Nuclear Components Targeted (What's Inside the Nucleus)")
add_bullet("**Nucleic acids:** DNA (double-stranded / single-stranded), RNA (transfer, mRNA, rRNA).")
add_bullet("**Histones:** Proteins around which DNA is wrapped (like thread on a spool).")
add_bullet("**Ribonucleoproteins (RNP):** Ro, La, U1 RNP, Sm (Smith).")
add_bullet("**Enzymes:** Helicase, DNA Topoisomerase (Scl-70), Histidyl-tRNA synthetase (Jo-1).")
add_bullet("**Centromere** — targeted in CREST syndrome.")

# >>> IMAGE: Nuclear antigens diagram (page 4)
add_image(4)

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Good morning doctors, how are you? Can you hear me? What's the news? How's the studying going? We've completed the previous lectures — we finished the entire Papulosquamous chapter. Today we're starting a completely new chapter: Autoimmune Connective Tissue Diseases. The lecture will be good, God willing, but it requires your focus. I want you to participate with me. What do we mean by Autoimmune Connective Tissue Diseases? What diseases are we talking about? In Psoriasis I told you it's an immunological disease, and in Lichen Planus it's also an immunological disease. So what's the difference between those and this entire chapter called Autoimmune Connective Tissue Diseases?")
add_normal("Student:  Excuse me Dr. Eman, isn't Eczema supposed to be today's lecture according to the schedule?")
add_normal("Dr. Eman:  There's no problem, we're following the curriculum but if there's a change in lecture order, it's still within the same schedule. Can you hear me?")
add_normal("Student (Dr. Doaa):  Yes, we can hear you. Immunological diseases are T-cell mediated, while autoimmune diseases are B-cell mediated because of the antibodies.")
add_normal("Dr. Eman:  Exactly — autoantibodies that attack things in our body and cause specific diseases. What are the Autoimmune CTD diseases?")
add_normal("Student (Dr. Sherifa):  Lupus, Scleroderma, Dermatomyositis.")
add_normal("Dr. Eman:  Very good. The Autoimmune Connective Tissue Diseases, as Dr. Sherifa said, are a group of autoimmune disorders characterized by production of circulating autoantibodies against multiple body systems. Meaning the skin, kidney, neurological system (I'm talking about Systemic LE for example) — so in every disease we'll find it affects the skin and another systemic part of the body.")
add_normal("Dr. Eman:  Here's a slide showing the types of CTD: Lupus, Rheumatoid Arthritis, Scleroderma, Sjögren's Syndrome, Dermatomyositis, Microscopic Polyangiitis, Churg-Strauss Syndrome, and Granulomatosis with Polyangiitis (the ANCA vasculitis ones).")
add_normal("Dr. Eman:  The Autoimmune CTD with cutaneous manifestations that we'll cover: First, LE (today's lecture). Second, Scleroderma — divided into Diffuse and Limited. The Limited is called what?")
add_normal("Student:  CREST.")
add_normal("Dr. Eman:  CREST syndrome, or Acrosclerosis separately and CREST separately. Then Morphea, Dermatomyositis (affects skin + muscle), Antiphospholipid Syndrome, Mixed CTD, Overlap Syndrome, Sjögren's, and ANCA-associated Vasculitis.")
add_normal("Dr. Eman:  All of these fall under Autoimmune CTD — meaning autoantibodies are attacking my body. What exactly do they attack in the cell? They attack the cells. What do they attack most in the cell? And based on that we created a major test we perform when we suspect any autoimmune disease. What is it?")
add_normal("Student:  The Nucleus.")
add_normal("Dr. Eman:  Yes, the Nucleus. So what investigation do we do to detect autoantibodies against the nucleus?")
add_normal("Student:  ANA.")
add_normal("Dr. Eman:  The ANA — Antinuclear Antibody. The autoantibodies are directed against: first, nuclear molecules (the ANA); second, cytoplasmic molecules (ANCA); third, cell membrane molecules. Can anyone give me a disease example for cell membrane molecules?")
add_normal("Student:  Antiphospholipid syndrome.")
add_normal("Dr. Eman:  Exactly! From its name — the antibody is directed against the phospholipid that forms the cell membrane. The disease itself is called Antiphospholipid Syndrome.")
add_normal("Dr. Eman:  Now, the nuclear antigens — what's inside the nucleus itself: First, nucleic acids (DNA or RNA). DNA can be double-stranded or single-stranded. RNA includes transfer, mRNA, and rRNA. Then proteins — I need DNA to be wrapped around something in the nucleus, it doesn't just exist as a strand...")
add_normal("Student:  It's wrapped around histones.")
add_normal("Dr. Eman:  Exactly! Remember from before — DNA is kilometers long in the body, so it's wrapped around protein cores called histones, like thread on a spool. There's also RNP (Ribonucleoprotein) — Ro, La, U1 RNP, and Sm (Smith). If an antibody forms against Ro and La, what disease does it give me?")
add_normal("Student:  Ro and La give Subacute Lupus, or drug-induced lupus.")
add_normal("Dr. Eman:  Yes, Subacute LE, correct. There are also enzymes in the nucleus: helicase, DNA topoisomerase, histidyl-tRNA synthetase. Anti-topoisomerase gives what disease? Scleroderma (Scl-70). Anti-helicase and histidyl-tRNA synthetase? That's Jo-1.")
add_normal("Student:  Dermatomyositis.")
add_normal("Dr. Eman:  Yes, Dermatomyositis — Jo-1 and Mi-2. And anti-centromere gives CREST syndrome. So the nucleus contains many components, and an antibody against each gives a different disease.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 2 – ANA Testing — Methodology & Patterns
# ═══════════════════════════════════════════════════════════════════════
add_h2("2. ANA Testing — Methodology & Patterns")
add_h3("Summary")

add_h4("What is ANA?")
add_bullet("**ANA = Antinuclear Antibodies** — a screening test, NOT specific for any single disease.")
add_bullet("Includes all autoantibodies reactive against nuclei of human cells.")
add_bullet("Primarily of the **IgG class**.")

add_h4("Detection Methods")
add_bullet("**Indirect Immunofluorescence (IIF)** — most common, most clinically efficient screening test.")
add_bullet("**ELISA** — gives numerical result; can request 'with pattern' for fluorescence pattern.")

add_h4("How Indirect Immunofluorescence Works")
add_bullet("Take patient **serum** (contains autoantibodies).")
add_bullet("Place on a **substrate** (tissue with nuclei) — the autoantibodies bind to the nucleus.")
add_bullet("Add a **secondary antibody** conjugated with fluorochrome — it binds to the patient's antibody and glows.")
add_bullet("View under fluorescence microscope → see WHERE it glows = the pattern.")

add_h4("Substrate")
add_bullet("**HEp-2 cells** = Human laryngeal/bronchogenic carcinoma cell line (current standard).")
add_bullet("Previously used **animal substrates** (rat kidney/liver) — REPLACED because **Anti-Ro antigen** is NOT present in animals, causing false negatives.")

# >>> IMAGE: ANA pattern images (page 8)
add_image(8)

add_h4("Direct Immunofluorescence = Lupus Band Test")
add_bullet("Patient **biopsy** + patient **serum** placed together.")
add_bullet("Antibodies bind at the **dermoepidermal junction (DEJ)** → fluorescence at the DEJ.")
add_bullet("Shows: **Granular deposition of IgG + complement (C3)** at the DEJ.")
add_bullet("Used when histopathology is equivocal (cannot clearly diagnose LE).")

add_h4("ANA Patterns & Associated Diseases")
add_bullet("**Homogeneous:** SLE, Discoid LE, Rheumatoid Arthritis.")
add_bullet("**Speckled:** SLE, Mixed CTD, Scleroderma.")
add_bullet("**Nucleolar:** Scleroderma.")
add_bullet("**Peripheral:** CREST, SLE.")
add_bullet("**Centromeric:** CREST syndrome (anti-centromere antibody).")

# >>> IMAGE: ANA flowchart (page 10)
add_image(10)

add_h4("Positive ANA Titer")
add_bullet("Positive: **≥1/160** on HEp-2 cells.")
add_bullet("1/160 is MORE positive than 1/80 — because it means at higher dilution the antibody is still detected (more concentrated antibody present).")
add_bullet("If positive → proceed to specific autoantibody testing (Anti-dsDNA, Anti-Sm, Anti-Ro, Anti-La, etc.).")

add_h4("Indications for ANA Testing")
add_bullet("Suspected **autoimmune CTD** — screening test.")
add_bullet("**Baseline for Discoid LE patients** — 5% convert to Systemic LE; 15% of SLE has Discoid as skin sign.")
add_bullet("**Workup of photosensitivity** — rule out SLE (malar erythema).")
add_bullet("**Baseline before phototherapy** (theoretically).")
add_bullet("**Chronic cutaneous vasculitis** — screen for ANCA-associated disease (Churg-Strauss, Wegener's).")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Now we need a test. If I suspect someone has Dermatomyositis or Systemic LE, do I go straight to Anti-dsDNA? Or is there something that tells me we're under the umbrella of Autoimmune CTD without being too specific? What if I suspect something autoimmune but don't know what exactly? What test do I start with?")
add_normal("Dr. Eman:  I start with the ANA. The ANA is not specific for any single disease — it tells me if it's positive, then there's a problem with one of these autoimmune conditions. What specifically in the nucleus is being attacked? I don't know yet — I start searching after that.")
add_normal("Dr. Eman:  The ANA is detected by indirect immunofluorescence or ELISA. What happens? I take the patient's serum — it contains autoantibodies. I place it on a substrate (tissue with nuclei). The autoantibodies go and bind to the nucleus. Then I add a secondary antibody that glows — it attaches to the patient's antibody like following it, tracking it. It binds and lights up, showing me where the antibodies are — peripheral around the nucleus, speckled, nucleolar, etc.")
add_normal("Dr. Eman:  The substrate should be rich in nuclear material. Two types: the human substrate (HEp-2 cell — human tumor cell line) and the animal substrate (rodent cell line — rat kidney/liver). They abandoned animal substrates because not all human nuclear antigens exist in animals. The main problem was Anti-Ro — it's not present in animals, so it caused false negatives.")
add_normal("Dr. Eman:  The classic ANA indirect immunofluorescence assay is still the most clinically efficient screening test. Screening means: I have 100 suspected cases, I do ANA. 50 come back positive — those have something. 50 negative — I exclude autoimmune disease in those. The 50 positive ones may or may not have disease depending on specificity, but I've screened the group.")
add_normal("Dr. Eman:  The patterns — Homogeneous means the fluorescence covers the entire nucleus uniformly: think SLE, Discoid LE, Rheumatoid Arthritis. Speckled (dotted): SLE, Mixed CTD, Scleroderma. Nucleolar: Scleroderma. Peripheral: CREST and SLE. Centromeric: CREST syndrome.")
add_normal("Dr. Eman:  The Lupus Band Test is the direct ANA. You take a biopsy from the patient and place the patient's own serum on it. The antibodies light up at the dermoepidermal junction. We use it when histopathology is equivocal — when we can't clearly diagnose LE from the biopsy alone.")
add_normal("Dr. Eman:  What does positive ANA titer mean? Greater than or equal to 1/160. Is 1/160 positive or 1/80? Which is more positive?")
add_normal("Student:  1/160.")
add_normal("Dr. Eman:  Yes! Why? Because it's a titer — meaning I'm diluting. At 1/80 the antibody is gone, but at 1/160 it's still there. That means the antibody concentration is higher. So ≥1/160 = autoimmune CTD confirmed. If negative or below that without systemic symptoms, the patient doesn't have anything.")
add_normal("Dr. Eman:  Indications for ANA: First, suspected autoimmune CTD screening. Second, baseline for Discoid LE patients (5% convert to SLE). Third, workup of photosensitivity — any patient with photosensitivity and malar rash needs ANA screening. Fourth, chronic cutaneous vasculitis — screening for ANCA (Churg-Strauss, Wegener's).")


# ═══════════════════════════════════════════════════════════════════════
# TOPIC 3 – Specific Autoantibodies & Disease Specificity
# ═══════════════════════════════════════════════════════════════════════
add_h2("3. Specific Autoantibodies & Disease Specificity")
add_h3("Summary")

add_h4("Highly Specific Autoantibodies")
add_bullet("**Anti-dsDNA (double-stranded DNA):** Highly specific for **Systemic LE**; correlates with **nephritis** (kidney involvement).")
add_bullet("**Anti-Sm (Smith):** Highly specific for **Systemic LE**.")
add_bullet("**Anti-U1 RNP:** **Mixed Connective Tissue Disease** (MCTD); mnemonic: 'Rabbit' = U1 RNP; correlates with **CNS involvement**.")
add_bullet("**Anti-Ro / Anti-La:** **Subacute CLE**, **Neonatal LE**, **Drug-induced lupus**.")
add_bullet("**Anti-Topoisomerase I (Scl-70):** **Diffuse/Generalized Scleroderma**.")
add_bullet("**Anti-Centromere:** **CREST syndrome**.")
add_bullet("**Anti-Mi-2:** **Classic Dermatomyositis**.")
add_bullet("**Anti-Jo-1 (histidyl-tRNA synthetase):** **Dermatomyositis** — specifically the **Anti-synthetase syndrome** (lung fibrosis, mechanic's hands).")
add_bullet("**ANCA (Anti-neutrophil cytoplasmic antibody):** **Churg-Strauss**, **Granulomatosis with Polyangiitis (Wegener's)**.")

# >>> IMAGE: Autoantibody specificity table (page 14)
add_image(14)

add_h4("Prognostic Value of Autoantibodies")
add_bullet("Autoantibodies don't just diagnose — they predict **organ involvement and prognosis**:")
add_bullet("**Anti-dsDNA** → Nephritis (kidney).", level=1)
add_bullet("**Anti-U1 RNP** → CNS involvement.", level=1)
add_bullet("**Anti-Jo-1** → Lung fibrosis.", level=1)
add_bullet("Clinical pearl: In fellowship exams, they ask which specific autoantibody predicts attack on kidney vs. CNS vs. lung.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Certain antibodies have significant disease specificity. Anti-dsDNA and Anti-Sm are highly specific for Systemic LE. Anti-U1 RNP is for Mixed Connective Tissue Disease — memorize it with the word 'Rabbit' (U1 RNP). Anti-Topoisomerase (Scl-70) is for generalized Scleroderma. Anti-Centromere is for CREST. Anti-Mi-2 is for classic Dermatomyositis. Anti-Jo-1 is the more aggressive Dermatomyositis — the anti-synthetase syndrome with lung fibrosis and mechanic's hands.")
add_normal("Dr. Eman:  Here's the important part: if I have positive dsDNA, I think SLE — but it also tells me the patient is more prone to nephritis. It doesn't just diagnose the disease, it tells me WHERE it's going in the body. Is it mild SLE or is it attacking the CNS? The autoantibody tells you. Anti-dsDNA = nephritis, Anti-RNP = CNS. For the lung, this was actually a fellowship exam question — they asked which autoantibody is specific for SLE attacking the lung vs. kidney vs. CNS. We found kidney (dsDNA) and CNS (RNP) but honestly couldn't find a definitive one for lung. These must be memorized!")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 4 – Lupus Erythematosus — Classification & Pathogenesis
# ═══════════════════════════════════════════════════════════════════════
add_h2("4. Lupus Erythematosus — Classification & Pathogenesis")
add_h3("Summary")

add_h4("Classification of Cutaneous LE")
add_bullet("**Chronic Cutaneous LE:**")
add_bullet("Discoid LE (localized / disseminated / hypertrophic).", level=1)
add_bullet("Tumid LE.", level=1)
add_bullet("Lupus Panniculitis.", level=1)
add_bullet("Lupus Profundus (= Discoid + Panniculitis together).", level=1)
add_bullet("Chilblain LE.", level=1)
add_bullet("**Subacute Cutaneous LE:**")
add_bullet("Classic (annular/polycyclic or psoriasiform/papulosquamous).", level=1)
add_bullet("Neonatal LE.", level=1)
add_bullet("**Acute Cutaneous LE:**")
add_bullet("Malar erythema (butterfly rash).", level=1)
add_bullet("Generalized — essentially = Systemic LE with skin features.", level=1)
add_bullet("Other: Bullous LE, Rowell Syndrome.", level=1)

# >>> IMAGE: LE classification (page 18)
add_image(18)

add_h4("Spectrum: Chronic → Subacute → Acute")
add_bullet("**Chronic** = MOST skin involvement, LEAST systemic.")
add_bullet("**Subacute** = Intermediate.")
add_bullet("**Acute** = LEAST skin (for dermatologist), MOST systemic.")
add_bullet("As you move from Chronic to Acute, systemic involvement increases and skin-only disease decreases.")

add_h4("Pathogenesis")
add_bullet("**Environmental trigger** + **Genetic susceptibility** → **Immune dysregulation** → Disease.")
add_bullet("Immunological outcome: **Immune complex** (antigen-antibody complex) formation and deposition in tissues.")

# >>> IMAGE: Pathogenesis diagram (page 20)
add_image(20)

add_h4("The KEY GENE: STAT4")
add_bullet("**STAT4** = Signal Transducer and Activator of Transcription.")
add_bullet("Responsible for **clearance of apoptotic bodies** (the body's 'garbage collection system').")
add_bullet("**Defect in STAT4** → Accumulation of apoptotic cells → Immune system recognizes them as foreign → Antibodies attach → Immune complex formation → Tissue deposition → Inflammation → Organ damage.")

add_h4("The 'Garbage Analogy' (Dr. Eman's Teaching Pearl)")
add_bullet("Our body produces 'garbage' daily (apoptotic cells, damaged nuclei from UV exposure, etc.).")
add_bullet("Normally, a **clearance system** removes this garbage every day — like garbage collection from a building.")
add_bullet("In LE patients: The 'garbage truck' (STAT4/clearance system) is **defective**.")
add_bullet("Result: Garbage accumulates → Attracts 'insects' (antibodies) → Antibodies bind the garbage (antigen-antibody complex) → Complexes deposit in organs (kidney, skin, lung, brain, heart) → Immune attack on surrounding tissue.")

# >>> IMAGE: STAT4/clearance mechanism (page 22)
add_image(22)

add_h4("Environmental Triggers")
add_bullet("**Ultraviolet radiation** — MOST IMPORTANT trigger.")
add_bullet("**Drugs:** Isoniazid, Griseofulvin, Penicillamine, Dapsone (must memorize).")
add_bullet("**Viral infections** (e.g., Herpes Zoster).")
add_bullet("**Trauma**, X-ray, Diathermy, Chemical burn.")
add_bullet("**Stress**, Seasonal exacerbation.")
add_bullet("**Cold exposure**.")
add_bullet("**Pregnancy**, Premenstrual period, Hormonal replacement therapy — hence hormonal component.")

add_h4("Genetic Evidence")
add_bullet("Family history present.")
add_bullet("Association with certain HLA types.")
add_bullet("Role of anti-Ro/anti-La passing from mother to infant (Neonatal LE) proves genetic component.")

add_h4("Immune Complex Deposition")
add_bullet("Antigen-antibody complexes deposit in: **Kidney, Skin, Lung, Brain, Heart**.")
add_bullet("Immune system attacks the complex AND surrounding tissue → inflammation → organ damage.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Cutaneous LE is divided into three: Chronic, Subacute, and Acute. Can Chronic convert to Systemic? Yes. Can Subacute? Yes. Can Acute? It essentially IS Systemic. The more you go from Chronic toward Acute, the more systemic involvement increases and skin involvement decreases. Chronic has the most skin disease, Acute has the most systemic disease.")
add_normal("Dr. Eman:  What causes LE? Environmental trigger in a genetically susceptible patient causes immunological response. The immunological product is immune complex (antigen-antibody complex). These complexes deposit in various organs and cause damage.")
add_normal("Dr. Eman:  The most important gene is STAT4 — Signal Transducer and Activator of Transcription. What happens when there's a defect in STAT4? It causes translocation of cellular and nuclear antigens to the cell surface, triggering immunity and inflammation. STAT4 must not be forgotten! I've been explaining for 10 minutes so that when I say 'STAT4' you'll understand it properly.")
add_normal("Dr. Eman:  Think of it like this — imagine a building where every day someone comes to collect the garbage. The garbage represents apoptotic cells and waste. Every day the clearance system removes it. Now imagine it stops working — the garbage accumulates, insects appear, everything is ruined. That's exactly what happens in the body. Apoptotic cells accumulate because there's no clearance system. Then antibodies come along and say 'what's this foreign thing?' and attach to it. Antigen-antibody complex forms, deposits in organs, and the immune system attacks both the complex and the surrounding tissue. That's the pathogenesis of SLE.")
add_normal("Dr. Eman:  Environmental factors: Ultraviolet radiation is the most important. Medications like Isoniazid, Griseofulvin, Penicillamine, Dapsone — these drugs must be memorized. Viral infections, trauma, cold exposure, pregnancy and hormonal factors. The evidence for genetic involvement: family history, HLA associations, and the fact that anti-Ro/La pass from mother to baby (Neonatal LE).")


# ═══════════════════════════════════════════════════════════════════════
# TOPIC 5 – Histopathology of LE & Lupus Band Test
# ═══════════════════════════════════════════════════════════════════════
add_h2("5. Histopathology of LE & Lupus Band Test")
add_h3("Summary")

add_h4("LE Histopathology Pattern")
add_bullet("All LE = **Vacuolar interface dermatitis** (NOT lichenoid band-like).")
add_bullet("Means: **Perivascular (patchy) lymphocytic infiltrate** — lymphocytes around blood vessels, not band-like.")
add_bullet("Interface dermatitis = Vacuolar degeneration of the basement membrane + inflammatory infiltrate.")

add_h4("Discoid LE Histopathology (Most Important)")
add_bullet("**Epidermis — Things that INCREASE:**")
add_bullet("Hyperkeratosis (in stratum corneum).", level=1)
add_bullet("Follicular plugging (keratin plugs in follicles).", level=1)
add_bullet("**Epidermis — Things that DECREASE:**")
add_bullet("Epidermal atrophy (thin epidermis despite hyperkeratosis above).", level=1)
add_bullet("Dyskeratosis (death of keratinocytes).", level=1)
add_bullet("**Dermis:**")
add_bullet("Patchy lymphocytic infiltrate (perivascular).", level=1)
add_bullet("Periadnexal infiltrate (lymphocytes around hair follicles — specific for LE).", level=1)
add_bullet("**BMZ (Basement Membrane Zone):**")
add_bullet("Thickening — due to immune complex deposition.", level=1)
add_bullet("Destruction → causes SCARRING (the BMZ is 'beaten to destruction').", level=1)

# >>> IMAGE: Discoid LE histopath (page 26)
add_image(26)

add_h4("Subacute LE Histopathology")
add_bullet("**Epidermis:** Normal — NO hyperkeratosis, NO follicular plugging.")
add_bullet("**Interface:** More intense inflammation — severe hydropic degeneration, may form subepidermal cleft.")
add_bullet("**Dermis:** Pronounced dermal edema (explains the edematous clinical appearance).")
add_bullet("**BMZ:** Not completely destroyed → **NO scar** (heals with hypo/hyperpigmentation only).")
add_bullet("Abundant patchy lymphocytic infiltrate present.")

add_h4("Lupus Band Test (Direct Immunofluorescence)")
add_bullet("**Indication:** When histopathology is equivocal (cannot diagnose LE from biopsy alone).")
add_bullet("**Method:** Patient biopsy + patient serum → fluorescence at DEJ.")
add_bullet("**Finding:** Granular deposition of **IgG + complement (C3)** at the dermoepidermal junction.")
add_bullet("NOT routinely done — only when diagnosis is uncertain.")

# >>> IMAGE: Lupus band test image (page 28)
add_image(28)

add_h4("'5L' — Patchy Lymphocytic Infiltrate Differential Diagnosis")
add_bullet("**L**upus Erythematosus.")
add_bullet("**L**ymphocytic lymphoma.")
add_bullet("**L**ymphocytoma cutis.")
add_bullet("**L**ight eruption (Polymorphous light eruption — PLE).")
add_bullet("**L**ymphocytic infiltration of Jessner.")
add_bullet("These 5 diseases share patchy lymphocytic infiltrate — must be memorized!")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  How does LE appear on pathology? What's the most important finding?")
add_normal("Student (Dr. Sherifa):  The dermoepidermal junction — interface dermatitis — and there's thickening of the basement membrane from immune complex deposition.")
add_normal("Dr. Eman:  Excellent! We're dealing with interface dermatitis — vacuolar degeneration of the basement membrane. Is it vacuolar or lichenoid? It's vacuolar interface dermatitis — meaning perivascular (patchy) lymphocytic infiltrate, not band-like.")
add_normal("Dr. Eman:  In Chronic (Discoid) LE, the basement membrane is completely destroyed — that's why it scars. In the epidermis, two things increase (Hyperkeratosis + Follicular plugging) and two things decrease (Epidermal atrophy + Dyskeratosis). The dermis shows patchy lymphocytic infiltrate that's also periadnexal (around follicles) — this is specific for LE.")
add_normal("Dr. Eman:  In Subacute LE: no hyperkeratosis, no follicular plugging — the epidermis is normal. There's more intense interface inflammation with severe hydropic degeneration and pronounced dermal edema. But importantly, the BMZ doesn't completely break down — so NO scar forms.")
add_normal("Dr. Eman:  The Lupus Band Test is direct immunofluorescence. I take a biopsy from the patient, put the patient's serum on it. If it lights up at the DEJ — granular IgG + complement deposition — it confirms LE. We only do this when histopathology can't clearly diagnose LE.")
add_normal("Dr. Eman:  Important differential for patchy lymphocytic infiltrate — the 5L: LE, Lymphocytic lymphoma, Lymphocytoma cutis, Light eruption (PLE), and Lymphocytic infiltration of Jessner. These must be memorized and you should look at their clinical and histopathology images.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 6 – Chronic CLE — Discoid LE
# ═══════════════════════════════════════════════════════════════════════
add_h2("6. Chronic CLE — Discoid LE")
add_h3("Summary")

add_h4("Clinical Appearance")
add_bullet("**Discoid** = coin-shaped (like placing a coin on the skin).")
add_bullet("**Erythematous telangiectatic plaque** with follicular plugging/stippling.")
add_bullet("**Scales** present on the surface.")
add_bullet("**Severe scarring + atrophy** — central depression.")
add_bullet("**Central hypopigmentation** + **Peripheral hyperpigmentation**.")

# >>> IMAGE: Discoid LE clinical photos (page 30)
add_image(30)

add_h4("Sites")
add_bullet("**Localized:** Head, neck, face — sun-exposed areas (confirms UV as main trigger).")
add_bullet("**Disseminated:** Sun-exposed + sun-protected areas — higher risk of Systemic LE.")

add_h4("Conversion to Systemic LE")
add_bullet("**5%** of Discoid LE converts to Systemic LE.")
add_bullet("**15%** of Systemic LE has Discoid as a cutaneous manifestation.")
add_bullet("**MUST do ANA** in ALL Discoid LE cases.")

add_h4("Complications")
add_bullet("Scarring → risk of **Squamous Cell Carcinoma (SCC)** in scar tissue.")

# >>> IMAGE: DDx comparison (page 32)
add_image(32)

add_h4("Differential Diagnosis")
add_bullet("**Tinea faciei** — fungal infection of the face.")
add_bullet("**Psoriasis** — annular or discoid type.")
add_bullet("**Sarcoidosis** — fleshy papules/plaques.")
add_bullet("**Lupus vulgaris** (TB) — shares the severe scarring feature; different disease entirely.")
add_bullet("**Granuloma faciale** — chronic vasculitis; plaque with follicular dilatation (visible dilated pores).")
add_bullet("**Jessner Lymphocytic Infiltrate** — T-cell lymphoproliferative disorder.")
add_bullet("**Polymorphous Light Eruption (PLE)**.")
add_bullet("**Dimorphic fungal infection**.")
add_bullet("**Discoid eczema**.")
add_bullet("Clinical pearl: Don't just memorize differentials — look at IMAGES of each disease. Photo-memory is the most important tool for clinical dermatology.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Let's talk about Discoid or Chronic Cutaneous LE. What does Discoid mean? It means coin-shaped — like placing a coin on the skin. It's an erythematous telangiectatic plaque with follicular plugging (you can see white dots on the follicles with the naked eye — keratin plugging). There's severe scarring and atrophy, central hypopigmentation with peripheral hyperpigmentation.")
add_normal("Dr. Eman:  Location: Localized to head, neck, face — sun-exposed areas. This confirms that UV is the main trigger for LE. There's also a disseminated type that appears in both sun-exposed and sun-protected areas — this should worry us more about Systemic LE.")
add_normal("Dr. Eman:  Complications: scarring that can develop SCC. Regarding systemic conversion: 5% of Discoid LE converts to Systemic LE, OR it could be Systemic from the start with Discoid as one of its skin signs (15%). So it is a MUST to do ANA in every Discoid LE case.")
add_normal("Student:  Tinea faciei for differential.")
add_normal("Dr. Eman:  Yes! Also Psoriasis, Sarcoidosis, Lupus vulgaris (TB — same name 'Lupus' meaning 'wolf' because it eats into the tissue), Granuloma faciale, Jessner lymphocytic infiltrate, PLE, and dimorphic fungal infection. But don't just memorize these — open each one and LOOK at the images. The differentials are about recognizing visual patterns, not memorizing lists.")


# ═══════════════════════════════════════════════════════════════════════
# TOPIC 7 – Chronic CLE — Tumid LE, Lupus Panniculitis/Profundus, Chilblain LE
# ═══════════════════════════════════════════════════════════════════════
add_h2("7. Chronic CLE — Tumid LE, Lupus Panniculitis/Profundus, Chilblain LE")
add_h3("Summary")

add_h4("Tumid LE")
add_bullet("Looks like **urticaria** — edematous plaque (like a wheal that doesn't fade in 24 hours).")
add_bullet("Only **edema + erythema** — nothing else.")
add_bullet("**NO scale, NO plugging, NO scar, NO atrophy** — no epidermal changes.")
add_bullet("Histopathology: No epidermal changes; intense **dermal lymphocytic infiltrate** (patchy).")
add_bullet("DDx: **Urticarial vasculitis**, 5L diseases.")

# >>> IMAGE: Tumid LE (page 36)
add_image(36)

add_h4("Lupus Panniculitis")
add_bullet("**Panniculitis** = inflammation in the subcutaneous fat.")
add_bullet("The skin surface is **NORMAL** — the problem is deep in the fat.")
add_bullet("Presents as: **Firm nodule or plaque under normal skin**.")
add_bullet("Fat is destroyed → leaves a **cup-shaped depression** after resolution.")
add_bullet("Sites: Face, upper arm, upper trunk, breast, buttocks.")

# >>> IMAGE: Lupus panniculitis/profundus (page 38)
add_image(38)

add_h4("Lupus Profundus")
add_bullet("**Lupus Profundus** = **Discoid LE** (surface) + **Lupus Panniculitis** (deep) TOGETHER.")
add_bullet("Has BOTH: surface scarring/atrophy/plugging (from Discoid) AND deep depression (from panniculitis).")

add_h4("Chilblain LE")
add_bullet("**Red-purple papules/plaques** on toes, fingers, nose, ears, knees, elbows.")
add_bullet("Triggered by **cold + pregnancy**.")
add_bullet("Same concept as idiopathic chilblains (vasoconstriction → inflammation) but **more severe and persistent**.")
add_bullet("Idiopathic chilblains: seasonal (winter only), family history, common in females.")
add_bullet("Chilblain LE: may persist beyond winter (summer too), more aggressive, associated with LE.")
add_bullet("**Screening:** If severe/persistent → do **ANA**. Positive = investigate for LE. Negative = idiopathic.")

# >>> IMAGE: Chilblain LE (page 40)
add_image(40)

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Tumid LE — what does it look like? It looks very much like urticaria! Like seeing a wheal. But of course it's not transient and doesn't fade within 24 hours. I see edema and erythema only — no atrophy, no hyperkeratosis, no scales, no plugging, nothing. Just an edematous plaque that looks urticarial.")
add_normal("Dr. Eman:  Lupus Panniculitis — 'panniculitis' means inflammation in the subcutaneous fat. So the skin surface is completely normal — the problem is underneath. When the fat gets inflamed and destroyed, it leaves a cup-shaped depression. It presents as a firm nodule or plaque under normal skin, then resolves leaving a depression. Sites: face, upper arm, trunk, breast, buttocks.")
add_normal("Dr. Eman:  Now if I have Lupus Panniculitis underneath AND Discoid LE on the surface at the same time — both together — what's that called?")
add_normal("Student:  Lupus Profundus.")
add_normal("Dr. Eman:  Exactly! Lupus Profundus means Discoid on the skin plus Panniculitis deep. So you'll see the depression AND the scar/atrophy/plugging.")
add_normal("Dr. Eman:  Chilblain LE — what's a chilblain? When we're exposed to cold, peripheral vessels vasoconstrict, causing inflammation that appears as itching, burning, and purple papules on fingers and toes. For idiopathic chilblains, treatment is keeping warm, topical steroids, Minoxidil (vasodilator), Trental (improves circulation), Vitamin C. But Chilblain LE is the same description but more aggressive — red-purple papules on toes, fingers, nose, ears, knees, elbows. Increased by cold and pregnancy. It persists beyond winter. If severe, do ANA screening. Positive = investigate for LE.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 8 – Subacute CLE & Neonatal LE
# ═══════════════════════════════════════════════════════════════════════
add_h2("8. Subacute CLE & Neonatal LE")
add_h3("Summary")

add_h4("Subacute CLE")
add_bullet("**Photosensitive eruption** on chest and back (sun-exposed areas).")
add_bullet("Morphology: **Annular/polycyclic** OR **Psoriasiform/papulosquamous**.")
add_bullet("**No scarring** — heals with hypo/hyperpigmentation only.")
add_bullet("Antibodies: **Anti-Ro + Anti-La positive**.")
add_bullet("**Drug triggers (MUST MEMORIZE):** Thiazides, Terbinafine.")

# >>> IMAGE: Subacute CLE photos (page 44)
add_image(44)

add_h4("Subacute CLE — Differential Diagnosis")
add_bullet("**Disseminated fungal infection** — active border resemblance.")
add_bullet("**Erythema annulare** (eruptive).")
add_bullet("**Pemphigus foliaceus** — MOST SIMILAR differential (erythema + minimal scales/crustation, doesn't look bullous).")
add_bullet("**Annular Psoriasis**.")
add_bullet("**Tinea corporis**.")
add_bullet("**Pityriasis rosea**.")
add_bullet("**Granuloma annulare**.")
add_bullet("**Annular eczema**.")

add_h4("Neonatal LE")
add_bullet("Mother passes **Anti-Ro antibodies** to infant (transplacental transfer).")
add_bullet("**3 Systems Affected:**")
add_bullet("**Skin:** Raccoon eyes (annular lesions around eyes), photosensitivity — SELF-LIMITED, no treatment needed.", level=1)
add_bullet("**Heart:** Congenital heart block — **20% mortality rate** — may need **pacemaker**. MOST FEARED complication.", level=1)
add_bullet("**Hepatobiliary + Thrombocytopenia:** Reversible, relatively mild.", level=1)

# >>> IMAGE: Neonatal LE - raccoon eyes (page 46)
add_image(46)

add_h4("Neonatal LE — Investigations & Treatment")
add_bullet("Investigations: **Echo** (cardiac), **CBC** (thrombocytopenia), **Liver function tests**.")
add_bullet("Treatment — Cutaneous: Self-limited, no treatment.")
add_bullet("Treatment — Cardiac: **Pacemaker** if heart block.")
add_bullet("Treatment — Hepatic/Hematological: Reversible, supportive care.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Subacute LE comes in two forms: the regular Subacute that affects females, and Neonatal LE. It appears on the chest and back as annular or psoriasiform lesions — a photosensitive eruption in sun-exposed areas. It's positive for anti-Ro and anti-La. Important drug triggers: Thiazides and Terbinafine — memorize these. No scarring — heals with hypopigmentation or hyperpigmentation only.")
add_normal("Dr. Eman:  Differential diagnosis for annular lesions: disseminated fungal infection (looks like active border), erythema annulare, Pemphigus foliaceus (the MOST similar — it's erythema with minimal scales, doesn't look bullous at all), annular psoriasis, tinea corporis, pityriasis rosea, granuloma annulare.")
add_normal("Dr. Eman:  Neonatal LE — the mother gives anti-Ro or anti-La to her baby. The baby gets Neonatal LE. What do I fear? What am I afraid of?")
add_normal("Student:  Heart block.")
add_normal("Dr. Eman:  Exactly! I'm afraid of heart block. I see a newborn with annular lesions on the face, especially around the eyes (Raccoon eyes). My fear is congenital heart block — 20% mortality rate. The three systems: Skin (raccoon eyes, photosensitivity — self-limited for us), Heart (congenital heart block — may need pacemaker), and Hepatobiliary plus Thrombocytopenia (reversible). Investigations: Echo, CBC, Liver function. The skin is our concern as dermatologists but it resolves on its own. The heart is the most feared complication.")


# ═══════════════════════════════════════════════════════════════════════
# TOPIC 9 – Acute CLE & Systemic LE Overview
# ═══════════════════════════════════════════════════════════════════════
add_h2("9. Acute CLE & Systemic LE Overview")
add_h3("Summary")

add_h4("Acute CLE = Multi-Organ Disease (Systemic LE)")
add_bullet("Acute CLE essentially equals Systemic LE — multi-organ disease with skin features.")

add_h4("Specific Cutaneous Signs")
add_bullet("**Malar erythema** (butterfly rash).")
add_bullet("**Poikiloderma**.")
add_bullet("**Bullous LE** — antibody against Collagen VII → separation → bullae.")
add_bullet("**Oral ulceration**.")

add_h4("Non-Specific Cutaneous Signs (Vascular)")
add_bullet("**Raynaud's phenomenon**.")
add_bullet("**Livedo reticularis**.")
add_bullet("**Periungual telangiectasia**.")
add_bullet("**Purpura**.")

add_h4("Other Cutaneous Features")
add_bullet("Alopecia.")
add_bullet("Palmar/plantar erythema.")
add_bullet("Scleroderma-like changes.")
add_bullet("Calcinosis cutis.")
add_bullet("Rheumatoid nodules.")

# >>> IMAGE: Malar erythema (page 50)
add_image(50)

add_h4("Systemic LE — The 11 Criteria")
add_list_number("**Malar rash** (butterfly distribution).")
add_list_number("**Discoid rash**.")
add_list_number("**Photosensitivity**.")
add_list_number("**Oral ulcers**.")
add_list_number("**Arthritis** (non-erosive, 2+ peripheral joints).")
add_list_number("**Serositis** (pleuritis or pericarditis).")
add_list_number("**Renal** — persistent proteinuria >0.5g/day OR cellular casts.")
add_list_number("**Neurological** — seizures OR psychosis.")
add_list_number("**Hematological** — hemolytic anemia, leukopenia, lymphopenia, OR thrombocytopenia.")
add_list_number("**Immunological** — anti-dsDNA, anti-Sm, or antiphospholipid antibodies.")
add_list_number("**ANA positive**.")

# >>> IMAGE: SLE criteria table (page 52)
add_image(52)

add_h4("Bullous LE — Special Variant")
add_bullet("Antibody against **Collagen VII** → basement membrane separation → bullae formation.")
add_bullet("NOT a separate bullous disease — it's part of Systemic LE process (autoantibodies targeting collagen VII just like they target kidney/lung).")
add_bullet("Histopathology: **Like Dermatitis Herpetiformis (DH) with neutrophils** — memorize this association.")

add_h4("Rowell Syndrome")
add_bullet("**Erythema multiforme-like lesions** + **Cutaneous LE**.")
add_bullet("EM-like lesions on face, neck, hands, chest.")
add_bullet("Important for SCE exams — frequently tested.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Acute LE is a multi-organ disease. Remember the SLE criteria we've learned throughout our years of study? That's what this is. From the skin perspective: specific signs include malar erythema, poikiloderma, bullous LE, and oral ulceration. Non-specific vascular signs: Raynaud's, livedo reticularis, periungual telangiectasia, purpura. Other findings: scleroderma, calcinosis cutis, alopecia, palmar erythema.")
add_normal("Dr. Eman:  The criteria — approximately 11: Malar rash, Discoid rash, Photosensitivity, Oral ulceration (these four are the skin ones). Then: Arthritis, Serositis, Renal, Neurological, Hematological, Immunological. Renal means persistent proteinuria >0.5 with cellular casts. Neurological means seizures or psychosis. Hematological: hemolytic anemia, leukopenia, lymphopenia. These come as cases in exams.")
add_normal("Dr. Eman:  Bullous LE — there's an antibody against collagen VII. What's collagen VII? It's part of the DEJ. If it's destroyed, separation occurs and bullae form. But this is part of Systemic LE — the same autoimmune process that attacks kidney and lung also attacks collagen VII. Histopathology: like DH with neutrophils. The Rowell Syndrome is LE with erythema multiforme-like lesions on face, neck, hands, chest — important for exams!")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 10 – Treatment of LE
# ═══════════════════════════════════════════════════════════════════════
add_h2("10. Treatment of LE")
add_h3("Summary")

add_h4("Cutaneous LE Treatment")
add_bullet("**Sun protection** — essential for all LE patients.")
add_bullet("**Topical/intralesional corticosteroids** (superpotent).")
add_bullet("**Topical calcineurin inhibitors**.")
add_bullet("**Topical retinoids**.")
add_bullet("**Systemic antimalarials:** Hydroxychloroquine, Chloroquine, Quinacrine 100mg daily.")
add_bullet("Combination antimalarial therapy if needed.")

add_h4("Antimalarial-Resistant Cutaneous LE")
add_bullet("**Retinoids** (systemic).")
add_bullet("**Thalidomide**.")
add_bullet("**Dapsone**.")
add_bullet("**Immunosuppressives:** Mycophenolate mofetil, Azathioprine.")

# >>> IMAGE: Treatment algorithm (page 56)
add_image(56)

add_h4("Systemic LE Treatment — By Severity")
add_bullet("**Mild** (no major organ involvement): NSAIDs.")
add_bullet("**Moderate:** Corticosteroids + immunosuppressive agents.")
add_bullet("**Moderate-Severe:** Cyclophosphamide + Corticosteroids.")

add_h4("Biologics")
add_bullet("**Rituximab** = Anti-CD20 / Anti-B cell monoclonal antibody.")
add_bullet("Used in ANY antibody-mediated disease (attacks B-cells that produce the pathogenic antibodies).")
add_bullet("Also used in: **Pemphigus** (bullous disease), other antibody-mediated conditions.")
add_bullet("Mechanism: Targets CD20 on B-lymphocytes → depletes B-cells → reduces autoantibody production.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Treatment for cutaneous LE: Sun protection, topical/intralesional steroids, topical calcineurin inhibitors, topical retinoids, and systemic antimalarials — Hydroxychloroquine, Chloroquine, Quinacrine 100mg daily. If antimalarial-resistant: retinoids, thalidomide, dapsone, immunosuppressives like Mycophenolate mofetil and Azathioprine.")
add_normal("Dr. Eman:  For Systemic LE — if mild with no major organ involvement: NSAIDs. Moderate: corticosteroids plus immunosuppressive. Moderate to severe: Cyclophosphamide with corticosteroids. What biologic do we use? Something against B-lymphocytes, against CD20. Has anyone heard of Rituximab?")
add_normal("Student:  Rituximab... I get confused between it and JAK inhibitors.")
add_normal("Dr. Eman:  No, JAK inhibitors are different (Tofacitinib). Rituximab is a monoclonal antibody against CD20, against B-cells. We use it in any disease related to antibodies attacking the body. It's used in Pemphigus as well — any disease related to B-cells and antibodies. It's called anti-CD20 or anti-B cell (Rituximab).")


# ═══════════════════════════════════════════════════════════════════════
# BONUS — End-of-Lecture MCQ Discussion
# ═══════════════════════════════════════════════════════════════════════
add_h2("Bonus: End-of-Lecture MCQ Discussion")
add_h3("Q1: Venous Ulcer Management")

add_normal("Dr. Eman:  Let's solve a few SCE questions. This one: A 49-year-old woman presented to the leg ulcer clinic with a painful ulcerated area over the left ankle, 4 months duration, dressed with iodine and honey by the practice nurse but gradually enlarging and becoming more painful. She was prescribed antibiotics after a swab grew Pseudomonas. History: diabetes, rheumatoid arthritis, ischemic heart disease. Medications: Ramipril, Simvastatin, Aspirin, Naproxen. Examination: ulcer on left lower leg with dull red indurated surrounding skin; right leg had edema and dull red discoloration with induration. Foot pulses palpable.")
add_normal("Dr. Eman:  What is the most appropriate treatment? Options: A) Advise elevating the left lower leg. B) Multi-layer compression bandage. C) Single layer compression bandage. D) Urgent incisional biopsy. E) Compression stocking.")
add_normal("Student:  Multi-layer compression bandage.")
add_normal("Dr. Eman:  Correct! They're all close to each other. The key is understanding that for varicose veins/venous ulcers, we need compression. But is it multi-layer or single layer? A large systematic review showed that chronic venous ulcers heal faster with compression — specifically with **multi-component systems containing elastic**, which are more effective than single-component systems. This level of detail won't come up in our lectures but appears in SCE questions. The answer: **Multi-layer compression bandage**.")

add_h3("Q2: Vitamin A and Isotretinoin")

add_normal("Dr. Eman:  A patient with nodulocystic acne on face, chest, and back is about to start Isotretinoin. He asks whether his vitamin supplements would interact. Which vitamin should he discontinue?")
add_normal("Student:  Vitamin A.")
add_normal("Dr. Eman:  Correct — straightforward question. Isotretinoin is a retinoid, chemically related to Vitamin A. Concomitant use of Vitamin A supplements causes potential danger of **hypervitaminosis A**. Simple question, nothing complicated.")

add_h3("Q3: Acneiform Eruption from Anti-TB Drugs")

add_normal("Dr. Eman:  A patient on TB medications develops inflammatory papules and pustules on cheeks, chin, upper chest, and upper back. No comedones. Which medication is responsible?")
add_normal("Student:  Isoniazid and Rifampicin?")
add_normal("Dr. Eman:  Let me explain the thought process. First identify WHAT the rash is: papules and pustules, no comedones = **acneiform eruption** (not true acne — no comedones). What drugs cause acneiform eruption? Steroids, Phenytoin, Barbiturates... and among anti-TB drugs: **Isoniazid and Rifampicin**. This is uncommon and not well-documented in many textbooks, but it's the answer. The full list of drugs causing acneiform eruption: Anabolic steroids, Androgenic hormones, Anticonvulsants, Antidepressants (Lithium), Corticosteroids, Isoniazid, Rifampicin — know this list.")

add_h3("Bonus Q&A: Drug-Triggered Psoriasis & Biologics")

add_normal("Student (Dr. Doaa):  Dr. Eman, about yesterday's psoriasis drug question — Lithium, Indomethacin, and Inderal. Why wasn't the answer 'all of the above'?")
add_normal("Dr. Eman:  Because Lithium is the strongest/most important trigger (50% trigger rate). Indomethacin (NSAID) is mentioned in some sources but isn't among the 'classic strong triggers' for psoriasis. The question logic: always pick the MOST correct or STRONGEST answer. For SCE specifically, the questions are case-based, so always choose the best single answer.")
add_normal("Student (Dr. Doaa):  And the Crohn's disease biologics question — why Ustekinumab over Infliximab?")
add_normal("Dr. Eman:  Ustekinumab (Stelara) is FDA-approved for Psoriatic Arthritis AND Crohn's disease. Secukinumab (Cosentyx) should NOT be used with Crohn's — it can actually trigger/worsen IBD. Infliximab has specific indications. Read the explanation — it clarifies why Ustekinumab was the correct answer for a patient with both Psoriatic Arthritis and Crohn's.")


# ══════════════════════════════════════════════════════════════════════════════
# ██  SAVE OUTPUT DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
print("\n" + "="*60)
print("  SAVING OUTPUT DOCUMENT")
print("="*60)

doc.save(OUT_PATH)
print(f"\n  ✓ Saved: {OUT_PATH}")
print(f"  ✓ Lecture 9 (Autoimmune CTD) appended successfully with images!")
print("="*60)
