"""Build a comprehensive Word document collating all 62 dermatology transcripts."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = '/projects/sandbox/mypdf'

# Each entry: (transcript_filename, display_title, summary_table_rows)
# summary_table_rows: list of (key, value) tuples
ENTRIES = [
    # -------- Inflammatory / Eczema --------
    ('Atopic_Dermatitis_transcript.txt', 'Atopic Dermatitis', [
        ('Atopic Triad', 'Atopic dermatitis + Allergic rhinitis + Asthma'),
        ('Epidemiology', 'Higher prevalence in high-income / urban areas'),
        ('Subtypes', 'Early-onset (most common, by 1-2y, 50% IgE+); Late-onset (after puberty); Senile (after 60y)'),
        ('Pathogenesis', 'Filaggrin mutations (strongest genetic risk; severe & persistent disease, hand dermatitis); SPINK5 mutations (LEKTI / Desmoglein-1 degradation); S. aureus colonisation; trans-epidermal water loss'),
        ('Acute AD Cytokines', 'Th2 dominant: IL-4, IL-5, IL-10, IL-13'),
        ('Chronic AD Cytokines', 'Th1 dominant: IL-1, IFN-gamma'),
        ('Associations', 'Food allergies, ichthyosis vulgaris (filaggrin)'),
    ]),
    # -------- Genodermatoses / Ichthyoses --------
    ('Ichthyosis_4_transcript.txt', 'Ichthyoses (Part 4) - Chanarin-Dorfman, Refsum, KID, EKV', [
        ('Chanarin-Dorfman Syndrome', 'ABHD5 mutation; neutral lipid storage disease; developmental delay, hepatomegaly, electron-lucent inclusions'),
        ('Refsum Disease', 'PEX7/PHYH mutation; cerebellar ataxia, deafness, salt-and-pepper retinitis pigmentosa, ↑ phytanic acid'),
        ('KID Syndrome', 'GJB2 (Connexin 26); Keratitis + Ichthyosis + Deafness; coarse leather skin; ↑ risk oral & cutaneous SCC'),
        ('Erythrokeratodermia Variabilis', 'GJB3/GJB4; transient erythematous patches + geographic hyperkeratotic plaques'),
    ]),
    ('Acrodermatitis_Enteropathica_transcript.txt', 'Acrodermatitis Enteropathica', [
        ('Inheritance', 'Autosomal recessive'),
        ('Gene', 'SLC39A4 → encodes ZIP4 zinc transporter'),
        ('Acquired form', 'Secondary zinc deficiency: alcoholism, anorexia, IBD, formula-fed infants, low maternal breastmilk zinc'),
        ('Clinical Triad', 'Erosive vesiculopustular eczematous lesions (diaper, face, acral) + diarrhoea + alopecia'),
        ('Histology', 'Cytoplasmic pallor of keratinocytes + ballooning in upper epidermis'),
        ('Labs', '↓ serum zinc, ↓ alkaline phosphatase'),
        ('Differentials', 'Biotin deficiency, cystic fibrosis (if zinc normal)'),
        ('Treatment', 'Lifelong zinc sulfate supplementation'),
    ]),
    ('Rothmund_Thomson_Syndrome_transcript.txt', 'Rothmund-Thomson Syndrome', [
        ('Inheritance', 'Autosomal recessive'),
        ('Gene', 'RECQL4 (DNA helicase)'),
        ('Skin', 'Erythema/edema → blistering on cheeks → poikiloderma; acral verrucous keratoses'),
        ('Skeletal', 'Short stature, absent/hypoplastic thumbs, radius, ulna'),
        ('Other', 'Juvenile cataracts, hypogonadism'),
        ('Cancer risk', 'Osteosarcoma & SCC'),
    ]),
    ('Trichothiodystrophy_transcript.txt', 'Trichothiodystrophy (TTD / Tay / PIBIDS)', [
        ('Inheritance', 'Autosomal recessive'),
        ('Hair', 'Brittle, low cysteine; "tiger tail" appearance on polarised microscopy'),
        ('Photosensitive form', 'ERCC2 (XPB), ERCC3 (XPD), GTF2H5 (TTDA); has ichthyosis'),
        ('Non-photosensitive form', 'No ichthyosis'),
        ('Key feature', 'NO increased skin cancer risk (unlike XP)'),
        ('PIBIDS', 'Photosensitivity, Ichthyosis, Brittle hair, Intellectual impairment, Decreased fertility, Short stature'),
    ]),
    ('Tuberous_Sclerosis_transcript.txt', 'Tuberous Sclerosis Complex', [
        ('Inheritance', 'Autosomal dominant'),
        ('Genes', 'TSC1 (hamartin) or TSC2 (tuberin) → mTOR pathway'),
        ('Skin', 'Ash-leaf macules, shagreen patches, facial angiofibromas (topical rapamycin), periungual fibromas (Koenen), café-au-lait, dental enamel pits'),
        ('Neurological', 'Cortical tubers, infantile spasms, seizures (#1 cause of mortality), periventricular calcifications'),
        ('Renal', 'Angiomyolipomas + cysts (#2 cause of mortality)'),
        ('Other', 'Retinal phacomas, cardiac rhabdomyomas, WPW, pulmonary LAM'),
    ]),
    ('Xeroderma_Pigmentosum_transcript.txt', 'Xeroderma Pigmentosum (XP)', [
        ('Defect', 'Nucleotide excision repair (NER) pathway'),
        ('Variants', 'XPA-XPG and XPV'),
        ('XPV', 'Defect in post-replication repair (DNA pol η)'),
        ('XPB, XPD, XPG', 'XP-Cockayne overlap'),
        ('XPB, XPD', 'Also associated with TTD (tiger tail hair)'),
        ('Skin', 'UV sensitivity, severe photodamage, 1000-fold ↑ skin cancer risk under age 20 (BCC, SCC, melanoma, lentigines)'),
        ('Neurological', 'XPA & XPD (20-30%): developmental delay, hyporeflexia, ataxia'),
        ('Severe variant', 'De Sanctis-Cacchione syndrome'),
    ]),
    # -------- Histiocytoses --------
    ('Histiocytoses_transcript.txt', 'Histiocytoses (Cell Types & Stains)', [
        ('Common origin', 'CD34+ progenitor cell in bone marrow'),
        ('Langerhans cells', 'Antigen-presenting; migrate to/from epidermis'),
        ('Langerhans stains', 'Langerin (CD207) - most specific; CD1a; S100'),
        ('Birbeck granules', 'Tennis-racket shaped on EM; stained by CD207'),
        ('Langerhans nuclei', 'Reniform (kidney-shaped)'),
        ('Macrophages', 'CD68+, HAM56+'),
        ('Dendritic cells Type 1', 'Factor XIIIa positive'),
        ('Dendritic cells Type 2', 'CD34 positive'),
    ]),
    ('Histiocytoses_2_transcript.txt', 'Langerhans Cell Histiocytosis (LCH)', [
        ('Cell type', 'Immature haematopoietic myeloid precursors'),
        ('Age', 'Children 1-3 years; more common in white patients'),
        ('Key mutation', 'BRAF V600E (50-60%) → MAP kinase pathway'),
        ('Stains', 'S100+, CD1a+, CD207+, CD68+, HAM56 negative'),
        ('Birbeck granules', 'Present'),
        ('Most common organ', 'Bone (especially skull) → then skin'),
        ('Risk organs', 'Liver, spleen, bone marrow → worse prognosis (75% 5-yr survival)'),
        ('Skin', 'Seborrheic dermatitis-like crusted papules on scalp; groin involvement'),
        ('Sequelae', 'Diabetes insipidus, facial asymmetry, vertebral collapse, hearing loss, neurodegeneration'),
        ('Treatment', 'Mild: clobetasol; Diffuse: prednisone + vinblastine'),
    ]),
    # -------- Metabolic --------
    ('Homocystinuria_transcript.txt', 'Homocystinuria', [
        ('Inheritance', 'Autosomal recessive'),
        ('Enzyme defect', 'Cystathionine beta-synthase (CBS)'),
        ('Skin', 'Hypopigmentation, livedo reticularis, leg ulcers'),
        ('Eyes', 'Ectopia lentis - downward (vs Marfan: upward)'),
        ('Skeletal', 'Marfanoid habitus'),
        ('Cardiac', 'Mitral valve prolapse'),
        ('Vascular', 'Thromboembolic events - venous, cardiac, cerebrovascular'),
        ('Treatment', 'Vitamin B6 + folate + B12; cystine supplementation; methionine restriction'),
    ]),
    ('Hunter_Syndrome_transcript.txt', 'Hunter Syndrome (MPS II)', [
        ('Inheritance', 'X-linked recessive'),
        ('Gene', 'IDS → encodes iduronate-2-sulfatase'),
        ('Skin', 'Pebbled ivory plaques between scapulae, upper arms, thighs; hypertrichosis; coarse facies'),
        ('Pigmentation', 'Extensive dermal melanosis (Mongolian spots)'),
        ('Family', 'Mucopolysaccharidoses - all have hypertrichosis + coarse facies'),
        ('Related', 'Hurler syndrome (MPS I) - gargoyle appearance + developmental delays'),
    ]),
    ('Incontinentia_Pigmenti_transcript.txt', 'Incontinentia Pigmenti', [
        ('Gene', 'NEMO (IKBKG) - prevents NF-κB activation → ↓ TNF-α apoptosis'),
        ('Inheritance', 'X-linked dominant; lethal in males (except Klinefelter)'),
        ('Pattern', 'Functional mosaicism via lyonisation → Blaschkoid distribution'),
        ('System', 'Neuro-ectodermal disorder'),
        ('Stage 1 (Vesicular)', 'Eosinophilic spongiosis, apoptotic keratinocytes'),
        ('Stage 2 (Verrucous)', 'Acanthosis, squamous eddies'),
        ('Stage 3 (Hyperpigmented)', 'Pigment incontinence, melanophages'),
        ('Stage 4 (Hypopigmented)', 'Epidermal atrophy, melanin loss'),
        ('Skin', 'Scarring alopecia (10-20%); subungual SCC-like tumours (10%)'),
        ('Teeth', 'Pegged, conical teeth, anodontia (~50%) - most common extracutaneous'),
        ('CNS', 'Seizures, developmental delays'),
        ('Eyes', '↓ visual acuity, retinal vascular changes'),
    ]),
    ('Lesch_Nyhan_Syndrome_transcript.txt', 'Lesch-Nyhan Syndrome', [
        ('Inheritance', 'X-linked recessive'),
        ('Gene', 'HGPRT → hypoxanthine-guanine phosphoribosyltransferase → uric acid accumulation'),
        ('Early sign', 'Orange uric acid crystals in nappy / haematuria'),
        ('Neurological', 'Severe developmental delays, intellectual impairment'),
        ('Behavioural', 'Self-mutilation (pathognomonic)'),
        ('Renal', 'Uric acid nephropathy'),
        ('Joints', 'Gout'),
        ('Treatment', 'Allopurinol'),
    ]),
    ('Alagille_Syndrome_transcript.txt', 'Alagille Syndrome', [
        ('Gene', 'JAG1 → encodes Jagged-1 protein'),
        ('Facies', 'Triangular face'),
        ('Skin/Lipids', 'Tuberous xanthomas, hypercholesterolaemia, hypertriglyceridaemia'),
        ('Liver', 'Congenital intrahepatic biliary hypoplasia → cholestasis, pruritus'),
        ('Growth', 'Failure to thrive'),
        ('Treatment', 'Liver transplantation'),
        ('Prognosis', 'Often fatal before age 5 if untreated'),
    ]),
    ('Alkaptonuria_transcript.txt', 'Alkaptonuria', [
        ('Inheritance', 'Autosomal recessive'),
        ('Enzyme defect', 'Homogentisate-1,2-dioxygenase (HGD)'),
        ('Pigmentation', 'Blue-grey ochronosis - face, nose, ear cartilage, sclera'),
        ('Body fluids', 'Dark urine, sweat, cerumen on standing'),
        ('Joints', 'Severe arthritis, intervertebral disc calcification'),
    ]),
    ('Biotinidase_Multiple_Carboxylase_Deficiency_transcript.txt', 'Biotinidase & Multiple Carboxylase Deficiency', [
        ('Inheritance', 'Autosomal recessive'),
        ('Mechanism', 'Biotin is cofactor for 4 carboxylases'),
        ('Biotinidase deficiency', 'Childhood onset'),
        ('Holocarboxylase synthetase deficiency', 'Early infancy; more severe; often fatal'),
        ('Skin', 'Perioral & generalised dermatitis, alopecia'),
        ('Eyes', 'Optic atrophy (biotinidase deficiency)'),
        ('Treatment', 'IV biotin replacement'),
    ]),
    ('Bloom_Syndrome_transcript.txt', 'Bloom Syndrome', [
        ('Inheritance', 'Autosomal recessive'),
        ('Gene', 'BLM (RECQL3) - DNA helicase'),
        ('Aka', 'Congenital telangiectatic erythema'),
        ('Growth', 'Pre/postnatal retardation; height ≤5 feet'),
        ('Skin', 'Photosensitivity, telangiectatic erythema malar (butterfly)'),
        ('Facies', 'Narrow face, prominent ears, bird-like nose, high-pitched voice'),
        ('Endocrine', 'Primary hypogonadism - males sterile'),
        ('Immunology', '↓ IgA and IgM'),
        ('Cytogenetics', '↑ sister chromatid exchange'),
        ('Cancer risk', '↑ Lymphoma, leukaemia (#1 cause of death)'),
    ]),
    ('Cockayne_Syndrome_transcript.txt', 'Cockayne Syndrome', [
        ('Inheritance', 'Autosomal recessive'),
        ('Defect', 'Defective transcription-coupled NER (vs XP = global genomic NER)'),
        ('Genes', 'CSA → ERCC8; CSB → ERCC6'),
        ('Skin', 'Photosensitivity, telangiectatic erythema; NO skin cancer risk'),
        ('Facies', 'Pinched, narrow, bird-like face with beak nose'),
        ('Neuro', 'Progressive demyelination CNS & PNS, ataxia'),
        ('Growth', 'Cachectic dwarfism'),
        ('Eyes', 'Salt and pepper retinopathy'),
        ('Prognosis', 'Death by 4th decade from neurological complications'),
    ]),
    ('Cutis_Laxa_transcript.txt', 'Cutis Laxa', [
        ('Aka', 'Generalised elastolysis - "hound dog" facies'),
        ('AD form', 'ELN gene; less common, early adulthood, mainly skin'),
        ('AR Type I', 'FBLN5 (fibulin-5); hypoplastic lungs/emphysema, GI/GU diverticula'),
        ('AR Type IIA', 'ATP6V0A2'),
        ('AR Type IIB', 'PYCR1'),
        ('AR Type III (De Barsy) IIIA', 'ALDH18A1; corneal clouding, cataracts'),
        ('AR Type IIIB', 'PYCR1'),
        ('X-linked recessive', 'Occipital Horn Syndrome - wedge-shaped occipital calcifications'),
        ('Acquired', 'Adults; drugs (penicillamine, INH); RA, sarcoidosis, lupus'),
    ]),
    ('Hartnup_Disease_transcript.txt', 'Hartnup Disease', [
        ('Inheritance', 'Autosomal recessive'),
        ('Gene', 'SLC6A19 - intestinal/renal neutral amino acid transporter'),
        ('Mechanism', '↓ Tryptophan absorption → ↓ Niacin'),
        ('Skin', 'Pellagra-like photosensitivity → erythema, blistering, scaling, scarring'),
        ('Onset', 'Acute photodermatitis in childhood'),
        ('Neuro', 'Cerebellar ataxia, seizures, intellectual disability, psychosis'),
        ('Treatment', 'Sun avoidance + oral nicotinamide'),
    ]),
    # -------- Connective tissue --------
    ('Marfan_Syndrome_transcript.txt', 'Marfan Syndrome', [
        ('Inheritance', 'Autosomal dominant'),
        ('Gene', 'FBN1 → encodes Fibrillin-1'),
        ('Habitus', 'Tall, slim, long extremities (Marfanoid)'),
        ('Skin', 'Striae; lack of subcutaneous fat'),
        ('Skin association', 'Elastosis perforans serpiginosa - lateral neck'),
        ('Skeletal', 'Arachnodactyly'),
        ('Eyes', 'Ectopia lentis - upward (~60%)'),
        ('Cardiovascular', 'Ascending aortic dilatation, MVP, dissection risk'),
        ('Pulmonary', '↑ risk spontaneous pneumothorax'),
    ]),
    # -------- Infections --------
    ('Molluscum_Contagiosum_transcript.txt', 'Molluscum Contagiosum', [
        ('Cause', 'Molluscum contagiosum virus 1 & 2 (poxvirus)'),
        ('Demographics', 'School-aged children'),
        ('Lesion', 'Pearly pink umbilicated papules'),
        ('Widespread in', 'Atopic dermatitis, ichthyosis (impaired barrier)'),
        ('Giant in', 'Immunosuppression, HIV/AIDS'),
        ('Histology', 'Molluscum (Henderson-Patterson) bodies in dermis'),
        ('Treatments', 'Cryotherapy, curettage, cantharidin, imiquimod, Candida antigen'),
        ('Course', 'Self-limiting'),
    ]),
    # -------- Connective tissue / Skeletal --------
    ('Osteogenesis_Imperfecta_transcript.txt', 'Osteogenesis Imperfecta', [
        ('Inheritance', 'Autosomal dominant'),
        ('Gene', 'COL1A1, COL1A2 (Type I collagen)'),
        ('Type 1', 'Most common (50%); mild'),
        ('Type 2', 'Most severe - fatal perinatal'),
        ('Type 3', 'Progressive deforming'),
        ('Type 4', 'Lacks blue sclera (unlike 1-3)'),
        ('Skin', 'Thin, atrophic, translucent; easy bruising'),
        ('Skeletal', 'Fragile bones, fractures, beaded ribs'),
        ('Eyes', 'Blue sclera (90%, except type 4)'),
        ('Ears', 'Hearing loss from otosclerosis'),
        ('Teeth', 'Fragile, discoloured (dentinogenesis imperfecta)'),
    ]),
    ('Pachyonychia_Congenita_transcript.txt', 'Pachyonychia Congenita (Part 1)', [
        ('Inheritance', 'Autosomal dominant'),
        ('Shared features', 'Onychodystrophy, pincer nails, plantar keratoderma + plantar pain, hyperhidrosis'),
        ('Type 1 genes', 'KRT6A, KRT16'),
        ('Type 1 features', 'Benign oral leukokeratosis; follicular hyperkeratosis'),
        ('Type 2 genes', 'KRT6B, KRT17'),
        ('Type 2 features', 'Natal teeth, steatocystoma multiplex'),
        ('Treatment', 'EGFR inhibitors (erlotinib) for plantar keratoderma'),
    ]),
    ('Phenylketonuria_transcript.txt', 'Phenylketonuria (PKU)', [
        ('Inheritance', 'Autosomal recessive'),
        ('Gene/Enzyme', 'PAH - phenylalanine hydroxylase'),
        ('Skin', 'Diffuse hypopigmentation, blonde hair, blue eyes; eczematous dermatitis; photosensitivity; sclerodermatous changes'),
        ('Classic', 'Mousey odour of urine'),
        ('Other', 'Short stature, microcephaly, developmental delay'),
        ('Treatment', 'Low phenylalanine diet; avoid aspartame'),
    ]),
    ('Prolidase_Deficiency_transcript.txt', 'Prolidase Deficiency', [
        ('Gene', 'PEPD - encodes prolidase'),
        ('Function', 'Ubiquitous enzyme in protein catabolism'),
        ('Skin', 'Severe progressive ulcerations of lower extremities'),
        ('Facies', 'Hypertelorism (wide-set eyes)'),
    ]),
    # -------- Cutaneous Lymphomas --------
    ('Cutaneous_B_Cell_Lymphomas_transcript.txt', 'Cutaneous B-Cell Lymphomas (PCBCL)', [
        ('General markers', 'CD20+, CD79a+'),
        ('BCL Score', 'BCL2 + BCL6: higher = worse prognosis'),
        ('PCFCL (Crosti)', 'Scalp/forehead/back; BCL6+, BCL2-; lacks t(14;18); excellent prognosis'),
        ('Marginal Zone Lymphoma', 'Upper extremities/trunk; BCL2+; monocytoid B cells; Dutcher bodies; excellent prognosis'),
        ('DLBCL Leg Type', 'Elderly females, leg; BCL2+, BCL6+, MUM1+; 5-yr survival ~50%'),
        ('Intravascular Large B-cell', 'Trunk/thighs; CNS involvement; CD20+ atypical B cells within vessels'),
    ]),
    ('Cutaneous_Lymphomas_Variants_transcript.txt', 'Adult T-Cell Leukaemia/Lymphoma (ATLL)', [
        ('Cause', 'HTLV-1'),
        ('Geography', 'Caribbean, Japan, Central Africa'),
        ('Clinical', 'Hypercalcaemia, lymphadenopathy'),
        ('Histology', 'MF-like + floret/cloverleaf-shaped malignant T cells'),
        ('Markers', 'PD1+, CD25+'),
    ]),
    ('Cutaneous_Lymphomas_Variants_2_transcript.txt', 'Lymphomatoid Papulosis (LyP)', [
        ('Demographics', 'Adults in 40s'),
        ('Clinical', 'Recurrent ulcerative red-brown papulonodules on trunk/extremities'),
        ('Course', 'Self-resolves over 1-2 months → atrophic varioliform scars'),
        ('Cancer association', '~25% develop lymphoma (MF > Hodgkin > ALCL)'),
        ('Markers', 'CD30+ (typically)'),
        ('Treatment', 'First-line: methotrexate (90% respond); refractory: brentuximab'),
    ]),
    ('Cutaneous_Lymphomas_Variants_3_transcript.txt', 'LyP Variants (Types A-E + 6p25.3)', [
        ('Type A', 'Wedge-shaped infiltrate; large CD30+ Reed-Sternberg-like cells; eos/neutrophils'),
        ('Type B', 'MF-like patch/plaque; cerebriform CD4+ cells; CD30 NEGATIVE (only type)'),
        ('Type C', 'Resembles ALCL with pandermal infiltrate'),
        ('Type D', 'CD8+ with epidermotropism'),
        ('Type E', 'Angioinvasive/angiodestructive; CD8+, CD30+'),
        ('6p25.3 variant', 'DUSP22-IRF4 rearrangement; biphasic histology; CD30+'),
    ]),
    ('Cutaneous_Lymphomas_Variants_4_transcript.txt', 'Primary Cutaneous Anaplastic Large Cell Lymphoma (pcALCL)', [
        ('Clinical', 'Solitary persistent lesion (vs LyP - comes and goes)'),
        ('Diagnosis', 'Must exclude systemic ALCL'),
        ('Prognosis', '~90% 5-yr survival'),
        ('Poor prognostic factors', 'Older age, generalised skin involvement'),
        ('Markers', 'CD30+ (sheets of CD30+ cells)'),
        ('Treatment', 'Surgery OR radiotherapy; refractory: brentuximab'),
    ]),
    ('Cutaneous_Lymphomas_Variants_5_transcript.txt', 'Subcutaneous Panniculitis-like T-Cell Lymphoma (SPTCL)', [
        ('Clinical', 'Generalised SC nodules legs/trunk; regression → lipoatrophy'),
        ('Markers', 'CD8+, TIA-1+, Granzyme B+, αβ T-cell phenotype'),
        ('Histology', 'Bean-bag cells; rimming of adipocytes'),
        ('Prognosis', '80% 5-yr survival'),
        ('Complication', '~15% develop HLH (HAVCR2 mutation); high mortality'),
        ('Treatment', 'Steroids, methotrexate, ciclosporin'),
    ]),
    ('Cutaneous_Lymphomas_Variants_6_transcript.txt', 'Primary Cutaneous γδ T-Cell Lymphoma', [
        ('Course', 'Aggressive, rapidly fatal'),
        ('Markers', 'CD4-, CD8-, CD56+, TIA-1+, Granzyme B+, Perforin+'),
        ('Histology', 'Lymphoid infiltrate w/ epidermotropism; lichenoid interface; fat rimming'),
    ]),
    ('Cutaneous_Lymphomas_Variants_7_transcript.txt', 'Extranodal NK/T-Cell Lymphoma, Nasal Type', [
        ('Phenotype', 'NK cell'),
        ('Clinical', 'Sudden ulcerated tumours of nasal region'),
        ('Aetiology', 'EBV strongly associated'),
        ('Histology', 'Prominent vascular destruction'),
        ('Markers', 'CD56+, CD3+, CD2+'),
        ('Prognosis', 'Usually fatal'),
    ]),
    ('Cutaneous_Lymphomas_Variants_8_transcript.txt', 'Primary Cutaneous CD4+ Small/Medium Pleomorphic T-Cell LPD', [
        ('Clinical', 'Solitary plaque/nodule on head or neck'),
        ('Prognosis', 'Excellent'),
        ('Markers', 'CD4+; MF-like immunophenotype'),
        ('Histology', 'Pleomorphic small-medium T cells'),
    ]),
    # -------- Cancer Predisposition Syndromes --------
    ('Carney_Complex_transcript.txt', 'Carney Complex (NAME / LAMB Syndrome)', [
        ('Inheritance', 'Autosomal dominant'),
        ('Gene', 'PRKAR1A (~50%)'),
        ('LAMB', 'Lentigines, Atrial Myxomas, Blue nevi'),
        ('Lentigines', 'Periorificial (~77%)'),
        ('Blue nevi', 'Epithelioid blue nevi - 2 melanocyte populations'),
        ('Cardiac', 'Atrial myxomas (life-threatening)'),
        ('Cutaneous myxomas', 'Eyelids, ear, nipple, breast'),
        ('Cancer risks', 'Breast, thyroid'),
        ('Endocrine', 'Primary pigmented adrenocortical disease (Cushing-like)'),
        ('Testicular', 'Sertoli cell tumours (33%)'),
        ('Other', 'Psammomatous melanotic schwannomas'),
    ]),
    ('CHH_and_CHILD_transcript.txt', 'Conradi-Hünermann-Happle (CHH) and CHILD Syndromes', [
        ('CHH gene', 'EBP (Emopamil-binding protein)'),
        ('CHILD gene', 'NSDHL (Now Sit Down Here Little child)'),
        ('Inheritance', 'X-linked dominant; lethal in males (both)'),
        ('CHH ichthyosis', 'Generalised along Blaschko lines'),
        ('CHILD ichthyosis', 'Hemilateral with sharp midline demarcation'),
        ('CHH features', 'Unilateral cataracts, ice-pick scarring, follicular atrophoderma'),
        ('CHILD features', 'Verrucous hyperkeratosis, hemidysplasia, hypoplastic limbs'),
        ('Shared', 'Asymmetric limb shortening, chondrodysplasia punctata, ichthyosiform erythroderma'),
    ]),
    ('Brooke_Spiegler_Syndrome_transcript.txt', 'Brooke-Spiegler Syndrome', [
        ('Gene', 'CYLD - tumour suppressor'),
        ('Function', 'Lysine deubiquitinase; inhibits NF-κB and NEMO'),
        ('Tumour spectrum', 'Cylindromas, Spiradenomas, Trichoepitheliomas, BCCs'),
        ('Cylindroma', '"Jigsaw / giraffe skin" pattern'),
        ('Trichoepithelioma', '"Swiss cheese" cribriform with mucin'),
        ('Spiradenoma', 'Big blue balls with slit-like holes'),
        ('BCC', 'Basaloid islands + mucinous stroma + peripheral palisading + retraction artefact'),
    ]),
    ('Ataxia_Telangiectasia_transcript.txt', 'Ataxia Telangiectasia (Louis-Bar Syndrome)', [
        ('Inheritance', 'Autosomal recessive'),
        ('Gene', 'ATM - serine-threonine protein kinase'),
        ('First sign', 'Cerebellar truncal ataxia when child begins to walk'),
        ('Telangiectasias', 'Onset ~age 6; periocular → perioral → auricular'),
        ('Skin/hair', 'Progeric changes (~90%)'),
        ('Pulmonary', 'Recurrent infections, bronchiectasis (#1 cause death, mean age 20)'),
        ('Cancer risks', 'Leukaemia, lymphoma (chr 7 & 14 rearrangements)'),
        ('Female carriers', 'Increased breast cancer risk'),
        ('Immunology', '↓ IgE, IgA, IgG; lymphopenia'),
    ]),
    ('Dyskeratosis_Congenita_transcript.txt', 'Dyskeratosis Congenita (DKC)', [
        ('Inheritance', 'X-linked recessive (DKC1) most common; AD (TERT, TERC)'),
        ('Mechanism', 'Defective telomere maintenance'),
        ('Classic Triad', 'Reticulated hyperpigmentation + Nail dystrophy + Leukoplakia'),
        ('Mucosa', 'Pre-malignant leukoplakia (vs benign in PC)'),
        ('Cancer risk', '↑ cutaneous SCC (annual screening); SCC oral, oesophageal, anus'),
        ('Nails', 'Pterygium → complete nail loss'),
        ('Bone marrow', 'Failure in 50-90% (major cause of mortality)'),
        ('Eyes', 'Excessive lacrimation (lacrimal duct atresia)'),
        ('Other', 'Liver cirrhosis, pulmonary fibrosis, premature greying'),
    ]),
    ('Birt_Hogg_Dube_Syndrome_transcript.txt', 'Birt-Hogg-Dubé Syndrome', [
        ('Inheritance', 'Autosomal dominant'),
        ('Gene', 'BHD/FLCN (folliculin) - tumour suppressor in mTOR pathway'),
        ('FAT HOG', 'Fibrofolliculomas, Acrochordons, Trichodiscomas'),
        ('Pulmonary', 'Cysts (~90%) → spontaneous pneumothorax'),
        ('Renal', 'Chromophobe renal cell carcinoma'),
        ('Thyroid', 'Medullary thyroid carcinoma; ~65% have nodules'),
    ]),
    ('Gorlin_Syndrome_transcript.txt', 'Gorlin Syndrome (Basal Cell Nevus Syndrome)', [
        ('Inheritance', 'Autosomal dominant; 50% new mutations'),
        ('Gene', 'PTCH1 - tumour suppressor in Sonic Hedgehog pathway'),
        ('Diagnostic criteria', '1 major + molecular OR 2 major OR 1 major + 2 minor'),
        ('Major criteria', 'Multiple BCCs, medulloblastoma (<3y), falx cerebri calcification, palmar/plantar pits, odontogenic keratocysts'),
        ('Minor criteria', 'Macrocephaly + frontal bossing, polydactyly, ocular anomalies, cardiac/ovarian fibromas, bifid/fused ribs'),
        ('Treatment', 'Surgical excision; vismodegib (smoothened inhibitor)'),
    ]),
    # -------- Hyperpigmentation --------
    ('Inherited_Disorders_Hyperpigmentation_transcript.txt', 'Inherited Disorders of Hyperpigmentation', [
        ('Dowling-Degos Disease', 'AD; KRT5; adult intertriginous reticulate pigmentation; antler-like pigment on histology'),
        ('McCune-Albright', 'GNAS1 mosaic; café-au-lait (coast of Maine) + endocrine + polyostotic fibrous dysplasia'),
        ('Reticulate Acropigmentation Kitamura', 'AD; ADAM10; Japanese; depressed hyperpigmented macules dorsal hands/feet; palmar pits'),
    ]),
    ('Waardenburg_and_Piebaldism_transcript.txt', 'Waardenburg Syndrome and Piebaldism', [
        ('Waardenburg Type 1', 'PAX3 (AD); poliosis, hypertelorism with synophrys, deafness, dystopia canthorum'),
        ('Waardenburg Type 2', 'MITF (AD); same as 1 but lacks dystopia canthorum'),
        ('Tietz Syndrome', 'AR MITF; diffuse pigment dilution, deafness'),
        ('Waardenburg Type 3', 'PAX3 (AD); type 1 + limb defects'),
        ('Waardenburg Type 4 (Shah)', 'SOX10/EDN3/EDNRB; defining: Hirschsprung disease'),
        ('Piebaldism', 'c-KIT; white forelock, leukoderma; NO deafness'),
    ]),
    ('Tay_Syndrome_transcript.txt', 'Tay Syndrome (TTD / PIBIDS)', [
        ('Inheritance', 'Autosomal recessive'),
        ('Genes', 'ERCC2 (XPD), ERCC3 (XPB) - same family as Cockayne'),
        ('Hair', 'Brittle, "tiger tail" banding on polarised microscopy'),
        ('PIBIDS', 'Photosensitivity, Ichthyosis, Brittle hair, Intellectual impairment, Decreased fertility, Short stature'),
        ('Key', 'Photosensitivity does NOT increase malignancy risk'),
        ('Other', 'Hypogammaglobulinaemia, PPK, KP, cataracts'),
    ]),
    ('NFJ_and_DPR_transcript.txt', 'NFJ Syndrome and Dermatopathia Pigmentosa Reticularis (DPR)', [
        ('Both genes', 'KRT14 (AD) - both ectodermal dysplasias'),
        ('NFJ pigmentation', 'Brown-grey reticulated, classically abdomen'),
        ('DPR pigmentation', 'Reticulate hyperpigmentation upper torso, proximal extremities'),
        ('NFJ course', 'Improves with age'),
        ('DPR course', 'Persistent'),
        ('NFJ distinctive', 'Dental anomalies (early loss of teeth)'),
        ('DPR distinctive', 'Diffuse non-scarring alopecia'),
        ('Shared', 'PPK, hypohidrosis, onychodystrophy, absent dermatoglyphics, poikiloderma'),
    ]),
    ('Sturge_Weber_Syndrome_transcript.txt', 'Sturge-Weber Syndrome', [
        ('Inheritance', 'Sporadic; somatic mosaic GNAQ → MAPK pathway'),
        ('Cutaneous', 'Port-wine stain in V1 (most common)'),
        ('CNS', 'Ipsilateral leptomeningeal vascular anomalies'),
        ('Neuro', 'Seizures (most common, <1y), stroke-like episodes, developmental delay'),
        ('Imaging', 'Tram-track cortical calcifications'),
        ('Ocular', 'Glaucoma (30-60%) - eyelid/forehead PWS involvement'),
        ('Worse prognosis', 'Bilateral V1 PWS'),
        ('Treatment', 'Neuroimaging ≥1y; prophylactic aspirin'),
    ]),
    ('Hailey_Hailey_transcript.txt', 'Hailey-Hailey Disease (Benign Familial Pemphigus)', [
        ('Inheritance', 'Autosomal dominant'),
        ('Gene', 'ATP2C1 - calcium ATPase in Golgi'),
        ('Onset', 'Teens/20s'),
        ('Distribution', 'Intertriginous (axilla, neck, inframammary, waistline)'),
        ('Lesions', 'Vesicles/bullae early → macerated eroded plaques'),
        ('Histology', 'Full-thickness "dilapidated brick wall" acantholysis; less dyskeratosis than Darier; NO follicular involvement'),
        ('Complication', 'Kaposi varicelliform eruption (HSV)'),
        ('Treatment', 'CO2 laser ablation (best); gentamicin; retinoids NOT effective'),
    ]),
    # -------- Pachyonychia 2 --------
    ('Pachyonychia_Congenita_2_transcript.txt', 'Pachyonychia Congenita (Part 2)', [
        ('Inheritance', 'Autosomal dominant - keratin defects'),
        ('Type 1 (Jadassohn-Lewandowsky)', 'KRT6A + KRT16'),
        ('Type 1 features', 'Benign oral leukokeratosis (NOT pre-malignant); follicular hyperkeratosis'),
        ('Type 2 (Jackson-Lawler)', 'KRT6B + KRT17'),
        ('Type 2 features', 'Natal teeth; steatocystoma multiplex (shark tooth cuticle)'),
        ('Shared', 'Painful PPK, hyperhidrosis, corneal dystrophy'),
        ('Key nail', 'Pincer nails (transverse curvature)'),
    ]),
    ('Muir_Torre_Syndrome_transcript.txt', 'Muir-Torre Syndrome', [
        ('Inheritance', 'Autosomal dominant - subtype of Lynch syndrome (HNPCC)'),
        ('Genes', 'DNA mismatch repair: MSH2 (~90%), MLH1, MSH6'),
        ('Cutaneous tumours', 'Sebaceous neoplasms; multiple keratoacanthomas'),
        ('Visceral malignancies', 'Colorectal (most common, <50, proximal) > GU (endometrial, ovarian, urinary) > lymphoma'),
        ('Surveillance', 'Colonoscopy q1-2y from 20-25; yearly UA cytology from 25-30; pelvic exam + TVUSS for women'),
    ]),
    ('HED_transcript.txt', 'Hypohidrotic Ectodermal Dysplasia (HED) - Christ-Siemens-Touraine', [
        ('Inheritance', 'X-linked recessive (most common) - EDA gene'),
        ('Other forms', 'AD/AR EDARADD'),
        ('Cardinal features', 'Eczema, Dyshidrosis, Dysplasia, Retarded dentition'),
        ('Skin', 'Hypohidrosis; no eccrine glands; no dermatoglyphics'),
        ('Hair', 'Hypotrichosis - sparse, fine'),
        ('Teeth', 'Conical/peg-shaped'),
        ('Facies', 'Frontal bossing, periorbital hyperpigmentation, large nostrils'),
        ('Major risk', 'Hyperthermia'),
        ('HED + Immunodef', 'IKBKG (NEMO); susceptible to pyogenic & atypical mycobacterial infections'),
    ]),
    ('Darier_Disease_transcript.txt', "Darier's Disease", [
        ('Inheritance', 'Autosomal dominant'),
        ('Gene', 'ATP2A2 → SERCA2 (Smooth ER Calcium ATPase)'),
        ('vs Hailey-Hailey', 'ATP2C1 (Golgi calcium ATPase)'),
        ('Onset', 'Usually <20 years'),
        ('Skin', 'Verrucous crusted papules in seborrhoeic distribution'),
        ('Histology', 'Acantholysis + dyskeratosis (corps ronds + grains)'),
        ('Nails', 'V-shaped nicking; red & white longitudinal lines'),
        ('Mucosa', 'Cobblestone oral mucosa'),
        ('Hands', 'Acrokeratosis verruciformis of Hopf, palmar pits'),
        ('Complications', 'Kaposi varicelliform eruption; bipolar, depression, schizophrenia'),
        ('Avoid', 'Lithium (worsens disease)'),
        ('Treatment', 'Topical steroids → systemic retinoids (90% effective)'),
    ]),
    ('Epidermolysis_Bullosa_transcript.txt', 'Epidermolysis Bullosa (EB)', [
        ('EBS Generalised Intermediate (Koebner)', 'KRT5/KRT14'),
        ('EBS Mottled pigmentation', 'KRT5 only'),
        ('EBS Localised (Weber-Cockayne)', 'KRT5/KRT14; hands & feet'),
        ('EBS with Muscular Dystrophy', 'PLECTIN; only AR EBS form'),
        ('EBS Generalised Severe (Dowling-Meara)', 'KRT5/KRT14; herpetiform; PPK'),
        ('JEB Severe (Herlitz)', 'Laminin 332; hoarse cry, perioral granulation, infant death'),
        ('JEB Intermediate', 'Laminin 332 + Collagen 17'),
        ('JEB with Pyloric Atresia', 'α6β4 integrin'),
        ('DEB Autosomal Dominant', 'Collagen 7; albopapuloid papules'),
        ('DEB Autosomal Recessive', 'Collagen 7; severe; SCC > renal failure; mitten deformity'),
        ('Kindler Syndrome', 'FERMT1; acral blistering, poikiloderma, photosensitivity'),
    ]),
    ('EDV_transcript.txt', 'Epidermodysplasia Verruciformis (EDV)', [
        ('Inheritance', 'Autosomal recessive'),
        ('Susceptibility', 'HPV genus β (esp HPV-5, HPV-8)'),
        ('Genes', 'EVER1 (TMC6) and EVER2 (TMC8)'),
        ('Skin findings', 'Flat-wart-like papules; tinea versicolor-like macules in seborrheic distribution; "Tree-man syndrome"'),
        ('Cancer risk', '50% develop SCC by age 30 (UV co-carcinogen with HPV)'),
        ('Acquired EDV-like', 'Immunosuppressed (AIDS, transplant, lymphoma)'),
        ('Histology', 'Vacuolated keratinocytes with large blue-grey nuclei'),
    ]),
    ('Clouston_Hidrotic_Ectodermal_Dysplasia_transcript.txt', "Clouston's Hidrotic Ectodermal Dysplasia", [
        ('Inheritance', 'Autosomal dominant'),
        ('Gene', 'GJB6 → Connexin 30'),
        ('Key vs CST', 'Normal sweating (hidrotic); can regulate temperature'),
        ('Clinical Triad', 'Onychodystrophy + PPK + Hypotrichosis'),
        ('Radiographic', 'Tufted distal phalanges (broccoli-like)'),
        ('Eyes', 'Conjunctivitis, strabismus, cataracts'),
    ]),
    ('Cowden_Syndrome_transcript.txt', "Cowden's Syndrome", [
        ('Inheritance', 'Autosomal dominant'),
        ('Gene', 'PTEN (tumour suppressor)'),
        ('Related disorders', 'Bannayan-Riley-Ruvalcaba (BRR)'),
        ('Skin', 'Trichilemmomas, sclerotic fibromas (almost exclusive), acrochordons, lipomas, PPK'),
        ('Pathognomonic', 'Lhermitte-Duclos disease (cerebellar dysplastic gangliocytoma)'),
        ('Mucosa', 'Cobblestone lips/oral mucosa'),
        ('Eyes', 'Myopia, angioid streaks (PEPSI LIGHT)'),
        ('Genitalia', 'Penile lentigines'),
        ('Other', 'Macrocephaly, hydrocephalus'),
        ('Cancer risks', 'Breast, Endometrial, Thyroid; GI hamartomatous polyps'),
    ]),
    # -------- Granulomatous & Lymphoid --------
    ('Granulomatous_Disorders_transcript.txt', 'Necrobiotic Xanthogranuloma (NXG)', [
        ('Clinical', 'Firm yellow xanthomatous plaques/nodules - classically periorbital'),
        ('Course', 'Often ulcerate → scarring'),
        ('Ophthalmic', '~50% of patients'),
        ('Cardiac', 'Endocardial involvement (majority)'),
        ('GI', 'Hepatosplenomegaly'),
        ('Haematology', 'IgG kappa monoclonal gammopathy (strong association)'),
        ('Histology', 'Palisading xanthogranulomas with necrobiotic collagen'),
        ('Other histology', 'Cholesterol clefts; bizarre foreign body giant cells; horseshoe arrangement of 25-50 nuclei'),
        ('Treatment', 'Treat underlying paraproteinaemia/malignancy'),
    ]),
    ('Granulomatous_Disorders_2_transcript.txt', 'Granulomatous Disorders Overview', [
        ('Sarcoidosis', 'Common in young African-American females'),
        ('Granuloma Annulare', 'Chronic ring-shaped lesions'),
        ('Necrobiosis Lipoidica', 'Diabetes-associated; yellow-brown patches lower legs'),
        ('Annular Elastolytic Giant Cell Granuloma', 'Annular sun-exposed lesions'),
        ('Cutaneous Crohn Disease', 'Skin manifestations of Crohn'),
        ('Rheumatoid Nodule', 'Firm SC nodules near joints; RA association'),
        ('Interstitial Granulomatous Dermatitis', 'Linear cords ("rope sign")'),
        ('PNGD', 'Umbilicated papules elbows, extensors, fingers'),
    ]),
    ('Granulomatous_Disorders_3_transcript.txt', "Annular Elastolytic Giant Cell Granuloma (Actinic Granuloma of O'Brien)", [
        ('Variant of', 'Granuloma Annulare'),
        ('Distribution', 'Sun-exposed (chronically photodamaged) skin'),
        ('Clinical', 'Pink papules merging into annular plaques'),
        ('Histology', 'Interstitial granulomatous infiltrate; ↑ multinucleated foreign body giant cells; elastophagocytosis'),
        ('Mucin', 'Absent (key differentiator from GA)'),
        ('Verhoeff stain', 'Negative for elastic fibres in lesional skin'),
        ('Course', 'Persistent; poor response to GA treatments'),
    ]),
    ('Leukemia_Cutis_transcript.txt', 'Leukaemia Cutis', [
        ('Most common', 'AML (especially myelomonocytic and monocytic)'),
        ('Histology', 'Grenz zone (papillary dermis spared); dermal myeloid blasts'),
        ('Cell pattern', 'Indian filing'),
        ('Stains', 'CD117 (c-KIT) +, MPO +'),
        ('Chloromas', 'Green nodules with AML; green from MPO activity'),
    ]),
    ('Pseudolymphoma_transcript.txt', 'Pseudolymphoma - Histology & Clinical', [
        ('B-cell origin', 'Polyclonal B cells'),
        ('Dermal infiltrate', 'Mixed: lymphocytes + eosinophils + plasma cells'),
        ('Grenz zone', 'Present'),
        ('Architecture', 'Large blue nodules/follicles in dermis & superficial fat'),
        ('Germinal centres', 'Pale-appearing'),
        ('Tingible body macrophages', 'Present (key feature)'),
        ('Mantle zone', 'Normal-appearing lymphocytes (key vs true B-cell lymphoma)'),
        ('IHC', 'Kappa + Gamma'),
        ('IgH gene rearrangement', 'Polyclonal (no clonality)'),
        ('Clinical', 'Red-to-plum firm plaques/nodules; usually above waist'),
    ]),
    ('Pseudolymphoma_2_transcript.txt', 'Pseudolymphoma - Causes', [
        ('Anticonvulsants', 'Phenytoin, Phenobarbital, Carbamazepine, Lamotrigine'),
        ('Neuroleptics', 'Promethazine, Chlorpromazine'),
        ('ARBs', 'Angiotensin Receptor Blockers'),
        ('Imatinib', 'Tyrosine kinase inhibitor'),
        ('Arthropod bites', 'Tick bites and other infestations'),
        ('Borrelia burgdorferi', 'Lyme disease'),
        ('Tattoo reactions', '-'),
        ('Post-zoster', 'Following Herpes Zoster infection'),
    ]),
    # -------- Papulosquamous / Inflammatory --------
    ('Pityriasis_Rosea_transcript.txt', 'Pityriasis Rosea', [
        ('Family', 'Papulosquamous disorders (with psoriasis, seb derm, PRP, granular parakeratosis)'),
        ('Aetiology', 'Linked to viral infection - HHV-7 (most common) and HHV-6'),
        ('Drug-induced PR triggers', 'ACE inhibitors (most common), gold, beta-blockers, isotretinoin, metronidazole, barbiturates'),
        ('Initial lesion', 'Herald patch (single larger plaque)'),
        ('Distribution', 'Christmas tree pattern along Langer lines, favours trunk'),
        ('Population', 'More extensive in African-American children'),
        ('Key feature', 'Trailing collarette of scale (also seen in EAC)'),
        ('Course', 'Self-limited; resolves over 6-8 weeks'),
        ('Treatment', 'Symptomatic; erythromycin if hastening clearance is needed'),
    ]),
    ('DIHS_transcript.txt', 'Drug Induced Hypersensitivity Syndrome (DIHS / DRESS)', [
        ('Mortality', '~10%'),
        ('Onset', '2-6 weeks after starting culprit drug'),
        ('Fever', '85% of patients'),
        ('Skin', 'Morbilliform eruption (75%); facial edema (classic clue)'),
        ('Multiorgan', 'Liver (most common, primary cause of death) > kidney'),
        ('Labs', 'Peripheral eosinophilia >1500; atypical lymphocytosis'),
        ('Late sequelae', 'Thyroiditis or Graves disease'),
        ('Risk factors', 'Inability to detoxify arene-oxide metabolites; slow acetylators'),
        ('Common culprits', 'Anticonvulsants, allopurinol, sulfonamides, dapsone, antibiotics'),
        ('Treatment', 'Drug withdrawal; systemic corticosteroids if severe'),
    ]),
    ('Figurate_Erythemas_transcript.txt', 'Figurate Erythemas', [
        ('Erythema Annulare Centrifugum (EAC)', 'Adults; trailing scale; associated with lymphoproliferative disorders, dermatophyte infections, pregnancy, medications'),
        ('EAC histology', 'Coat-sleeve perivascular lymphocytic infiltrate'),
        ('Erythema Marginatum', 'Sequela of untreated Group A strep pharyngitis (rheumatic fever)'),
        ('JONES criteria', 'Joints, heart [O], Nodules, Erythema marginatum, Sydenham chorea'),
        ('Erythema Migrans', 'Lyme disease (Borrelia burgdorferi); Ixodes tick (>24h attachment)'),
        ('EM treatment', 'Doxycycline (>8y, non-pregnant); amoxicillin (pregnant + <8y)'),
        ('Erythema Gyratum Repens', 'Paraneoplastic - strongly associated with lung adenocarcinoma; classic wood-grain pattern'),
    ]),
    ('Lichenoid_Histology_transcript.txt', 'Lichenoid Histology', [
        ('Shared features', 'Band of lymphocytes at DEJ; ragged BMZ; vacuolar interface; Civatte bodies; pigment incontinence'),
        ('Lichen Planus (LP)', 'Saw-toothed rete ridges; NO eosinophils; NO parakeratosis (compact orthokeratosis)'),
        ('Drug-Induced Lichenoid', 'LP-like + may have eosinophils + parakeratosis (permissible)'),
        ('Hypertrophic Lupus', 'Superficial + deep infiltrate; mucin in dermis; follicular plugging'),
        ('Lupus DIF', 'Full house (IgG, IgA, IgM, C3 at BMZ)'),
        ('Distinguishing features', 'Follicular plugging, eosinophils, parakeratosis'),
    ]),
    ('Neonatal_Rashes_transcript.txt', 'Neonatal Rashes', [
        ('TNPM (Transient Neonatal Pustular Melanosis)', 'At birth; black/Hispanic full-term babies; 3 stages: pustules → collarette of scale → hyperpigmented macules'),
        ('TNPM sites', 'Forehead, ears, back, fingers, toes; pustules contain neutrophils'),
        ('SCPS DDx (CAT TIPS)', 'Candida, Acropustulosis of infancy, TNPM, Sneddon-Wilkinson, Impetigo, Pustular psoriasis, Staph scalded skin'),
        ('ETN (Erythema Toxicum Neonatorum)', 'Onset in first 2 days; >2.5kg full-term; spares palms/soles; macules, papules, pustules, wheals'),
        ('ETN cytology', 'Eosinophils on Wright stain'),
        ('Miliaria Crystallina', 'Onset ~7 days; sweat duct occlusion at stratum corneum'),
        ('Miliaria distribution', 'Pinpoint clear "crystal-like" vesicles; upper extremities, head'),
        ('Miliaria risk factor', 'Overheating/fever'),
    ]),
    ('PLEVA_transcript.txt', 'PLEVA / PLC', [
        ('Full name', 'Pityriasis Lichenoides et Varioliformis Acuta (PLEVA) / Chronica (PLC)'),
        ('Histology - Low power', 'Lichenoid reaction; wedge-shaped lymphocytic infiltrate (heaviest at centre)'),
        ('Cell type', 'Cytotoxic lymphocytes; often monoclonal'),
        ('Key features', 'Vacuolar interface; spongiosis with "lymph in every hole"'),
        ('Other features', 'RBC extravasation, ulceration, parakeratosis'),
        ('DDx with "lymph in every hole"', 'PLEVA/PLC and other lichenoid conditions'),
        ('Treatment', 'Erythromycin, tetracycline, phototherapy'),
    ]),
    ('PRP_transcript.txt', 'Pityriasis Rubra Pilaris (PRP)', [
        ('Aka', 'Devergie disease'),
        ('Classic features', 'Wheat-bran-like scale; perifollicular red papules (nutmeg grater); salmon patches with islands of sparing; thick yellow-orange waxy PPK'),
        ('Nails', 'Thick hyperkeratotic, yellow-brown discolouration; NO nail pits (vs psoriasis)'),
        ('Photosensitivity', 'Some patients flare with phototherapy - photo-test before NB-UVB'),
        ('Type 1', 'Classic adult; can cause erythroderma; good prognosis (resolves <3y)'),
        ('Type 2', 'Atypical adult; ichthyosiform leg lesions; coarse PPK; alopecia; chronic'),
        ('Type 3', 'Like Type 1 but in children; good prognosis'),
        ('Type 4', 'Most common pediatric; follicular papules on elbows/knees'),
        ('Type 5', 'Atypical juvenile; sclerodermoid hands/feet; CARD14 mutation; prominent facial involvement'),
        ('Type 6', 'HIV-associated; PRP + hidradenitis suppurativa + acne conglobata + lichen spinulosus'),
    ]),
    ('Psoriasis_transcript.txt', 'Psoriasis', [
        ('Susceptibility loci', 'PSORS1-9; PSORS1 most important'),
        ('Strongest HLA association', 'HLA-Cw6 (15-fold increased risk)'),
        ('HLA-B17', 'Erythrodermic & guttate psoriasis'),
        ('HLA-B13', 'Guttate psoriasis'),
        ('HLA-B27', 'Sacroiliitis-associated, psoriatic arthritis, pustular psoriasis'),
        ('Köbner phenomenon', 'Lesions develop in trauma areas; lag 2-6 weeks'),
        ('Pregnancy trigger', 'Impetigo herpetiformis'),
        ('Rapid-onset triggers (<4w)', 'NSAIDs, terbinafine'),
        ('Intermediate triggers (4-12w)', 'Hydroxychloroquine, ACE inhibitors'),
        ('Long-latency triggers (>12w)', 'Beta-blockers, lithium'),
        ('Paradoxical', 'TNF-α inhibitors can both treat AND trigger psoriasis (especially palmoplantar pustulosis)'),
        ('Increased cytokines', 'IL-22, IL-23, IL-17'),
        ('Pathway', 'IL-23 (DCs) → Th17 → IL-17 + IL-22 → keratinocyte proliferation'),
        ('Decreased cytokine', 'IL-10 (anti-inflammatory)'),
    ]),
    ('SJS_TEN_transcript.txt', 'Stevens-Johnson Syndrome / Toxic Epidermal Necrolysis (SJS/TEN)', [
        ('Onset', '1-2 weeks after culprit drug (8 weeks for anticonvulsants)'),
        ('Mucosal involvement', '92-100% of cases (oral, ocular → blindness, genital → strictures)'),
        ('Consults', 'Ophthalmology + Urology + Dermatology'),
        ('Skin', 'Atypical targetoid lesions; painful (not pruritic); spares distal extremities'),
        ('Signs', 'Positive Nikolsky and Asboe-Hansen signs'),
        ('Common culprits', 'Allopurinol, cephalosporins, minocycline, TMP-SMX, NSAIDs'),
        ('HIV NNRTIs', 'Abacavir, efavirenz, nevirapine'),
        ('Aromatic anticonvulsants', 'Phenobarbital, phenytoin, lamotrigine, carbamazepine; valproate is safe alternative'),
        ('Pathogenesis', 'Drug binds MHC1 → CD8+ T cells → keratinocyte apoptosis'),
        ('Mediators', 'Granzyme B (major); Fas ligand (CD95L) → caspase cascade'),
        ('SCORTEN (TAME BUGS)', 'Tachycardia, Age, Malignancy, Epidermal loss >10%, Bicarbonate <20, Urea, Glucose >210'),
        ('Most important mortality factor', 'Serum bicarbonate <20'),
        ('Common cause of death', 'Secondary infection - Staph aureus, Pseudomonas'),
    ]),
    # -------- Vasculitis & Neutrophilic Dermatoses --------
    ('Pyoderma_Gangrenosum_transcript.txt', 'Pyoderma Gangrenosum (PG)', [
        ('Definition', 'Neutrophilic dermatosis - sterile inflammatory ulceration'),
        ('Pathergy', 'Lesions develop/worsen at sites of trauma'),
        ('Classic ulcer', 'Painful with violaceous/undermined borders'),
        ('Ulcerative form', 'Most common; classic painful ulcer; lower extremities'),
        ('Bullous form', 'Hemorrhagic bullae; associated with haematologic malignancies (AML)'),
        ('Pustular form', 'Pustules; associated with IBD'),
        ('Vegetative form', 'Verrucous/exophytic; less aggressive; minimal systemic associations'),
        ('Peristomal PG', 'Around ileostomy/colostomy sites'),
        ('Systemic associations', 'IBD (especially UC), arthritis, haematologic disorders, monoclonal gammopathy'),
        ('Histology', 'Dense neutrophilic infiltrate; non-specific; rule out infection/vasculitis'),
        ('Diagnosis', 'Diagnosis of exclusion'),
        ('Treatment', 'Topical/systemic corticosteroids, ciclosporin, biologics (TNF-α inhibitors, especially infliximab for IBD-associated)'),
        ('Important', 'Avoid surgical debridement (can worsen due to pathergy)'),
    ]),
    ('Behcet_Disease_transcript.txt', "Behcet's Disease", [
        ('Geographic distribution', 'Silk Road - Mediterranean, Middle East, East Asia'),
        ('HLA association', 'HLA-B51'),
        ('Pathergy', 'Positive (papule/pustule at site of needle prick after 48h)'),
        ('Major criterion', 'Recurrent oral ulcers (≥3x/year)'),
        ('Minor criteria', 'Genital ulcers, eye lesions (uveitis), skin lesions, positive pathergy test'),
        ('Oral ulcers', 'Aphthous - shallow/deep with white-yellow base; painful'),
        ('Genital ulcers', 'Scrotum (men), vulva (women); often scar'),
        ('Skin', 'Erythema nodosum-like lesions, papulopustular lesions, acneiform'),
        ('Eye', 'Posterior uveitis (retinal vasculitis) - leading cause of blindness'),
        ('Vascular', 'Both arterial and venous; thrombosis, aneurysms'),
        ('Neuro-Behcet', 'Brainstem involvement; mortality risk'),
        ('GI', 'Ileocecal ulcers (mimics Crohn)'),
        ('Treatment', 'Colchicine for mucocutaneous; corticosteroids; biologics (anti-TNF) for severe'),
    ]),
    ('Cryoglobulinemia_transcript.txt', 'Cryoglobulinemia', [
        ('Definition', 'Immunoglobulins that precipitate at cold temperatures'),
        ('Type I', 'Monoclonal Ig (IgM > IgG); associated with multiple myeloma, Waldenstrom macroglobulinemia, lymphoma'),
        ('Type I features', 'Cold-induced acrocyanosis, livedo, ulcers, retiform purpura'),
        ('Type II (Mixed)', 'Monoclonal IgM with rheumatoid factor activity + polyclonal IgG'),
        ('Type II association', 'Hepatitis C (>90% of cases)'),
        ('Type III (Mixed)', 'Polyclonal IgM + polyclonal IgG'),
        ('Type III association', 'Connective tissue diseases, infections'),
        ('Mixed cryoglobulinemia features', 'Palpable purpura, arthralgia, weakness (Meltzer triad), glomerulonephritis, peripheral neuropathy'),
        ('Histology', 'Leukocytoclastic vasculitis; intravascular hyaline thrombi (Type I)'),
        ('Treatment', 'Treat underlying cause; rituximab; plasmapheresis for severe disease'),
    ]),
    ('HSP_transcript.txt', 'Henoch-Schonlein Purpura (HSP / IgA Vasculitis)', [
        ('Aka', 'IgA Vasculitis'),
        ('Demographics', 'Most common vasculitis in children (4-11 years)'),
        ('Trigger', 'Often follows URI (Group A strep, viral infections)'),
        ('Classic tetrad', 'Palpable purpura + arthritis/arthralgia + abdominal pain + renal involvement'),
        ('Skin', 'Palpable purpura - lower extremities and buttocks (dependent areas)'),
        ('Joints', 'Arthralgia/arthritis - knees, ankles'),
        ('GI', 'Abdominal pain, intussusception (ileoileal), GI bleeding'),
        ('Renal', 'IgA nephropathy - haematuria, proteinuria; major long-term morbidity'),
        ('Histology', 'Leukocytoclastic vasculitis with IgA deposits on DIF'),
        ('Course', 'Usually self-limited (4-6 weeks)'),
        ('Treatment', 'Supportive; steroids for severe abdominal pain or renal involvement'),
        ('Adult HSP', 'More severe with higher rate of renal involvement'),
    ]),
    ('Kawasaki_Disease_transcript.txt', 'Kawasaki Disease', [
        ('Demographics', 'Children <5 years; Asian descent; M > F'),
        ('Aetiology', 'Unknown; possibly infectious trigger'),
        ('Diagnostic criteria', 'Fever ≥5 days + 4 of 5 features'),
        ('Criterion 1', 'Bilateral non-exudative conjunctivitis'),
        ('Criterion 2', 'Mucosal changes - strawberry tongue, cracked red lips, pharyngeal erythema'),
        ('Criterion 3', 'Cervical lymphadenopathy (>1.5cm, usually unilateral)'),
        ('Criterion 4', 'Polymorphous rash'),
        ('Criterion 5', 'Extremity changes - palmar/plantar erythema, edema, periungual desquamation (late)'),
        ('Cardiac complication', 'Coronary artery aneurysms (~25% if untreated)'),
        ('Labs', 'Elevated ESR/CRP, leukocytosis, thrombocytosis (later)'),
        ('Treatment', 'IVIG (2g/kg) within 10 days + high-dose aspirin'),
        ('Refractory cases', 'Repeat IVIG, infliximab, corticosteroids'),
        ('Long-term', 'Echocardiography to monitor coronary arteries'),
    ]),
    ('Urticarial_Vasculitis_transcript.txt', 'Urticarial Vasculitis', [
        ('Clinical', 'Urticarial wheals lasting >24 hours (vs typical urticaria <24h)'),
        ('Symptoms', 'Burning/painful (vs pruritic in typical urticaria)'),
        ('Resolution', 'Heals with hyperpigmentation/purpura'),
        ('Hypocomplementemic UV (HUV)', 'Low C1q, C3, C4; more severe; systemic involvement'),
        ('HUV systemic features', 'Arthralgia, GI, renal, pulmonary (COPD), uveitis'),
        ('Normocomplementemic UV', 'Normal complement; usually idiopathic; less severe'),
        ('Associations', 'SLE, Sjogren syndrome, hepatitis C, drug reactions, malignancy'),
        ('Histology', 'Leukocytoclastic vasculitis - perivascular neutrophils with karyorrhexis, fibrinoid necrosis'),
        ('Workup', 'CBC, complement levels (C1q, C3, C4), ANA, anti-C1q antibody, hepatitis serologies'),
        ('Treatment', 'NSAIDs, antihistamines, dapsone, colchicine, systemic steroids; rituximab/MMF for severe HUV'),
    ]),
    ('Small_Vessel_Vasculitis_transcript.txt', 'Small Vessel Vasculitis Overview', [
        ('Definition', 'Vasculitis affecting capillaries, post-capillary venules, arterioles'),
        ('Histology', 'Leukocytoclastic vasculitis - neutrophilic infiltrate, karyorrhexis, fibrinoid necrosis, RBC extravasation'),
        ('Cutaneous Small Vessel Vasculitis (CSVV)', 'Skin-limited; idiopathic or post-infectious/drug-induced'),
        ('Common drug triggers', 'Antibiotics (penicillins, cephalosporins, sulfonamides), NSAIDs, diuretics'),
        ('Common infectious triggers', 'Strep, viral hepatitis, HIV'),
        ('IgA Vasculitis (HSP)', 'IgA deposits; children; renal involvement'),
        ('Cryoglobulinemic Vasculitis', 'Type II/III cryoglobulins; Hep C association'),
        ('Urticarial Vasculitis', 'Wheals lasting >24h; can have systemic involvement'),
        ('ANCA-Associated Vasculitis (AAV)', 'GPA, MPA, EGPA - small to medium vessels'),
        ('Granulomatosis with Polyangiitis (GPA)', 'c-ANCA (PR3); upper/lower respiratory + renal; saddle nose'),
        ('Microscopic Polyangiitis (MPA)', 'p-ANCA (MPO); pulmonary-renal syndrome'),
        ('Eosinophilic GPA (EGPA)', 'p-ANCA; asthma, eosinophilia, sinusitis (Churg-Strauss)'),
        ('Workup', 'CBC, comprehensive metabolic panel, urinalysis, ESR/CRP, ANA, ANCA, complement, cryoglobulins, hepatitis serologies, RF'),
        ('Treatment', 'Address trigger; corticosteroids; immunosuppressants for systemic disease'),
    ]),
    ('Sweet_Syndrome_transcript.txt', "Sweet's Syndrome (Acute Febrile Neutrophilic Dermatosis)", [
        ('Demographics', 'Women aged 30-50; commonly post-infectious or paraneoplastic'),
        ('Clinical', 'Tender erythematous papules/plaques/nodules; pseudovesicular appearance'),
        ('Distribution', 'Face, neck, upper extremities (asymmetric)'),
        ('Systemic features', 'Fever, leukocytosis with neutrophilia, malaise, arthralgia'),
        ('Diagnostic criteria - Major', 'Abrupt onset of tender plaques/nodules + characteristic histology'),
        ('Diagnostic criteria - Minor', 'Fever >38°C, association (malignancy/IBD/pregnancy/preceding infection), excellent response to steroids, abnormal labs (3 of 4: ↑ESR, ↑CRP, ↑WBC, neutrophilia)'),
        ('Classical Sweet', 'Post-URI; women 30-60'),
        ('Malignancy-associated', '~20% - especially AML; can precede malignancy diagnosis'),
        ('Drug-induced', 'G-CSF (most common), all-trans retinoic acid (ATRA), TMP-SMX'),
        ('Histology', 'Dense dermal neutrophilic infiltrate; papillary dermal edema; NO vasculitis'),
        ('Histiocytoid variant', 'Immature myelomonocytic cells; associated with myelodysplastic syndrome'),
        ('Treatment', 'Systemic corticosteroids (rapid response); colchicine, dapsone, potassium iodide'),
        ('Prognosis', 'Excellent for classic; depends on underlying disease for paraneoplastic'),
    ]),
    # -------- Immunobullous Diseases --------
    ('Bullous_Pemphigoid_Basic_Science_transcript.txt', 'Bullous Pemphigoid - Basic Science & Basement Membrane', [
        ('Cell adhesion', 'Hemidesmosomes connect basal keratinocytes to basement membrane zone (BMZ)'),
        ('Target antigen 1', 'BPAG1 (BP230) - intracellular plakin family protein'),
        ('Target antigen 2', 'BPAG2 (BP180 / Collagen XVII) - transmembrane; NC16A domain (extracellular)'),
        ('NC16A domain', 'Most immunogenic epitope; primary target in BP'),
        ('Lamina lucida', 'Above lamina densa; integrin α6β4 + laminin 332 located here'),
        ('Lamina densa', 'Type IV collagen + nidogen + perlecan'),
        ('Sub-lamina densa', 'Anchoring fibrils of Type VII collagen (target in EBA, DEB)'),
        ('BP split level', 'Sub-epidermal, within lamina lucida'),
        ('Pathogenesis', 'IgG (and IgE) autoantibodies bind BP180/BP230 → complement activation → eosinophil/neutrophil recruitment → release of proteases → blister formation'),
        ('Salt-split skin DIF', 'Roof binding (epidermal side) - distinguishes BP from EBA (floor binding)'),
        ('Mnemonic', 'BP = "Big Plaques" with autoantibodies to BIG (180kDa & 230kDa) proteins'),
    ]),
    ('Bullous_Pemphigoid_Clinical_transcript.txt', 'Bullous Pemphigoid - Clinical Features & Treatments', [
        ('Demographics', 'Most common autoimmune blistering disease; elderly (>60y)'),
        ('Drug triggers', 'DPP-4 inhibitors (gliptins - vildagliptin, linagliptin), PD-1/PD-L1 inhibitors, diuretics (furosemide), spironolactone, NSAIDs'),
        ('Prodromal phase', 'Pruritic urticarial/eczematous lesions weeks-months before blisters'),
        ('Bullous phase', 'Tense, fluid-filled bullae on erythematous base; flexural distribution'),
        ('Sites', 'Lower abdomen, inner thighs, flexor forearms; mucosa rare (<20%)'),
        ('Nikolsky sign', 'NEGATIVE (vs pemphigus: positive)'),
        ('Histology', 'Sub-epidermal blister with eosinophil-rich infiltrate'),
        ('DIF (perilesional)', 'Linear IgG and C3 at BMZ (n-serrated pattern)'),
        ('IIF on salt-split skin', 'IgG binds epidermal (roof) side'),
        ('ELISA', 'Anti-BP180 (NC16A), anti-BP230 antibodies'),
        ('Treatment - Localised/mild', 'Potent topical steroids (clobetasol) - first-line per Joly trial'),
        ('Treatment - Moderate/severe', 'Systemic prednisolone + steroid-sparing agents (azathioprine, MMF, methotrexate, dapsone, doxycycline + nicotinamide)'),
        ('Refractory', 'Rituximab, omalizumab, IVIG, dupilumab'),
        ('Variants', 'Pemphigoid gestationis (pregnancy, PG), mucous membrane pemphigoid, lichen planus pemphigoides, p200 pemphigoid (laminin γ1)'),
        ('Mortality', '20-40% at 1 year (especially elderly with comorbidities)'),
    ]),
    ('Pemphigus_Vulgaris_transcript.txt', 'Pemphigus Vulgaris and Variants', [
        ('Pathogenesis', 'IgG autoantibodies against desmogleins → loss of cell-cell adhesion → intraepidermal acantholysis'),
        ('PV target antigens', 'Desmoglein 3 (mucosal-only) ± Desmoglein 1 (mucocutaneous)'),
        ('Mucosal-dominant PV', 'Anti-Dsg3 only - oral lesions only'),
        ('Mucocutaneous PV', 'Anti-Dsg3 + anti-Dsg1 - oral + skin lesions'),
        ('Demographics', 'Adults 40-60; ↑ in Ashkenazi Jewish, Mediterranean, Indian populations'),
        ('HLA associations', 'HLA-DR4, HLA-DR14'),
        ('Clinical', 'Painful flaccid blisters → erosions; mucosal involvement (>90%) often first'),
        ('Sites', 'Oral mucosa (most common), scalp, face, chest, axilla, groin'),
        ('Nikolsky sign', 'POSITIVE (lateral pressure → epidermal sloughing)'),
        ('Asboe-Hansen sign', 'POSITIVE (pressure on bulla extends it laterally)'),
        ('Histology', 'Suprabasal acantholysis; "tombstoning" of basal cells'),
        ('DIF', 'Intercellular IgG and C3 ("chicken wire" / "fishnet" pattern)'),
        ('IIF', 'Monkey oesophagus substrate'),
        ('ELISA', 'Anti-Dsg3 ± anti-Dsg1 antibodies (correlate with disease activity)'),
        ('Pemphigus Foliaceus', 'Anti-Dsg1 only; NO mucosal involvement; superficial blisters; sub-corneal acantholysis; "cornflakes" appearance'),
        ('Fogo Selvagem', 'Endemic PF in Brazil; black-fly vector'),
        ('IgA Pemphigus', 'Anti-desmocollin 1 (subcorneal pustular type) or anti-Dsg3 (intraepidermal type)'),
        ('Paraneoplastic Pemphigus (PNP)', 'Antibodies against multiple plakin family proteins (envoplakin, periplakin, desmoplakin); polymorphous mucocutaneous; assoc. with NHL, CLL, Castleman disease, thymoma'),
        ('PNP feature', 'Severe stomatitis + lichenoid changes; bronchiolitis obliterans (often fatal)'),
        ('Treatment', 'Systemic corticosteroids + rituximab (FIRST-LINE per ritux PV trial); azathioprine, MMF, IVIG; cyclophosphamide for severe'),
    ]),
    # -------- Infections (Bacterial / Mycobacterial) --------
    ('Atypical_Mycobacteria_transcript.txt', 'Atypical Mycobacteria', [
        ('Definition', 'Non-tuberculous mycobacteria (NTM); environmental organisms'),
        ('Runyon classification', 'Group I-IV based on growth rate and pigment production'),
        ('M. marinum (Group I)', 'Photochromogen; "fish tank/swimming pool granuloma"; sporotrichoid distribution; treatment: clarithromycin + ethambutol or rifampin + ethambutol'),
        ('M. kansasii (Group I)', 'Photochromogen; can mimic TB; pulmonary disease'),
        ('M. scrofulaceum (Group II)', 'Scotochromogen; cervical lymphadenitis in children (scrofula)'),
        ('M. avium-intracellulare (MAC, Group III)', 'Non-chromogen; disseminated disease in HIV/AIDS (CD4 <50)'),
        ('M. ulcerans (Group III)', 'Buruli ulcer; mycolactone toxin; tropical regions; large painless ulcers'),
        ('M. fortuitum / chelonae / abscessus (Group IV)', 'Rapid growers (<7 days); skin/soft tissue infection after trauma, surgery, tattoos, pedicures'),
        ('M. leprae', 'Causes leprosy (covered separately)'),
        ('Histology', 'Granulomatous inflammation; AFB stain may show acid-fast bacilli'),
        ('Diagnosis', 'Tissue culture (gold standard; specify mycobacterial culture); PCR'),
        ('Treatment', 'Combination therapy 3-6+ months; depends on species; macrolides (clarithromycin) often backbone'),
    ]),
    ('HPV_transcript.txt', 'Human Papillomavirus (HPV)', [
        ('Family', 'Papillomaviridae; double-stranded circular DNA virus; >200 types'),
        ('Common warts (verruca vulgaris)', 'HPV 2, 4, 7'),
        ('Plantar warts (myrmecia)', 'HPV 1'),
        ('Flat warts (verruca plana)', 'HPV 3, 10'),
        ('Anogenital warts (condyloma acuminata)', 'HPV 6, 11 (low-risk)'),
        ('Cervical/anogenital cancer', 'HPV 16, 18 (high-risk; oncogenic E6 inhibits p53, E7 inhibits Rb)'),
        ('Bowenoid papulosis', 'HPV 16, 18; multiple red-brown papules; SCC in situ histology'),
        ('Epidermodysplasia verruciformis (EDV)', 'HPV 5, 8 (β-papillomaviruses); 50% develop SCC by age 30'),
        ('Heck disease (focal epithelial hyperplasia)', 'HPV 13, 32; oral mucosa; common in Native Americans, Inuits'),
        ('Histology', 'Hyperkeratosis, papillomatosis, koilocytes (perinuclear halo + raisinoid nuclei)'),
        ('Treatment - Common warts', 'Cryotherapy, salicylic acid, imiquimod, intralesional Candida antigen, cantharidin'),
        ('Treatment - Genital warts', 'Imiquimod, podophyllotoxin, cryotherapy, sinecatechins, surgery'),
        ('Vaccines', 'Gardasil 9 (covers HPV 6, 11, 16, 18, 31, 33, 45, 52, 58); recommended ages 9-26 (up to 45)'),
        ('Latency', 'Virus persists in basal keratinocytes; reactivation possible'),
    ]),
    ('Leprosy_transcript.txt', 'Leprosy (Hansen Disease)', [
        ('Pathogen', 'Mycobacterium leprae (also M. lepromatosis); acid-fast bacillus; cannot be cultured in vitro'),
        ('Reservoir', 'Humans; armadillos (zoonotic in southern US)'),
        ('Transmission', 'Respiratory droplets; prolonged contact'),
        ('Cellular target', 'Schwann cells (peripheral nerves) and macrophages (skin)'),
        ('Spectrum (Ridley-Jopling)', 'TT (tuberculoid) ↔ BT ↔ BB ↔ BL ↔ LL (lepromatous)'),
        ('Tuberculoid (TT)', 'Strong Th1 response; 1-3 hypopigmented anesthetic plaques; well-defined; few bacilli (paucibacillary); negative lepromin test'),
        ('Lepromatous (LL)', 'Weak Th2 response; symmetric multiple lesions; "leonine facies", madarosis, saddle nose; many bacilli (multibacillary); positive lepromin test'),
        ('Borderline (BT, BB, BL)', 'Intermediate features; unstable forms'),
        ('Indeterminate', 'Earliest form; single hypopigmented macule'),
        ('Lepromin test', 'Skin reaction to killed M. leprae - positive in TT, negative in LL'),
        ('Histology TT', 'Well-formed granulomas around nerves; few/no bacilli'),
        ('Histology LL', 'Foamy macrophages (Virchow cells/lepra cells); many bacilli; Grenz zone'),
        ('Type 1 reaction (reversal)', 'Upgrading reaction in BT/BB; existing lesions become red/swollen; nerve damage; treat with prednisone'),
        ('Type 2 reaction (ENL)', 'Erythema nodosum leprosum; tender SC nodules + fever in LL; treat with thalidomide, prednisone'),
        ('Lucio phenomenon', 'Necrotizing reaction in untreated LL'),
        ('Multibacillary treatment', 'Rifampin + dapsone + clofazimine x 12 months'),
        ('Paucibacillary treatment', 'Rifampin + dapsone x 6 months'),
    ]),
    ('Syphilis_transcript.txt', 'Syphilis', [
        ('Pathogen', 'Treponema pallidum (spirochete)'),
        ('Transmission', 'Sexual contact, vertical (congenital), blood'),
        ('Primary syphilis', 'Painless chancre (3 weeks post-exposure); indurated ulcer with rolled border; regional lymphadenopathy'),
        ('Secondary syphilis', '4-10 weeks after chancre; "great mimicker"'),
        ('Secondary cutaneous', 'Generalized maculopapular rash including PALMS AND SOLES (key feature)'),
        ('Condylomata lata', 'Moist papules in intertriginous areas (highly infectious)'),
        ('Mucous patches', 'Oral/genital mucosal lesions'),
        ('Other secondary', 'Patchy "moth-eaten" alopecia; constitutional symptoms; lymphadenopathy'),
        ('Latent syphilis', 'Asymptomatic; positive serology; early latent (<1y) vs late latent (>1y)'),
        ('Tertiary syphilis', '3-15 years post-infection'),
        ('Gummas', 'Granulomatous lesions of skin/bone/organs'),
        ('Cardiovascular syphilis', 'Aortitis, aortic regurgitation, aneurysm'),
        ('Neurosyphilis', 'Tabes dorsalis (dorsal column), general paresis, Argyll Robertson pupils'),
        ('Congenital syphilis - early', 'Hutchinson teeth, snuffles, saddle nose, mulberry molars, palms/soles rash'),
        ('Congenital syphilis - late', 'Hutchinson triad: interstitial keratitis + Hutchinson teeth + 8th nerve deafness'),
        ('Non-treponemal tests', 'RPR, VDRL (screening; quantitative; can have biological false positives)'),
        ('Treponemal tests', 'FTA-ABS, TP-PA (specific; remain positive for life)'),
        ('Histology', 'Lichenoid + plasma cell-rich infiltrate; endothelial swelling; can use Warthin-Starry stain or IHC for spirochetes'),
        ('Treatment - Primary/Secondary/Early Latent', 'Benzathine penicillin G 2.4 million units IM x 1 dose'),
        ('Treatment - Late Latent/Tertiary', 'Benzathine penicillin G 2.4 million units IM weekly x 3 weeks'),
        ('Treatment - Neurosyphilis', 'IV aqueous penicillin G x 10-14 days'),
        ('Penicillin allergy', 'Doxycycline (not in pregnancy); pregnant must be desensitized'),
        ('Jarisch-Herxheimer reaction', 'Fever, chills within 24h of treatment due to spirochete lysis'),
    ]),
    ('Syphilis_and_Zoon_transcript.txt', 'Syphilis Histology and Zoon Balanitis', [
        ('Syphilis histology - all stages', 'Lichenoid + plasma cell-rich superficial and deep perivascular infiltrate; endothelial swelling'),
        ('Primary chancre histology', 'Ulcerated epidermis; dense plasma cell-rich dermal infiltrate; endothelial swelling'),
        ('Secondary syphilis histology', 'Psoriasiform/lichenoid pattern; plasma cells; vacuolar interface; can mimic many diseases'),
        ('Tertiary gumma histology', 'Granulomatous inflammation with central necrosis; plasma cells'),
        ('Special stains for syphilis', 'Warthin-Starry silver stain; immunohistochemistry (IHC) for treponemes (most sensitive)'),
        ('Plasma cell infiltrate DDx', 'Syphilis, rosacea (rare), arthropod bite, Borrelia infection, Zoon balanitis/vulvitis, lymphoplasmacellular plaque, sclerosing lipogranuloma'),
        ('Zoon Balanitis (Plasma Cell Balanitis)', 'Idiopathic chronic balanitis; uncircumcised middle-aged/elderly men'),
        ('Zoon clinical', 'Solitary shiny red-orange "cayenne pepper" plaque on glans/prepuce'),
        ('Zoon histology', 'Atrophic epidermis; lozenge-shaped keratinocytes; band-like plasma cell infiltrate; dilated capillaries; siderophages'),
        ('Zoon female counterpart', 'Plasma cell vulvitis (Zoon vulvitis)'),
        ('Zoon treatment', 'Circumcision (definitive); topical steroids, calcineurin inhibitors, imiquimod'),
        ('Zoon prognosis', 'Benign but chronic; no malignant potential'),
    ]),
    ('Tinea_Capitis_transcript.txt', 'Tinea Capitis', [
        ('Definition', 'Dermatophyte infection of scalp and hair shafts'),
        ('Demographics', 'Predominantly children (3-7y); rare in adults'),
        ('Most common cause (US/Europe)', 'Trichophyton tonsurans (anthropophilic; endothrix)'),
        ('Most common cause worldwide', 'Microsporum canis (zoophilic - cats/dogs; ectothrix; fluoresces under Wood lamp)'),
        ('Endothrix pattern', 'Spores INSIDE hair shaft - T. tonsurans, T. violaceum'),
        ('Ectothrix pattern', 'Spores OUTSIDE hair shaft - M. canis, M. audouinii (fluoresce on Wood lamp); T. mentagrophytes, T. verrucosum (do NOT fluoresce)'),
        ('Favus', 'T. schoenleinii; cup-shaped scutula; hair shaft full of air spaces; permanent scarring alopecia'),
        ('Clinical patterns', 'Black dot tinea (broken hairs); grey patch (scaling); kerion (boggy inflammatory plaque); favus'),
        ('Kerion', 'Severe T-cell mediated inflammatory response; boggy purulent plaque; scarring alopecia; treat aggressively to prevent scarring'),
        ('Wood lamp', 'Bright green-yellow fluorescence with M. canis, M. audouinii (ectothrix); does NOT fluoresce with T. tonsurans'),
        ('Diagnosis', 'KOH microscopy; fungal culture (gold standard, but slow ~4 weeks); PCR'),
        ('Treatment - First line', 'ORAL antifungals required (topicals do not penetrate hair shaft)'),
        ('Griseofulvin', 'FDA-approved; 8-12 weeks; preferred for Microsporum'),
        ('Terbinafine', 'FDA-approved (>4y); 4-6 weeks; preferred for Trichophyton (esp. T. tonsurans)'),
        ('Itraconazole / Fluconazole', 'Alternative options'),
        ('Adjunct', 'Selenium sulfide or ketoconazole shampoo to reduce spore shedding'),
        ('Asymptomatic carriers', 'Treat household contacts with antifungal shampoo to prevent reinfection'),
        ('Id reaction', 'Hypersensitivity to fungal antigen; sterile vesicular eruption distant from infection site'),
    ]),
    ('Tuberculosis_transcript.txt', 'Cutaneous Tuberculosis', [
        ('Pathogen', 'Mycobacterium tuberculosis; M. bovis; BCG vaccine strain (M. bovis BCG)'),
        ('Classification', 'True cutaneous TB (organisms present) vs Tuberculids (hypersensitivity reactions)'),
        ('Inoculation forms', 'Primary chancre (TB inoculation); TB verrucosa cutis (warty plaques in immunized hosts)'),
        ('Lupus vulgaris', 'Most common form; chronic progressive; "apple jelly" nodules on diascopy; head/neck; reactivation; Mantoux strongly positive'),
        ('Scrofuloderma', 'Direct extension from underlying lymph node/bone TB; cervical/inguinal/axillary; cold abscess → fistula → ulceration'),
        ('Orificial TB', 'Mucocutaneous junctions (mouth, anus); painful ulcers; advanced systemic TB; poor prognosis'),
        ('Acute miliary TB', 'Disseminated; severe immunosuppression; fatal'),
        ('Metastatic TB abscess', 'Single/multiple SC abscesses; immunocompromised'),
        ('Tuberculids - definition', 'Hypersensitivity reaction to TB; cultures NEGATIVE; PCR may be positive; respond to anti-TB therapy'),
        ('Lichen scrofulosorum', 'Tuberculid; lichenoid follicular papules in children with TB'),
        ('Papulonecrotic tuberculid', 'Tuberculid; symmetric crusted papules on extensor surfaces; heals with scarring'),
        ('Erythema induratum (Bazin disease)', 'Tuberculid; lobular panniculitis on calves (women); ulcerating SC nodules'),
        ('Mantoux/PPD', 'Positive in most forms (except orificial/disseminated where anergy may occur)'),
        ('IGRA (QuantiFERON, T-SPOT)', 'More specific than PPD; not affected by BCG'),
        ('Histology', 'Caseating granulomas with Langhans giant cells; AFB stain may be positive'),
        ('Diagnosis', 'Culture (gold standard, slow); PCR; histology + clinical correlation'),
        ('Treatment', 'Standard 4-drug therapy: RIPE (Rifampin + Isoniazid + Pyrazinamide + Ethambutol) x 2 months → RI x 4 months'),
    ]),
]


def get_transcript_text(filename):
    """Read a transcript file and return only the body text (skip header)."""
    path = os.path.join(BASE_DIR, filename)
    if not os.path.exists(path):
        return f"[Transcript file not found: {filename}]"
    with open(path, 'r') as f:
        content = f.read()
    # Skip the title line and the ===== separator
    lines = content.split('\n')
    body = []
    skip = True
    for line in lines:
        if skip:
            if line.startswith('===') or line.strip() == '':
                continue
            elif line.startswith('TRANSCRIPT:'):
                continue
            else:
                skip = False
        body.append(line)
    # Trim leading/trailing empty lines
    return '\n'.join(body).strip()


def build_doc():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title page
    title = doc.add_heading('Dermatology Transcripts - Compiled Study Notes', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f'A collection of {len(ENTRIES)} dermatology audio transcripts with summary tables.')
    sub_run.italic = True

    doc.add_paragraph()
    doc.add_paragraph('Source: Dermatographics audio episodes. Transcribed using OpenAI Whisper.')
    doc.add_page_break()

    # Add each entry
    for i, (filename, title, summary_rows) in enumerate(ENTRIES, start=1):
        # Heading 2
        doc.add_heading(f'{i}. {title}', level=2)

        # Transcript subheading
        doc.add_heading('Transcript', level=3)
        transcript_text = get_transcript_text(filename)
        # Add as paragraph (Whisper output is one big paragraph; keep as-is)
        p = doc.add_paragraph(transcript_text)
        p.paragraph_format.space_after = Pt(6)

        # Summary table subheading
        doc.add_heading('Summary Table', level=3)

        # Create table
        table = doc.add_table(rows=len(summary_rows) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        # Header row
        hdr = table.rows[0].cells
        hdr[0].text = 'Topic'
        hdr[1].text = 'Details'
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        # Data rows
        for j, (key, value) in enumerate(summary_rows, start=1):
            row = table.rows[j].cells
            row[0].text = key
            row[1].text = value

        # Bold the first column
        for j in range(1, len(summary_rows) + 1):
            for paragraph in table.rows[j].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add separator
        doc.add_paragraph()
        if i < len(ENTRIES):
            doc.add_page_break()

    # Save
    out_path = os.path.join(BASE_DIR, 'Dermatology_Transcripts_Compiled.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    print(f'Total entries: {len(ENTRIES)}')

    # Quick check that all transcripts are accounted for
    actual_files = sorted([f for f in os.listdir(BASE_DIR) if f.endswith('_transcript.txt')])
    listed_files = sorted([entry[0] for entry in ENTRIES])
    missing_in_doc = set(actual_files) - set(listed_files)
    missing_files = set(listed_files) - set(actual_files)
    if missing_in_doc:
        print(f'WARNING: Transcripts in repo NOT included in doc: {missing_in_doc}')
    if missing_files:
        print(f'WARNING: Entries in doc with missing files: {missing_files}')


if __name__ == '__main__':
    build_doc()
