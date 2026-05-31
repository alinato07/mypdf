"""
md_to_docx.py — Convert any study-guide markdown file to a formatted Word .docx
Usage:  python3 md_to_docx.py <input.md> <output.docx>
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colour palette ────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x39, 0x64)   # title / H1
MID_BLUE    = RGBColor(0x27, 0x5E, 0xA0)   # H2
LIGHT_BLUE  = RGBColor(0x37, 0x86, 0xC0)   # H3
TEAL        = RGBColor(0x1F, 0x70, 0x8A)   # H4
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREY        = RGBColor(0x50, 0x50, 0x50)

# ── XML helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_para_bg(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "4472C4")
    pBdr.append(bot); pPr.append(pBdr)

# ── inline formatting ─────────────────────────────────────────────────────────
_INLINE = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')

def add_runs(para, text: str):
    last = 0
    for m in _INLINE.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1) is not None:
            r = para.add_run(m.group(1)); r.bold = True
        elif m.group(2) is not None:
            r = para.add_run(m.group(2)); r.italic = True
        else:
            r = para.add_run(m.group(3))
            r.font.name = "Courier New"; r.font.size = Pt(9)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])

# ── table parser ──────────────────────────────────────────────────────────────
_SEP = re.compile(r'^\|[-| :]+\|$')

def flush_table(doc, rows):
    if not rows: return
    cols = max(len(r) for r in rows)
    rows = [r + ['']*(cols-len(r)) for r in rows]
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell_txt in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_runs(p, cell_txt.strip())
            if ri == 0:
                set_cell_bg(cell, '1F3964')
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = WHITE
            elif ri % 2 == 0:
                set_cell_bg(cell, 'DCE6F1')
    doc.add_paragraph()

# ── code-block pass-through ───────────────────────────────────────────────────
def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Cm(1)
        set_para_bg(p, 'F2F2F2')
        r = p.add_run(line)
        r.font.name = 'Courier New'; r.font.size = Pt(9)

# ── main ──────────────────────────────────────────────────────────────────────
def convert(md_path: str, docx_path: str):
    doc = Document()

    # margins
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(1)
        sec.top_margin  = sec.bottom_margin = Inches(1)

    # base style
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # heading colours
    for style, colour, size in [
        ('Title',     DARK_BLUE,  22),
        ('Heading 1', DARK_BLUE,  16),
        ('Heading 2', MID_BLUE,   13),
        ('Heading 3', LIGHT_BLUE, 11),
        ('Heading 4', TEAL,       11),
    ]:
        s = doc.styles[style]
        s.font.size  = Pt(size)
        s.font.bold  = True
        s.font.color.rgb = colour

    with open(md_path, encoding='utf-8') as f:
        lines = f.read().splitlines()

    i = 0
    table_rows = []
    code_lines = []
    in_code = False

    while i < len(lines):
        raw = lines[i]; s = raw.strip()

        # ---- fenced code block
        if s.startswith('```'):
            if not in_code:
                in_code = True; i += 1; continue
            else:
                add_code_block(doc, code_lines)
                code_lines = []; in_code = False; i += 1; continue
        if in_code:
            code_lines.append(raw); i += 1; continue

        # ---- flush pending table when non-table line encountered
        if table_rows and not s.startswith('|'):
            flush_table(doc, table_rows); table_rows = []

        # ---- empty line
        if not s:
            i += 1; continue

        # ---- horizontal rule
        if s in ('---', '***', '___'):
            add_hr(doc); i += 1; continue

        # ---- headings
        if s.startswith('#### '):
            doc.add_paragraph(s[5:].strip(), style='Heading 4'); i += 1; continue
        if s.startswith('### '):
            doc.add_paragraph(s[4:].strip(), style='Heading 3'); i += 1; continue
        if s.startswith('## ') and not s.startswith('### '):
            doc.add_paragraph(s[3:].strip(), style='Heading 2'); i += 1; continue
        if s.startswith('# ') and not s.startswith('## '):
            doc.add_paragraph(s[2:].strip(), style='Title');     i += 1; continue

        # ---- blockquote (clinical pearl / note)
        if s.startswith('> '):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            set_para_bg(p, 'E8F0FE')
            add_runs(p, s[2:]); i += 1; continue

        # ---- markdown table row
        if s.startswith('|'):
            if _SEP.match(s): i += 1; continue   # skip separator row
            cells = [c for c in s.strip('|').split('|')]
            table_rows.append(cells); i += 1; continue

        # ---- sub-bullet (2+ spaces then - or *)
        if re.match(r'^\s{2,}[-*]\s', raw):
            p = doc.add_paragraph(style='List Bullet 2')
            add_runs(p, re.sub(r'^\s+[-*]\s', '', s)); i += 1; continue

        # ---- numbered list
        if re.match(r'^\d+\.\s', s):
            p = doc.add_paragraph(style='List Number')
            add_runs(p, re.sub(r'^\d+\.\s', '', s)); i += 1; continue

        # ---- bullet
        if re.match(r'^[-*]\s', s):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, re.sub(r'^[-*]\s', '', s)); i += 1; continue

        # ---- italic-only line (notes)
        if s.startswith('*') and s.endswith('*') and not s.startswith('**'):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run(s.strip('*')); r.italic = True
            r.font.color.rgb = GREY; i += 1; continue

        # ---- normal paragraph
        p = doc.add_paragraph(style='Normal')
        add_runs(p, s); i += 1

    # flush any remaining table
    if table_rows:
        flush_table(doc, table_rows)

    doc.save(docx_path)
    print(f'Saved: {docx_path}')

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python3 md_to_docx.py <input.md> <output.docx>')
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
