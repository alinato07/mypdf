"""Build a comprehensive Word document collating all 63 dermatology Word transcript files."""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = '/projects/sandbox/mypdf'

# Each entry: (docx_filename, display_title, summary_table_rows)
# summary_table_rows: list of (key, value) tuples
ENTRIES = [
    # -------- Drug-Induced Lupus Erythematosus --------
    ('DILE-wolfs-drug-surplus-website-file.doc', 'Drug-Induced Lupus Erythematosus (DILE)', [
        ('Onset', 'Usually >1 year on the medication before symptoms appear'),
        ('Cutaneous findings', 'NOT prominent in typical DILE; TNF-alpha inhibitor DILE = exception (photosensitivity, malar rash)'),
        ('Most common symptoms', 'Arthralgias and myalgias (#1); also fever, weight loss, serositis'),
        ('Typical drugs (anti-histone Ab+)', 'Procainamide, Isoniazid, Hydralazine, Methyldopa, Penicillamine (PIHMP)'),
        ('Penicillamine distinction', 'Can UNMASK true underlying SLE in addition to causing DILE'),
        ('Minocycline', 'Atypical DILE; NOT anti-histone antibodies; instead p-ANCA positive'),
        ('TNF-alpha inhibitors', 'Atypical DILE; cutaneous lupus features (photosensitivity, malar rash); positive dsDNA antibodies'),
        ('Slow acetylators', 'Higher risk for typical DILE (especially procainamide, hydralazine, isoniazid)'),
        ('Resolution', 'Symptoms resolve within 4-6 weeks of stopping the offending agent'),
    ]),

    ('A-day-to-pie-for-script.docx', 'Pyoderma Gangrenosum (PG)', [
        ('Classification', 'Neutrophilic dermatosis; sterile inflammatory ulceration'),
        ('Pathergy', 'Lesions develop/worsen at sites of trauma; avoid unnecessary debridement'),
        ('Classic ulcer', 'Painful; violaceous undermined/overhanging border; anterior lower leg most common site'),
        ('Vegetative form', 'Least aggressive; verrucous/exophytic plaques; responds to conservative therapy; often post-surgical'),
        ('Bullous form', 'Haemorrhagic bullae; associated with haematologic malignancies (AML)'),
        ('Pustular form', 'Multiple sterile pustules; associated with IBD'),
        ('Systemic associations', 'IBD (especially UC), rheumatoid arthritis, haematologic disorders, monoclonal gammopathy'),
        ('Peristomal PG', 'Around ileostomy/colostomy sites; pathergy important'),
        ('Histology', 'Dense neutrophilic infiltrate; non-specific; must exclude infection/vasculitis'),
        ('Diagnosis', 'Diagnosis of exclusion; Paracelsus score can assist'),
        ('Treatment', 'Topical/systemic corticosteroids; ciclosporin; infliximab (especially IBD-associated); dapsone'),
        ('Scarring', 'Atrophic cribriform scarring on healing'),
    ]),
    # -------- Acne --------
    ('Acne.docx', 'Acne Vulgaris - Clinical Features, Pathogenesis & Treatment', [
        ('Pathogenesis', 'Abnormal follicular keratinization + androgens (DHT/testosterone) → ↑ sebaceous gland growth → ↑ sebum'),
        ('C. acnes role', 'Gram-positive anaerobic rod; activates TLR2 on macrophages → IL-1α → neutrophil recruitment → NLRP3 inflammasome'),
        ('Wood lamp', 'C. acnes produces coproporphyrin III → red-orange fluorescence'),
        ('Dietary triggers', 'Skim milk, whey protein, high glycemic foods (evidence moderate)'),
        ('Female flares', 'Premenstrual flares ~1 week before menstruation'),
        ('Topical retinoids', 'First-line for comedonal acne; tretinoin, adapalene, tazarotene'),
        ('Benzoyl peroxide', 'Bactericidal; reduces C. acnes resistance; combine with antibiotics'),
        ('Oral antibiotics', 'Doxycycline, minocycline (moderate-severe inflammatory)'),
        ('Isotretinoin', 'Severe nodulocystic/scarring acne; iPLEDGE program; teratogenic; causes xerosis, cheilitis, pseudotumor cerebri, hyperlipidaemia'),
        ('Spironolactone', 'Anti-androgen; women with hormonal acne'),
        ('Comedones', 'Open (blackhead) = oxidized melanin; Closed (whitehead) = follicular plugging'),
        ('Gram-negative folliculitis', 'Complication of prolonged oral antibiotic use; treat with isotretinoin'),
    ]),

    # -------- Acrodermatitis Enteropathica --------
    ('Acrodermatitis_Enteropathica.docx', 'Acrodermatitis Enteropathica', [
        ('Inheritance', 'Autosomal recessive; SLC39A4 gene → encodes ZIP4 zinc transporter'),
        ('Clinical triad', 'Erosive vesiculopustular eczematous lesions (diaper/face/acral) + diarrhoea + alopecia'),
        ('Acquired form', 'Secondary zinc deficiency: alcoholism, anorexia, malabsorption (IBD), formula-fed infants, low maternal breastmilk zinc'),
        ('Histology', 'Cytoplasmic pallor of keratinocytes + ballooning in upper epidermis'),
        ('Labs', '↓ serum zinc; ↓ alkaline phosphatase (zinc-dependent enzyme)'),
        ('Treatment', 'Lifelong oral zinc sulfate supplementation'),
    ]),
    # -------- Actinic Keratosis --------
    ('Actinic_Keratosis.docx', 'Actinic Keratosis & Squamous Cell Carcinoma In Situ', [
        ('Pathogenesis', 'UVB → thymidine dimers → p53 mutations → impaired apoptosis'),
        ('Histology', 'Basal layer atypia limited to lower 1/3 epidermis; budding finger-like projections into dermis'),
        ('Flag sign', 'Parakeratosis (pink) alternating with orthokeratosis (blue) on histology'),
        ('SCC transformation', '<0.1% per AK per year'),
        ('Treatment - single lesion', 'Cryotherapy (liquid nitrogen)'),
        ('Treatment - field', '5-fluorouracil (5-FU), imiquimod, photodynamic therapy, diclofenac gel'),
    ]),
    # -------- Congenital Ichthyoses Part 1 --------
    ('All_the_Fish_in_the_Sea-_Congenital_Ichthyoses.docx', 'Congenital Ichthyoses (Part 1) - EI, Lamellar, Netherton, SSD', [
        ('Epidermolytic Ichthyosis (EI)', 'AD; KRT1 + KRT10; bullous at birth → cobblestone hyperkeratosis over joints; malodorous; clumped keratin suprabasal cells on histology'),
        ('Ichthyosis hystrix (Curth-Macklin)', 'KRT1 only; related to EI'),
        ('KID Syndrome', 'GJB2 (Connexin 26); Keratitis + Ichthyosis + Deafness; ↑ risk oral/cutaneous SCC'),
        ('Superficial EI (Ichthyosis bullosa of Siemens)', 'KRT2; superficial blistering; Mauserung phenomenon (peeling)'),
        ('Lamellar Ichthyosis', 'AR; TGM1 (most common); collodion baby at birth; large dark plate-like scales; ectropion + eclabion; no erythema'),
        ('Non-bullous CIE', 'AR; varied genes; erythroderma + fine white scales'),
        ('Netherton Syndrome', 'SPINK5 (LEKTI); ichthyosis linearis circumflexa (double-edged scale); trichorrhexis invaginata (bamboo hair); atopic features; ↑ IgE'),
        ('Steroidal Sulfatase Deficiency (SSD)', 'X-linked recessive; STS gene; large brown scales sparing flexures; retained vernix caseosa at birth; cryptorchidism; corneal opacities'),
        ('Ichthyosis Vulgaris', 'AD; FLG (filaggrin); variable penetrance; fine white scales; keratosis pilaris association'),
    ]),

    # -------- Pityriasis Rosea --------
    ('America-is-Rosea-Press-Release-Pityriasis-Rosea.docx', 'Pityriasis Rosea', [
        ('Family', 'Papulosquamous disorders (psoriasis, seb derm, PRP, granular parakeratosis)'),
        ('Aetiology', 'HHV-7 (most common) and HHV-6'),
        ('Drug-induced PR', 'ACE inhibitors (most common), gold, beta-blockers, barbiturates, isotretinoin, metronidazole'),
        ('Herald patch', 'Single larger plaque precedes generalised eruption by 1-2 weeks'),
        ('Distribution', 'Christmas-tree pattern along Langer lines; trunk; spares face'),
        ('Key feature', 'Trailing collarette of scale (peripheral); also seen in EAC'),
        ('Populations', 'More extensive in African-American children'),
        ('Course', 'Self-limited; resolves in 6-8 weeks'),
        ('Treatment', 'Symptomatic; erythromycin if earlier clearance needed'),
    ]),
    # -------- Amyloidosis (script) --------
    ('Amyloidosis-script-.docx', 'Amyloidosis - Primary Systemic & Cutaneous Forms', [
        ('Primary Systemic (AL)', 'IgG lambda light chains; underlying plasma cell dyscrasia/multiple myeloma; check SPEP/UPEP'),
        ('Cutaneous signs - primary', 'Sclerodermoid changes, pinch purpura (periocular), waxy papules/nodules, alopecia, macroglossia'),
        ('Cardiac involvement', 'Worst prognosis in primary systemic amyloidosis'),
        ('Lichen Amyloidosis', 'Keratin (K5); lichenified papules on shins; MEN type 2A association'),
        ('Macular Amyloidosis', 'Keratin (K5); rippled hyperpigmented patches on interscapular upper back'),
        ('Nodular Amyloidosis', 'AL; skin-limited; ~7% progress to systemic; associated with CTDs'),
        ('Histology stains', 'Congo red (apple-green birefringence under polarised light), Crystal violet, Thioflavin T'),
        ('Secondary Amyloidosis', 'AA protein from chronic inflammation (ankylosing spondylitis, TB, dystrophic EB, Muckle-Wells)'),
    ]),
    # -------- Ataxia Telangiectasia / DKC --------
    ('Ataxia_Telangiectasia_Final.docx', 'Dyskeratosis Congenita (DKC)', [
        ('Inheritance', 'X-linked recessive (DKC1) most common; also AD (TERT, TERC)'),
        ('Mechanism', 'Defective telomere maintenance → progressive telomere shortening → chromosomal instability'),
        ('Classic Triad', 'Reticulated hyperpigmentation (neck/arms/upper torso) + Nail dystrophy (pterygium → nail loss) + Leukoplakia (pre-malignant)'),
        ('Bone marrow', 'Failure in 50-90%; major cause of mortality'),
        ('Cancer risk', '↑ SCC oral/oesophageal/anal/cutaneous; annual skin screening'),
        ('Eyes', 'Excessive lacrimation (lacrimal duct atresia)'),
        ('Other', 'Liver cirrhosis, pulmonary fibrosis, premature greying'),
    ]),

    # -------- Atopic Dermatitis --------
    ('Atopic_Dermatitis.docx', 'Atopic Dermatitis', [
        ('Atopic Triad', 'Atopic dermatitis + Allergic rhinitis + Asthma'),
        ('Pathogenesis', 'Filaggrin mutations (FLG) → barrier dysfunction + trans-epidermal water loss; SPINK5 (LEKTI); S. aureus colonisation'),
        ('Acute cytokines', 'Th2 dominant: IL-4, IL-5, IL-10, IL-13'),
        ('Chronic cytokines', 'Th1 dominant: IL-1, IFN-gamma'),
        ('Epidemiology', 'Higher prevalence in high-income/urban areas; onset by 1-2 years in early-onset form'),
        ('Subtypes', 'Early-onset (most common, 50% IgE+); Late-onset (post-puberty); Senile (>60y)'),
        ('Treatment', 'Emollients + mild-potent topical steroids; tacrolimus/pimecrolimus; dupilumab (anti-IL-4Rα); tralokinumab (anti-IL-13); JAK inhibitors (upadacitinib, abrocitinib)'),
    ]),
    # -------- Biotinidase --------
    ('Biotinidase_Multiple_Carboxylase_Deficiency_transcript.docx', 'Biotinidase & Multiple Carboxylase Deficiency', [
        ('Inheritance', 'Autosomal recessive'),
        ('Mechanism', 'Biotin cofactor required for 4 carboxylases (ACC, PC, MCC, PCC)'),
        ('Biotinidase deficiency', 'Childhood onset; milder'),
        ('Holocarboxylase synthetase deficiency', 'Early infancy; more severe; often fatal without treatment'),
        ('Skin', 'Perioral and generalised dermatitis, alopecia'),
        ('Eyes', 'Optic atrophy (biotinidase type)'),
        ('Neuro', 'Seizures, hypotonia, developmental delay'),
        ('Treatment', 'IV or oral biotin replacement; dramatic response'),
    ]),
    # -------- Bites & Stings --------
    ('Bites_&_Stings.docx', 'Arthropod Bites & Stings - DEET, Ticks, Spiders, Scabies', [
        ('DEET', 'Most effective insect repellent; safe in children >2 months; repels mosquitoes, ticks, flies'),
        ('Permethrin', 'Synthetic pyrethroid; applied to clothing/gear (not skin); kills on contact'),
        ('Tick removal', 'Fine-tipped tweezers; grasp close to skin; steady upward pull; clean with alcohol'),
        ('Rocky Mountain Spotted Fever', 'Rickettsia rickettsii; Dermacentor tick; centripetal rash (starts wrists/ankles, spreads centrally); treat: doxycycline'),
        ('Lyme Disease', 'Borrelia burgdorferi; Ixodes tick (>36h attachment needed); erythema migrans; treat: doxycycline'),
        ('Brown Recluse', 'Loxosceles reclusa; necrotic arachnidism; violin-shaped mark on cephalothorax; dapsone for treatment'),
        ('Black Widow', 'Latrodectus; neurotoxic (alpha-latrotoxin); muscle cramps, diaphoresis; treat with calcium gluconate/antivenom'),
        ('Scabies', 'Sarcoptes scabiei; intensely pruritic; burrows in web spaces/wrists/genitalia; Norwegian (crusted) scabies in immunosuppressed; treat with permethrin 5% cream'),
        ('Pediculosis', 'Pediculus humanus capitis (head), corporis (body), Phthirus pubis (pubic); treat with permethrin, malathion, or spinosad'),
    ]),

    # -------- Bloom Syndrome --------
    ('Bloom_Syndrome.docx', 'Bloom Syndrome', [
        ('Inheritance', 'Autosomal recessive'),
        ('Gene', 'BLM (RECQL3) - DNA helicase'),
        ('Aka', 'Congenital telangiectatic erythema'),
        ('Growth', 'Pre- and postnatal growth retardation; height ≤5 feet'),
        ('Skin', 'Photosensitivity; telangiectatic malar butterfly erythema; café-au-lait macules'),
        ('Facies', 'Narrow face, prominent ears, bird-like nose, high-pitched voice'),
        ('Endocrine', 'Primary hypogonadism; males sterile'),
        ('Immunology', '↓ IgA and IgM'),
        ('Cytogenetics', '↑ sister chromatid exchanges (hallmark)'),
        ('Cancer risk', '↑ lymphoma and leukaemia (#1 cause of death)'),
    ]),
    # -------- CH-fashion (CHH & CHILD) --------
    ('CH-fashion-script.docx', 'Conradi-Hünermann-Happle (CHH) & CHILD Syndromes', [
        ('CHH gene', 'EBP (emopamil-binding protein); X-linked dominant; lethal in males'),
        ('CHILD gene', 'NSDHL; X-linked dominant; lethal in males'),
        ('CHH ichthyosis', 'Generalised along Blaschko lines; unilateral cataracts; ice-pick scarring; follicular atrophoderma'),
        ('CHILD ichthyosis', 'Hemilateral with sharp midline demarcation; verrucous hyperkeratosis; hemidysplasia; hypoplastic limbs'),
        ('Shared features', 'Chondrodysplasia punctata; asymmetric limb shortening; ichthyosiform erythroderma'),
        ('Metabolism', 'Both involve cholesterol biosynthesis defects'),
    ]),
    # -------- Cockayne Syndrome --------
    ('Cockayne_Syndrome.docx', 'Cockayne Syndrome', [
        ('Inheritance', 'Autosomal recessive'),
        ('Defect', 'Defective transcription-coupled NER (vs XP = global genomic NER)'),
        ('Genes', 'CSA → ERCC8; CSB → ERCC6'),
        ('Skin', 'Photosensitivity, telangiectatic erythema; NO increased skin cancer risk'),
        ('Facies', 'Pinched, narrow, bird-like face with beak nose'),
        ('Neuro', 'Progressive demyelination CNS + PNS; ataxia'),
        ('Growth', 'Cachectic dwarfism (postnatal growth failure)'),
        ('Eyes', 'Salt-and-pepper retinopathy; cataracts'),
        ('Prognosis', 'Death by 4th decade from neurological complications'),
    ]),
    # -------- Congenital Ichthyosis Part 2 --------
    ('congenital-ichthyosis-part-2-script.docx', 'Congenital Ichthyoses (Part 2) - Vulgaris, KID, Lamellar & Netherton', [
        ('Ichthyosis Vulgaris', 'AD; FLG (filaggrin); variable penetrance; fine white scales; KP association'),
        ('KID Syndrome', 'GJB2 (Connexin 26); Keratitis + Ichthyosis + Deafness; ↑ SCC risk'),
        ('Ichthyosis bullosa of Siemens', 'KRT2; superficial; Mauserung peeling phenomenon'),
        ('Lamellar Ichthyosis', 'AR; TGM1; collodion baby; large dark plate-like scales; ectropion/eclabion'),
        ('Steroidal Sulfatase Deficiency', 'X-linked recessive; large brown scales sparing flexures; cryptorchidism; corneal opacities; associated with Kallmann syndrome'),
        ('Netherton Syndrome', 'SPINK5/LEKTI; ichthyosis linearis circumflexa; bamboo hair (trichorrhexis invaginata); atopic features; ↑ IgE'),
    ]),

    # -------- Cutis Laxa --------
    ('Cutis_Laxa.docx', 'Cutis Laxa', [
        ('Aka', 'Generalised elastolysis; "hound dog" facies'),
        ('AD form', 'ELN gene; less severe; mainly skin; early adulthood'),
        ('AR Type I', 'FBLN5 (fibulin-5); hypoplastic lungs/emphysema; GI/GU diverticula'),
        ('AR Type IIA', 'ATP6V0A2; wrinkly skin syndrome'),
        ('AR Type IIB', 'PYCR1'),
        ('AR Type III De Barsy (IIIA)', 'ALDH18A1; corneal clouding + cataracts'),
        ('X-linked recessive', 'Occipital Horn Syndrome (ATP7A); wedge-shaped occipital calcifications; copper metabolism defect (same gene as Menkes)'),
        ('Acquired', 'Adults; drugs (penicillamine, isoniazid); RA, sarcoidosis, lupus, post-inflammatory'),
        ('Histology', 'Fragmented, clumped, shortened elastic fibres; Verhoeff-van Gieson stain'),
    ]),
    # -------- DKC (rough) - same content as Ataxia_Telangiectasia_Final --------
    ('DKC_(rough).docx', 'Dyskeratosis Congenita (DKC) - Detailed Review', [
        ('Inheritance', 'X-linked recessive (DKC1) most common; AD (TERT, TERC); AR forms exist'),
        ('Mechanism', 'Defective telomerase → shortened telomeres → genomic instability'),
        ('Classic Triad', 'Reticulated hyperpigmentation + Nail dystrophy (pterygium → nail loss) + Leukoplakia (pre-malignant oral)'),
        ('Bone marrow failure', '50-90%; major mortality cause; Fanconi anaemia overlap'),
        ('Cancer risk', 'SCC (skin, oral, oesophageal, anus); annual surveillance'),
        ('Poikiloderma', 'Common skin finding alongside hyperpigmentation'),
        ('Other', 'Lacrimal duct atresia (excessive tearing), pulmonary fibrosis, liver cirrhosis, premature greying'),
        ('Hoyeraal-Hreidarsson', 'Severe variant; cerebellar hypoplasia + IUGR + immunodeficiency'),
    ]),
    # -------- Dermatomyositis --------
    ('Dermatomyositis-script-.docx', 'Dermatomyositis - Clinical Features & Autoantibodies', [
        ('Classic skin features', 'Heliotrope rash (periorbital violaceous + oedema); Gottron papules (knuckles); Gottron sign (elbows/knees)'),
        ('Other skin', 'Shawl sign, V-sign, holster sign, mechanic\'s hands, periungual erythema/telangiectasias, flagellate erythema (bleomycin-like)'),
        ('Calcinosis cutis', 'Especially childhood DM; NXP-2 antibody association'),
        ('Anti-Mi-2', 'Classic DM features; GOOD prognosis'),
        ('Anti-TIF1-gamma (p155/140)', 'Amyopathic DM; cancer-associated DM (screen for malignancy)'),
        ('Anti-NXP-2', 'Juvenile DM; calcinosis cutis'),
        ('Anti-MDA5', 'Clinically amyopathic DM; mechanic\'s hands; rapidly progressive ILD (worst pulmonary prognosis)'),
        ('Anti-SRP', 'Severe necrotising myopathy; cardiac involvement; very poor prognosis'),
        ('Anti-Jo-1 (anti-synthetase)', 'ILD + arthritis + myositis + mechanic\'s hands + Raynaud\'s'),
        ('Malignancy screening', 'Required especially for TIF1-gamma; ovarian most common associated'),
        ('Treatment', 'Systemic corticosteroids; methotrexate; IVIG; hydroxychloroquine for skin features'),
    ]),

    # -------- Diffuse PPKs --------
    ('Diffuse-PPKs-An-Uma-Thurman-Action-Film.docx', 'Diffuse Palmoplantar Keratodermas (PPKs)', [
        ('Unna-Thost (Diffuse non-epidermolytic PPK)', 'AD; KRT1; diffuse waxy yellow PPK; no epidermolysis on histology'),
        ('Vörner PPK (Diffuse epidermolytic PPK)', 'AD; KRT1 or KRT9; identical to Unna-Thost clinically but has epidermolysis on histology'),
        ('Mal de Meleda', 'AR; SLURP1; transgrediens PPK (extends to dorsal hands/feet); constricting bands; pseudoainhum risk'),
        ('Papillon-Lefèvre Syndrome', 'AR; CTSC (cathepsin C); PPK + severe periodontitis (tooth loss by age 5); ↑ pyogenic infections (Papillon-Lefèvre + Haim-Munk)'),
        ('Haim-Munk Syndrome', 'AR; CTSC; PPK + periodontitis + arachnodactyly + onychogryphosis + flat feet'),
        ('Vohwinkle Syndrome', 'GJB2 (connexin 26); honeycomb PPK + starfish-shaped dorsal keratoderma + pseudoainhum + deafness'),
        ('Loricrin Keratoderma', 'LOR gene; Vohwinkle-like without deafness; ichthyosis'),
        ('Howel-Evans Syndrome (TOC)', 'AD; RHBDF2; PPK + oesophageal SCC; tylosis'),
        ('Bart-Pumphrey Syndrome', 'GJB2; PPK + leukonychia + knuckle pads + deafness'),
    ]),
    # -------- EED & Acute Haemorrhagic Oedema of Infancy --------
    ('EED-and-Acute-hemorrhagic-edema-of-infancy.docx', 'Erythema Elevatum Diutinum (EED) & Haemorrhagic Oedema of Infancy (HEI)', [
        ('EED clinical', 'Red-brown to violaceous papules/nodules over extensor surfaces and joints; arthralgias; ocular symptoms'),
        ('EED associations', 'HIV (most important); lupus; IBD; TB; lymphoma; IgA paraproteinaemia'),
        ('EED histology', 'Perivascular onion-skin fibrosis (early: LCV; late: fibrosing)'),
        ('EED treatment', 'Dapsone (treatment of choice); colchicine'),
        ('HEI demographics', 'Ages 6-24 months; bright red-violaceous annular/cockade (targetoid) plaques; upper extremities, head, neck; facial oedema; child appears well despite alarming appearance'),
        ('HEI triggers', 'Hepatitis A; antibiotics (cephalosporins, beta-lactams, TMP-SMX); streptococcal pharyngitis; NSAIDs; viral infections (coxsackie, EBV, HSV, VZV); vaccinations'),
        ('HEI course', 'Self-limiting; resolves without sequelae'),
    ]),
    # -------- Epidermolysis Bullosa --------
    ('Epidermolysis_Bullosa.docx', 'Epidermolysis Bullosa (EB)', [
        ('EBS Generalised Severe (Dowling-Meara)', 'KRT5/KRT14; herpetiform blistering; PPK'),
        ('EBS Generalised Intermediate (Koebner)', 'KRT5/KRT14'),
        ('EBS Localised (Weber-Cockayne)', 'KRT5/KRT14; hands and feet only'),
        ('EBS Mottled Pigmentation', 'KRT5 only'),
        ('EBS with Muscular Dystrophy', 'PLEC (plectin); only AR EBS form'),
        ('JEB Severe (Herlitz)', 'LAMA3/LAMB3/LAMC2 (laminin 332); hoarse cry; perioral granulation tissue; infant death'),
        ('JEB Intermediate', 'Laminin 332 + Collagen XVII (COL17A1)'),
        ('JEB with Pyloric Atresia', 'ITGA6/ITGB4 (α6β4 integrin)'),
        ('DEB Dominant (Pasini/Cockayne-Touraine)', 'COL7A1; albopapuloid papules (Pasini)'),
        ('DEB Recessive Severe (Hallopeau-Siemens)', 'COL7A1; severe; SCC > renal failure; mitten deformity'),
        ('Kindler Syndrome', 'FERMT1 (kindlin-1); acral blistering, poikiloderma, photosensitivity'),
        ('Split level mnemonic', 'EBS = intra-epidermal (above BMZ); JEB = lamina lucida; DEB = sub-lamina densa'),
    ]),

    # -------- Grandpa's Big Party (Bullous Pemphigoid clinical) --------
    ('Grandpas-Big-Party.docx', 'Bullous Pemphigoid - Clinical Features & Treatment', [
        ('Demographics', 'Most common autoimmune blistering disease; elderly >60 years'),
        ('Drug triggers', 'DPP-4 inhibitors (gliptins), PD-1/PD-L1 inhibitors, furosemide, spironolactone, NSAIDs'),
        ('Prodromal phase', 'Pruritic urticarial/eczematous plaques weeks-months before blisters'),
        ('Bullous phase', 'Tense fluid-filled bullae on erythematous base; flexural distribution; mucosa rare (<20%)'),
        ('Sites', 'Inner thighs, lower abdomen, flexor forearms'),
        ('Nikolsky sign', 'NEGATIVE'),
        ('Histology', 'Sub-epidermal blister; eosinophil-rich infiltrate'),
        ('DIF (perilesional)', 'Linear IgG and C3 at BMZ (n-serrated pattern)'),
        ('IIF salt-split skin', 'IgG binds epidermal (roof) side'),
        ('ELISA', 'Anti-BP180 (NC16A domain) + anti-BP230'),
        ('Treatment mild', 'Topical clobetasol (Joly trial - first-line)'),
        ('Treatment moderate/severe', 'Systemic prednisolone + steroid-sparing (azathioprine, MMF, doxycycline + nicotinamide)'),
        ('Refractory', 'Rituximab, omalizumab, dupilumab, IVIG'),
        ('Mortality', '20-40% at 1 year in elderly with comorbidities'),
    ]),
    # -------- Griscelli Syndrome --------
    ('Griscelli-syndrome.docx', 'Griscelli Syndrome & Related Hypopigmentation Disorders', [
        ('Griscelli Type 1', 'MYO5A; neurological impairment; silvery-grey hair; large clumped melanosomes on hair microscopy'),
        ('Griscelli Type 2', 'RAB27A; haemophagocytic lymphohistiocytosis (HLH) + silvery-grey hair; immunodeficiency'),
        ('Griscelli Type 3', 'MLPH (melanophilin); only hypopigmentation; NO neurological or immune defects'),
        ('Chediak-Higashi Syndrome', 'LYST; partial albinism + giant lysosomal granules + recurrent pyogenic infections + HLH accelerated phase'),
        ('Hermansky-Pudlak Syndrome', 'AP3B1 (Type 2); HPS1, HPS4 (Types 1/4 - Puerto Rican); OCA + platelet storage pool deficiency (bleeding) + pulmonary fibrosis + ceroid accumulation'),
        ('Elejalde Syndrome', 'MYO5A mutation; severe neurology; similar to GS Type 1'),
        ('Cross-McKusick-Breen', 'GS + oculocerebral syndrome'),
    ]),
    # -------- HSP --------
    ('HSP-script.docx', 'Henoch-Schönlein Purpura (HSP / IgA Vasculitis)', [
        ('Aka', 'IgA Vasculitis; most common vasculitis in children (4-11 years)'),
        ('Trigger', 'URI (Group A strep, viral) preceding by 1-3 weeks'),
        ('Classic tetrad', 'Palpable purpura (dependent areas) + arthritis/arthralgia + abdominal pain + renal involvement'),
        ('Skin', 'Palpable purpura on lower extremities and buttocks'),
        ('GI', 'Abdominal pain; intussusception (ileoileal); GI bleeding'),
        ('Renal', 'IgA nephropathy → haematuria/proteinuria; major long-term morbidity'),
        ('Histology', 'Leukocytoclastic vasculitis with IgA deposits on DIF'),
        ('Course', 'Self-limited 4-6 weeks; recurrence in ~30%'),
        ('Adult HSP', 'More severe; higher rate of renal involvement'),
        ('Treatment', 'Supportive; steroids for severe abdominal pain or renal involvement'),
    ]),

    # -------- Hermansky-Pudlak --------
    ('Hermansky-Pudlak.docx', 'Hermansky-Pudlak Syndrome (HPS)', [
        ('Inheritance', 'Autosomal recessive'),
        ('Classic Type 1', 'HPS1 gene; most severe; Puerto Rican population; OCA + bleeding + pulmonary fibrosis'),
        ('Type 2', 'AP3B1; similar severity; ↑ infections'),
        ('Type 4', 'HPS4; Puerto Rican; severe'),
        ('OCA features', 'Variable hypopigmentation; nystagmus; reduced visual acuity'),
        ('Platelet defect', 'Storage pool deficiency (absent dense granules) → mucocutaneous bleeding, easy bruising'),
        ('Pulmonary fibrosis', 'Leading cause of death (Types 1, 4); ceroid deposition'),
        ('Ceroid accumulation', 'Also in GI tract, kidney, bone marrow'),
        ('Distinguish from Chediak-Higashi', 'HPS: no giant granules; no recurrent infections; no HLH'),
    ]),
    # -------- Human Herpesviruses --------
    ('Human_Herpesviruses.docx', 'Human Herpesviruses - HSV, VZV, EBV, CMV, HHV-6/7/8', [
        ('HSV-1', 'Orolabial herpes; herpes labialis; herpetic whitlow; herpes gladiatorum; keratoconjunctivitis; Kaposi varicelliform eruption'),
        ('HSV-2', 'Genital herpes; neonatal herpes (most dangerous - encephalitis)'),
        ('HSV treatment', 'Aciclovir/valaciclovir/famciclovir; foscarnet (aciclovir-resistant)'),
        ('VZV (HHV-3)', 'Varicella (chickenpox); reactivation = zoster; Hutchinson sign (nasociliary → herpes zoster ophthalmicus); Ramsay Hunt (CN VII)'),
        ('Zoster treatment', 'Valaciclovir 1g TDS x 7 days within 72h; post-herpetic neuralgia: gabapentin, tricyclics, capsaicin'),
        ('EBV (HHV-4)', 'Infectious mononucleosis; ampicillin rash (90% if given amoxicillin); EBV-associated: Burkitt lymphoma, NPC, hairy leukoplakia, oral hairy leukoplakia, PTLD'),
        ('CMV (HHV-5)', 'Immunocompromised: perianal/oral ulcers, retinitis; "owl-eye" intranuclear inclusions; treat: ganciclovir/valganciclovir'),
        ('HHV-6 & HHV-7', 'Roseola infantum (exanthem subitum); high fever then maculopapular rash; also linked to pityriasis rosea'),
        ('HHV-8 (KSHV)', 'Kaposi sarcoma; PEL (primary effusion lymphoma); Castleman disease; CD34+ spindle cells; HHV-8 LNA-1 IHC'),
        ('Kaposi Sarcoma types', 'Classic (elderly Mediterranean men); Endemic (Africa); Iatrogenic (post-transplant); AIDS-related (CD4 <200)'),
        ('DIF patterns', 'Herpesvirus family: direct immunofluorescence not routinely used; Tzanck smear shows multinucleated giant cells'),
    ]),
    # -------- Hyperpigmentation --------
    ('Hyperpigmentation.docx', 'Hyperpigmentation - Melasma, Post-Inflammatory & Drug-Induced', [
        ('Melasma', 'UV, hormones (OCP, pregnancy), genetics; brown patches on cheeks/forehead/upper lip; dermal vs epidermal (Wood lamp: epidermal enhances)'),
        ('Melasma treatment', 'Triple cream (hydroquinone 4% + tretinoin 0.05% + fluocinolone 0.01%); tranexamic acid (oral/topical); azelaic acid; chemical peels; laser (risk of PIH)'),
        ('Post-Inflammatory Hyperpigmentation (PIH)', 'Melanin deposition following inflammation; darker skin types more susceptible; treat: hydroquinone, kojic acid, azelaic acid, sunscreen'),
        ('Drug-induced hyperpigmentation', 'Amiodarone (slate-grey, face/sun-exposed); minocycline (blue-grey, scars/shins/oral); antimalarials (blue-grey); clofazimine (red-brown); silver/argyria (blue-grey periorbital)'),
        ('Minocycline pigmentation types', 'Type I: facial acne scars (iron chelates); Type II: shins/forearms (haemosiderin); Type III: diffuse sun-exposed (melanin)'),
        ('Ochronosis (exogenous)', 'Hydroquinone overuse; paradoxical blue-black pigmentation; banana-shaped deposits on histology'),
    ]),

    # -------- Ichthyosis (large multi-part) --------
    ('Ichthyosis.docx', 'Ichthyoses - Comprehensive Multi-Part Review', [
        ('Ichthyosis Vulgaris', 'AD; FLG (filaggrin); fine white scales; spares flexures; KP; hyperlinear palms'),
        ('X-linked Recessive Ichthyosis (SSD)', 'STS gene; large brown scales; spares flexures; cryptorchidism; corneal opacities; associated with Kallmann syndrome'),
        ('Lamellar Ichthyosis', 'AR; TGM1 (most common); collodion baby; large dark scales; ectropion; eclabion'),
        ('Non-bullous CIE', 'AR; fine white scales + erythroderma'),
        ('Epidermolytic Ichthyosis', 'AD; KRT1/KRT10; bullous at birth; cobblestone hyperkeratosis; malodorous'),
        ('Netherton Syndrome', 'SPINK5; ichthyosis linearis circumflexa; bamboo hair; atopic; ↑ IgE'),
        ('Chanarin-Dorfman', 'ABHD5; neutral lipid storage; ↑ LFTs; hepatomegaly; EM: lipid vacuoles'),
        ('Refsum Disease', 'PEX7/PHYH; ↑ phytanic acid; ichthyosis + cerebellar ataxia + retinitis pigmentosa + deafness'),
        ('KID Syndrome', 'GJB2 (Cx26); keratitis + ichthyosis + deafness; ↑ SCC risk'),
        ('Erythrokeratodermia Variabilis', 'GJB3/GJB4; transient erythematous patches + fixed hyperkeratotic plaques'),
        ('Collodion baby DDx', 'Lamellar ichthyosis, non-bullous CIE, Sjögren-Larsson, CHILD, Gaucher, Netherton'),
        ('Treatment', 'Emollients; keratolytics (urea, lactic acid, salicylic acid); retinoids (acitretin) for severe forms'),
    ]),
    # -------- Incontinentia Pigmenti --------
    ('Incontinentia_Pigmenti.docx', 'Incontinentia Pigmenti', [
        ('Gene', 'NEMO (IKBKG); prevents NF-κB activation → ↓ TNF-α-mediated apoptosis'),
        ('Inheritance', 'X-linked dominant; lethal in males (except Klinefelter 47,XXY)'),
        ('Pattern', 'Functional mosaicism via X-inactivation (lyonisation) → Blaschkoid distribution'),
        ('Stage 1 Vesicular', 'Eosinophilic spongiosis; apoptotic keratinocytes; peripheral eosinophilia'),
        ('Stage 2 Verrucous', 'Acanthosis; squamous eddies'),
        ('Stage 3 Hyperpigmented', 'Pigment incontinence; melanophages ("splashed paint" along Blaschko lines)'),
        ('Stage 4 Hypopigmented/Atrophic', 'Epidermal atrophy; melanin loss; alopecia'),
        ('Teeth', 'Pegged/conical teeth; anodontia (~50%); MOST common extracutaneous feature'),
        ('CNS', 'Seizures; developmental delay'),
        ('Eyes', '↓ visual acuity; retinal vascular abnormalities'),
        ('Nails', 'Subungual SCC-like tumours (~10%)'),
    ]),

    # -------- Localized Cutaneous Amyloidosis --------
    ('Localized_Cutaneous_Amyloidosis.docx', 'Localised Cutaneous Amyloidosis', [
        ('Lichen Amyloidosis', 'Keratin (K5) deposition; intensely pruritic lichenified papules/plaques on shins; associated with MEN 2A; Darier-Roussy bodies on EM'),
        ('Macular Amyloidosis', 'Keratin (K5); rippled hyperpigmented patches on interscapular upper back; mast cells in papillary dermis'),
        ('Nodular Amyloidosis', 'AL (immunoglobulin light chain); single/multiple waxy nodules; ~7% progress to systemic amyloidosis; associated with CTDs'),
        ('Primary localised cutaneous', 'Epidermis-derived keratin; NOT associated with systemic disease (lichen/macular forms)'),
        ('Histology', 'Eosinophilic amorphous deposits in papillary dermis; Congo red + (apple-green birefringence); crystal violet +; thioflavin T +'),
        ('Treatment', 'Topical steroids; DMSO; narrowband UVB; dermabrasion; laser for lichen amyloidosis'),
    ]),
    # -------- Lupus Autoantibodies --------
    ('Lupus_Erythematosus_Autoantibodies.docx', 'Lupus Erythematosus - Autoantibodies & Subtypes', [
        ('ANA', 'Most sensitive (>95% SLE); homogeneous IIF pattern most common; also speckled, peripheral, nucleolar'),
        ('Anti-dsDNA', 'Highly specific for SLE (~70%); correlates with disease activity + lupus nephritis'),
        ('Anti-Sm', 'Highly specific for SLE (~30%); does not correlate with activity'),
        ('Anti-Ro/SSA', 'SCLE (~90%); neonatal lupus (heart block); Sjögren; HLA-B8/DR3'),
        ('Anti-La/SSB', 'Associated with Ro antibodies; Sjögren; neonatal lupus (lesser degree)'),
        ('Anti-histone', 'Drug-induced lupus (procainamide, hydralazine, isoniazid, methyldopa); >90% sensitivity'),
        ('Anti-phospholipid (aPL)', 'APS: arterial/venous thrombosis, recurrent miscarriages; livedo reticularis; anti-cardiolipin, anti-β2-GPI, lupus anticoagulant'),
        ('Anti-Scl-70', 'Diffuse systemic sclerosis + ILD (NOT lupus)'),
        ('Complement', '↓ C3, C4 in active SLE (consumption); ↓ C1q in HUV'),
        ('SLE criteria (ACR/EULAR 2019)', 'ANA ≥1:80 required entry criterion; 10 domains; ≥10 points = SLE'),
    ]),
    # -------- Melanoma --------
    ('Melanoma.docx', 'Melanoma - Subtypes, Genetics & Treatment', [
        ('Superficial Spreading Melanoma (SSM)', 'Most common (70%); radial growth phase; BRAF V600E most frequent mutation'),
        ('Nodular Melanoma', 'Most aggressive; vertical growth only; no radial phase; frequently amelanotic'),
        ('Lentigo Maligna Melanoma', 'Sun-damaged skin (face); elderly; in-situ phase = lentigo maligna; NRAS mutations more common'),
        ('Acral Lentiginous Melanoma (ALM)', 'Palms/soles/nail beds; KIT mutations; most common in Asian + African populations; Hutchinson sign (periungual pigment)'),
        ('Desmoplastic Melanoma', 'Fibromatous stroma; perineural invasion; spindled cells; S100+ (most reliable marker); HMB-45 often negative'),
        ('Breslow thickness', 'Most important prognostic factor; T1 <1mm; T4 >4mm'),
        ('Ulceration', 'Adverse prognostic factor; upstages T category'),
        ('Sentinel lymph node biopsy', 'For tumours ≥0.8mm or ulcerated T1b; melanoma ≥1mm (T2+)'),
        ('BRAF V600E', '40-50% of melanoma; treated with BRAF inhibitors (vemurafenib, dabrafenib) ± MEK inhibitors (trametinib)'),
        ('Immunotherapy', 'PD-1 inhibitors (pembrolizumab, nivolumab); CTLA-4 inhibitor (ipilimumab); combo first-line advanced disease'),
        ('IHC markers', 'S100 (most sensitive); SOX10; MART-1/Melan-A; HMB-45 (specific but less sensitive)'),
    ]),

    # -------- Modern Men (Non-scarring alopecias) --------
    ('Modern-Men-Script.docx', 'Androgenetic Alopecia & Non-Scarring Alopecias', [
        ('Androgenetic Alopecia (AGA) - Male', 'Hamilton-Norwood scale; bitemporal recession → vertex thinning; DHT-mediated miniaturisation; AR polygenic'),
        ('Androgenetic Alopecia - Female', 'Ludwig scale; diffuse crown thinning; preserved frontal hairline; check androgens + DHEA-S + ferritin'),
        ('AGA treatment', 'Minoxidil (topical/oral); finasteride (5-alpha reductase inhibitor - males); dutasteride; platelet-rich plasma'),
        ('Telogen Effluvium', 'Shedding 2-3 months after trigger (surgery, illness, delivery, weight loss, hypothyroidism, iron deficiency)'),
        ('Alopecia Areata', 'Autoimmune; CD8+ T cell attack on hair follicle; exclamation mark hairs; nail pitting; totalis/universalis; JAK inhibitors (baricitinib, ritlecitinib)'),
        ('Traction Alopecia', 'Chronic tension; marginal/temporal hairline; reversible if caught early'),
        ('Trichotillomania', 'Compulsive hair pulling; irregular patches; broken hairs of different lengths; psychiatric comorbidity'),
    ]),
    # -------- Mycosis Fungoides --------
    ('Mycosis_Fungoides.docx', 'Mycosis Fungoides (MF) & Sézary Syndrome', [
        ('Cell of origin', 'CD4+ memory T helper cells (CD45RO+)'),
        ('Typical phenotype', 'CD3+, CD4+; loss of CD7 (most common) and CD26'),
        ('Stages', 'Patch → Plaque → Tumour → Erythrodermic'),
        ('Patch stage', 'Erythematous scaly patches in sun-protected/bathing trunk distribution'),
        ('Histology', 'Epidermotropism; Pautrier microabscesses; cerebriform nuclei'),
        ('Folliculotropic MF', 'Head/neck; alopecia; follicular plugging; deeper infiltrate → PUVA less effective'),
        ('Pagetoid Reticulosis (Woringer-Kolopp)', 'Solitary indolent patch/plaque; excellent prognosis'),
        ('Granulomatous Slack Skin', 'Axilla/groin; progressive laxity; associated with lymphoma'),
        ('Sézary Syndrome', 'Erythroderma + lymphadenopathy + Sézary cells ≥1000/µL or ≥10% of lymphocytes'),
        ('Staging (TNMB)', 'T (skin extent) + N (nodes) + M (visceral) + B (blood)'),
        ('Prognosis', 'IA-IIA: 80-90% 5-yr survival; Stage IV: <20%'),
        ('Treatment', 'Skin-directed (early): topical steroids, mechlorethamine, bexarotene, PUVA, TSEBT; Systemic (advanced): bexarotene, mogamulizumab, romidepsin, ECP'),
    ]),
    # -------- Nail-it-down --------
    ('Nail-it-down-script.docx', 'Nail Disorders & Nail Signs', [
        ('Yellow Nail Syndrome', 'Yellow, slow-growing, thickened nails + lymphoedema + pleural effusions/bronchiectasis; FOXC2 gene'),
        ('Terry nails', 'Proximal white + distal pink/brown band; hepatic cirrhosis, CHF, diabetes'),
        ('Half-and-half nails (Lindsay)', 'Proximal white + distal brown; chronic renal failure'),
        ('Muehrcke lines', 'Paired transverse white bands (leukonychia); hypoalbuminaemia'),
        ('Beau lines', 'Transverse ridges; systemic illness, chemotherapy'),
        ('Onycholysis', 'Distal/lateral nail plate separation; psoriasis, thyroid disease, trauma, tetracyclines + UV'),
        ('Subungual melanoma', 'Hutchinson sign (periungual pigmentation); biopsy rule: any single dark band >3mm, new in adult, proximal widening'),
        ('Psoriatic nails', 'Pitting (#1 nail psoriasis finding), oil drop, onycholysis, subungual hyperkeratosis, splinter haemorrhages'),
        ('Lichen Planus nails', 'Pterygium unguis (pathognomonic); twenty-nail dystrophy'),
    ]),

    # -------- Neurocutaneous Dermatoses --------
    ('Neurocutaneous_Dermatoses.docx', 'Neurocutaneous Dermatoses (Phakomatoses)', [
        ('NF1', 'AD; chr 17; neurofibromin (Ras-GAP); ≥6 café-au-lait (≥5mm prepuberty); axillary/groin freckling (Crowe sign); Lisch nodules; neurofibromas; optic gliomas; learning disability'),
        ('NF1 cancers', 'MPNST (malignant peripheral nerve sheath tumour); gastrointestinal stromal tumour (GIST); leukaemia; optic glioma'),
        ('NF2', 'AD; chr 22; merlin/schwannomin; bilateral vestibular schwannomas; meningiomas; ependymomas; few skin findings; juvenile cataracts'),
        ('Tuberous Sclerosis', 'AD; TSC1/TSC2; mTOR pathway; ash-leaf macules (earliest), shagreen patch, facial angiofibromas, periungual fibromas (Koenen); seizures; cardiac rhabdomyomas; renal AML; pulmonary LAM'),
        ('Sturge-Weber', 'Sporadic; GNAQ somatic; port-wine stain V1 + leptomeningeal vascular anomaly; tram-track calcifications; glaucoma; seizures; aspirin prophylaxis'),
        ('VHL', 'AD; VHL gene; retinal/cerebellar haemangioblastomas; clear cell RCC; phaeochromocytoma; pancreatic cysts'),
        ('Gorlin Syndrome', 'AD; PTCH1; multiple BCCs; odontogenic keratocysts; bifid ribs; falx calcification; palmar/plantar pits; medulloblastoma (<3y)'),
        ('Cowden Syndrome', 'AD; PTEN; trichilemmomas; sclerotic fibromas; Lhermitte-Duclos; breast/thyroid/endometrial cancers'),
    ]),
    # -------- Non-Scarring Alopecias --------
    ('Non-Scarring_Alopecias.docx', 'Non-Scarring Alopecias - Mind Map', [
        ('Androgenetic Alopecia', 'Most common; DHT-mediated miniaturisation; Hamilton-Norwood (M); Ludwig (F)'),
        ('Alopecia Areata', 'Autoimmune (CD8+ T); exclamation mark hairs; nail pitting; totalis/universalis; JAK inhibitors'),
        ('Telogen Effluvium', 'Diffuse shedding 2-3m after trigger; self-limiting; correct underlying cause'),
        ('Anagen Effluvium', 'Chemotherapy; radiation; rapid onset; regrowth after cessation'),
        ('Trichotillomania', 'Compulsive pulling; irregular patches; varying hair lengths'),
        ('Traction Alopecia', 'Chronic tension; temporal/marginal; reversible if early'),
        ('Tinea Capitis', 'Dermatophyte; children; oral antifungals required; griseofulvin (Microsporum); terbinafine (Trichophyton)'),
        ('Loose Anagen Syndrome', 'Children; hair pulled painlessly; anagen hairs lack root sheaths on trichogram'),
        ('Secondary syphilis', 'Moth-eaten alopecia; associated with non-scarring patchy hair loss'),
    ]),
    # -------- Paraneoplastic Syndromes --------
    ('Paraneoplastic_Syndromes_-_Cutaneous_Metastases.docx', 'Paraneoplastic Syndromes & Cutaneous Metastases', [
        ('Acanthosis Nigricans (malignant)', 'Rapid onset + extensive; stomach adenocarcinoma most common; TGF-α/EGF-related'),
        ('Tripe Palms', 'Rugose velvety palms; lung > gastric cancer; often coexists with AN'),
        ('Sign of Leser-Trelat', 'Sudden eruption of multiple seborrhoeic keratoses; gastric/colorectal/lung cancer; GI most common'),
        ('Bazex Syndrome (Acrokeratosis paraneoplastica)', 'Psoriasiform plaques on acral sites (ears, nose, fingers); squamous cell carcinoma of upper aerodigestive tract'),
        ('Necrolytic Migratory Erythema (NME)', 'Glucagonoma; 4Ds: Diabetes, DVT, Depression, Diarrhoea; periorificial/intertriginous crusted erosions'),
        ('Sweet Syndrome (paraneoplastic)', '~20% malignancy (AML most common); can precede diagnosis'),
        ('Pyoderma Gangrenosum', 'IBD; haematologic malignancies (bullous form = AML)'),
        ('Cutaneous Metastases', 'Most common primaries: breast (women), lung (men); Sister Mary Joseph nodule = periumbilical metastasis (gastric/ovarian)'),
        ('Inflammatory carcinoma', 'Breast; peau d\'orange; dermal lymphatic invasion; not a true panniculitis'),
        ('Carcinoid syndrome', 'Episodic flushing (non-pruritic); 5-HT; small bowel NET most common'),
    ]),

    # -------- Bullous Pemphigoid basic science --------
    ('Pemphigoid-basic-science-script.docx', 'Bullous Pemphigoid - Basic Science & Basement Membrane Zone', [
        ('Target antigen 1', 'BPAG2 (BP180 / Collagen XVII); transmembrane; NC16A domain = most immunogenic epitope'),
        ('Target antigen 2', 'BPAG1 (BP230); intracellular plakin family protein'),
        ('BMZ layers', 'Lamina lucida (above lamina densa; α6β4 integrin + laminin 332); Lamina densa (type IV collagen + nidogen); Sub-lamina densa (type VII collagen = target in EBA/DEB)'),
        ('Split level in BP', 'Sub-epidermal within lamina lucida'),
        ('Pathogenesis', 'IgG (+ IgE) against NC16A of BP180 → complement activation → eosinophil/neutrophil recruitment → protease release → blister'),
        ('Salt-split skin DIF', 'BP: roof (epidermal) binding; EBA: floor (dermal) binding'),
        ('ELISA', 'Anti-BP180 NC16A correlates with disease activity; anti-BP230'),
        ('IIF substrate', 'Normal human skin or monkey oesophagus'),
    ]),
    # -------- Pemphigus --------
    ('Pemphigus-script-corrected.docx', 'Pemphigus Vulgaris & Variants', [
        ('Pathogenesis', 'IgG autoantibodies against desmogleins → loss of intercellular adhesion → intraepidermal acantholysis'),
        ('PV antigens', 'Mucosal-dominant: anti-Dsg3 only; Mucocutaneous: anti-Dsg3 + anti-Dsg1'),
        ('Demographics', 'Adults 40-60; ↑ Ashkenazi Jewish, Mediterranean, Indian; HLA-DR4, HLA-DR14'),
        ('Clinical', 'Painful flaccid blisters → erosions; mucosal involvement >90% (often first)'),
        ('Signs', 'Positive Nikolsky sign; positive Asboe-Hansen sign'),
        ('Histology', 'Suprabasal acantholysis; "tombstoning" of basal cells'),
        ('DIF', 'Intercellular IgG + C3 ("chicken wire/fishnet" pattern)'),
        ('Pemphigus Foliaceus', 'Anti-Dsg1 only; NO mucosal involvement; superficial blisters; Fogo Selvagem = endemic PF in Brazil'),
        ('IgA Pemphigus', 'Anti-desmocollin 1 (subcorneal) or anti-Dsg3 (intraepidermal); neutrophilic spongiosis'),
        ('PNP', 'Antibodies to multiple plakins (envoplakin, periplakin, desmoplakin); NHL/CLL/Castleman/thymoma; bronchiolitis obliterans'),
        ('Treatment', 'Systemic steroids + rituximab (first-line per RCT); azathioprine, MMF, IVIG, cyclophosphamide'),
    ]),
    # -------- Porphyrias --------
    ('Porphyrias.docx', 'Porphyrias - Comprehensive Review', [
        ('Rate-limiting enzyme', 'ALA synthetase (ALAS)'),
        ('PCT (Porphyria Cutanea Tarda)', 'UROD deficiency; most common; adults; hypertrichosis + vesicles/bullae dorsal hands + milia + hyperpigmentation; triggers: Hep C, iron, HIV, EtOH, oestrogen'),
        ('PCT treatment', 'Phlebotomy (primary); low-dose hydroxychloroquine; photoprotection'),
        ('PCT histology', 'Cell-poor sub-epidermal blister; festooning; caterpillar bodies; linear Ig/C3 at DEJ'),
        ('HEP', 'Homozygous UROD; childhood; severe PCT-like + sclerodermoid changes'),
        ('CEP (Gunther disease)', 'UROS3; erythrodontia (red teeth, Wood lamp); red urine; severe scarring/mutilation; haemolytic anaemia; splenomegaly'),
        ('EPP', 'Ferrochelatase deficiency; most common childhood porphyria; burning pain 3-30 min after sun; linear facial pits; photo-onycholysis; treat: beta-carotene'),
        ('AIP (Acute Intermittent Porphyria)', 'HMBS; acute attacks: abdominal pain + neuropathy + psychiatric (5-P); avoid: EtOH, OCP, sulfonamides, barbiturates; treat: hemin + glucose'),
    ]),

    # -------- Primary Systemic Amyloidosis --------
    ('Primary_Systemic_Amyloidosis.docx', 'Primary Systemic Amyloidosis', [
        ('Protein', 'AL (amyloid light chain); IgG lambda most common'),
        ('Underlying disease', 'Plasma cell dyscrasia; multiple myeloma; Waldenström macroglobulinaemia'),
        ('Workup', 'SPEP, UPEP, serum free light chains, bone marrow biopsy'),
        ('Skin findings', 'Pinch purpura (periocular/perioral); waxy papules/plaques/nodules; sclerodermoid changes; macroglossia; alopecia'),
        ('Worst prognosis', 'Cardiac involvement (amyloid cardiomyopathy → restrictive)'),
        ('Other organs', 'Peripheral neuropathy; carpal tunnel syndrome; nephrotic syndrome; hepatomegaly'),
        ('Treatment', 'Daratumumab + bortezomib + dexamethasone; autologous stem cell transplant in eligible patients'),
    ]),
    # -------- Pruritus --------
    ('Pruritus.docx', 'Pruritus - Mechanisms, Mediators & Treatment', [
        ('Definition', 'Unpleasant sensation provoking the desire to scratch; itch ≠ pain (separate C-fibres)'),
        ('Key mediator', 'Histamine (mast cells); also substance P, IL-31, IL-4/13, TSLP, PAR-2, serotonin, opioids'),
        ('IL-31', 'Key cytokine in atopic itch; nemolizumab (anti-IL-31RA) approved for AD + PN'),
        ('Itch pathways', 'Peripheral C-fibres → spinal cord → thalamus → somatosensory cortex; GRPR (gastrin-releasing peptide receptor) in spinal cord'),
        ('Systemic causes', 'Renal failure (uraemic pruritus), cholestasis, polycythaemia vera (aquagenic), iron deficiency, thyroid disease, lymphoma, HIV'),
        ('Aquagenic pruritus', 'PV (most important); water triggers; prostacyclin/prostaglandin D2 hypothesis'),
        ('Prurigo Nodularis', 'Chronic scratching → lichenified nodules; ↑ IL-31; dupilumab; nemolizumab approved'),
        ('Cholestatic itch', 'Bile salts + LPA + lysophosphatidic acid + endogenous opioids; rifampicin, naltrexone, sertraline, cholestyramine'),
        ('Uraemic itch', 'Dialysis patients; gabapentin/pregabalin (first-line); difelikefalin (kappa opioid agonist - FDA approved)'),
        ('Treatment ladder', 'Emollients → topical (steroids, calcineurin inhibitors, menthol) → antihistamines (limited) → gabapentin/pregabalin → dupilumab/nemolizumab'),
    ]),
    # -------- Psoriasis --------
    ('Psoriasis.docx', 'Psoriasis - Genetics, Immunopathology & Treatment', [
        ('Strongest HLA', 'HLA-Cw6 (15-fold ↑ risk); PSORS1 locus most important'),
        ('HLA subtypes', 'HLA-B17 (erythrodermic/guttate); HLA-B13 (guttate); HLA-B27 (PsA/pustular)'),
        ('Cytokine pathway', 'DC → IL-23 → Th17 → IL-17 + IL-22 → keratinocyte proliferation'),
        ('Decreased cytokine', 'IL-10 (anti-inflammatory)'),
        ('Köbner phenomenon', 'Lesions at trauma sites; 2-6 week lag'),
        ('Drug triggers', 'NSAIDs/terbinafine (fast <4w); HCQ/ACEi (4-12w); beta-blockers/lithium (>12w)'),
        ('Paradoxical psoriasis', 'TNF-α inhibitors → psoriasis (especially palmoplantar pustulosis)'),
        ('Histology', 'Munro microabscesses (parakeratosis); Kogoj spongiform pustules (spinous layer); dilated dermal vessels'),
        ('Biologic targets', 'IL-17A: secukinumab, ixekizumab; IL-17R: brodalumab; IL-23 p19: guselkumab, risankizumab, tildrakizumab; IL-12/23 p40: ustekinumab; TNF-α: adalimumab, etanercept'),
        ('Pustular psoriasis', 'CARD14 (Type 5 PRP-like); IL36RN; GPP → systemic illness; treat with spesolimab (anti-IL-36R)'),
    ]),

    # -------- Psychodermatoses --------
    ('Psychodermatoses.docx', 'Psychodermatoses', [
        ('Primary (psychiatric → skin)', 'Delusions of parasitosis (Ekbom syndrome); trichotillomania; dermatitis artefacta; onychotillomania; neurotic excoriations'),
        ('Delusions of Parasitosis', 'Monosymptomatic hypochondriacal psychosis; Matchbox sign (specimens brought to clinician); treat: pimozide (1st gen antipsychotic) or risperidone/olanzapine (2nd gen)'),
        ('Secondary (skin disease → psychiatric)', 'Depression in severe psoriasis/eczema; social anxiety in vitiligo/alopecia; PTSD from burns'),
        ('Tertiary (coexisting)', 'Psoriasis + depression; atopic dermatitis + anxiety (bidirectional neuro-immune axis)'),
        ('Dermatitis Artefacta', 'Self-inflicted; angular/geometric lesions; patient denies; on accessible sites; "hollow history"'),
        ('Body Dysmorphic Disorder (BDD)', 'Perceived skin flaw; common in cosmetic dermatology patients; SSRIs first-line; avoid surgical intervention'),
        ('Neurotic Excoriations', 'OCD spectrum; repetitive picking; treat with SSRIs, NAC (N-acetylcysteine), habit reversal therapy'),
    ]),
    # -------- Rothmund-Thomson --------
    ('Rothmund-Thomson_Syndrome.docx', 'Rothmund-Thomson Syndrome', [
        ('Inheritance', 'Autosomal recessive'),
        ('Gene', 'RECQL4 (DNA helicase)'),
        ('Skin', 'Erythema/oedema cheeks at 3-6 months → blistering → poikiloderma (atrophy + telangiectasia + hypo/hyperpigmentation); acral verrucous keratoses'),
        ('Skeletal', 'Short stature; absent/hypoplastic thumbs; radial aplasia'),
        ('Eyes', 'Juvenile bilateral cataracts (by age 5); characteristic feature'),
        ('Endocrine', 'Hypogonadism; sparse hair'),
        ('Cancer risk', 'Osteosarcoma (#1); SCC'),
    ]),
    # -------- Running with the Bulls (Vasculitis) --------
    ('Running-with-the-bulls-script.docx', 'Medium and Large Vessel Vasculitis', [
        ('PAN (Polyarteritis Nodosa)', 'Medium vessel; skin: livedo reticularis + palpable purpura + nodules + ulcers; internal organ infarction; HBV association; ANCA NEGATIVE; treat: cyclophosphamide + steroids'),
        ('Kawasaki Disease', 'Children <5 years; coronary artery aneurysms (untreated 25%); fever ≥5d + 4 of 5 criteria; IVIG + aspirin'),
        ('Giant Cell Arteritis (GCA)', 'Women >50y; temporal artery; jaw claudication; visual loss (anterior ischaemic optic neuropathy); ↑ ESR + CRP; temporal artery biopsy; treat: high-dose prednisolone + tocilizumab'),
        ('Takayasu Arteritis', 'Young Asian women; aorta + major branches; pulseless disease; arm BP discrepancy; treat: steroids'),
        ('Buerger Disease (TAO)', 'Young male smokers; distal vessel occlusion; digital ischaemia; smoking cessation critical'),
        ('Behçet Disease', 'Silk Road; HLA-B51; oral ulcers (≥3/year) + genital ulcers + uveitis + skin; pathergy test'),
        ('EGPA (Churg-Strauss)', 'Asthma + eosinophilia + sinusitis; p-ANCA (MPO); skin: palpable purpura + subcutaneous nodules'),
        ('GPA (Granulomatosis with Polyangiitis)', 'c-ANCA (PR3); saddle nose; upper + lower respiratory + renal; treat: rituximab'),
    ]),

    # -------- Seborrhoeic Dermatitis --------
    ('Seborrheic_Dermatitis.docx', 'Seborrhoeic Dermatitis', [
        ('Aetiology', 'Malassezia furfur (globosa/restricta); abnormal host response; sebum-rich areas'),
        ('Distribution', 'Scalp, nasolabial folds, glabella, retroauricular, central chest, intertriginous areas'),
        ('HIV association', 'Severe/refractory seb derm = marker of HIV; flares with CD4 <100'),
        ('Parkinsons association', 'Sebum overproduction; seborrhoeic dermatitis common'),
        ('Histology', 'Psoriasiform hyperplasia; follicular plugging; shoulder parakeratosis (at follicular ostia); sparse perivascular lymphocytes; neutrophils in stratum corneum'),
        ('Infantile (cradle cap)', 'Thick yellow-brown waxy scale; scalp + flexures; self-limiting'),
        ('Treatment adults', 'Antifungals: ketoconazole shampoo/cream, selenium sulfide, zinc pyrithione; topical steroids (mild-moderate); tacrolimus; ciclopirox'),
        ('Dandruff', 'Mild form; ketoconazole or pyrithione zinc shampoo'),
    ]),
    # -------- Secondary Systemic Amyloidosis --------
    ('Secondary_Systemic_Amyloidosis.docx', 'Secondary Systemic Amyloidosis (AA Amyloidosis)', [
        ('Protein', 'AA (serum amyloid A); acute phase reactant produced by liver in chronic inflammation'),
        ('Causes', 'Rheumatoid arthritis (most common in developed world); ankylosing spondylitis; Castleman disease; TB; osteomyelitis; IBD; FMF; Muckle-Wells syndrome; dystrophic EB'),
        ('Skin findings', 'Less prominent than primary; macroglossia rare; petechiae/purpura possible'),
        ('Main organs affected', 'Kidney (nephrotic syndrome → renal failure, most common complication); liver; spleen'),
        ('Distinguish from AL', 'AA: no monoclonal protein; Congo red staining of rectum/abdominal fat; SAA level elevated'),
        ('Treatment', 'Treat underlying disease; colchicine (FMF); canakinumab (anti-IL-1β for CAPS/FMF); eprodisate (investigational)'),
    ]),
    # -------- STIs --------
    ('Sexually_Transmitted_Bacterial_Infections.docx', 'Sexually Transmitted Bacterial Infections', [
        ('Syphilis (Primary)', 'Treponema pallidum; painless indurated chancre; heals in 3-6 weeks'),
        ('Syphilis (Secondary)', 'Maculopapular rash palms/soles; condylomata lata; moth-eaten alopecia; mucous patches'),
        ('Syphilis (Tertiary)', 'Gummas; cardiovascular (aortitis); neurosyphilis (Argyll Robertson pupils, tabes dorsalis)'),
        ('Syphilis serology', 'Non-treponemal: RPR/VDRL (quantitative, follow activity); Treponemal: FTA-ABS, TP-PA (remain positive for life)'),
        ('Syphilis treatment', 'Benzathine penicillin G; IV aqueous penicillin for neurosyphilis; doxycycline if penicillin-allergic (except pregnancy)'),
        ('Gonorrhoea (DGI)', 'Neisseria gonorrhoeae; disseminated gonococcal infection: petechial/pustular skin lesions + tenosynovitis + polyarthralgia; treat: ceftriaxone'),
        ('Chlamydia (LGV)', 'L1-L3 serovars; inguinal lymphadenopathy; groove sign; treat: doxycycline 21 days'),
        ('Chancroid', 'Haemophilus ducreyi; painful soft chancre + inguinal lymphadenopathy (bubo); treat: azithromycin or ceftriaxone'),
        ('Granuloma Inguinale (Donovanosis)', 'Klebsiella granulomatis; painless beefy-red ulcers; Donovan bodies (intracellular rod-shaped bacteria in macrophages); treat: azithromycin/doxycycline'),
    ]),

    # -------- Small Vessel Vasculitis (script) --------
    ('Small-vessel-vasculitis-script.docx', 'Small Vessel Vasculitis - Narrative Review', [
        ('Definition', 'Vasculitis affecting capillaries, post-capillary venules, arterioles'),
        ('Cutaneous findings', 'Palpable purpura (non-blanching); necrosis; urticarial lesions; livedo reticularis'),
        ('Leukocytoclastic vasculitis (LCV)', 'Neutrophilic perivascular infiltrate + nuclear dust (karyorrhexis) + fibrinoid necrosis + RBC extravasation'),
        ('Common triggers', 'Infections (strep, viral hep B/C, HIV); drugs (antibiotics, NSAIDs, diuretics); autoimmune CTDs'),
        ('IgA Vasculitis (HSP)', 'IgA deposits on DIF; children; renal involvement important'),
        ('Cryoglobulinemic vasculitis', 'Hep C (Type II mixed); purpura + arthralgia + neuropathy + GN (Meltzer triad)'),
        ('Urticarial vasculitis', 'Wheals >24h; burns/stings; heals with pigmentation; HUV: low complement'),
        ('ANCA vasculitis', 'GPA (c-ANCA/PR3); MPA (p-ANCA/MPO); EGPA (p-ANCA/MPO + asthma + eosinophilia)'),
        ('Treatment', 'Remove trigger; systemic steroids; dapsone for LCV; immunosuppressants for systemic ANCA vasculitis (rituximab or cyclophosphamide)'),
    ]),
    # -------- Small Vessel Vasculitis Mind Map --------
    ('Small_Vessel_Vasculitis_Mind_Map.docx', 'Small Vessel Vasculitis - Mind Map Overview', [
        ('Classification', 'Immune complex (IgA, cryo, CSVV) vs ANCA-mediated (GPA, MPA, EGPA) vs Direct invasion vs Paraneoplastic'),
        ('GPA', 'c-ANCA (PR3+); upper + lower airway granulomas; necrotising GN; saddle nose; subglottic stenosis'),
        ('MPA', 'p-ANCA (MPO+); pulmonary-renal syndrome; alveolar haemorrhage; pauci-immune GN'),
        ('EGPA (Churg-Strauss)', 'p-ANCA (MPO+, 40%); asthma (mandatory) + eosinophilia + sinusitis; skin nodules on scalp/elbows'),
        ('LCV histology grades', 'Early: neutrophils; Late: mononuclear cells; fibrinoid necrosis throughout'),
        ('Workup', 'CBC, CMP, UA, ESR/CRP, ANA, ANCA, complement, cryoglobulins, hepatitis serologies, RF'),
        ('Skin biopsy timing', 'Within 24-48 hours of new lesion for best DIF/histology yield'),
    ]),
    # -------- SCC --------
    ('Squamous_Cell_Carcinoma.docx', "Squamous Cell Carcinoma & Bowen's Disease", [
        ("Bowen's Disease (SCC in situ)", 'Full-thickness epidermal atypia; p53 mutation; HPV 16/18 association (genital); arsenical keratoses; treat: 5-FU, PDT, cryotherapy, imiquimod'),
        ('Erythroplasia of Queyrat', "Bowen's disease on glans penis; higher malignant potential"),
        ('SCC Risk factors', 'UV exposure; immunosuppression (transplant: 65-250x risk); HPV; chronic wounds (Marjolin ulcer); arsenic; ionising radiation; PUVA'),
        ('Marjolin Ulcer', 'SCC arising in chronic wounds/scars/burn scars; aggressive; poor prognosis'),
        ('High-risk features', 'Diameter >2cm; depth >2mm; perineural invasion; location (lip, ear, scalp); poor differentiation; immunosuppression'),
        ('Staging', 'Brigham & Women Hospital system; T3 = ≥3 high-risk factors; worst prognosis'),
        ('Treatment', 'Wide excision; Mohs for high-risk/head-neck; cemiplimab/pembrolizumab for advanced/metastatic'),
        ('Keratoacanthoma', 'Rapidly growing crateriform tumour; self-involuting; histologically resembles well-differentiated SCC; treat like SCC'),
    ]),

    # -------- Sweet's Arcade --------
    ('Sweets-arcade-narration-.docx', "Sweet's Syndrome (Acute Febrile Neutrophilic Dermatosis)", [
        ('Definition', 'Neutrophilic dermatosis; sterile inflammatory papules/plaques/nodules with fever and leukocytosis'),
        ('Demographics', 'Women 30-60; commonly post-infectious or paraneoplastic'),
        ('Clinical', 'Tender erythematous papules/plaques; pseudovesicular appearance; face/neck/upper extremities'),
        ('Malignancy-associated', '~20%; AML most common; can precede diagnosis by months'),
        ('Drug-induced', 'G-CSF (most common); ATRA; TMP-SMX; azathioprine; furosemide'),
        ('Major diagnostic criteria', 'Abrupt onset tender plaques/nodules + characteristic histology (dense neutrophilic infiltrate, no vasculitis, papillary dermal oedema)'),
        ('Minor criteria', 'Fever >38°C; association (malignancy/IBD/pregnancy/infection); rapid steroid response; 3/4 abnormal labs (↑ESR, CRP, WBC, neutrophilia)'),
        ('Histiocytoid variant', 'Immature myeloid cells; associated with MDS'),
        ('Treatment', 'Systemic corticosteroids (rapid, dramatic response); alternatives: colchicine, dapsone, potassium iodide'),
    ]),
    # -------- The Amazing Amazon River Ride (Scleroderma / SSc) --------
    ('The-Amazing-Amazon-River-Ride-script.docx', 'Systemic Sclerosis (Scleroderma) & Related Conditions', [
        ('Types', 'Diffuse SSc (dcSSc): proximal skin thickening; Limited SSc (lcSSc): CREST syndrome; Sine scleroderma: systemic without skin fibrosis'),
        ('CREST syndrome', 'Calcinosis, Raynaud, Oesophageal dysmotility, Sclerodactyly, Telangiectasias'),
        ('Anti-Scl-70 (topoisomerase I)', 'Diffuse SSc; ↑ ILD risk; poor prognosis'),
        ('Anti-centromere (ACA)', 'Limited SSc (CREST); ↑ PAH risk; better prognosis'),
        ('Anti-RNA polymerase III', 'Diffuse SSc; ↑ renal crisis risk; ↑ malignancy (especially gastric)'),
        ('Anti-U1RNP', 'Mixed CTD (MCTD); overlap syndrome'),
        ('Raynaud phenomenon', 'White → Blue → Red (vasospasm → cyanosis → reperfusion); nifedipine first-line'),
        ('Renal crisis', 'Hypertension + AKI; anti-RNA Pol III; treat: ACE inhibitors (captopril)'),
        ('Pulmonary', 'ILD (dcSSc + anti-Scl-70); PAH (lcSSc + ACA); nintedanib/mycophenolate for ILD'),
        ('Skin', 'Salt-and-pepper pigmentation; mat telangiectasias; calcinosis cutis; sclerodactyly; puffy fingers (early); "hidebound" skin (late)'),
        ('Morphea', 'Localised scleroderma; no systemic involvement; linear (en coup de sabre), generalised, deep forms'),
    ]),
    # -------- Old and the Restless (Werner/Bloom) --------
    ('the-old-and-the-restless-script.docx', 'Premature Ageing Syndromes - Werner & Bloom', [
        ('Werner Syndrome', 'AR; WRN (RECQL2A) DNA helicase; accelerated ageing from ~30y'),
        ('Werner - Skin', 'Sclerodermatous/atrophic changes (face, acral); leg ulcers; bird-like facies; mottled pigmentation (poikiloderma)'),
        ('Werner - Other', 'Primary hypogonadism; bilateral cataracts; diabetes; atherosclerosis; osteoporosis'),
        ('Werner - Cancer', 'Sarcomas (especially osteosarcoma, fibrosarcoma) + melanoma; UNUSUAL spectrum'),
        ('Bloom Syndrome', 'AR; BLM (RECQL3) DNA helicase; ↑ sister chromatid exchanges'),
        ('Bloom - Skin', 'Telangiectatic malar butterfly erythema; photosensitivity; café-au-lait macules; short stature'),
        ('Bloom - Cancer', 'Lymphoma + leukaemia (#1); broad spectrum (GI, skin)'),
        ('Shared features', 'DNA helicase mutations; premature ageing; ↑ malignancy risk; growth retardation'),
    ]),

    # -------- Trichothiodystrophy --------
    ('Trichothiodystrophy_(TTD).docx', 'Trichothiodystrophy (TTD / PIBIDS)', [
        ('Inheritance', 'Autosomal recessive'),
        ('Photosensitive form genes', 'ERCC2 (XPD), ERCC3 (XPB), GTF2H5 (TTDA)'),
        ('Hair finding', 'Brittle hair; low sulphur (cysteine); "tiger tail" banding on polarised microscopy'),
        ('PIBIDS mnemonic', 'Photosensitivity, Ichthyosis, Brittle hair, Intellectual impairment, Decreased fertility, Short stature'),
        ('Key distinction from XP', 'NO increased skin cancer risk in TTD'),
        ('Non-photosensitive form', 'TTDN1 (C7orf11); no ichthyosis; less severe neuro'),
        ('Tay syndrome', 'Ichthyosiform erythroderma variant of TTD'),
    ]),
    # -------- Tuberous Sclerosis --------
    ('Tuberous_Sclerosis.docx', 'Tuberous Sclerosis Complex', [
        ('Inheritance', 'Autosomal dominant; ~60% sporadic'),
        ('Genes', 'TSC1 (chr 9q34, hamartin) or TSC2 (chr 16p13, tuberin); TSC2 more severe'),
        ('Pathway', 'mTOR activation → tissue overgrowth and hamartomas'),
        ('Earliest skin sign', 'Ash-leaf macules (hypomelanotic, ≥3, ≥5mm; present at birth; Wood lamp)'),
        ('Other skin', 'Facial angiofibromas (central face, 2-5y; treat: topical rapamycin); shagreen patch (lower back); periungual fibromas (Koenen, adolescence); forehead fibrous plaque; confetti macules'),
        ('Neurological', 'Cortical tubers; seizures (infantile spasms earliest, ≥80%); SEGA; intellectual disability; autism'),
        ('Cardiac', 'Rhabdomyomas (antenatal; usually regress; WPW risk)'),
        ('Renal', 'Angiomyolipomas (>4cm → haemorrhage risk; ↑ RCC risk; #2 mortality)'),
        ('Pulmonary', 'LAM (women 20-40; cystic lung disease; pneumothorax)'),
        ('Treatment', 'Topical sirolimus (angiofibromas); everolimus/sirolimus (SEGA, AML, LAM, refractory seizures); vigabatrin (infantile spasms)'),
    ]),
    # -------- Ulcers --------
    ('Ulcers.docx', 'Leg Ulcers - Venous, Arterial, Neuropathic & Lymphatic', [
        ('Venous ulcers', 'Most common (70-80%); medial gaiter area; shallow with irregular edges; lipodermatosclerosis; haemosiderin (brown); champagne bottle leg; treat: compression (first-line)'),
        ('Atrophie blanche', 'Ivory-white porcelain scars with telangiectasias; surrounds healed venous ulcers; vasculopathy'),
        ('Arterial ulcers', 'Painful; punched-out; distal (toes/heel); pale + cold limb; absent pulses; ABI <0.9; treat: revascularisation'),
        ('Diabetic/Neuropathic ulcers', 'Pressure points (metatarsal heads, heels); painless; callus surrounding; treat: offloading + wound care + glycaemic control'),
        ('Marjolin ulcer', 'SCC in chronic wound/scar; aggressive'),
        ('Pyoderma Gangrenosum', 'Pathergy; violaceous undermined borders; diagnose by exclusion; treat: steroids + ciclosporin'),
        ('Calciphylaxis', 'CKD/dialysis; uraemic small vessel calcification; retiform purpura → necrosis; treat: sodium thiosulfate + wound care; avoid warfarin'),
        ('ABI (Ankle Brachial Index)', '>1.0 = normal; 0.7-0.9 = mild arterial; <0.4 = severe ischaemia; >1.4 = calcified (diabetic/CKD)'),
    ]),

    # -------- Urticarial Vasculitis --------
    ('Urticarial-vasculitis-script.docx', 'Urticarial Vasculitis', [
        ('Definition', 'Urticarial wheals lasting >24 hours; burns/stings rather than pruritic; heals with purpura/pigmentation'),
        ('Normocomplementemic UV', 'Normal C3/C4; usually idiopathic; less systemic involvement'),
        ('Hypocomplementaemic UV (HUV)', 'Low C1q, C3, C4; anti-C1q antibodies; systemic: arthralgia, GI, renal, pulmonary (COPD), uveitis/episcleritis'),
        ('Associations', 'SLE; Sjögren; Hep C; serum sickness; drug reactions; malignancy (especially HUV + plasma cell dyscrasia)'),
        ('Histology', 'Leukocytoclastic vasculitis (neutrophils + karyorrhexis + fibrinoid necrosis); eosinophils common'),
        ('Workup', 'CBC; complement (C1q, C3, C4); anti-C1q Ab; ANA; hepatitis serologies; SPEP'),
        ('Treatment', 'NSAIDs; antihistamines; dapsone; colchicine; systemic steroids; rituximab/MMF for HUV'),
    ]),
    # -------- Vitiligo --------
    ('Vitiligo.docx', 'Vitiligo', [
        ('Pathogenesis', 'Autoimmune destruction of melanocytes; CD8+ T cells; IFN-γ/CXCL10 axis; JAK-STAT pathway'),
        ('Associations', 'Thyroid disease (most common; check TFTs); T1DM; pernicious anaemia; Addison disease; alopecia areata'),
        ('Segmental vitiligo', 'Unilateral; follows Blaschko lines; earlier onset; more stable; less autoimmune overlap'),
        ('Non-segmental vitiligo', 'Symmetric; progressive; periocular, perioral, acral, genitalia; Koebner phenomenon'),
        ('Repigmentation pattern', 'Perifollicular (from outer root sheath melanocytes) → marginal → diffuse'),
        ('Wood lamp', 'Enhances depigmentation (white fluorescence); useful in light skin types'),
        ('Treatment - first-line', 'Topical steroids or topical tacrolimus/pimecrolimus; narrowband UVB (phototherapy)'),
        ('Systemic', 'Oral mini-pulse steroids (arrest spread); oral ruxolitinib (JAK inhibitor - approved for extensive non-segmental)'),
        ('Topical ruxolitinib (Opzelura)', 'FDA-approved 2022; JAK1/2 inhibitor; repigmentation especially face/sun-exposed'),
        ('Surgical', 'Stable disease only; suction blister grafting; melanocyte transfer'),
    ]),
    # -------- Xeroderma Pigmentosum --------
    ('Xeroderma_Pigmentosum.docx', 'Xeroderma Pigmentosum (XP)', [
        ('Defect', 'Nucleotide excision repair (NER) pathway'),
        ('Variants', 'XPA to XPG and XPV (7 complementation groups + variant)'),
        ('XPV', 'DNA polymerase η (post-replication repair); milder UV sensitivity; still ↑ skin cancer'),
        ('XPB, XPD', 'Also cause TTD (tiger tail hair) and XP-Cockayne overlap (CS + cancer risk)'),
        ('Skin', 'UV sensitivity; severe photodamage; 1000-fold ↑ skin cancer risk under age 20 (BCC, SCC, melanoma)'),
        ('Neurological (20-30%)', 'XPA + XPD; developmental delay; hyporeflexia; ataxia; De Sanctis-Cacchione syndrome (severe variant)'),
        ('Management', 'Strict UV protection; annual skin surveillance; oral retinoids (chemopreventive); sun-protective clothing/film; vismodegib for BCCs'),
    ]),
]



import copy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor


def copy_paragraph_with_formatting(src_para, out_doc):
    """
    Copy a source paragraph into out_doc preserving:
      - run-level colour, bold, italic, underline, font size
      - inline images (drawings / blipFill elements)
    Returns the new paragraph object.
    """
    new_para = out_doc.add_paragraph()
    # Copy paragraph-level alignment
    if src_para.alignment is not None:
        new_para.alignment = src_para.alignment

    src_xml = src_para._element

    for child in src_xml:
        tag = child.tag

        # --- inline run (w:r) ---
        if tag == qn('w:r'):
            # Check if this run contains an image (drawing or pict)
            drawing = child.find('.//' + qn('w:drawing'))
            pict    = child.find('.//' + qn('w:pict'))

            if drawing is not None or pict is not None:
                # Deep-copy the entire run XML (image + rPr) into new paragraph
                new_run_elem = copy.deepcopy(child)
                new_para._element.append(new_run_elem)
                continue

            # Plain text run — copy text + formatting
            new_run = new_para.add_run()
            # Copy run properties (rPr)
            rpr = child.find(qn('w:rPr'))
            if rpr is not None:
                new_run._element.insert(0, copy.deepcopy(rpr))

            # Copy w:t text nodes
            for t_elem in child.findall(qn('w:t')):
                new_run.text += (t_elem.text or '')

            # Apply formatting directly to run object as well (belt-and-braces)
            if rpr is not None:
                # Bold
                b = rpr.find(qn('w:b'))
                if b is not None:
                    new_run.bold = True
                # Italic
                it = rpr.find(qn('w:i'))
                if it is not None:
                    new_run.italic = True
                # Underline
                u = rpr.find(qn('w:u'))
                if u is not None and u.get(qn('w:val'), '') not in ('none', ''):
                    new_run.underline = True
                # Colour
                color_elem = rpr.find(qn('w:color'))
                if color_elem is not None:
                    val = color_elem.get(qn('w:val'))
                    if val and val.lower() not in ('auto', '000000', ''):
                        try:
                            r = int(val[0:2], 16)
                            g = int(val[2:4], 16)
                            b_val = int(val[4:6], 16)
                            new_run.font.color.rgb = RGBColor(r, g, b_val)
                        except Exception:
                            pass
                # Font size
                sz = rpr.find(qn('w:sz'))
                if sz is not None:
                    half_pts = sz.get(qn('w:val'))
                    if half_pts:
                        try:
                            new_run.font.size = Pt(int(half_pts) / 2)
                        except Exception:
                            pass

        # --- hyperlink (w:hyperlink) — treat as plain text preserving colour ---
        elif tag == qn('w:hyperlink'):
            for sub_run in child.findall(qn('w:r')):
                new_run = new_para.add_run()
                rpr = sub_run.find(qn('w:rPr'))
                if rpr is not None:
                    new_run._element.insert(0, copy.deepcopy(rpr))
                for t_elem in sub_run.findall(qn('w:t')):
                    new_run.text += (t_elem.text or '')
                if rpr is not None:
                    color_elem = rpr.find(qn('w:color'))
                    if color_elem is not None:
                        val = color_elem.get(qn('w:val'))
                        if val and val.lower() not in ('auto', '000000', ''):
                            try:
                                r = int(val[0:2], 16)
                                g = int(val[2:4], 16)
                                b_val = int(val[4:6], 16)
                                new_run.font.color.rgb = RGBColor(r, g, b_val)
                            except Exception:
                                pass

        # --- bookmark start/end, proofErr, etc. — skip ---

    return new_para


def copy_docx_into_doc(src_path, out_doc):
    """
    Read a .docx source file and copy all its paragraphs (with colour + images)
    into out_doc, preserving run-level formatting.
    """
    src = Document(src_path)
    for para in src.paragraphs:
        copy_paragraph_with_formatting(para, out_doc)


def get_doc_text_fallback(path):
    """Fallback for legacy .doc files: extract readable ASCII text via olefile."""
    try:
        import olefile, re as _re
        ole = olefile.OleFileIO(path)
        word_stream = ole.openstream('WordDocument').read()
        chunks = _re.findall(b'[ -~]{15,}', word_stream)
        lines = []
        for c in chunks:
            try:
                t = c.decode('latin-1').strip()
                if t:
                    lines.append(t)
            except Exception:
                pass
        return '\n'.join(lines)
    except Exception as e:
        return f"[Error reading {path}: {e}]"


def build_doc():
    out_doc = Document()

    # Set default font
    style = out_doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title page
    title = out_doc.add_heading('Dermatology Word Transcripts - Compiled Study Notes', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = out_doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f'A collection of {len(ENTRIES)} dermatology audio transcripts (Word source files) with summary tables.'
    )
    sub_run.italic = True

    out_doc.add_paragraph()
    out_doc.add_paragraph('Source: Dermatographics / Dermnemonics audio episodes. Original Word transcript documents.')
    out_doc.add_page_break()

    for i, (filename, display_title, summary_rows) in enumerate(ENTRIES, start=1):
        src_path = os.path.join(BASE_DIR, filename)

        # ---- Section heading ----
        out_doc.add_heading(f'{i}. {display_title}', level=2)

        # ---- Transcript subheading ----
        out_doc.add_heading('Transcript', level=3)

        if not os.path.exists(src_path):
            out_doc.add_paragraph(f'[File not found: {filename}]')
        elif filename.endswith('.doc'):
            # Legacy .doc — plain text fallback (no colour/image support)
            text = get_doc_text_fallback(src_path)
            p = out_doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
        else:
            # .docx — copy paragraphs preserving colour + images
            copy_docx_into_doc(src_path, out_doc)

        # ---- Summary table ----
        out_doc.add_heading('Summary Table', level=3)

        table = out_doc.add_table(rows=len(summary_rows) + 1, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header row
        hdr = table.rows[0].cells
        hdr[0].text = 'Topic'
        hdr[1].text = 'Details'
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for j, (key, value) in enumerate(summary_rows, start=1):
            row = table.rows[j].cells
            row[0].text = key
            row[1].text = value

        # Bold first column
        for j in range(1, len(summary_rows) + 1):
            for paragraph in table.rows[j].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        out_doc.add_paragraph()
        if i < len(ENTRIES):
            out_doc.add_page_break()

    out_path = os.path.join(BASE_DIR, 'Dermatology_Word_Transcripts_Compiled.docx')
    out_doc.save(out_path)
    print(f'Saved: {out_path}')
    print(f'Total entries: {len(ENTRIES)}')

    # Sanity check
    actual = set(
        f for f in os.listdir(BASE_DIR)
        if (f.endswith('.docx') or f.endswith('.doc')) and not f.startswith('~')
        and f not in ('Dermatology_Transcripts_Compiled.docx', 'Dermatology_Word_Transcripts_Compiled.docx')
    )
    listed = set(e[0] for e in ENTRIES)
    not_included = actual - listed
    missing_files = listed - actual
    if not_included:
        print(f'WARNING: Docx files in repo NOT included: {sorted(not_included)}')
    if missing_files:
        print(f'WARNING: Entries referencing missing files: {sorted(missing_files)}')


if __name__ == '__main__':
    build_doc()
