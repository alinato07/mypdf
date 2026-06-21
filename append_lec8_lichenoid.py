"""
Append Lecture 8: Other Lichenoid Dermatoses to SCE_ConnectAC_Study_Guide_Lec1-7.docx
Follows the exact same formatting structure as Lectures 1-7.

Style mapping:
  Heading 1  : lecture title
  Heading 2  : numbered section (## N. Title)
  Heading 3  : "Summary" / "Translated & Tidied Audio Transcript"
  Heading 4  : sub-sub heading
  Normal     : transcript lines; "Dr. Eman:" bold run + rest normal
  List Bullet  : bullet, inline bold for key terms
  List Bullet 2: sub-bullet, inline bold
  List Number  : numbered items
"""

import re, os
from docx import Document
from docx.shared import Inches

DOCX_PATH = r"C:\Users\ahaww\Downloads\mypdf\SCE_ConnectAC_Study_Guide_Lec1-7.docx"
OUT_PATH  = r"C:\Users\ahaww\Downloads\mypdf\SCE_ConnectAC_Study_Guide_Lec1-8.docx"
PDF_PATH  = r"C:\Users\ahaww\Downloads\mypdf\ConnectAC_Slides\Lec8_Other lichenoid dermatosis.pdf"

doc = Document(DOCX_PATH)

# ── Try to extract images from the PDF ───────────────────────────────────────
extracted_images = []
try:
    import fitz  # PyMuPDF
    pdf_doc = fitz.open(PDF_PATH)
    img_dir = r"C:\Users\ahaww\Downloads\mypdf\ConnectAC_Slides\lec8_images"
    os.makedirs(img_dir, exist_ok=True)
    img_count = 0
    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            # Only keep reasonably sized images (skip tiny icons)
            if len(image_bytes) > 5000:
                img_path = os.path.join(img_dir, f"lec8_p{page_num+1}_{img_index}.{image_ext}")
                with open(img_path, "wb") as f:
                    f.write(image_bytes)
                extracted_images.append(img_path)
                img_count += 1
    pdf_doc.close()
    print(f"  Extracted {img_count} images from PDF.")
except Exception as e:
    print(f"  Image extraction skipped: {e}")

# ── inline-markdown parser  **bold** *italic* ────────────────────────────────

def add_inline_runs(para, text: str):
    tokens = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = para.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            run = para.add_run(tok[1:-1])
            run.italic = True
        else:
            para.add_run(tok)

# ── paragraph adders ─────────────────────────────────────────────────────────

def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.add_run(text)

def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.add_run(text)

def add_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.add_run(text)

def add_h4(text):
    p = doc.add_paragraph(style='Heading 4')
    p.add_run(text)

def add_normal(text: str):
    p = doc.add_paragraph(style='Normal')
    speaker_pat = re.match(
        r'^(\*\*)?((Dr\.|Student|Professor)[^:]*:)(\*\*)?\s*(.*)',
        text, re.DOTALL
    )
    if speaker_pat:
        label = speaker_pat.group(2).strip()
        rest  = speaker_pat.group(5).strip()
        run_l = p.add_run(label + '  ')
        run_l.bold = True
        add_inline_runs(p, rest)
    else:
        add_inline_runs(p, text.strip())

def add_bullet(text: str, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text.strip())

def add_list_number(text: str):
    p = doc.add_paragraph(style='List Number')
    add_inline_runs(p, text.strip())

def add_image(img_path, width_inches=4.5):
    """Try to add an image; skip silently on failure."""
    try:
        doc.add_picture(img_path, width=Inches(width_inches))
    except Exception:
        pass

# ══════════════════════════════════════════════════════════════════════════════
# LECTURE 8: OTHER LICHENOID DERMATOSES
# ══════════════════════════════════════════════════════════════════════════════

add_h1("Lecture 8: Other Lichenoid Dermatoses")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 1 – Lichen Nitidus
# ═══════════════════════════════════════════════════════════════════════
add_h2("1. Lichen Nitidus")
add_h3("Summary")

add_h4("Definition & Clinical Features")
add_bullet("**Lichen Nitidus** = Multiple tiny, discrete, skin-colored, dome-shaped, asymptomatic papules.")
add_bullet("Most common site: **Dorsum of the hand**, upper extremities, flexors, genitalia.")
add_bullet("More common in **children and young adults**, especially **females**.")
add_bullet("The papules are **pinpoint-sized** and **skin-colored** (not violaceous like LP).")

add_h4("Variants")
add_bullet("**Classic** (most common)")
add_bullet("**Palmoplantar**")
add_bullet("**Linear**")
add_bullet("**Follicular**")
add_bullet("**Generalized**")
add_bullet("**Perforating**")

add_h4("Etiopathogenesis")
add_bullet("Etiology is **unknown/debated**.")
add_bullet("Theory 1: May be an **unusual variant of Lichen planus**.")
add_bullet("Theory 2: May be related to **tuberculosis (TB)** — because the histopathology shows **tuberculoid granuloma** formation.")
add_bullet("A tuberculoid granuloma = well-circumscribed dermal infiltrate of epithelioid histiocytes surrounded by lymphocytes.")

add_h4("Histopathology — \"Ball and Claw Appearance\"")
add_bullet("**\"Ball and claw appearance\"** = pathognomonic histological finding.")
add_bullet("The **\"Ball\"** = Well-circumscribed dermal infiltrate of **epithelioid histiocytes + lymphocytes** (granuloma), confined to the width of a single dermal papilla.")
add_bullet("The **\"Claw\"** = Elongated rete ridges (dermal papillae) flanking and gripping the granuloma like claws holding a ball.")
add_bullet("Also shows: **Lichenoid interface dermatitis** + **Epidermal atrophy** (unlike LP which has acanthosis).")
add_bullet("**Parakeratosis** is common (vs rare in LP).")
add_bullet("Granular cell layer is **thin to absent** (vs focally increased in LP).")

add_h4("Nail & Oral Involvement")
add_bullet("Nail involvement: **10%** of patients — **pitting** and **longitudinal ridging**.")
add_bullet("Key difference from LP: LP has longitudinal ridging and pterygium but **NO pitting**; Lichen Nitidus **has pitting**.")
add_bullet("Oral involvement: **Very rare**.")

add_h4("Associations — Mnemonic \"DMC\"")
add_bullet("**D** = Down syndrome")
add_bullet("**M** = Megacolon")
add_bullet("**C** = Crohn's disease")

add_h4("Lymphocytes")
add_bullet("Lymphocytes in Lichen Nitidus are **CD4+**.")
add_bullet("Key difference: LP lymphocytes are **CD8+**.")

add_h4("Differential Diagnosis")
add_bullet("**Plane warts (flat warts)** — most important DDx.")
add_bullet("**Frictional lichenoid dermatitis**.")

add_h4("Treatment")
add_bullet("Usually **no treatment needed** — asymptomatic and self-limited.")
add_bullet("If symptomatic: **Topical steroids**, **emollients**, **antihistamines**.")
add_bullet("Spontaneous clearing occurs.")

add_h4("Key Differences from Lichen Planus (Summary Table)")
add_bullet("**Size:** Pinpoint (LN) vs Pinhead to large plaques (LP).")
add_bullet("**Color:** Skin-colored (LN) vs Violaceous (LP).")
add_bullet("**Wickham striae:** Absent (LN) vs Present (LP) — because LN has no hypergranulosis.")
add_bullet("**Post-inflammatory hyperpigmentation:** None (LN) vs Present (LP).")
add_bullet("**Histology — Epidermis:** Epidermal atrophy (LN) vs Acanthosis (LP).")
add_bullet("**Histology — Dermal infiltrate:** Granuloma present (LN) vs No granuloma (LP).")
add_bullet("**Histology — Parakeratosis:** Common (LN) vs Rare (LP).")
add_bullet("**Lymphocytes:** CD4+ (LN) vs CD8+ (LP).")
add_bullet("**DIF — Fibrin deposition:** Absent (LN) vs Present/shaggy fibrinogen (LP).")
add_bullet("**Koebner phenomenon:** Common in both.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Good morning doctors, how are you all? How is the studying going? Let me check in with everyone. Have you started solving questions yet?")
add_normal("Student (Dr. Naima):  Good morning Dr. Eman. I've started reading the first lectures to connect the new material with the old, thank God.")
add_normal("Dr. Eman:  The first two lectures are genuinely difficult — they're like a compilation of the entire dermatology base. Try to listen to them multiple times. But the remaining lectures are easier because each one discusses a separate topic — you'll understand the disease and that's it.")
add_normal("Student (Dr. Naima):  That's what I'm doing — I listen to them almost every day. With work and Ramadan I'm trying to push myself to listen to a section and read it simultaneously.")
add_normal("Dr. Eman:  Excellent, and try to incorporate question-solving too. You must start solving questions because solving is part of studying — it's not just reviewing information you've already learned, there's also new information in the questions themselves.")
add_normal("Student (Dr. Doaa):  I'm doing well, thank God. It's Ramadan and time is short, but I feel I'm getting into the mood a bit.")
add_normal("Dr. Eman:  May God make it easy for everyone. Let's begin today's lecture. We finished Psoriasis, Other Papulosquamous diseases, and Lichen Planus. Today we'll finish the Other Lichenoid Dermatoses — it should be a light and pleasant lecture, God willing.")
add_normal("Dr. Eman:  Last time we finished all of Lichen Planus. We said there are many diseases under the umbrella of Interface dermatitis. There are two types of interface dermatitis — who can tell me what Interface dermatitis means and what are its two types?")
add_normal("Student:  Degeneration of the basement membrane.")
add_normal("Dr. Eman:  Yes, degeneration of the basement membrane. What are the two types?")
add_normal("Student:  Lichenoid and Vacuolar.")
add_normal("Dr. Eman:  The Lichenoid type is a band-like inflammatory infiltrate of lymphocytes beneath the dermoepidermal junction. Diseases in this category include: Lichen planus, Lichen nitidus, Lichen striatus, and Lichenoid drug eruption. The Vacuolar interface dermatitis includes: Erythema multiforme, SLE, Dermatomyositis, Fixed drug eruption, and Graft-versus-host disease. This is very important because it encompasses diseases from many different chapters.")
add_normal("Dr. Eman:  Let's start with Lichen Nitidus. It's quite different from Lichen planus. It presents as multiple tiny discrete skin-colored papules, as you can see on the hand here. It comes mainly on the dorsum of the hand. Will this confuse us with Lichen planus? No — it looks distinctly different. It's asymptomatic, tiny, discrete, dome-shaped, skin-colored papules on the dorsum of the hand and elbows, more common in children and young adults, more common in females.")
add_normal("Dr. Eman:  The most common sites are the dorsal aspect of the hand, upper extremities, flexors, and genitalia. It has variants: the classic form we're seeing, plus palmoplantar, linear, follicular, generalized, and perforating.")
add_normal("Dr. Eman:  Now what is Lichen Nitidus pathologically? Part of it — because it's a 'Lichen' — will have interface dermatitis and band-like inflammatory infiltrate. But they also found something unusual in its pathology: the shape of a Granuloma. A granuloma is a collection of histiocytes with lymphocytes around them.")
add_normal("Dr. Eman:  That's why there's debate about Lichen Nitidus — some say it's a rare variant of LP, and others say it's related to TB. Why TB? Because it forms a Tuberculoid granuloma: a well-circumscribed dermal infiltrate of epithelioid histiocytes and lymphocytes.")
add_normal("Dr. Eman:  The pathology of Lichen Nitidus is very specific and recognizable — we call it the \"Ball and claw appearance.\" The ball is the granuloma collection surrounded by dermal papillae (the claws). Like a ball being held by claws. The ball is the well-circumscribed dermal infiltrate of epithelioid histiocytes and lymphocytes. There's also lichenoid interface dermatitis and epidermal atrophy — unlike LP which has acanthosis.")
add_normal("Dr. Eman:  Does it have oral lesions? Very rare. Nail involvement? Yes, 10% of patients have nail abnormalities like pitting and longitudinal ridging. Remember: LP had longitudinal ridging but NO pitting. Lichen Nitidus HAS pitting.")
add_normal("Dr. Eman:  The associations of Lichen Nitidus — memorize them with the mnemonic DMC: Down syndrome, Megacolon, and Crohn's disease. Most cases I've seen were normal with no associations, but you must know DMC for the exam.")
add_normal("Dr. Eman:  The differential diagnosis — what could this look like? Plane warts (flat warts) are the number one DDx. Also frictional lichenoid dermatitis.")
add_normal("Dr. Eman:  The histopathology — what's the key word?")
add_normal("Student:  Ball and claw.")
add_normal("Dr. Eman:  Ball and claw. Ball from the Granuloma (epithelioid histiocytes with lymphocytes), plus lichenoid interface dermatitis and epidermal atrophy. This is extremely important — the pathology of Lichen Nitidus is very high-yield.")
add_normal("Dr. Eman:  Treatment? No treatment needed — it's asymptomatic. If needed: topical steroids, emollient, antihistamines if itchy. Spontaneous clearing occurs.")
add_normal("Dr. Eman:  Now here's a comparison table between Lichen Nitidus and Lichen Planus. Size: LN is pinpoint, LP is pinhead to large plaques. Color: LN is skin-colored, LP is violaceous. Distribution: LN favors upper extremities, flexors, genitalia, chest; LP favors wrists, forearms, peri-sacral area. Oral involvement: LN very rare, LP common — over 50%. Nail: LN has pitting (5-10%), LP does NOT have pitting.")
add_normal("Dr. Eman:  Wickham striae — will we see them in LN? No. Why do we see them in LP? Because of the focal hypergranulosis. Hyperpigmentation after healing — does LN leave it? No. LP leaves it because of basement membrane damage causing melanocytes to fall into the dermis as melanophages.")
add_normal("Dr. Eman:  Histopathology: Parakeratosis is common in LN but rare in LP. Granular cell layer is thin to absent in LN but focally increased in LP. Dermal infiltrate in LN is lymphocytes + epithelioid cells = granuloma! LP has no granuloma — just band-like lymphocytes. In LN the infiltrate is confined to the width of dermal papillae (the ball in the claws), in LP it's a band. Fibrin deposition: absent in LN, present in LP (shaggy fibrinogen on DIF).")
add_normal("Dr. Eman:  The lymphocytes in Lichen Nitidus are CD4+. In Lichen Planus they were CD8+.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 2 – Lichen Striatus
# ═══════════════════════════════════════════════════════════════════════
add_h2("2. Lichen Striatus")
add_h3("Summary")

add_h4("Definition & Clinical Features")
add_bullet("**Lichen Striatus** = Uncommon, asymptomatic, self-limited linear dermatosis following **Blaschko's lines**.")
add_bullet("Appearance: Pink/erythematous lichenoid papules in a **linear, unilateral** distribution, mainly on **extremities**.")
add_bullet("More common in **children** and **females**.")
add_bullet("Seasonal predilection: **Spring and Summer** (supports viral infection theory).")
add_bullet("Heals with **hypopigmentation** (NOT hyperpigmentation like LP).")
add_bullet("**Spontaneous resolution** within 6 months to years.")

add_h4("Etiopathogenesis (Theories)")
add_bullet("**Somatic mosaicism** — different cell populations along Blaschko's lines.")
add_bullet("**Viral infection** — supported by seasonal increase in spring/summer.")
add_bullet("**Cell-mediated CD8+ response**.")
add_bullet("**Atopy**.")
add_bullet("**Trauma / Koebnerization**.")

add_h4("Histopathology")
add_bullet("**Eczema-like features** in early stages: acanthosis, spongiosis, exocytosis.")
add_bullet("**Vacuolar interface dermatitis** also present.")
add_bullet("**KEY DISTINGUISHING FEATURE:** Inflammatory infiltrate **around hair follicles and eccrine sweat glands**.")
add_bullet("This perifollicular and periadnexal infiltrate differentiates Lichen Striatus from simple eczema or vacuolar interface dermatitis alone.", level=1)

add_h4("Differential Diagnosis")
add_bullet("**Linear Lichen Planus** — older age, pruritic, heals with hyperpigmentation (vs LS: younger, asymptomatic, heals with hypopigmentation).")
add_bullet("**Linear GVHD**.")
add_bullet("**Blaschkitis** (inflammation along Blaschko's lines).")
add_bullet("**ILVEN** (Inflammatory Linear Verrucous Epidermal Nevus) — most similar to Lichen Striatus.")

add_h4("Treatment")
add_bullet("**No treatment needed** — asymptomatic and self-resolving (resolves within 6 months).")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Now we'll move to a completely different topic: Lichen Striatus. It presents as a linear dermatosis — erythematous papules following Blaschko's lines or a linear pattern.")
add_normal("Dr. Eman:  Some people classified it as a type of eczema. Others said it might be caused by viral infection because it increases in summer and spring. Some proposed somatic mosaicism. Others suggested Koebnerization or trauma. All these theories exist because the exact cause is unknown.")
add_normal("Dr. Eman:  Lichen Striatus is an uncommon, asymptomatic, self-limited disease. A streak of erythematous papules appears in a linear pattern, stays for a while, then heals with hypopigmentation. It's an asymptomatic streak of pink lichenoid papules that's unilateral, on the extremities, along Blaschko's lines.")
add_normal("Dr. Eman:  Has anyone seen Lichen Striatus before? It's more common in children and females — same as Lichen Nitidus. It has seasonal variation — more common in summer and spring, which is why they proposed the viral infection theory.")
add_normal("Dr. Eman:  It appears suddenly and resolves spontaneously within months to years. The etiopathogenesis is either unknown, or one of these theories: somatic mosaicism (different cells along Blaschko's lines), viral infection (seasonal), cell-mediated CD8 response, atopy, or trauma.")
add_normal("Dr. Eman:  The histopathology of Lichen Striatus: First, it has vacuolar interface dermatitis, but not strongly. It also shows eczema-like features — acanthosis, spongiosis, and exocytosis in the early stages.")
add_normal("Dr. Eman:  But the KEY distinguishing feature is that the infiltrate surrounds hair follicles and eccrine sweat glands. The lymphocytes encircle the hair follicle and the sweat glands. This is crucial — it's not just eczema-like or vacuolar interface dermatitis on the surface. You have lymphocytes also surrounding the hair follicles and eccrine sweat glands deep in the dermis.")
add_normal("Dr. Eman:  Differential diagnosis: Linear lichen planus — but that comes in older age, is very pruritic, and heals with hyperpigmentation. Linear Graft-versus-host disease. Blaschkitis — inflammation of the Blaschko's lines. And ILVEN (Inflammatory Linear Verrucous Epidermal Nevus) — this is the most similar to Lichen Striatus.")
add_normal("Dr. Eman:  Treatment? Asymptomatic, resolves within 6 months, nothing needed. Just like Lichen Nitidus — nothing needed. These two topics (Lichen Striatus and Lichen Nitidus) are quick, easy topics on the side.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 3 – Graft-Versus-Host Disease (GVHD)
# ═══════════════════════════════════════════════════════════════════════
add_h2("3. Graft-Versus-Host Disease (GVHD)")
add_h3("Summary")

add_h4("Definition & Concept")
add_bullet("**GVHD** = The transplanted immune cells (graft) attack the recipient's body (host).")
add_bullet("Occurs ONLY after **hematopoietic stem cell (bone marrow) transplantation** — NOT after solid organ transplants.")
add_bullet("Why not solid organs? Solid organs (liver, kidney, heart) do NOT contain immune cells capable of attacking the host. In solid organ transplants, it's the host that rejects the organ (tissue rejection).", level=1)

add_h4("Mechanism")
add_bullet("The recipient is **immunosuppressed** (to prevent rejection of the transplant) → immune system is essentially zero.")
add_bullet("Donor bone marrow contains **T-cells** that are transplanted into the host.")
add_bullet("The donor T-cells recognize host tissues as **foreign** → activate → attack the host's organs.")
add_bullet("The host cannot defend itself because it is immunocompromised.")

add_h4("Target Organs")
add_bullet("**Skin** (most common)")
add_bullet("**GI tract** (nausea, vomiting, diarrhea, abdominal pain)")
add_bullet("**Liver** (hepatitis)")

add_h4("Acute GVHD")
add_bullet("**Mechanism:** Recipient tissue damage → T-cell activation → Cytotoxic T-cells cause tissue injury to skin, GI, and liver.")
add_bullet("**Skin manifestation:** **Morbilliform exanthem** (maculopapular rash).")
add_bullet("**Pathognomonic sign:** **Peripheral desquamation** — bullae and desquamation of the acral area (palms and soles).")
add_bullet("GI: Nausea, vomiting, diarrhea, abdominal pain.", level=1)
add_bullet("Liver: Hepatitis.", level=1)
add_bullet("Also: Generalized pruritus, cicatricial alopecia, erythroderma (in severe cases).", level=1)
add_bullet("**Histopathology:** Vacuolar interface dermatitis + keratinocyte necrosis.")
add_bullet("**DDx:** Viral exanthem, Drug eruption (differentiated by peripheral desquamation and clinical history of bone marrow transplantation).")
add_bullet("**Treatment:** Topical and systemic steroids, Tacrolimus, Mycophenolate mofetil, TNF-alpha inhibitors.")

add_h4("Chronic GVHD")
add_bullet("**Mechanism:** T-cell activation (same as acute) **PLUS B-cell activation** → production of **autoantibodies** (ANA positive, Anti-dsDNA positive).")
add_bullet("The autoantibodies attack host tissues, producing autoimmune-like disease manifestations.")
add_bullet("**Skin manifestations mimic:**")
add_bullet("**Lichen planus**-like", level=1)
add_bullet("**Lichen striatus**-like", level=1)
add_bullet("**Morphea**-like", level=1)
add_bullet("**Scleroderma**-like", level=1)
add_bullet("**Poikiloderma**-like", level=1)
add_bullet("**Eosinophilic fasciitis**-like", level=1)
add_bullet("**Systemic features:** Dry mouth, dry eyes, oral sores, rash, jaundice, shortness of breath, nausea, vomiting.")
add_bullet("Can affect **any body system**.")
add_bullet("**Histopathology:** Variable — depends on which disease it mimics (LP pathology if LP-like, morphea pathology if morphea-like, etc.).")
add_bullet("**Treatment:** Sun protection, phototherapy, systemic steroids.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Now let's move to a completely new topic: Graft-versus-host disease (GVHD). Does anyone have any information about this disease or can explain what it means?")
add_normal("Student:  In transplantation, rejection of the graft occurs.")
add_normal("Dr. Eman:  Rejection happens, yes. But who does the rejection? Does my body reject what was transplanted, or is it the opposite?")
add_normal("Student:  I think the opposite.")
add_normal("Dr. Eman:  Exactly the opposite! This is NOT tissue rejection. In tissue rejection, the host rejects the organ — that's what happens in any normal organ transplantation. But here the disease is called Graft VERSUS Host disease — the graft attacks the host.")
add_normal("Dr. Eman:  Let me explain from scratch. Imagine someone has a problem with their hematopoietic stem cells (bone marrow) — cancer, radiation damage, anything. They need a bone marrow transplant. When we transplant bone marrow, we're transplanting immune cells and blood cells.")
add_normal("Dr. Eman:  Where do these immune cells come from? From another person whose HLA is somewhat matched. We take their bone marrow and transplant it. Now, for the transplant to succeed, does the recipient's immunity need to be strong or weak?")
add_normal("Student:  Obviously weak.")
add_normal("Dr. Eman:  Obviously weak — why? So they don't reject the transplant! I need to suppress their immune system to near zero. Then the new immune cells (the graft full of functional T-cells) enter and what happens? These new immune cells find themselves in a foreign body! It's not their own body. So they start attacking the host's body!")
add_normal("Dr. Eman:  That's why it's called Graft versus Host. The immune cells I transplanted (the Graft) attack the patient's body (the Host). The host can't defend itself because it's immunocompromised. Understand the concept?")
add_normal("Dr. Eman:  This only happens in hematopoietic stem cell transplantation. Does it happen with other organ transplants? No — because solid organs don't contain immune cells that can attack the body. If I transplant a liver, the liver won't attack the body! The body attacks the liver and rejects it. But here I'm transplanting immune cells (bone marrow/stem cells), so THEY attack the body, and the host can't defend itself.")
add_normal("Dr. Eman:  GVHD is divided into Acute and Chronic, and the target organs are: Skin, GI tract, and Liver.")
add_normal("Dr. Eman:  In Acute GVHD, the pathognomonic sign is bullae and desquamation of the acral area — desquamation of the palms and soles. On the skin you get a morbilliform eruption (maculopapular rash), and desquamation in the acral area. There's also pruritus and cicatricial alopecia. GI presents with nausea, vomiting, diarrhea, and abdominal pain. Liver presents as hepatitis.")
add_normal("Dr. Eman:  So a typical case: an older patient with a history of bone marrow transplantation, presenting with a full-body exanthem that looks like a viral infection or drug eruption, PLUS peripheral desquamation of the palms and soles, PLUS acute abdomen and hepatitis. The peripheral desquamation is what differentiates it from viral or drug causes.")
add_normal("Dr. Eman:  The Acute form has 3 phases: Recipient tissue damage → Activation of T-cells → Cytotoxic T-cells cause tissue injury to skin, liver, and GIT. It manifests as morbilliform exanthem, erythroderma, bullae and desquamation in the acral area, generalized pruritus, and cicatricial alopecia.")
add_normal("Dr. Eman:  Now the Chronic form. The acute was T-cell activation — the donor T-cells get activated against the host. In Chronic GVHD, the same T-cell activation occurs BUT ALSO B-cell activation. What does B-cell activation mean? It means autoantibodies are produced against the host's body. These antibodies attack the host.")
add_normal("Dr. Eman:  What do these autoantibodies produce? ANA positive and Anti-double stranded DNA positive. So the chronic form manifests as: morphea-like, scleroderma-like, poikiloderma-like, eosinophilic fasciitis-like, and lichen planus-like disease. All because autoantibodies are attacking.")
add_normal("Dr. Eman:  Chronic GVHD causes dry mouth, dry eyes, mouth sores, rash, jaundice, shortness of breath, nausea, vomiting. On the skin it appears as lichen planus, lichen striatus, morphea-like, scleroderma-like, and poikiloderma-like. It can affect any organ system.")
add_normal("Dr. Eman:  Histopathology: Acute GVHD shows vacuolar interface dermatitis and keratinocyte necrosis — that's why it's classified under lichenoid dermatoses. Chronic GVHD is variable — if it looks like LP, it has LP pathology; if it looks like morphea, it has morphea pathology.")
add_normal("Dr. Eman:  DDx for Acute: viral exanthem and drug eruption. Treatment for Acute: topical and systemic steroids, Tacrolimus, Mycophenolate Mofetil, and TNF-alpha inhibitors. Treatment for Chronic: sun protection, phototherapy, and systemic steroids.")
add_normal("Student:  Doctor, the recipient who receives the graft — they get immunosuppression, correct? So what causes the donor T-cells to activate? They already have no immunity.")
add_normal("Dr. Eman:  What activates the donor T-cells? Any stimulus — bacteria on the skin, any antigen presenting cell activation. Why are we transplanting in the first place? To give this person immunity! Their bone marrow had cancer, and I transplanted new bone marrow so they'd have a functioning immune system to protect them from infections. But that immune system turns against them. The T-cells aren't 'theirs' — so they start attacking the host tissue itself.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 4 – Erythroderma (Exfoliative Dermatitis)
# ═══════════════════════════════════════════════════════════════════════
add_h2("4. Erythroderma (Exfoliative Dermatitis)")
add_h3("Summary")

add_h4("Definition")
add_bullet("**Erythroderma** = Erythema + scaling involving **>90% body surface area**.")
add_bullet("Also called **Exfoliative Dermatitis**.")
add_bullet("NOT a disease itself — it is an **endpoint/manifestation** of many different diseases.")
add_bullet("Course: **Chronic** (months to years), relapsing-remitting.")
add_bullet("Main cause of death: **Pneumonia** (secondary infection due to loss of skin barrier).")

add_h4("Causes in Adults — Mnemonic \"BAD\"")
add_bullet("**Most common causes:**")
add_bullet("**B** = (B)ig two: **Psoriasis** and **Atopic dermatitis** (most common overall)", level=1)
add_bullet("**A** = (A)ssorted inflammatory conditions", level=1)
add_bullet("**D** = **Drugs** — mnemonic **\"ASBAGAK\"**: Allopurinol, Sulfasalazine/Sulfonamides, Phenytoin, Beta-blockers, Gold, Carbamazepine", level=1)
add_bullet("**Less common causes:** Mycosis Fungoides (MF)/Sézary Syndrome, Pityriasis Rubra Pilaris (PRP), Eczema, Paraneoplastic Pemphigus, Pemphigus Foliaceus, Bullous Pemphigoid.")
add_bullet("**Rare causes:** Crusted Scabies, Congenital Ichthyosiform Erythroderma, Lichen Planus, GVHD.")
add_bullet("**Idiopathic (25%)** = **\"Red Man Syndrome\"** — erythroderma with no identifiable cause.")

add_h4("Causes in Children/Neonates")
add_bullet("**Genetic/Ichthyosis:** Congenital Ichthyosiform Erythroderma (bullous and non-bullous), Netherton Syndrome.")
add_bullet("**Immunodeficiency:** Wiskott-Aldrich Syndrome.")
add_bullet("**Inflammatory:** Atopic dermatitis, Seborrheic dermatitis, Psoriasis.")
add_bullet("**Infections:** Staphylococcal Scalded Skin Syndrome (SSSS), Congenital Cutaneous Candidiasis.")
add_bullet("**Others:** PRP, GVHD, Mastocytosis.")

add_h4("Drug Mnemonics Summary")
add_bullet("**Erythroderma drugs** = \"ASBAGAK\" (Allopurinol, Sulfasalazine/Sulfonamides, Phenytoin, Beta-blockers, Gold, Carbamazepine).")
add_bullet("**Psoriasis exacerbation drugs** = \"LIPA\" (Lithium, Interferon, Beta-blockers, Antimalarials).")
add_bullet("**Lichenoid drug eruption/LP drugs** = \"DAMAT\" (Diuretics, Antihypertensives, Antimalarials, Amalgam/Metals, Tetracyclines).")

add_h4("Clinical Features")
add_bullet("**Erythema appears first** → scaling develops 2–8 days later.")
add_bullet("Skin is **bright red, hot, and dry**.")
add_bullet("**Intense pruritus** — worst with **Sézary syndrome** and **Atopic dermatitis**.")

add_h4("Complications")
add_bullet("**Skin complications:** Lichenification, depigmentation, palmoplantar keratoderma, nail changes, alopecia, subungual hyperkeratosis, *Staphylococcus aureus* infection (no skin barrier).")
add_bullet("**Systemic complications:** Fever, malaise, **generalized lymphadenopathy (50%)**, **peripheral edema (50%)**, hepatomegaly (20%), splenomegaly.")
add_bullet("**Cardiovascular cascade:** Vasodilatation → fluid loss → electrolyte disturbance → tachycardia → hypothermia → **High-output cardiac failure** → cachexia, anemia.")
add_bullet("High-output cardiac failure is the reason for **mandatory hospitalization**!", level=1)

add_h4("Sézary Syndrome — Key Distinction")
add_bullet("Sézary syndrome is **NOT the same** as erythrodermic Mycosis Fungoides.")
add_bullet("It is a **separate leukemic variant of Cutaneous T-cell Lymphoma (CTCL)**.")
add_bullet("Starts with **erythroderma from the beginning** — no preceding patches or plaques (unlike MF which progresses through patches → plaques → tumors → erythroderma).")
add_bullet("The malignant T-cell in Sézary is **different** from the malignant T-cell in MF.")
add_bullet("Features: **Sézary cells in blood** + **Generalized lymphadenopathy**.")

add_h4("Histopathology")
add_bullet("**Non-specific:** Acanthosis, dilated blood vessels, edema.")
add_bullet("**Acute:** Spongiosis + parakeratosis.")
add_bullet("**Chronic:** Acanthosis.")
add_bullet("The underlying disease may show its own specific pathology (e.g., psoriasis plaques may still be identifiable).")

add_h4("Treatment")
add_bullet("**Hospitalization required** (for monitoring and managing cardiovascular complications).")
add_bullet("**Stop suspected drugs** immediately.")
add_bullet("**Close monitoring:** Fluid balance, electrolytes, temperature, cardiac function.")
add_bullet("**Systemic antibiotics** — prevent secondary bacterial infection (no functioning skin barrier).")
add_bullet("**Sedating antihistamines** — for severe pruritus.")
add_bullet("**Topical care:**")
add_bullet("Wet dressings", level=1)
add_bullet("Emollients", level=1)
add_bullet("**LOW-potency topical steroids ONLY** — high absorption risk over >90% BSA → systemic toxicity if high-potency steroids used", level=1)
add_bullet("**NEVER cover the entire body with topical steroids**", level=1)

add_h4("Specific Treatment by Cause")
add_bullet("**Psoriasis →** Methotrexate, Acitretin, or Biologics. **AVOID systemic steroids** (causes rebound flare or pustular psoriasis).")
add_bullet("**Atopic dermatitis →** Systemic steroids OR Cyclosporine.")
add_bullet("**Drug-induced →** Stop the offending drug + systemic steroids.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  We've now finished the Lichenoid dermatoses. Many diseases end up causing erythroderma — including GVHD and Psoriasis. Let's discuss Erythroderma now.")
add_normal("Dr. Eman:  What is the definition of Erythroderma?")
add_normal("Student:  Generalized erythema and scaling of the skin more than 90% body surface area.")
add_normal("Dr. Eman:  Yes — erythema and scales covering more than 90% of the BSA. Is Erythroderma a disease in itself?")
add_normal("Student:  No, it's the result of many diseases.")
add_normal("Dr. Eman:  Exactly. It can be chronic, lasting months to years, with a relapsing-remitting course. The main cause of death in erythroderma is pneumonia — secondary infection.")
add_normal("Dr. Eman:  Causes divide into primary (idiopathic, 25% — called Red Man Syndrome) and secondary. Secondary divides into adult causes and pediatric causes.")
add_normal("Dr. Eman:  What are the most important causes of erythroderma in adults?")
add_normal("Student:  Psoriasis, Atopic Dermatitis, Drugs.")
add_normal("Dr. Eman:  Correct — the three most important adult causes, summarized as \"BAD.\" The drugs causing erythroderma are memorized with \"ASBAGAK\": Allopurinol, Sulfasalazine/Sulfonamides, Phenytoin, Beta-blockers, Gold, Carbamazepine.")
add_normal("Dr. Eman:  Less common causes include Mycosis Fungoides, Pityriasis Rubra Pilaris, Eczema, Paraneoplastic Pemphigus, Pemphigus Foliaceus, and Bullous Pemphigoid. Very rare causes: Crusted Scabies, Congenital Ichthyosiform Erythroderma, Lichen Planus, and GVHD. You must know all of these.")
add_normal("Dr. Eman:  In children and neonates, the most important cause is genetic/ichthyosis — a newborn isn't going to get psoriasis or take drugs. Congenital Ichthyosiform Erythroderma (bullous and non-bullous) and Netherton Syndrome are the key genetic causes. Then immunodeficiency syndromes like Wiskott-Aldrich. Then inflammatory causes like Atopic dermatitis, Seborrheic dermatitis, and Psoriasis. Infections like SSSS (Staphylococcal Scalded Skin Syndrome) and Congenital Cutaneous Candidiasis. And others: PRP, GVHD, Mastocytosis.")
add_normal("Dr. Eman:  Let me consolidate the drug mnemonics: Erythroderma drugs = \"ASBAGAK\". Psoriasis exacerbation drugs = \"LIPA\" (Lithium, Interferon, Beta-blockers, Antimalarials). Lichen Planus/Lichenoid drug eruption drugs = \"DAMAT\" (Diuretics, Antihypertensives, Antimalarials, Amalgam/Metals, Tetracyclines).")
add_normal("Dr. Eman:  The clinical picture: Erythema appears first, then scaling develops 2 to 8 days later. The body is bright red, hot, and dry. Is there pruritus? Most common — it's very intense, especially with Sézary syndrome and Atopic dermatitis.")
add_normal("Dr. Eman:  Complications — the skin: Lichenification (increased skin markings), depigmentation, palmoplantar keratoderma, nail changes, alopecia, subungual hyperkeratosis, and Staph aureus infection because there's no skin barrier.")
add_normal("Dr. Eman:  Systemic complications: Fever, malaise, generalized lymphadenopathy in 50% of cases, peripheral edema in 50%, hepatomegaly in 20%, splenomegaly.")
add_normal("Dr. Eman:  The most important systemic complication — due to massive vasodilatation in the skin: fluid loss, electrolyte disturbance, tachycardia, hypothermia, and ultimately High-output cardiac failure, leading to cachexia and anemia. Can you imagine that erythroderma can cause cardiac failure? That's why these patients require hospitalization.")
add_normal("Dr. Eman:  Histopathology is non-specific — acanthosis, dilated blood vessels, edema. Nothing specific to help you identify the cause. In acute erythroderma: spongiosis and parakeratosis. In chronic: acanthosis.")
add_normal("Dr. Eman:  Now about Sézary Syndrome — does anyone know what it is?")
add_normal("Student:  It's on the spectrum of MF.")
add_normal("Dr. Eman:  Actually, it's NOT on the MF spectrum. It's different from MF. Mycosis Fungoides progresses from patches to plaques to tumors to erythroderma. Sézary Syndrome is a SEPARATE entity — a leukemic variant of Cutaneous T-cell Lymphoma. It starts with erythroderma from the very beginning, with no preceding patches or plaques. The malignant T-cell itself is different from MF. It has Sézary cells in the blood at a certain percentage, plus generalized lymphadenopathy.")
add_normal("Dr. Eman:  Treatment of Erythroderma: First — hospitalization. Second — stop any suspected drugs. Third — close monitoring for fluid disturbance and cardiac issues; adjust body temperature, electrolyte balance, and fluid replacement. Fourth — systemic antibiotics to prevent secondary bacterial infection since there's no skin barrier. Fifth — sedating antihistamines for severe pruritus. Sixth — topical care: wet dressings, emollients, and low-potency topical steroids.")
add_normal("Dr. Eman:  IMPORTANT: You cannot cover the entire body with topical steroids in erythroderma — the absorption will be extremely high over 90% BSA and will cause systemic toxicity.")
add_normal("Dr. Eman:  Specific treatment depends on the cause. For Psoriasis: Methotrexate, Acitretin, or Biologics — AVOID systemic steroids because they cause rebound flare or pustular psoriasis. For Atopic dermatitis: systemic steroids or Cyclosporine. For Drug-induced: stop the offending drug and give systemic steroids.")
add_normal("Student:  Doctor, regarding erythroderma, when do we use corticosteroids in the acute phase?")
add_normal("Dr. Eman:  I would give systemic steroids to control severe inflammation if I know the cause is NOT psoriasis — for example, drug-induced erythroderma. If I'm unsure it's psoriasis and give systemic steroids, there'll be a rebound. If the patient is critically ill, I might give steroids and then determine the underlying disease once things calm down. But if I'm certain it's psoriasis, I'll use Methotrexate or another agent instead. Meanwhile, I treat all the other symptoms — fluid balance, electrolytes, etc.")
add_normal("Dr. Eman:  That's it for today's lecture, doctors. We covered Lichen Nitidus, Lichen Striatus, Graft-Versus-Host Disease, and Erythroderma. Any questions?")
add_normal("Dr. Eman:  I apologize we won't solve questions today due to time. I'll try to post some questions on the group after iftar for us to solve together. Happy Ramadan to everyone.")

# ══════════════════════════════════════════════════════════════════════════════
# ADD IMAGES (if extracted)
# ══════════════════════════════════════════════════════════════════════════════
# Images are not inserted inline (to avoid breaking flow) but could be added
# at specific points if desired. For now, we've included them in extraction only.

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

doc.save(OUT_PATH)
print(f"✓ Lecture 8 (Other Lichenoid Dermatoses) appended successfully!")
print(f"  Saved to: {OUT_PATH}")
