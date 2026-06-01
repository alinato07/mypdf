"""
Append Lecture 2 (Immunology) to Lec1-5_Derm_Study_Guide.docx
matching its exact formatting style.

Style mapping (from reference doc inspection):
  Heading 1  : 16pt bold #1F3964  — lecture title
  Heading 2  : 13pt bold #275EA0  — numbered section  (## N. Title)
  Heading 3  : bold    #3786C0  — "Summary" / "Translated & Tidied Audio Transcript"
  Heading 4  : bold    #4F81BD  — sub-sub heading (#### ...)
  Normal     : Calibri auto      — transcript lines; "Speaker:" bold run + rest normal
  List Bullet  : bullet, inline bold for key terms
  List Bullet 2: sub-bullet, "- " prefix, inline bold
  List Number  : numbered, first term bold
  Tables     : Table Grid, header #1F3964 white bold, alt rows #DCE6F1, first-col bold
"""

import re
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = '/projects/sandbox/mypdf/Lec1-5_Derm_Study_Guide.docx'
MD   = '/projects/sandbox/mypdf/Lec2_StudyGuide_Immunology.md'
OUT  = '/projects/sandbox/mypdf/Lec1-5_Derm_Study_Guide.docx'

doc = Document(SRC)

# ── tiny XML helpers ─────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_fill: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove old shd if present
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    tcPr.append(shd)

# ── inline-markdown parser  **bold** *italic* ────────────────────────────────

def add_inline_runs(para, text: str):
    """
    Split text on **bold** and *italic* markers and add runs
    with appropriate formatting.  Plain text gets no extra formatting
    (inherits the paragraph style).
    """
    # tokenise: bold first, then italic
    tokens = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = para.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            run = para.add_run(tok[1:-1])
            run.italic = True
        else:
            para.add_run(tok)

# ── paragraph adders ─────────────────────────────────────────────────────────

def add_h1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    p.add_run(text)

def add_h2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.add_run(text)

def add_h3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    p.add_run(text)

def add_h4(doc, text):
    p = doc.add_paragraph(style='Heading 4')
    p.add_run(text)

def add_normal(doc, text: str):
    """
    Normal paragraph.  If the text starts with a speaker label like
    "Dr. Eman:" or "Student (...):" make that label a bold run.
    """
    p = doc.add_paragraph(style='Normal')
    m = re.match(r'^(\*\*)?([^:]+:)(\*\*)?\s*(.*)', text, re.DOTALL)
    # detect transcript speaker lines
    speaker_pat = re.match(
        r'^(\*\*)?((Dr\.|Student|Professor)[^:]*:)(\*\*)?\s*(.*)',
        text, re.DOTALL
    )
    if speaker_pat:
        label = speaker_pat.group(2).strip()
        rest  = speaker_pat.group(5).strip()
        run_l = p.add_run(label + '  ')
        run_l.bold = True
        add_inline_runs(p, rest)
    else:
        add_inline_runs(p, text.strip())

def add_bullet(doc, text: str, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text.strip())

def add_list_number(doc, text: str):
    p = doc.add_paragraph(style='List Number')
    add_inline_runs(p, text.strip())

# ── table builder ────────────────────────────────────────────────────────────

def parse_md_table(md_lines):
    """Return list-of-lists (rows × cols), skipping separator lines."""
    rows = []
    for line in md_lines:
        line = line.strip()
        if re.match(r'^\|[-|: ]+\|$', line):
            continue
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            rows.append(cells)
    return rows

def add_table(doc, md_lines):
    rows = parse_md_table(md_lines)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'

    for r_idx, row in enumerate(rows):
        for c_idx in range(ncols):
            cell_text = row[c_idx] if c_idx < len(row) else ''
            cell = tbl.cell(r_idx, c_idx)

            # shading
            if r_idx == 0:
                set_cell_shading(cell, '1F3964')   # dark navy header
            elif r_idx % 2 == 0:
                set_cell_shading(cell, 'DCE6F1')   # light blue alt row
            # else: white (no shading needed)

            para = cell.paragraphs[0]
            # strip markdown bold for cell text, then re-apply
            clean = re.sub(r'\*\*(.+?)\*\*', r'__B__\1__B__', cell_text)
            clean = re.sub(r'\*(.+?)\*',     r'\1',           clean)
            parts = clean.split('__B__')

            first_col = (c_idx == 0)

            for i, part in enumerate(parts):
                if not part:
                    continue
                run = para.add_run(part)
                is_bold_segment = (i % 2 == 1)
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif first_col or is_bold_segment:
                    run.bold = True

    doc.add_paragraph()   # spacer after table

# ── code-block renderer (as indented Normal paragraphs) ─────────────────────

def add_code_block(doc, lines):
    """Render a fenced code block as Normal paragraphs with Courier New."""
    for line in lines:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(line.rstrip())
        run.font.name = 'Courier New'

# ── blockquote / callout notes ───────────────────────────────────────────────

def add_note(doc, text: str):
    """
    Render > blockquote lines as a List Bullet with a bullet char,
    inline bold, matching the reference style for clinical pearls.
    Uses List Bullet so it inherits the document's bullet styling.
    """
    # strip leading > markers
    clean = re.sub(r'^>+\s*', '', text).strip()
    if not clean:
        return
    # nested blockquote table rows that start with | — treat as table row note
    if clean.startswith('|'):
        # skip — these are inner table-within-blockquote; handled elsewhere
        return
    p = doc.add_paragraph(style='List Bullet')
    add_inline_runs(p, clean)

# ── main parser ──────────────────────────────────────────────────────────────

def build():
    with open(MD, encoding='utf-8') as f:
        raw = f.read()

    lines = raw.split('\n')
    n     = len(lines)
    i     = 0

    in_code       = False
    code_buf      = []
    in_table      = False
    table_buf     = []
    in_transcript = False
    transcript_buf= []

    # ── flush helpers ──────────────────────────────────────────────────────
    def flush_code():
        nonlocal in_code, code_buf
        if code_buf:
            add_code_block(doc, code_buf)
        in_code  = False
        code_buf = []

    def flush_table():
        nonlocal in_table, table_buf
        if table_buf:
            add_table(doc, table_buf)
        in_table  = False
        table_buf = []

    def flush_transcript():
        nonlocal in_transcript, transcript_buf
        if transcript_buf:
            for tline in transcript_buf:
                tline = tline.strip()
                if tline:
                    add_normal(doc, tline)
        in_transcript  = False
        transcript_buf = []

    # Lines to skip entirely (metadata header, TOC heading, TOC link list, etc.)
    SKIP_PATTERNS = [
        r'^\*\*Course:\*\*',
        r'^\*\*Lecturer:\*\*',
        r'^\*\*Source Materials:\*\*',
        r'^## Table of Contents',
        r'^\d+\.\s+\[.*\]\(#.*\)',   # TOC numbered links like "1. [Title](#anchor)"
        r'^>\s+\*\*Study Note',       # the blockquote study note at the very top
        r'^\*End of Lecture',
        r'^\*Prepared from',
        r'^\*"Understanding',
        r'^\*\*Study Note',
        r'^---$',                    # standalone HR — no visible element needed
    ]

    # ── line-by-line ───────────────────────────────────────────────────────
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip metadata / TOC lines
        skip = False
        for pat in SKIP_PATTERNS:
            if re.match(pat, stripped):
                skip = True
                break
        if skip:
            i += 1
            continue

        # ── code fences ───────────────────────────────────────────────────
        if stripped.startswith('```'):
            if in_code:
                flush_code()
            else:
                flush_table()
                flush_transcript()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── markdown table rows ───────────────────────────────────────────
        if stripped.startswith('|'):
            flush_transcript()
            in_table = True
            table_buf.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── headings ──────────────────────────────────────────────────────
        # H1: document title — skip (already there) except if it's the
        # "Lecture 2" line which we DO want as a new Heading 1
        if re.match(r'^# ', line) and not re.match(r'^## ', line):
            flush_transcript()
            title = line[2:].strip()
            # Skip the big document title and TOC entries
            if title.startswith('Lecture 2') or title.startswith('Dermatology Lecture 2'):
                add_h1(doc, title)
            i += 1
            continue

        if re.match(r'^## ', line):
            flush_transcript()
            heading = line[3:].strip()
            add_h2(doc, heading)
            i += 1
            continue

        if re.match(r'^### ', line):
            flush_transcript()
            heading = line[4:].strip()
            add_h3(doc, heading)
            # If this is the transcript heading, enter transcript mode
            if 'Transcript' in heading or 'Translated' in heading:
                in_transcript = True
            i += 1
            continue

        if re.match(r'^#### ', line):
            flush_transcript()
            add_h4(doc, line[5:].strip())
            i += 1
            continue

        # ── horizontal rule — already handled by SKIP_PATTERNS above ─────

        # ── blockquote (clinical pearls / notes) ─────────────────────────
        if stripped.startswith('>'):
            flush_transcript()
            block_lines = []
            while i < n and lines[i].strip().startswith('>'):
                inner = lines[i].strip().lstrip('> ').strip()
                block_lines.append(inner)
                i += 1
            # Each non-table, non-empty line becomes its own List Bullet
            current = []
            for bl in block_lines:
                # separator table rows — skip
                if re.match(r'^\|[-|: ]+\|$', bl):
                    continue
                # an inner markdown table row — treat as plain text stripped of pipes
                if bl.startswith('|') and bl.endswith('|'):
                    cells = [c.strip() for c in bl[1:-1].split('|')]
                    plain = '  |  '.join(c for c in cells if c)
                    current.append(plain)
                elif bl == '':
                    # blank line inside blockquote → flush current accumulation
                    if current:
                        add_note(doc, ' '.join(current))
                        current = []
                else:
                    current.append(bl)
            if current:
                add_note(doc, ' '.join(current))
            continue

        # ── transcript mode ───────────────────────────────────────────────
        if in_transcript:
            if stripped == '---':
                flush_transcript()
            else:
                transcript_buf.append(line)
            i += 1
            continue

        # ── bullet points ─────────────────────────────────────────────────
        # sub-bullet:  "  - " or "  * " (indented)
        if re.match(r'^\s{2,}[-*] ', line):
            text = re.sub(r'^\s+[-*]\s+', '', line)
            add_bullet(doc, text, level=1)
            i += 1
            continue

        # top-level bullet: "- " or "* "
        if re.match(r'^[-*] ', stripped):
            text = stripped[2:]
            add_bullet(doc, text, level=0)
            i += 1
            continue

        # numbered list:  "1. " or "  1. "
        if re.match(r'^\s*\d+\.\s', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            add_list_number(doc, text)
            i += 1
            continue

        # ── blank lines ───────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── everything else → Normal paragraph ───────────────────────────
        # (this catches the intro study-note block-quote text etc.)
        add_normal(doc, stripped)
        i += 1

    # flush anything left open
    flush_code()
    flush_table()
    flush_transcript()

    doc.save(OUT)
    print(f"Saved → {OUT}")


if __name__ == '__main__':
    build()
