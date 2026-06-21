"""
Rebuild SCE_ConnectAC_Study_Guide with Lectures 7 and 8 (WITH IMAGES).

Starts from: SCE_ConnectAC_Study_Guide_Lec1-6.docx (base)
Appends: Lecture 7 (Lichen Planus) with 16 inline images
Then:    Lecture 8 (Other Lichenoid Dermatoses) with 11 inline images
Saves:   SCE_ConnectAC_Study_Guide_Lec1-8.docx (overwrite)

Uses Pillow re-save technique to fix JPEG header issues.
"""

import re, os, io
from docx import Document
from docx.shared import Inches
from PIL import Image

# ── Paths ────────────────────────────────────────────────────────────────────
BASE_PATH = r"C:\Users\ahaww\Downloads\mypdf\SCE_ConnectAC_Study_Guide_Lec1-6.docx"
OUT_PATH  = r"C:\Users\ahaww\Downloads\mypdf\SCE_ConnectAC_Study_Guide_Lec1-8.docx"
LEC7_IMG_DIR = r"C:\Users\ahaww\Downloads\mypdf\ConnectAC_Slides\lec7_images"
LEC8_IMG_DIR = r"C:\Users\ahaww\Downloads\mypdf\ConnectAC_Slides\lec8_images"

# ── Load base document ───────────────────────────────────────────────────────
doc = Document(BASE_PATH)

# ── Image helper (Pillow re-save technique) ──────────────────────────────────

def add_image(page_num, lecture='lec7', width_inches=5.0):
    """Add an image from extracted slides, re-saving via Pillow to fix JPEG headers."""
    if lecture == 'lec7':
        img_dir = LEC7_IMG_DIR
    else:
        img_dir = LEC8_IMG_DIR
    fname = f"{lecture}_p{page_num}_0.jpeg"
    path = os.path.join(img_dir, fname)
    if os.path.exists(path):
        img = Image.open(path)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        stream = io.BytesIO()
        img.save(stream, format='JPEG', quality=90)
        stream.seek(0)
        try:
            doc.add_picture(stream, width=Inches(width_inches))
            print(f"  ✓ Inserted image: {fname}")
        except Exception as e:
            print(f"  ✗ Failed to insert {fname}: {e}")
    else:
        print(f"  ✗ Image not found: {path}")

# ── inline-markdown parser  **bold** *italic* ────────────────────────────────

def add_inline_runs(para, text: str):
    tokens = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = para.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            run = para.add_run(tok[1:-1])
            run.italic = True
        else:
            para.add_run(tok)

# ── paragraph adders ─────────────────────────────────────────────────────────

def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.add_run(text)

def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.add_run(text)

def add_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.add_run(text)

def add_h4(text):
    p = doc.add_paragraph(style='Heading 4')
    p.add_run(text)

def add_normal(text: str):
    p = doc.add_paragraph(style='Normal')
    speaker_pat = re.match(
        r'^(\*\*)?((Dr\.|Student|Professor)[^:]*:)(\*\*)?\s*(.*)',
        text, re.DOTALL
    )
    if speaker_pat:
        label = speaker_pat.group(2).strip()
        rest  = speaker_pat.group(5).strip()
        run_l = p.add_run(label + '  ')
        run_l.bold = True
        add_inline_runs(p, rest)
    else:
        add_inline_runs(p, text.strip())

def add_bullet(text: str, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text.strip())

def add_list_number(text: str):
    p = doc.add_paragraph(style='List Number')
    add_inline_runs(p, text.strip())


# ══════════════════════════════════════════════════════════════════════════════
# ██  LECTURE 7: LICHEN PLANUS (WITH IMAGES)
# ══════════════════════════════════════════════════════════════════════════════
print("\n" + "="*60)
print("  APPENDING LECTURE 7: LICHEN PLANUS (WITH IMAGES)")
print("="*60)

add_h1("Lecture 7: Lichen Planus")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 1 – Interface Dermatitis & Lichenoid Tissue Reaction
# ═══════════════════════════════════════════════════════════════════════
add_h2("1. Interface Dermatitis & Lichenoid Tissue Reaction (Introduction)")
add_h3("Summary")

add_h4("Definition of Interface Dermatitis")
add_bullet("**Interface dermatitis** = Vacuolar degeneration of the basement membrane (dermoepidermal junction) + inflammatory infiltrate in the dermis.")
add_bullet("The basement membrane zone (BMZ) is damaged/destroyed, and lymphocytic inflammatory infiltrate is present.")

add_h4("Two Subtypes Based on Pattern of Lymphocytic Infiltrate")
add_bullet("**Lichenoid interface dermatitis:** Band-like (dense) lymphocytic infiltrate immediately beneath the DEJ.")
add_bullet("Diseases in this group: **Lichen planus**, **Lichen striatus**, **Lichen nitidus**, **Lichenoid drug eruption**.", level=1)
add_bullet("**Vacuolar interface dermatitis:** Perivascular lymphocytic infiltrate (lymphocytes around blood vessels rather than band-like).")
add_bullet("Diseases in this group: **GVHD**, **PLEVA/PLC**, **Erythema multiforme**, **Connective tissue diseases** (SLE, Dermatomyositis), **Fixed drug eruption**, **Secondary syphilis**, **Paraneoplastic pemphigus**.", level=1)

# >>> IMAGE: Interface dermatitis diagram (page 2)
add_image(2, 'lec7')

add_h4("Key Concept: Spectrum Not Dichotomy")
add_bullet("These two patterns are a **spectrum**, not strict separate entities.")
add_bullet("Mild inflammation → lymphocytes remain perivascular (vacuolar pattern).", level=1)
add_bullet("Intense/increasing inflammation → lymphocytes become band-like (lichenoid pattern).", level=1)
add_bullet("Understanding this shared histopathology saves time: many diseases share the same BMZ destruction + lymphocytic infiltrate foundation.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Good morning doctors, how are you? Ramadan Kareem, may these blessed days return to us all, God willing.")
add_normal("Student:  God is most generous, happy Ramadan to you too doctor.")
add_normal("Dr. Eman:  Have you adjusted to the fasting schedule yet? I don't want anyone sleeping through Ramadan — we still need to study! Tell me, how is the revision going? Have you started doing practice questions yet?")
add_normal("Student:  It's the start of Ramadan so we're still settling in.")
add_normal("Dr. Eman:  Of course, everything stops in the first week of Ramadan. Does anyone have questions from previous lectures? Last time we finished Psoriasis and all the Papulosquamous diseases. Today we'll enter Lichenoid dermatoses and specifically Lichen planus.")
add_normal("Dr. Eman:  What are the Lichenoid dermatoses? Last time we briefly touched on this when we discussed PLEVA and PLC. We talked about something called Vacuolar interface dermatitis. What does Interface dermatitis mean?")
add_normal("Dr. Eman:  The Lichenoid dermatoses are a heterogeneous group of conditions that clinically resemble Lichen planus — pruritic, violaceous, flat-topped papules with variations. But histopathologically they share a common feature: the Lichenoid tissue reaction, which is vacuolar degeneration of the basement membrane plus a band-like lymphocytic inflammatory infiltrate in the dermis.")
add_normal("Dr. Eman:  Interface dermatitis means the basement membrane is destroyed — vacuolar degeneration of the dermoepidermal junction — with inflammatory infiltrate. Based on the pattern of the inflammatory infiltrate, we divide interface dermatitis into two categories: Lichenoid interface dermatitis (band-like lymphocytes beneath the DEJ) or Vacuolar interface dermatitis (perivascular lymphocytes).")
add_normal("Dr. Eman:  Don't get confused between the term 'vacuolar interface dermatitis' as a subtype, and the fact that interface dermatitis itself involves vacuolar degeneration. In lichenoid interface dermatitis, the infiltrate is band-like — lined up like a stripe beneath the DEJ. In vacuolar interface, it's perivascular. They used to classify them as two completely separate entities, but recent evidence shows they are actually one spectrum. When the inflammatory infiltrate is mild, it stays around blood vessels. When it becomes very intense, it forms a band-like pattern.")
add_normal("Dr. Eman:  Under the band-like (lichenoid) group, anything with 'Lichen' in the name goes here: Lichen planus, Lichen striatus, Lichen nitidus, Lichenoid drug eruption. Under the vacuolar (perivascular) group: GVHD, PLEVA and PLC, Erythema multiforme, Connective tissue diseases like SLE and Dermatomyositis, Fixed drug eruption, Secondary syphilis, and Paraneoplastic pemphigus. These are diseases from completely different chapters, but they all share the same histopathology!")
add_normal("Student (Dr. Narges):  Vacuolar interface dermatitis, doctor — which is vacuolar degeneration of the basement membrane and perivascular inflammatory infiltrate.")
add_normal("Dr. Eman:  Exactly! This interface dermatitis classification is extremely important and will save you a lot of time. Imagine how many diseases are vacuolar interface dermatitis? So when you're studying SLE and remember it has BMZ destruction with perivascular infiltrate, you've already covered half its histopathology. And when you see a histopathology slide with vacuolar degeneration and lymphocytic infiltrate, you know it's interface dermatitis — now just narrow it down among these diseases.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 2 – Definition, Epidemiology & Etiopathogenesis
# ═══════════════════════════════════════════════════════════════════════
add_h2("2. Lichen Planus — Definition, Epidemiology & Etiopathogenesis")
add_h3("Summary")

add_h4("Definition")
add_bullet("**Lichen planus (LP)** = Idiopathic inflammatory disease affecting the skin, mucous membranes, hair, and nails.")
add_bullet("Name derivation: **Lichen** = papule; **Planus** = flat → flat-topped violaceous papules.")
add_bullet("Characterized clinically by the **7 Ps** and histopathologically by dense band-like lymphocytic infiltrate with destruction of the basal cell layer (lichenoid interface dermatitis).")

# >>> IMAGE: LP clinical photo - violaceous papules (page 4)
add_image(4, 'lec7')

add_h4("Epidemiology")
add_bullet("Worldwide prevalence: **0.2–1%** of the adult population (cutaneous LP).")
add_bullet("Oral LP: **1–4%** of the population.")
add_bullet("Age of onset: **5th–6th decade** of life.")

add_h4("Etiopathogenesis")
add_bullet("Follows the standard model: **Environmental trigger** + **Genetic susceptibility** → **Immune dysregulation** → Disease.")
add_bullet("This is not a hereditary disease, but the patient must be genetically predisposed.")

add_h4("Environmental Triggers / Risk Factors")
add_bullet("**Hepatitis C virus** (most important — must screen all LP patients, especially in endemic areas).")
add_bullet("**Drugs:** antimalarials, D-penicillamine, diuretics (these cause Lichenoid drug eruption).")
add_bullet("**Amalgam dental fillings** (especially relevant for oral LP).")

add_h4("Immune Mechanism")
add_bullet("**CD8+ cytotoxic T-cells** attack and destroy basal keratinocytes.")
add_bullet("Same mechanism as Erythema multiforme (EM) and Toxic Epidermal Necrolysis (TEN), but LP is less severe.")
add_bullet("EM/TEN: same CD8+ cytotoxic attack but much more aggressive → complete separation of dermis and epidermis.", level=1)

# >>> IMAGE: Etiopathogenesis diagram (page 6)
add_image(6, 'lec7')

add_h4("Associations")
add_bullet("**Hepatitis C & B** (risk factors).")
add_bullet("**Liver disease**, **Diabetes mellitus**.")
add_bullet("**Autoimmune diseases:** Alopecia areata, Vitiligo (LP itself is considered autoimmune, hence association with other autoimmune conditions).")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Today's entire lecture will be about Lichen planus. From its name — Lichen means papule and Planus means flat. So it's a flat-topped violaceous papule. Lichen planus is an idiopathic inflammatory disease of the skin, mucous membranes, hair, and nails. Clinically characterized by the 7 Ps: pruritic, violaceous papules that favor the extremities. Histopathologically: dense band-like lymphocytic infiltrate and destruction of the basal cell layer — which is lichenoid interface dermatitis.")
add_normal("Dr. Eman:  Have we all seen Lichen planus before? Anyone who has worked in dermatology will know it. We see violaceous, purple papules that are very itchy, commonly on the wrists and hands — those are the most characteristic sites.")
add_normal("Student:  Hepatitis C.")
add_normal("Dr. Eman:  Yes, Virus C. It's a risk factor. Are there other risk factors? We'll learn them now.")
add_normal("Dr. Eman:  The epidemiology: cutaneous LP represents 0.2% to 1% of the adult population. Oral LP represents 1 to 4%. Age of onset is the 5th and 6th decade.")
add_normal("Dr. Eman:  What causes Lichen planus? The same arrow diagram I showed you during Psoriasis — Environmental trigger in a genetically susceptible patient causes immune dysregulation, which manifests as disease. This is NOT hereditary — it means the patient is genetically predisposed, and when exposed to the environmental factor, immune problems develop and the disease appears.")
add_normal("Student:  Sometimes it's associated with thyroid disease, maybe diabetes. The antidiabetic drugs...")
add_normal("Dr. Eman:  Those are associations, not environmental factors. Since LP is an autoimmune disease, it associates with other autoimmune diseases. I'm asking about the environmental trigger specifically.")
add_normal("Dr. Eman:  What's the immune dysregulation in LP? Is the problem in T-helper or T-cytotoxic cells?")
add_normal("Student:  Probably cytotoxic.")
add_normal("Dr. Eman:  Correct — CD8+ cytotoxic T-cells. Does anyone know another dermatological disease with the same mechanism?")
add_normal("Student:  Toxic epidermal necrolysis.")
add_normal("Dr. Eman:  Excellent! EM and TEN use the same mechanism. But of course in TEN it's much more aggressive — to the point where the dermis and epidermis completely separate from each other.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 3 – Clinical Features: The 7 Ps & Classification System
# ═══════════════════════════════════════════════════════════════════════
add_h2("3. Clinical Features — The 7 Ps & Classification System")
add_h3("Summary")

add_h4("The 7 Ps of Classic Lichen Planus")
add_list_number("**Papule**")
add_list_number("**Pruritic**")
add_list_number("**Polygonal**")
add_list_number("**Purple** (violaceous)")
add_list_number("**Planar** (flat-topped)")
add_list_number("**Polished** (shiny surface)")
add_list_number("**Plentiful**")

# >>> IMAGE: Classic LP photo (page 8)
add_image(8, 'lec7')

add_h4("Wickham Striae")
add_bullet("**Wickham striae** = Whitish streaks on the surface of LP lesions.")
add_bullet("Best seen with **dermoscopy** (sometimes visible clinically).")
add_bullet("Caused by **focal wedge-shaped hypergranulosis** in the epidermis.")

add_h4("Most Characteristic Sites")
add_bullet("**Flexor surfaces**, especially the **wrists** and **forearms**.")
add_bullet("**Dorsum of the hands**.")

add_h4("Classification System: '4 Groups of 4' = 16 Variants Total")
add_bullet("**Group 1 — By Site:**")
add_bullet("Skin, Mucous membrane (oral/genital), Hair (Lichen planopilaris), Nail.", level=1)
add_bullet("**Group 2 — Special Types:**")
add_bullet("Actinic, Pigmented (LP Pigmentosus), Palmoplantar, Inverse.", level=1)
add_bullet("**Group 3 — By Configuration:**")
add_bullet("Annular, Linear, Ulcerative/Guttate.", level=1)
add_bullet("**Group 4 — Pairs:**")
add_bullet("Hypertrophic/Atrophic, Bullous LP/LP Pemphigoides, LP-LE Overlap, Eruptive.", level=1)
add_bullet("Important: In any LP case, always examine the **oral cavity**, **hair**, and **nails** — exam tip!")

# >>> IMAGE: Variants overview slide (page 10)
add_image(10, 'lec7')

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  The clinical picture of classic Lichen planus: the 7 Ps. Papule, Pruritic, Polygonal, Purple, Planar, Polished, Plentiful. And there's no Wickham striae in the 7 Ps list — Wickham striae are whitish streaks on the lesion surface. We see them sometimes clinically, but more commonly with the dermoscope. The most characteristic site is the flexor surfaces — wrists, forearms, and dorsum of the hands.")
add_normal("Student:  Doctor, regarding environmental factors — when a drug like a diuretic or antimalarial causes it, do you call this Lichen planus or Lichenoid drug eruption?")
add_normal("Dr. Eman:  Drug-induced cases are Lichenoid drug eruption. The clinical presentation is different — it tends to be more generalized, affects older patients, and lacks Wickham striae. We'll discuss this at the end of the lecture.")
add_normal("Dr. Eman:  How many variants of LP are there? Remember Psoriasis was 4, 7, 2, 8? Here it's 4 groups of 4 — so 16 types total. First group by site: Skin, Mucous membrane (oral or genital), Hair, and Nail. In any LP case you suspect, you MUST examine oral mucosa, hair, and nails. If you're in a clinical exam and don't check these, you've essentially failed the case.")
add_normal("Dr. Eman:  Second group: Actinic, Pigmented, Palmoplantar, Inverse. Third group by configuration: Annular, Linear, Ulcerative/Guttate. Last group — pairs: Hypertrophic or Atrophic, Bullous LP or LP Pemphigoides, LP-LE Overlap, and Eruptive LP. So we have 16 types total. Note that each one is a separate variant — the pairs are four variants, not two.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 4 – Variants: Mucosal, Hair & Nail
# ═══════════════════════════════════════════════════════════════════════
add_h2("4. Lichen Planus Variants — Mucosal, Hair & Nail")
add_h3("Summary")

add_h4("Oral Lichen Planus")
add_bullet("Three main forms: **Erosive/Ulcerative** (associated with HCV), **Reticular** (white lace-like lines), **Leukoplakia** (whitish patch/plaque).")
add_bullet("May be the **only site** of LP involvement.")
add_bullet("**50%** of patients with cutaneous LP also have oral involvement.")
add_bullet("Differential diagnosis: **Pemphigus vulgaris**, **Behçet's disease**, **Aphthous stomatitis**.")
add_bullet("**Risk of SCC** (squamous cell carcinoma) — one of the highest-risk LP sites for malignancy.")

# >>> IMAGE: Oral LP photos (page 12)
add_image(12, 'lec7')

add_h4("Genital Lichen Planus")
add_bullet("May occur alone or with oral/cutaneous LP.")
add_bullet("**Glans penis:** Annular violaceous papules; DDx: Syphilis.")
add_bullet("**Vulvovaginal LP:** Always erosive → scarring → **dyspareunia** → **malignancy risk**.")
add_bullet("Differential diagnosis: **Lichen sclerosus et atrophicus**, **Cicatricial pemphigoid**.")
add_bullet("Oral and genital LP carry the **highest risk of SCC** among all LP variants.")

add_h4("Lichen Planopilaris (Hair LP)")
add_bullet("LP affecting the hair follicle = **Lichen planopilaris**.")
add_bullet("Appearance: **Keratotic plug surrounded by violaceous rim** — characteristic even without dermoscopy.")
add_bullet("Causes **lymphocytic cicatricial (scarring) alopecia**.")
add_bullet("Perifollicular scale and hair casts visible on dermoscopy.")
add_bullet("Differential diagnosis: **Discoid lupus erythematosus** (also causes cicatricial alopecia on scalp).")

# >>> IMAGE: Lichen planopilaris photo (page 15)
add_image(15, 'lec7')

add_h4("Variants of Lichen Planopilaris")
add_bullet("**Frontal Fibrosing Alopecia (FFA):**")
add_bullet("Recession of the frontal hairline.", level=1)
add_bullet("Predominantly affects **postmenopausal females**.", level=1)
add_bullet("More common in **high socioeconomic** groups.", level=1)
add_bullet("Suspected link to **sunscreens and skin care products** (no confirmed single cause).", level=1)
add_bullet("Skin-colored papules on forehead and chin may be present.", level=1)
add_bullet("**Graham-Little Syndrome:**")
add_bullet("Triad: Cicatricial alopecia of the scalp + Keratotic (spine-like) papules on skin + Non-cicatricial alopecia of body hair (axilla/groin).", level=1)
add_bullet("DDx of Graham-Little: Keratosis pilaris, Follicular mucinosis.", level=1)

add_h4("Nail Lichen Planus")
add_bullet("**Longitudinal lines/striations** — most characteristic and earliest sign (what you'll actually see clinically).")
add_bullet("Thinning and shedding of the nail plate.")
add_bullet("**Pterygium** (dorsal) — end-stage sign: dorsal nail fold adheres to nail bed after nail plate destruction.")
add_bullet("**Twenty-nail dystrophy** (all 20 nails affected) — specific to LP and atopic dermatitis.")
add_bullet("**Melanonychia** may occur.")

# >>> IMAGE: Nail LP - pterygium (page 17)
add_image(17, 'lec7')

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Let's start with Oral LP. It's one of the most important variants. I told you — if you see skin LP, you must examine the oral cavity. Oral LP comes in several forms: Erosive or Ulcerative (this one is always associated with Hepatitis C virus), Reticular (the white lines you can see), and Leukoplakia (whitish patch or plaque). It may be the only site of involvement, and 50% of cutaneous LP patients have oral involvement too.")
add_normal("Dr. Eman:  The differential diagnosis for oral LP: Pemphigus vulgaris is the most important, then Behçet's disease, and Aphthous stomatitis.")
add_normal("Dr. Eman:  Genital LP — it may occur alone or with oral or cutaneous disease. On the glans penis, you see annular violaceous papules. The DDx is syphilis. Vulvovaginal LP is always erosive, causes scarring, leads to dyspareunia, and ultimately carries malignancy risk. The sites with highest malignancy risk after LP are oral and genital. DDx: Lichen sclerosus et atrophicus and Cicatricial pemphigoid.")
add_normal("Dr. Eman:  Hair LP is called Lichen planopilaris. It presents as a keratotic plug surrounded by a violaceous rim — this appearance is so characteristic you can diagnose it even before putting on the dermoscope. It causes cicatricial alopecia. The DDx is Discoid LE.")
add_normal("Dr. Eman:  Frontal Fibrosing Alopecia is a variant of Lichen planopilaris. Who does it affect?")
add_normal("Student:  Females, postmenopausal, high socioeconomic class.")
add_normal("Dr. Eman:  Correct. The suspected cause is sunscreens and skin products — those in higher socioeconomic groups tend to use more of these products. But have they found a definitive cause? No, but this is the strongest association. There's regression of the forehead hairline with skin-colored papules on the forehead and chin.")
add_normal("Dr. Eman:  The second variant is Graham-Little syndrome: cicatricial alopecia of the scalp plus keratotic spine-like papules on the skin plus non-cicatricial alopecia of body hair in the axilla and groin.")
add_normal("Dr. Eman:  Now the nail. What's the appearance we see here?")
add_normal("Student:  Pterygium.")
add_normal("Dr. Eman:  Yes, pterygium. But does only pterygium occur? No! Pterygium is the end stage. First: thinning, shedding, and most importantly longitudinal lines — that's the most characteristic sign you'll actually see because patients come early. The pterygium is dorsal pterygium — the dorsal nail fold adheres to the nail bed after the nail plate is destroyed. Twenty-nail dystrophy is also specific to LP. Remember: longitudinal lines, pterygium, and twenty-nail dystrophy are the key nail findings.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 5 – Variants: Skin Types
# ═══════════════════════════════════════════════════════════════════════
add_h2("5. Lichen Planus Variants — Skin Types")
add_h3("Summary")

add_h4("Actinic Lichen Planus")
add_bullet("Occurs on **sun-exposed areas**, especially the **forehead above the eyebrows** (most characteristic site).")
add_bullet("Appearance: **Annular hyperpigmented** lesion **surrounded by a hypopigmented rim**.")
add_bullet("Colour is violaceous/dark rather than brown (differentiates from melasma).")
add_bullet("**NOT pruritic**.")
add_bullet("Variants: Classic (annular), Melasma-like, Dyschromic.")
add_bullet("Differential diagnosis: **Melasma**.")

# >>> IMAGE: Actinic LP photo (page 19)
add_image(19, 'lec7')

add_h4("LP Pigmentosus")
add_bullet("Hyperpigmented LP appearing in **both sun-exposed AND sun-protected areas** (key difference from Actinic LP).")
add_bullet("**IS pruritic** (differentiates from Actinic LP which is non-pruritic).")

add_h4("Palmoplantar Lichen Planus")
add_bullet("**Yellowish, firm papules/plaques** surrounded by a **violaceous rim** — very characteristic.")
add_bullet("**Non-itchy**.")
add_bullet("Differential diagnosis: **Syphilis**, **Psoriasis**, **Callosity**, **Warts**.")

# >>> IMAGE: Yellow papules photo (page 21)
add_image(21, 'lec7')

add_h4("Inverse Lichen Planus")
add_bullet("Intertriginous areas: **Axilla** (more common) and **groin**.")
add_bullet("Pink to violaceous papules and plaques.")
add_bullet("DDx: **LP Pigmentosus** (because inverse LP heals with post-inflammatory hyperpigmentation, mimicking LP pigmentosus).")

add_h4("Annular Lichen Planus")
add_bullet("Most commonly found on the **penis** and **axilla**.")

add_h4("Linear Lichen Planus")
add_bullet("Follows **Blaschko's lines** or has a **zosteriform** distribution.")
add_bullet("Not truly dermatomal — follows embryological developmental lines.")

add_h4("Ulcerative Lichen Planus")
add_bullet("Characteristically on the **sole** of the foot.")
add_bullet("**Risk of SCC** (squamous cell carcinoma).")

add_h4("Hypertrophic Lichen Planus")
add_bullet("Location: **Shin of the tibia** (anterior lower leg).")
add_bullet("**Highly pruritic** and **chronic**.")
add_bullet("**Risk of SCC** with long-standing lesions.")
add_bullet("Differential diagnosis: **Lichen simplex chronicus**, **Amyloidosis**, **Prurigo nodularis**.")

# >>> IMAGE: Hypertrophic LP on shin (page 23)
add_image(23, 'lec7')

add_h4("Why LP Leaves Persistent Hyperpigmentation")
add_bullet("BMZ destruction → **melanocytes fall from the basal layer into the dermis**.")
add_bullet("Melanocytes are engulfed by macrophages → **melanophages**.")
add_bullet("Melanophages persist indefinitely in the dermis → pigmentation does not fade easily.")
add_bullet("This explains why LP always leaves stubborn post-inflammatory hyperpigmentation.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Actinic LP — from its name, it appears on sun-exposed skin, mainly the face. It doesn't come as typical papules; rather you see pigmentation that looks violaceous, almost purple-black rather than brown like melasma. It's annular hyperpigmented surrounded by a hypopigmented rim. The most characteristic area is above the eyebrows. It is NOT pruritic. DDx: Melasma.")
add_normal("Dr. Eman:  LP Pigmentosus is hyperpigmented LP that appears in both sun-exposed AND sun-protected areas. And it IS pruritic. That's how you differentiate it from Actinic LP — Actinic is non-pruritic and only sun-exposed; Pigmentosus is pruritic and involves both areas.")
add_normal("Dr. Eman:  Palmoplantar LP — the most characteristic feature is yellowish papules with a rim of violaceous or hyperpigmentation around them. Yellowish, firm, non-itchy papules or plaques surrounded by a violaceous rim. Very characteristic and we see it frequently. DDx: Syphilis, Psoriasis, Callosity, Warts.")
add_normal("Dr. Eman:  Inverse LP occurs in the axilla (more common) and groin — intertriginous areas. Pink to violaceous papules and plaques. The DDx is LP Pigmentosus — why? Because when inverse LP heals, it leaves hyperpigmentation. And LP always leaves hyperpigmentation. So when I see that residual hyperpigmentation in the axilla, is it LP Pigmentosus or post-inflammatory from resolved Inverse LP? That's why they're DDx of each other.")
add_normal("Dr. Eman:  Now my question — why does LP leave such stubborn hyperpigmentation that barely fades?")
add_normal("Student:  Because of degeneration — the melanocytes get deposited deep in the dermis.")
add_normal("Dr. Eman:  Excellent! The basement membrane is destroyed, so melanocytes from the basal layer fall down into the dermis, get engulfed by macrophages — we call them melanophages — and these melanophages just sit there permanently. That's why the pigmentation is so persistent.")
add_normal("Dr. Eman:  Annular LP is always on the penis and axilla. Linear LP follows Blaschko's lines or a zosteriform distribution — not truly dermatomal but linear along developmental lines. Ulcerative LP occurs on the sole and can cause SCC. Hypertrophic LP — elevated above the skin, on the shin of the tibia, highly pruritic, chronic, and can also cause SCC. DDx: Lichen simplex chronicus, Amyloidosis, Prurigo nodularis.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 6 – Bullous LP vs LP Pemphigoides & LP-LE Overlap
# ═══════════════════════════════════════════════════════════════════════
add_h2("6. Bullous LP vs LP Pemphigoides & LP-LE Overlap")
add_h3("Summary")

add_h4("Bullous Lichen Planus")
add_bullet("Bullae develop **ON existing LP lesions** (the bullae are within the LP papules/plaques).")
add_bullet("Cause: Intense inflammatory infiltrate → severe BMZ destruction → Max Joseph space (cleft at DEJ) → fluid accumulates → bullae form.")
add_bullet("This is the same LP pathology, just more severe — no autoantibodies involved.")

# >>> IMAGE: Bullous LP photo (page 25)
add_image(25, 'lec7')

add_h4("LP Pemphigoides")
add_bullet("Bullae develop on **uninvolved (normal) skin** — NOT on LP lesions.")
add_bullet("Cause: **Autoantibodies** (against BP antigen 2 / Collagen XVII) attack skin anywhere, independent of LP lesions.")
add_bullet("**DIF (Direct Immunofluorescence) positive:** IgG on uninvolved skin.")
add_bullet("Key differentiator: Bullous LP = bullae ON LP lesions; LP Pemphigoides = bullae on NORMAL skin.")

add_h4("LP-LE Overlap")
add_bullet("Features of **both LP and Lupus Erythematosus** combined.")
add_bullet("Appearance: Violaceous to erythematous patches/plaques with **atrophy**, mild **scaling**, and **telangiectasia**.")
add_bullet("**Acral distribution** (hands and feet).")
add_bullet("Oral lesions are common.")

add_h4("Eruptive Lichen Planus")
add_bullet("**Generalized** distribution.")
add_bullet("Tends to affect **older** patients.")
add_bullet("**No Wickham striae** (differentiates from classic LP).")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Bullous LP versus LP Pemphigoides — who can tell me the difference?")
add_normal("Student:  Bullous LP is the bullous variant — the pathology is still Lichen planus. The bullae arise from intense inflammation at the DEJ and Max Joseph space formation.")
add_normal("Dr. Eman:  Exactly. The intense inflammatory infiltrate destroys the basal cell layer, causes separation — naturally when the dermis separates from the epidermis, a space forms (Max Joseph space), fills with fluid, and you get bullae. The key point: these bullae appear ON the Lichen planus lesions themselves.")
add_normal("Dr. Eman:  LP Pemphigoides is different. Here autoantibodies against BP antigen 2 (Collagen XVII) attack the skin. The bullae appear on uninvolved skin — not necessarily where LP lesions exist. LP did its job destroying the BMZ, which then triggered autoantibody production, and these antibodies attack skin anywhere. DIF shows IgG on uninvolved skin.")
add_normal("Dr. Eman:  LP-LE Overlap has features of both LP and Lupus erythematosus — violaceous to erythematous patches and plaques with atrophy, mild scaling, and telangiectasia, usually in an acral distribution. Oral lesions from LP are present alongside LE features.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 7 – Histopathology & Diagnosis
# ═══════════════════════════════════════════════════════════════════════
add_h2("7. Histopathology & Diagnosis")
add_h3("Summary")

add_h4("Histopathology Features of Lichen Planus")
add_list_number("**Compact orthokeratosis** — thickened stratum corneum with NO parakeratosis (therefore no significant clinical scaling).")
add_list_number("**Focal wedge-shaped hypergranulosis** — cause of Wickham striae; appears like 'eyes' beneath the stratum corneum.")
add_list_number("**Irregular acanthosis** with **sawtooth appearance** (not regular like psoriasis).")
add_list_number("**Vacuolar degeneration of the BMZ** — basement membrane zone destroyed.")
add_list_number("**Max Joseph spaces** — clefts at the DEJ from severe BMZ destruction; when filled with fluid → bullous LP.")
add_list_number("**Band-like (lichenoid) lymphocytic infiltrate** — dense infiltrate that obscures the DEJ.")
add_list_number("**Melanophages** in the upper dermis (explain persistent hyperpigmentation).")
add_list_number("**Civatte bodies** — apoptotic/dyskeratotic keratinocytes (dead keratinocytes from T-cell attack).")

# >>> IMAGE: Histopathology slide (page 27)
add_image(27, 'lec7')

add_h4("What Creates the Papule?")
add_bullet("The elevated papule results from: **Inflammatory infiltrate below** (pushing up) + **Compact orthokeratosis above** (thickened stratum corneum).")

add_h4("Direct Immunofluorescence (DIF)")
add_bullet("**Linear fibrin deposits** at the DEJ (from keratinocyte destruction).")
add_bullet("**Globular IgM deposits** below the DEJ (red dot-like pattern).")

add_h4("Dermoscopy")
add_bullet("**Wickham striae** = whitish streaks on a violaceous/erythematous background.")
add_bullet("Very characteristic and highly visible under dermoscopy.")

# >>> IMAGE: Wickham striae dermoscopy (page 29)
add_image(29, 'lec7')

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  The histopathology of LP. Look at how dissolved this looks — there's barely a basement membrane left from all the destruction. First: vacuolar degeneration of the BMZ, plus the inflammatory infiltrate — see how dense it is? Band-like inflammatory infiltrate.")
add_normal("Dr. Eman:  Where does the papule come from? The papule elevation comes from the inflammatory infiltrate pushing up from below and the compact orthokeratosis thickening above. Important: compact orthokeratosis with NO parakeratosis. Underline this 100 times — no parakeratosis! What does this mean clinically? It means no significant scaling in LP. Correct?")
add_normal("Dr. Eman:  Then there's focal hypergranulosis in a wedge shape — very characteristic of LP. It looks like eyes beneath the stratum corneum. And this is what causes the Wickham striae! The whitish streaks we see clinically and on dermoscopy are caused by this focal hypergranulosis.")
add_normal("Dr. Eman:  Irregular acanthosis with a sawtooth appearance — not regular like psoriasis, but irregular jagged projections. Vacuolar degeneration of the BMZ. Max Joseph spaces — when the infiltrate destroys the BMZ severely, separation occurs and clefts form at the DEJ. If these fill with fluid, that's what gives you bullous LP.")
add_normal("Dr. Eman:  Band-like lymphocytic infiltrate. Melanophages — the melanocytes that fell into the dermis and were engulfed by macrophages. And Civatte bodies — these are dyskeratotic (apoptotic) keratinocytes. The keratinocytes die from the T-cell attack and form these Civatte bodies.")
add_normal("Dr. Eman:  Direct immunofluorescence in LP shows linear fibrin deposits at the DEJ — from keratinocyte destruction releasing fibrin. And globular IgM deposits below the DEJ — those red dots you see. You must know this for the exam.")
add_normal("Dr. Eman:  Dermoscopy — the Wickham striae I've been telling you about all morning. Whitish streaks dividing the papules. You put the dermatoscope on and they light up in your field of view. So LP dermoscopy = Wickham striae = whitish streaks on an erythematous or violaceous background.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 8 – Treatment
# ═══════════════════════════════════════════════════════════════════════
add_h2("8. Treatment")
add_h3("Summary")

add_h4("General Measures")
add_bullet("**Exclude risk factors:** amalgam dental fillings, causative drugs, screen for HCV (mandatory, especially in endemic areas like Egypt).")
add_bullet("Screen for **Hepatitis C and B**.")
add_bullet("**Antihistamines** for pruritus (LP is inherently pruritic).")

add_h4("Topical Treatment")
add_bullet("**Potent/superpotent topical corticosteroids** — first-line treatment (e.g., Betamethasone valerate / Betnovate).")
add_bullet("Can be used under occlusion for enhanced efficacy.")
add_bullet("**Topical calcineurin inhibitors** (Tacrolimus/Pimecrolimus) — for areas unsuitable for long-term steroids (face, genitals, oral mucosa).")
add_bullet("**Inhaled corticosteroids** — for oral LP.")
add_bullet("Intralesional corticosteroid injection (rarely needed as topical steroids work very well for LP).")

add_h4("Systemic Treatment")
add_bullet("**Prednisolone 20 mg/day** for 2–6 weeks (lower dose than the usual 0.5–1 mg/kg used for other conditions).")
add_bullet("**Metronidazole** (Flagyl) — newer treatment, effective.")
add_bullet("**Griseofulvin** — also used effectively.")
add_bullet("**Antimalarials** — paradoxically can both treat LP AND cause Lichenoid drug eruption.")
add_bullet("Systemic retinoids (Acitretin) — less commonly used.")

add_h4("Phototherapy")
add_bullet("**Narrowband UVB** — for resistant/long-standing cases.")
add_bullet("Especially useful when **LP coexists with HCV** (avoids systemic immunosuppression).")

add_h4("Biologics & New Agents")
add_bullet("**Basiliximab** (anti-IL-2 / IL-2 antagonist) — newest biologic for LP; appears in SCE questions.")
add_bullet("**Apremilast** (Otezla) — PDE4 (phosphodiesterase 4) inhibitor; anti-inflammatory used in psoriasis and LP.")
add_bullet("**JAK inhibitors** (Tofacitinib, Upadacitinib) — not first-line for LP but may appear in questions; FDA-approved for Alopecia areata and Vitiligo.")

# >>> IMAGE: Treatment table (page 31)
add_image(31, 'lec7')

add_h4("LP with Hepatitis C")
add_bullet("**Treat the HCV** (antiviral therapy) + **Narrowband UVB** for the LP.")
add_bullet("Be cautious with systemic corticosteroids in HCV patients.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Student:  Topical corticosteroids.")
add_normal("Dr. Eman:  Exactly — that's the first line. Topical, but give potent steroids appropriate to the site. The most important. Second: topical calcineurin inhibitors for areas where you can't give steroids. Then of course systemic steroids.")
add_normal("Dr. Eman:  Let me organize the treatment for you systematically. First: general measures. You must look for risk factors — does the patient have amalgam dental fillings? Are they taking any causative drugs? You absolutely must screen for Hepatitis C, especially in Egypt where HCV was so prevalent. No LP case should pass without checking HCV — and you can add HBV screening too. Give antihistamines for the pruritus since LP is inherently itchy.")
add_normal("Dr. Eman:  Topical: steroids are essential — topical or intralesional, but topical is the mainstay. For oral LP, we use inhaled steroids. Tacrolimus for areas unsuitable for steroids. Systemic: steroids — the dose for LP is usually 20 mg per day for 2 to 6 weeks. That's lower than the usual 0.5 to 1 mg/kg. Metronidazole and Griseofulvin are both effective newer options. Antimalarials can be used despite also being a potential cause of lichenoid drug eruption — it's a paradox.")
add_normal("Dr. Eman:  Phototherapy — narrowband UVB for resistant, long-standing cases. If you have LP with Hepatitis C, treat the virus and give narrowband UVB. Be careful with steroids in HCV patients.")
add_normal("Dr. Eman:  Now for biologics — what's the newest one for LP?")
add_normal("Student:  Anti-TNF?")
add_normal("Dr. Eman:  No — it's Basiliximab, which targets Interleukin-2. Write this down because it appears in SCE questions. Basiliximab against IL-2 — the newest treatment for LP. Then there's Apremilast (Otezla), a PDE4 inhibitor — anti-inflammatory used in psoriasis, LP, and other autoimmune conditions. And JAK inhibitors like Tofacitinib — not first-line for LP but useful for Alopecia areata and Vitiligo. They can appear in exam questions as immunomodulators applicable across multiple conditions.")
add_normal("Student:  What is the duration of treatment? For how long do you treat, especially oral treatments like Prednisolone?")
add_normal("Dr. Eman:  Clinically, I give patients treatment for a maximum of 6 weeks to 2 months, with follow-up every 2 weeks. The textbooks don't specify a maximum endpoint — it's according to clinical response, risk factors, and side effects. If you're giving oral steroids, consider the dose, whether the patient is diabetic, has HCV, or other comorbidities.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 9 – Complications & Prognosis
# ═══════════════════════════════════════════════════════════════════════
add_h2("9. Complications & Prognosis")
add_h3("Summary")

add_h4("Complications by Site")
add_bullet("**Skin:** Post-inflammatory hyperpigmentation (persistent due to melanophages), scarring, atrophy.")
add_bullet("**Mucosa:** Scarring, stenosis, **SCC** (especially oral and genital).")
add_bullet("**Hair:** Cicatricial (permanent, scarring) alopecia.")
add_bullet("**Nail:** Permanent damage — pterygium (irreversible).")

# >>> IMAGE: Complications overview (page 33)
add_image(33, 'lec7')

add_h4("SCC Risk")
add_bullet("**Highest risk:** Oral LP and Genital LP (especially vulvovaginal).")
add_bullet("Also elevated with: Ulcerative LP (sole) and Hypertrophic LP (shin).")

add_h4("Prognosis")
add_bullet("LP is **NOT a chronic disease** like psoriasis.")
add_bullet("Responds well to treatment and can resolve completely.")
add_bullet("Especially good prognosis when risk factors are identified and addressed.")
add_bullet("May recur but is NOT typically a lifelong condition requiring indefinite treatment.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Complications of LP. We've been saying it repeatedly — this type causes SCC, that type causes SCC. SCC is very important, especially with oral and genital LP. What else? In the skin: post-inflammatory hyperpigmentation (which I've explained why it persists), scarring, and atrophy. In the mucosa: scarring and stenosis. In hair: cicatricial alopecia — that's permanent. In nails: permanent damage — the pterygium I showed you.")
add_normal("Dr. Eman:  Associations — important for the exam: Hepatitis C and B (risk factors), liver disease, diabetes mellitus, and autoimmune diseases like Alopecia areata and Vitiligo. LP itself is considered autoimmune, which is why it associates with other autoimmune conditions.")
add_normal("Student:  I've seen LP go on and off, improving then relapsing. Is it a chronic disease?")
add_normal("Dr. Eman:  We need to search for the cause. LP is NOT a chronic disease in the same way psoriasis is. It's not something where the patient will be on treatment for life. LP responds well to treatment, improves, and can completely resolve. We look for the risk factor, address it, and the disease goes away. Can it recur? Possibly, but it's not like psoriasis where you're managing it forever. From our clinical experience, we see LP very frequently, and it does improve and may not come back.")
add_normal("Student:  Thank you, doctor.")
add_normal("Dr. Eman:  You're welcome. Is the material making sense to everyone? Dermatology requires studying, and we all need to put in the work. That's it for today's lecture — I hope it was easy and enjoyable. Any questions? May everyone have a blessed Ramadan.")

print("\n  ✓ Lecture 7 (Lichen Planus) appended with 16 images.")


# ══════════════════════════════════════════════════════════════════════════════
# ██  LECTURE 8: OTHER LICHENOID DERMATOSES (WITH IMAGES)
# ══════════════════════════════════════════════════════════════════════════════
print("\n" + "="*60)
print("  APPENDING LECTURE 8: OTHER LICHENOID DERMATOSES (WITH IMAGES)")
print("="*60)

add_h1("Lecture 8: Other Lichenoid Dermatoses")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 1 – Lichen Nitidus
# ═══════════════════════════════════════════════════════════════════════
add_h2("1. Lichen Nitidus")
add_h3("Summary")

add_h4("Definition & Clinical Features")
# >>> IMAGE: Clinical photo of tiny papules (page 3)
add_image(3, 'lec8')

add_bullet("**Lichen Nitidus** = Multiple tiny, discrete, skin-colored, dome-shaped, asymptomatic papules.")
add_bullet("Most common site: **Dorsum of the hand**, upper extremities, flexors, genitalia.")
add_bullet("More common in **children and young adults**, especially **females**.")
add_bullet("The papules are **pinpoint-sized** and **skin-colored** (not violaceous like LP).")

add_h4("Variants")
add_bullet("**Classic** (most common)")
add_bullet("**Palmoplantar**")
add_bullet("**Linear**")
add_bullet("**Follicular**")
add_bullet("**Generalized**")
add_bullet("**Perforating**")

add_h4("Etiopathogenesis")
add_bullet("Etiology is **unknown/debated**.")
add_bullet("Theory 1: May be an **unusual variant of Lichen planus**.")
add_bullet("Theory 2: May be related to **tuberculosis (TB)** — because the histopathology shows **tuberculoid granuloma** formation.")
add_bullet("A tuberculoid granuloma = well-circumscribed dermal infiltrate of epithelioid histiocytes surrounded by lymphocytes.")

add_h4("Histopathology — \"Ball and Claw Appearance\"")
# >>> IMAGE: Histopathology image (page 9)
add_image(9, 'lec8')

add_bullet("**\"Ball and claw appearance\"** = pathognomonic histological finding.")
add_bullet("The **\"Ball\"** = Well-circumscribed dermal infiltrate of **epithelioid histiocytes + lymphocytes** (granuloma), confined to the width of a single dermal papilla.")
add_bullet("The **\"Claw\"** = Elongated rete ridges (dermal papillae) flanking and gripping the granuloma like claws holding a ball.")
add_bullet("Also shows: **Lichenoid interface dermatitis** + **Epidermal atrophy** (unlike LP which has acanthosis).")
add_bullet("**Parakeratosis** is common (vs rare in LP).")
add_bullet("Granular cell layer is **thin to absent** (vs focally increased in LP).")

add_h4("Nail & Oral Involvement")
add_bullet("Nail involvement: **10%** of patients — **pitting** and **longitudinal ridging**.")
add_bullet("Key difference from LP: LP has longitudinal ridging and pterygium but **NO pitting**; Lichen Nitidus **has pitting**.")
add_bullet("Oral involvement: **Very rare**.")

add_h4("Associations — Mnemonic \"DMC\"")
add_bullet("**D** = Down syndrome")
add_bullet("**M** = Megacolon")
add_bullet("**C** = Crohn's disease")

add_h4("Lymphocytes")
add_bullet("Lymphocytes in Lichen Nitidus are **CD4+**.")
add_bullet("Key difference: LP lymphocytes are **CD8+**.")

add_h4("Differential Diagnosis")
add_bullet("**Plane warts (flat warts)** — most important DDx.")
add_bullet("**Frictional lichenoid dermatitis**.")

add_h4("Treatment")
add_bullet("Usually **no treatment needed** — asymptomatic and self-limited.")
add_bullet("If symptomatic: **Topical steroids**, **emollients**, **antihistamines**.")
add_bullet("Spontaneous clearing occurs.")

add_h4("Key Differences from Lichen Planus (Summary Table)")
add_bullet("**Size:** Pinpoint (LN) vs Pinhead to large plaques (LP).")
add_bullet("**Color:** Skin-colored (LN) vs Violaceous (LP).")
add_bullet("**Wickham striae:** Absent (LN) vs Present (LP) — because LN has no hypergranulosis.")
add_bullet("**Post-inflammatory hyperpigmentation:** None (LN) vs Present (LP).")
add_bullet("**Histology — Epidermis:** Epidermal atrophy (LN) vs Acanthosis (LP).")
add_bullet("**Histology — Dermal infiltrate:** Granuloma present (LN) vs No granuloma (LP).")
add_bullet("**Histology — Parakeratosis:** Common (LN) vs Rare (LP).")
add_bullet("**Lymphocytes:** CD4+ (LN) vs CD8+ (LP).")
add_bullet("**DIF — Fibrin deposition:** Absent (LN) vs Present/shaggy fibrinogen (LP).")
add_bullet("**Koebner phenomenon:** Common in both.")

# >>> IMAGE: Comparison table slide (page 12)
add_image(12, 'lec8')

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Good morning doctors, how are you all? How is the studying going? Let me check in with everyone. Have you started solving questions yet?")
add_normal("Student (Dr. Naima):  Good morning Dr. Eman. I've started reading the first lectures to connect the new material with the old, thank God.")
add_normal("Dr. Eman:  The first two lectures are genuinely difficult — they're like a compilation of the entire dermatology base. Try to listen to them multiple times. But the remaining lectures are easier because each one discusses a separate topic — you'll understand the disease and that's it.")
add_normal("Student (Dr. Naima):  That's what I'm doing — I listen to them almost every day. With work and Ramadan I'm trying to push myself to listen to a section and read it simultaneously.")
add_normal("Dr. Eman:  Excellent, and try to incorporate question-solving too. You must start solving questions because solving is part of studying — it's not just reviewing information you've already learned, there's also new information in the questions themselves.")
add_normal("Student (Dr. Doaa):  I'm doing well, thank God. It's Ramadan and time is short, but I feel I'm getting into the mood a bit.")
add_normal("Dr. Eman:  May God make it easy for everyone. Let's begin today's lecture. We finished Psoriasis, Other Papulosquamous diseases, and Lichen Planus. Today we'll finish the Other Lichenoid Dermatoses — it should be a light and pleasant lecture, God willing.")
add_normal("Dr. Eman:  Last time we finished all of Lichen Planus. We said there are many diseases under the umbrella of Interface dermatitis. There are two types of interface dermatitis — who can tell me what Interface dermatitis means and what are its two types?")
add_normal("Student:  Degeneration of the basement membrane.")
add_normal("Dr. Eman:  Yes, degeneration of the basement membrane. What are the two types?")
add_normal("Student:  Lichenoid and Vacuolar.")
add_normal("Dr. Eman:  The Lichenoid type is a band-like inflammatory infiltrate of lymphocytes beneath the dermoepidermal junction. Diseases in this category include: Lichen planus, Lichen nitidus, Lichen striatus, and Lichenoid drug eruption. The Vacuolar interface dermatitis includes: Erythema multiforme, SLE, Dermatomyositis, Fixed drug eruption, and Graft-versus-host disease. This is very important because it encompasses diseases from many different chapters.")
add_normal("Dr. Eman:  Let's start with Lichen Nitidus. It's quite different from Lichen planus. It presents as multiple tiny discrete skin-colored papules, as you can see on the hand here. It comes mainly on the dorsum of the hand. Will this confuse us with Lichen planus? No — it looks distinctly different. It's asymptomatic, tiny, discrete, dome-shaped, skin-colored papules on the dorsum of the hand and elbows, more common in children and young adults, more common in females.")
add_normal("Dr. Eman:  The most common sites are the dorsal aspect of the hand, upper extremities, flexors, and genitalia. It has variants: the classic form we're seeing, plus palmoplantar, linear, follicular, generalized, and perforating.")
add_normal("Dr. Eman:  Now what is Lichen Nitidus pathologically? Part of it — because it's a 'Lichen' — will have interface dermatitis and band-like inflammatory infiltrate. But they also found something unusual in its pathology: the shape of a Granuloma. A granuloma is a collection of histiocytes with lymphocytes around them.")
add_normal("Dr. Eman:  That's why there's debate about Lichen Nitidus — some say it's a rare variant of LP, and others say it's related to TB. Why TB? Because it forms a Tuberculoid granuloma: a well-circumscribed dermal infiltrate of epithelioid histiocytes and lymphocytes.")
add_normal("Dr. Eman:  The pathology of Lichen Nitidus is very specific and recognizable — we call it the \"Ball and claw appearance.\" The ball is the granuloma collection surrounded by dermal papillae (the claws). Like a ball being held by claws. The ball is the well-circumscribed dermal infiltrate of epithelioid histiocytes and lymphocytes. There's also lichenoid interface dermatitis and epidermal atrophy — unlike LP which has acanthosis.")
add_normal("Dr. Eman:  Does it have oral lesions? Very rare. Nail involvement? Yes, 10% of patients have nail abnormalities like pitting and longitudinal ridging. Remember: LP had longitudinal ridging but NO pitting. Lichen Nitidus HAS pitting.")
add_normal("Dr. Eman:  The associations of Lichen Nitidus — memorize them with the mnemonic DMC: Down syndrome, Megacolon, and Crohn's disease. Most cases I've seen were normal with no associations, but you must know DMC for the exam.")
add_normal("Dr. Eman:  The differential diagnosis — what could this look like? Plane warts (flat warts) are the number one DDx. Also frictional lichenoid dermatitis.")
add_normal("Dr. Eman:  The histopathology — what's the key word?")
add_normal("Student:  Ball and claw.")
add_normal("Dr. Eman:  Ball and claw. Ball from the Granuloma (epithelioid histiocytes with lymphocytes), plus lichenoid interface dermatitis and epidermal atrophy. This is extremely important — the pathology of Lichen Nitidus is very high-yield.")
add_normal("Dr. Eman:  Treatment? No treatment needed — it's asymptomatic. If needed: topical steroids, emollient, antihistamines if itchy. Spontaneous clearing occurs.")
add_normal("Dr. Eman:  Now here's a comparison table between Lichen Nitidus and Lichen Planus. Size: LN is pinpoint, LP is pinhead to large plaques. Color: LN is skin-colored, LP is violaceous. Distribution: LN favors upper extremities, flexors, genitalia, chest; LP favors wrists, forearms, peri-sacral area. Oral involvement: LN very rare, LP common — over 50%. Nail: LN has pitting (5-10%), LP does NOT have pitting.")
add_normal("Dr. Eman:  Wickham striae — will we see them in LN? No. Why do we see them in LP? Because of the focal hypergranulosis. Hyperpigmentation after healing — does LN leave it? No. LP leaves it because of basement membrane damage causing melanocytes to fall into the dermis as melanophages.")
add_normal("Dr. Eman:  Histopathology: Parakeratosis is common in LN but rare in LP. Granular cell layer is thin to absent in LN but focally increased in LP. Dermal infiltrate in LN is lymphocytes + epithelioid cells = granuloma! LP has no granuloma — just band-like lymphocytes. In LN the infiltrate is confined to the width of dermal papillae (the ball in the claws), in LP it's a band. Fibrin deposition: absent in LN, present in LP (shaggy fibrinogen on DIF).")
add_normal("Dr. Eman:  The lymphocytes in Lichen Nitidus are CD4+. In Lichen Planus they were CD8+.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 2 – Lichen Striatus
# ═══════════════════════════════════════════════════════════════════════
add_h2("2. Lichen Striatus")
add_h3("Summary")

add_h4("Definition & Clinical Features")
# >>> IMAGE: Clinical photo of linear lesion (page 15)
add_image(15, 'lec8')

add_bullet("**Lichen Striatus** = Uncommon, asymptomatic, self-limited linear dermatosis following **Blaschko's lines**.")
add_bullet("Appearance: Pink/erythematous lichenoid papules in a **linear, unilateral** distribution, mainly on **extremities**.")
add_bullet("More common in **children** and **females**.")
add_bullet("Seasonal predilection: **Spring and Summer** (supports viral infection theory).")
add_bullet("Heals with **hypopigmentation** (NOT hyperpigmentation like LP).")
add_bullet("**Spontaneous resolution** within 6 months to years.")

add_h4("Etiopathogenesis (Theories)")
add_bullet("**Somatic mosaicism** — different cell populations along Blaschko's lines.")
add_bullet("**Viral infection** — supported by seasonal increase in spring/summer.")
add_bullet("**Cell-mediated CD8+ response**.")
add_bullet("**Atopy**.")
add_bullet("**Trauma / Koebnerization**.")

add_h4("Histopathology")
# >>> IMAGE: Histopathology image (page 18)
add_image(18, 'lec8')

add_bullet("**Eczema-like features** in early stages: acanthosis, spongiosis, exocytosis.")
add_bullet("**Vacuolar interface dermatitis** also present.")
add_bullet("**KEY DISTINGUISHING FEATURE:** Inflammatory infiltrate **around hair follicles and eccrine sweat glands**.")
add_bullet("This perifollicular and periadnexal infiltrate differentiates Lichen Striatus from simple eczema or vacuolar interface dermatitis alone.", level=1)

add_h4("Differential Diagnosis")
add_bullet("**Linear Lichen Planus** — older age, pruritic, heals with hyperpigmentation (vs LS: younger, asymptomatic, heals with hypopigmentation).")
add_bullet("**Linear GVHD**.")
add_bullet("**Blaschkitis** (inflammation along Blaschko's lines).")
add_bullet("**ILVEN** (Inflammatory Linear Verrucous Epidermal Nevus) — most similar to Lichen Striatus.")

add_h4("Treatment")
add_bullet("**No treatment needed** — asymptomatic and self-resolving (resolves within 6 months).")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Now we'll move to a completely different topic: Lichen Striatus. It presents as a linear dermatosis — erythematous papules following Blaschko's lines or a linear pattern.")
add_normal("Dr. Eman:  Some people classified it as a type of eczema. Others said it might be caused by viral infection because it increases in summer and spring. Some proposed somatic mosaicism. Others suggested Koebnerization or trauma. All these theories exist because the exact cause is unknown.")
add_normal("Dr. Eman:  Lichen Striatus is an uncommon, asymptomatic, self-limited disease. A streak of erythematous papules appears in a linear pattern, stays for a while, then heals with hypopigmentation. It's an asymptomatic streak of pink lichenoid papules that's unilateral, on the extremities, along Blaschko's lines.")
add_normal("Dr. Eman:  Has anyone seen Lichen Striatus before? It's more common in children and females — same as Lichen Nitidus. It has seasonal variation — more common in summer and spring, which is why they proposed the viral infection theory.")
add_normal("Dr. Eman:  It appears suddenly and resolves spontaneously within months to years. The etiopathogenesis is either unknown, or one of these theories: somatic mosaicism (different cells along Blaschko's lines), viral infection (seasonal), cell-mediated CD8 response, atopy, or trauma.")
add_normal("Dr. Eman:  The histopathology of Lichen Striatus: First, it has vacuolar interface dermatitis, but not strongly. It also shows eczema-like features — acanthosis, spongiosis, and exocytosis in the early stages.")
add_normal("Dr. Eman:  But the KEY distinguishing feature is that the infiltrate surrounds hair follicles and eccrine sweat glands. The lymphocytes encircle the hair follicle and the sweat glands. This is crucial — it's not just eczema-like or vacuolar interface dermatitis on the surface. You have lymphocytes also surrounding the hair follicles and eccrine sweat glands deep in the dermis.")
add_normal("Dr. Eman:  Differential diagnosis: Linear lichen planus — but that comes in older age, is very pruritic, and heals with hyperpigmentation. Linear Graft-versus-host disease. Blaschkitis — inflammation of the Blaschko's lines. And ILVEN (Inflammatory Linear Verrucous Epidermal Nevus) — this is the most similar to Lichen Striatus.")
add_normal("Dr. Eman:  Treatment? Asymptomatic, resolves within 6 months, nothing needed. Just like Lichen Nitidus — nothing needed. These two topics (Lichen Striatus and Lichen Nitidus) are quick, easy topics on the side.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 3 – Graft-Versus-Host Disease (GVHD)
# ═══════════════════════════════════════════════════════════════════════
add_h2("3. Graft-Versus-Host Disease (GVHD)")
add_h3("Summary")

add_h4("Definition & Concept")
# >>> IMAGE: Mechanism diagram (page 21)
add_image(21, 'lec8')

add_bullet("**GVHD** = The transplanted immune cells (graft) attack the recipient's body (host).")
add_bullet("Occurs ONLY after **hematopoietic stem cell (bone marrow) transplantation** — NOT after solid organ transplants.")
add_bullet("Why not solid organs? Solid organs (liver, kidney, heart) do NOT contain immune cells capable of attacking the host. In solid organ transplants, it's the host that rejects the organ (tissue rejection).", level=1)

add_h4("Mechanism")
add_bullet("The recipient is **immunosuppressed** (to prevent rejection of the transplant) → immune system is essentially zero.")
add_bullet("Donor bone marrow contains **T-cells** that are transplanted into the host.")
add_bullet("The donor T-cells recognize host tissues as **foreign** → activate → attack the host's organs.")
add_bullet("The host cannot defend itself because it is immunocompromised.")

add_h4("Target Organs")
add_bullet("**Skin** (most common)")
add_bullet("**GI tract** (nausea, vomiting, diarrhea, abdominal pain)")
add_bullet("**Liver** (hepatitis)")

add_h4("Acute GVHD")
# >>> IMAGE: Acute GVHD skin manifestation (page 26)
add_image(26, 'lec8')

add_bullet("**Mechanism:** Recipient tissue damage → T-cell activation → Cytotoxic T-cells cause tissue injury to skin, GI, and liver.")
add_bullet("**Skin manifestation:** **Morbilliform exanthem** (maculopapular rash).")
add_bullet("**Pathognomonic sign:** **Peripheral desquamation** — bullae and desquamation of the acral area (palms and soles).")
add_bullet("GI: Nausea, vomiting, diarrhea, abdominal pain.", level=1)
add_bullet("Liver: Hepatitis.", level=1)
add_bullet("Also: Generalized pruritus, cicatricial alopecia, erythroderma (in severe cases).", level=1)
add_bullet("**Histopathology:** Vacuolar interface dermatitis + keratinocyte necrosis.")
add_bullet("**DDx:** Viral exanthem, Drug eruption (differentiated by peripheral desquamation and clinical history of bone marrow transplantation).")
add_bullet("**Treatment:** Topical and systemic steroids, Tacrolimus, Mycophenolate mofetil, TNF-alpha inhibitors.")

add_h4("Chronic GVHD")
# >>> IMAGE: Chronic GVHD manifestations (page 31)
add_image(31, 'lec8')

add_bullet("**Mechanism:** T-cell activation (same as acute) **PLUS B-cell activation** → production of **autoantibodies** (ANA positive, Anti-dsDNA positive).")
add_bullet("The autoantibodies attack host tissues, producing autoimmune-like disease manifestations.")
add_bullet("**Skin manifestations mimic:**")
add_bullet("**Lichen planus**-like", level=1)
add_bullet("**Lichen striatus**-like", level=1)
add_bullet("**Morphea**-like", level=1)
add_bullet("**Scleroderma**-like", level=1)
add_bullet("**Poikiloderma**-like", level=1)
add_bullet("**Eosinophilic fasciitis**-like", level=1)
add_bullet("**Systemic features:** Dry mouth, dry eyes, oral sores, rash, jaundice, shortness of breath, nausea, vomiting.")
add_bullet("Can affect **any body system**.")
add_bullet("**Histopathology:** Variable — depends on which disease it mimics (LP pathology if LP-like, morphea pathology if morphea-like, etc.).")
add_bullet("**Treatment:** Sun protection, phototherapy, systemic steroids.")

# >>> IMAGE: Comparison table - Acute vs Chronic GVHD (page 36)
add_image(36, 'lec8')

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  Now let's move to a completely new topic: Graft-versus-host disease (GVHD). Does anyone have any information about this disease or can explain what it means?")
add_normal("Student:  In transplantation, rejection of the graft occurs.")
add_normal("Dr. Eman:  Rejection happens, yes. But who does the rejection? Does my body reject what was transplanted, or is it the opposite?")
add_normal("Student:  I think the opposite.")
add_normal("Dr. Eman:  Exactly the opposite! This is NOT tissue rejection. In tissue rejection, the host rejects the organ — that's what happens in any normal organ transplantation. But here the disease is called Graft VERSUS Host disease — the graft attacks the host.")
add_normal("Dr. Eman:  Let me explain from scratch. Imagine someone has a problem with their hematopoietic stem cells (bone marrow) — cancer, radiation damage, anything. They need a bone marrow transplant. When we transplant bone marrow, we're transplanting immune cells and blood cells.")
add_normal("Dr. Eman:  Where do these immune cells come from? From another person whose HLA is somewhat matched. We take their bone marrow and transplant it. Now, for the transplant to succeed, does the recipient's immunity need to be strong or weak?")
add_normal("Student:  Obviously weak.")
add_normal("Dr. Eman:  Obviously weak — why? So they don't reject the transplant! I need to suppress their immune system to near zero. Then the new immune cells (the graft full of functional T-cells) enter and what happens? These new immune cells find themselves in a foreign body! It's not their own body. So they start attacking the host's body!")
add_normal("Dr. Eman:  That's why it's called Graft versus Host. The immune cells I transplanted (the Graft) attack the patient's body (the Host). The host can't defend itself because it's immunocompromised. Understand the concept?")
add_normal("Dr. Eman:  This only happens in hematopoietic stem cell transplantation. Does it happen with other organ transplants? No — because solid organs don't contain immune cells that can attack the body. If I transplant a liver, the liver won't attack the body! The body attacks the liver and rejects it. But here I'm transplanting immune cells (bone marrow/stem cells), so THEY attack the body, and the host can't defend itself.")
add_normal("Dr. Eman:  GVHD is divided into Acute and Chronic, and the target organs are: Skin, GI tract, and Liver.")
add_normal("Dr. Eman:  In Acute GVHD, the pathognomonic sign is bullae and desquamation of the acral area — desquamation of the palms and soles. On the skin you get a morbilliform eruption (maculopapular rash), and desquamation in the acral area. There's also pruritus and cicatricial alopecia. GI presents with nausea, vomiting, diarrhea, and abdominal pain. Liver presents as hepatitis.")
add_normal("Dr. Eman:  So a typical case: an older patient with a history of bone marrow transplantation, presenting with a full-body exanthem that looks like a viral infection or drug eruption, PLUS peripheral desquamation of the palms and soles, PLUS acute abdomen and hepatitis. The peripheral desquamation is what differentiates it from viral or drug causes.")
add_normal("Dr. Eman:  The Acute form has 3 phases: Recipient tissue damage → Activation of T-cells → Cytotoxic T-cells cause tissue injury to skin, liver, and GIT. It manifests as morbilliform exanthem, erythroderma, bullae and desquamation in the acral area, generalized pruritus, and cicatricial alopecia.")
add_normal("Dr. Eman:  Now the Chronic form. The acute was T-cell activation — the donor T-cells get activated against the host. In Chronic GVHD, the same T-cell activation occurs BUT ALSO B-cell activation. What does B-cell activation mean? It means autoantibodies are produced against the host's body. These antibodies attack the host.")
add_normal("Dr. Eman:  What do these autoantibodies produce? ANA positive and Anti-double stranded DNA positive. So the chronic form manifests as: morphea-like, scleroderma-like, poikiloderma-like, eosinophilic fasciitis-like, and lichen planus-like disease. All because autoantibodies are attacking.")
add_normal("Dr. Eman:  Chronic GVHD causes dry mouth, dry eyes, mouth sores, rash, jaundice, shortness of breath, nausea, vomiting. On the skin it appears as lichen planus, lichen striatus, morphea-like, scleroderma-like, and poikiloderma-like. It can affect any organ system.")
add_normal("Dr. Eman:  Histopathology: Acute GVHD shows vacuolar interface dermatitis and keratinocyte necrosis — that's why it's classified under lichenoid dermatoses. Chronic GVHD is variable — if it looks like LP, it has LP pathology; if it looks like morphea, it has morphea pathology.")
add_normal("Dr. Eman:  DDx for Acute: viral exanthem and drug eruption. Treatment for Acute: topical and systemic steroids, Tacrolimus, Mycophenolate Mofetil, and TNF-alpha inhibitors. Treatment for Chronic: sun protection, phototherapy, and systemic steroids.")
add_normal("Student:  Doctor, the recipient who receives the graft — they get immunosuppression, correct? So what causes the donor T-cells to activate? They already have no immunity.")
add_normal("Dr. Eman:  What activates the donor T-cells? Any stimulus — bacteria on the skin, any antigen presenting cell activation. Why are we transplanting in the first place? To give this person immunity! Their bone marrow had cancer, and I transplanted new bone marrow so they'd have a functioning immune system to protect them from infections. But that immune system turns against them. The T-cells aren't 'theirs' — so they start attacking the host tissue itself.")

# ═══════════════════════════════════════════════════════════════════════
# TOPIC 4 – Erythroderma (Exfoliative Dermatitis)
# ═══════════════════════════════════════════════════════════════════════
add_h2("4. Erythroderma (Exfoliative Dermatitis)")
add_h3("Summary")

add_h4("Definition")
# >>> IMAGE: Erythroderma clinical photo (page 39)
add_image(39, 'lec8')

add_bullet("**Erythroderma** = Erythema + scaling involving **>90% body surface area**.")
add_bullet("Also called **Exfoliative Dermatitis**.")
add_bullet("NOT a disease itself — it is an **endpoint/manifestation** of many different diseases.")
add_bullet("Course: **Chronic** (months to years), relapsing-remitting.")
add_bullet("Main cause of death: **Pneumonia** (secondary infection due to loss of skin barrier).")

add_h4("Causes in Adults — Mnemonic \"BAD\"")
# >>> IMAGE: Causes table/list (page 40)
add_image(40, 'lec8')

add_bullet("**Most common causes:**")
add_bullet("**B** = (B)ig two: **Psoriasis** and **Atopic dermatitis** (most common overall)", level=1)
add_bullet("**A** = (A)ssorted inflammatory conditions", level=1)
add_bullet("**D** = **Drugs** — mnemonic **\"ASBAGAK\"**: Allopurinol, Sulfasalazine/Sulfonamides, Phenytoin, Beta-blockers, Gold, Carbamazepine", level=1)
add_bullet("**Less common causes:** Mycosis Fungoides (MF)/Sézary Syndrome, Pityriasis Rubra Pilaris (PRP), Eczema, Paraneoplastic Pemphigus, Pemphigus Foliaceus, Bullous Pemphigoid.")
add_bullet("**Rare causes:** Crusted Scabies, Congenital Ichthyosiform Erythroderma, Lichen Planus, GVHD.")
add_bullet("**Idiopathic (25%)** = **\"Red Man Syndrome\"** — erythroderma with no identifiable cause.")

add_h4("Causes in Children/Neonates")
add_bullet("**Genetic/Ichthyosis:** Congenital Ichthyosiform Erythroderma (bullous and non-bullous), Netherton Syndrome.")
add_bullet("**Immunodeficiency:** Wiskott-Aldrich Syndrome.")
add_bullet("**Inflammatory:** Atopic dermatitis, Seborrheic dermatitis, Psoriasis.")
add_bullet("**Infections:** Staphylococcal Scalded Skin Syndrome (SSSS), Congenital Cutaneous Candidiasis.")
add_bullet("**Others:** PRP, GVHD, Mastocytosis.")

add_h4("Drug Mnemonics Summary")
add_bullet("**Erythroderma drugs** = \"ASBAGAK\" (Allopurinol, Sulfasalazine/Sulfonamides, Phenytoin, Beta-blockers, Gold, Carbamazepine).")
add_bullet("**Psoriasis exacerbation drugs** = \"LIPA\" (Lithium, Interferon, Beta-blockers, Antimalarials).")
add_bullet("**Lichenoid drug eruption/LP drugs** = \"DAMAT\" (Diuretics, Antihypertensives, Antimalarials, Amalgam/Metals, Tetracyclines).")

add_h4("Clinical Features")
add_bullet("**Erythema appears first** → scaling develops 2–8 days later.")
add_bullet("Skin is **bright red, hot, and dry**.")
add_bullet("**Intense pruritus** — worst with **Sézary syndrome** and **Atopic dermatitis**.")

add_h4("Complications")
add_bullet("**Skin complications:** Lichenification, depigmentation, palmoplantar keratoderma, nail changes, alopecia, subungual hyperkeratosis, *Staphylococcus aureus* infection (no skin barrier).")
add_bullet("**Systemic complications:** Fever, malaise, **generalized lymphadenopathy (50%)**, **peripheral edema (50%)**, hepatomegaly (20%), splenomegaly.")
add_bullet("**Cardiovascular cascade:** Vasodilatation → fluid loss → electrolyte disturbance → tachycardia → hypothermia → **High-output cardiac failure** → cachexia, anemia.")
add_bullet("High-output cardiac failure is the reason for **mandatory hospitalization**!", level=1)

add_h4("Sézary Syndrome — Key Distinction")
add_bullet("Sézary syndrome is **NOT the same** as erythrodermic Mycosis Fungoides.")
add_bullet("It is a **separate leukemic variant of Cutaneous T-cell Lymphoma (CTCL)**.")
add_bullet("Starts with **erythroderma from the beginning** — no preceding patches or plaques (unlike MF which progresses through patches → plaques → tumors → erythroderma).")
add_bullet("The malignant T-cell in Sézary is **different** from the malignant T-cell in MF.")
add_bullet("Features: **Sézary cells in blood** + **Generalized lymphadenopathy**.")

add_h4("Histopathology")
add_bullet("**Non-specific:** Acanthosis, dilated blood vessels, edema.")
add_bullet("**Acute:** Spongiosis + parakeratosis.")
add_bullet("**Chronic:** Acanthosis.")
add_bullet("The underlying disease may show its own specific pathology (e.g., psoriasis plaques may still be identifiable).")

add_h4("Treatment")
add_bullet("**Hospitalization required** (for monitoring and managing cardiovascular complications).")
add_bullet("**Stop suspected drugs** immediately.")
add_bullet("**Close monitoring:** Fluid balance, electrolytes, temperature, cardiac function.")
add_bullet("**Systemic antibiotics** — prevent secondary bacterial infection (no functioning skin barrier).")
add_bullet("**Sedating antihistamines** — for severe pruritus.")
add_bullet("**Topical care:**")
add_bullet("Wet dressings", level=1)
add_bullet("Emollients", level=1)
add_bullet("**LOW-potency topical steroids ONLY** — high absorption risk over >90% BSA → systemic toxicity if high-potency steroids used", level=1)
add_bullet("**NEVER cover the entire body with topical steroids**", level=1)

add_h4("Specific Treatment by Cause")
add_bullet("**Psoriasis →** Methotrexate, Acitretin, or Biologics. **AVOID systemic steroids** (causes rebound flare or pustular psoriasis).")
add_bullet("**Atopic dermatitis →** Systemic steroids OR Cyclosporine.")
add_bullet("**Drug-induced →** Stop the offending drug + systemic steroids.")

add_h3("Translated & Tidied Audio Transcript")
add_normal("Dr. Eman:  We've now finished the Lichenoid dermatoses. Many diseases end up causing erythroderma — including GVHD and Psoriasis. Let's discuss Erythroderma now.")
add_normal("Dr. Eman:  What is the definition of Erythroderma?")
add_normal("Student:  Generalized erythema and scaling of the skin more than 90% body surface area.")
add_normal("Dr. Eman:  Yes — erythema and scales covering more than 90% of the BSA. Is Erythroderma a disease in itself?")
add_normal("Student:  No, it's the result of many diseases.")
add_normal("Dr. Eman:  Exactly. It can be chronic, lasting months to years, with a relapsing-remitting course. The main cause of death in erythroderma is pneumonia — secondary infection.")
add_normal("Dr. Eman:  Causes divide into primary (idiopathic, 25% — called Red Man Syndrome) and secondary. Secondary divides into adult causes and pediatric causes.")
add_normal("Dr. Eman:  What are the most important causes of erythroderma in adults?")
add_normal("Student:  Psoriasis, Atopic Dermatitis, Drugs.")
add_normal("Dr. Eman:  Correct — the three most important adult causes, summarized as \"BAD.\" The drugs causing erythroderma are memorized with \"ASBAGAK\": Allopurinol, Sulfasalazine/Sulfonamides, Phenytoin, Beta-blockers, Gold, Carbamazepine.")
add_normal("Dr. Eman:  Less common causes include Mycosis Fungoides, Pityriasis Rubra Pilaris, Eczema, Paraneoplastic Pemphigus, Pemphigus Foliaceus, and Bullous Pemphigoid. Very rare causes: Crusted Scabies, Congenital Ichthyosiform Erythroderma, Lichen Planus, and GVHD. You must know all of these.")
add_normal("Dr. Eman:  In children and neonates, the most important cause is genetic/ichthyosis — a newborn isn't going to get psoriasis or take drugs. Congenital Ichthyosiform Erythroderma (bullous and non-bullous) and Netherton Syndrome are the key genetic causes. Then immunodeficiency syndromes like Wiskott-Aldrich. Then inflammatory causes like Atopic dermatitis, Seborrheic dermatitis, and Psoriasis. Infections like SSSS (Staphylococcal Scalded Skin Syndrome) and Congenital Cutaneous Candidiasis. And others: PRP, GVHD, Mastocytosis.")
add_normal("Dr. Eman:  Let me consolidate the drug mnemonics: Erythroderma drugs = \"ASBAGAK\". Psoriasis exacerbation drugs = \"LIPA\" (Lithium, Interferon, Beta-blockers, Antimalarials). Lichen Planus/Lichenoid drug eruption drugs = \"DAMAT\" (Diuretics, Antihypertensives, Antimalarials, Amalgam/Metals, Tetracyclines).")
add_normal("Dr. Eman:  The clinical picture: Erythema appears first, then scaling develops 2 to 8 days later. The body is bright red, hot, and dry. Is there pruritus? Most common — it's very intense, especially with Sézary syndrome and Atopic dermatitis.")
add_normal("Dr. Eman:  Complications — the skin: Lichenification (increased skin markings), depigmentation, palmoplantar keratoderma, nail changes, alopecia, subungual hyperkeratosis, and Staph aureus infection because there's no skin barrier.")
add_normal("Dr. Eman:  Systemic complications: Fever, malaise, generalized lymphadenopathy in 50% of cases, peripheral edema in 50%, hepatomegaly in 20%, splenomegaly.")
add_normal("Dr. Eman:  The most important systemic complication — due to massive vasodilatation in the skin: fluid loss, electrolyte disturbance, tachycardia, hypothermia, and ultimately High-output cardiac failure, leading to cachexia and anemia. Can you imagine that erythroderma can cause cardiac failure? That's why these patients require hospitalization.")
add_normal("Dr. Eman:  Histopathology is non-specific — acanthosis, dilated blood vessels, edema. Nothing specific to help you identify the cause. In acute erythroderma: spongiosis and parakeratosis. In chronic: acanthosis.")
add_normal("Dr. Eman:  Now about Sézary Syndrome — does anyone know what it is?")
add_normal("Student:  It's on the spectrum of MF.")
add_normal("Dr. Eman:  Actually, it's NOT on the MF spectrum. It's different from MF. Mycosis Fungoides progresses from patches to plaques to tumors to erythroderma. Sézary Syndrome is a SEPARATE entity — a leukemic variant of Cutaneous T-cell Lymphoma. It starts with erythroderma from the very beginning, with no preceding patches or plaques. The malignant T-cell itself is different from MF. It has Sézary cells in the blood at a certain percentage, plus generalized lymphadenopathy.")
add_normal("Dr. Eman:  Treatment of Erythroderma: First — hospitalization. Second — stop any suspected drugs. Third — close monitoring for fluid disturbance and cardiac issues; adjust body temperature, electrolyte balance, and fluid replacement. Fourth — systemic antibiotics to prevent secondary bacterial infection since there's no skin barrier. Fifth — sedating antihistamines for severe pruritus. Sixth — topical care: wet dressings, emollients, and low-potency topical steroids.")
add_normal("Dr. Eman:  IMPORTANT: You cannot cover the entire body with topical steroids in erythroderma — the absorption will be extremely high over 90% BSA and will cause systemic toxicity.")
add_normal("Dr. Eman:  Specific treatment depends on the cause. For Psoriasis: Methotrexate, Acitretin, or Biologics — AVOID systemic steroids because they cause rebound flare or pustular psoriasis. For Atopic dermatitis: systemic steroids or Cyclosporine. For Drug-induced: stop the offending drug and give systemic steroids.")
add_normal("Student:  Doctor, regarding erythroderma, when do we use corticosteroids in the acute phase?")
add_normal("Dr. Eman:  I would give systemic steroids to control severe inflammation if I know the cause is NOT psoriasis — for example, drug-induced erythroderma. If I'm unsure it's psoriasis and give systemic steroids, there'll be a rebound. If the patient is critically ill, I might give steroids and then determine the underlying disease once things calm down. But if I'm certain it's psoriasis, I'll use Methotrexate or another agent instead. Meanwhile, I treat all the other symptoms — fluid balance, electrolytes, etc.")
add_normal("Dr. Eman:  That's it for today's lecture, doctors. We covered Lichen Nitidus, Lichen Striatus, Graft-Versus-Host Disease, and Erythroderma. Any questions?")
add_normal("Dr. Eman:  I apologize we won't solve questions today due to time. I'll try to post some questions on the group after iftar for us to solve together. Happy Ramadan to everyone.")

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

doc.save(OUT_PATH)
print(f"\n{'='*60}")
print(f"  ✓ COMPLETE! Study Guide rebuilt with Lectures 7 & 8 (WITH IMAGES)")
print(f"  Saved to: {OUT_PATH}")
print(f"  Lecture 7 images: 16 insertion points")
print(f"  Lecture 8 images: 11 insertion points")
print(f"{'='*60}")
