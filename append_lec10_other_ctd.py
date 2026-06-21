"""
Append Lecture 10 (Other Connective Tissue Diseases) WITH IMAGES
to SCE_ConnectAC_Study_Guide_Lec1-9.docx

Saves as: SCE_ConnectAC_Study_Guide_Lec1-10.docx

Uses Pillow re-save technique to fix JPEG header issues.
"""

import re, os, io
from docx import Document
from docx.shared import Inches
from PIL import Image

# ── Paths ────────────────────────────────────────────────────────────────────
BASE_PATH = r"C:\Users\ahaww\Downloads\mypdf\SCE_ConnectAC_Study_Guide_Lec1-9.docx"
OUT_PATH  = r"C:\Users\ahaww\Downloads\mypdf\SCE_ConnectAC_Study_Guide_Lec1-10.docx"
LEC10_IMG_DIR = r"C:\Users\ahaww\Downloads\mypdf\ConnectAC_Slides\lec10_images"

# ── Load base document ───────────────────────────────────────────────────────
doc = Document(BASE_PATH)

# ── Image helper (Pillow re-save technique) ──────────────────────────────────

def add_image(page_num, width_inches=5.0):
    """Add an image from extracted slides, re-saving via Pillow to fix JPEG headers."""
    fname = f"lec10_p{page_num}_0.jpeg"
    path = os.path.join(LEC10_IMG_DIR, fname)
    if os.path.exists(path):
        img = Image.open(path)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        stream = io.BytesIO()
        img.save(stream, format='JPEG', quality=90)
        stream.seek(0)
        try:
            doc.add_picture(stream, width=Inches(width_inches))
            print(f"  ✓ Inserted image: {fname}")
        except Exception as e:
            print(f"  ✗ Failed to insert {fname}: {e}")
    else:
        print(f"  ✗ Image not found: {path}")

# ── inline-markdown parser  **bold** *italic* ────────────────────────────────

def add_inline_runs(para, text: str):
    tokens = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = para.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            run = para.add_run(tok[1:-1])
            run.italic = True
        else:
            para.add_run(tok)

# ── paragraph adders ─────────────────────────────────────────────────────────

def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.add_run(text)

def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.add_run(text)

def add_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.add_run(text)

def add_h4(text):
    p = doc.add_paragraph(style='Heading 4')
    p.add_run(text)

def add_normal(text: str):
    p = doc.add_paragraph(style='Normal')
    speaker_pat = re.match(
        r'^(\*\*)?((Dr\.|Student|Professor)[^:]*:)(\*\*)?\s*(.*)',
        text, re.DOTALL
    )
    if speaker_pat:
        label = speaker_pat.group(2).strip()
        rest  = speaker_pat.group(5).strip()
        run_l = p.add_run(label + '  ')
        run_l.bold = True
        add_inline_runs(p, rest)
    else:
        add_inline_runs(p, text.strip())

def add_bullet(text: str, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text.strip())

def add_list_number(text: str):
    p = doc.add_paragraph(style='List Number')
    add_inline_runs(p, text.strip())

# ══════════════════════════════════════════════════════════════════════════════
#  LECTURE 10 CONTENT — Other Connective Tissue Diseases
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_h1("Lecture 10: Other Connective Tissue Diseases")

# ─── Topic 1: Dermatomyositis — Definition, Classification & Pathogenesis ────
add_h2("Dermatomyositis — Definition, Classification & Pathogenesis")

add_bullet("**Definition:** Multisystemic autoimmune disease affecting muscle, skin, and other organs")
add_bullet("Dermato = skin, Myositis = muscle inflammation", level=1)
add_bullet("Rare disease with **bimodal distribution**: Juvenile DM and Adult DM", level=0)
add_bullet("Adult DM is more dangerous — associated with **malignancy**", level=1)

add_h3("Classification")
add_bullet("**Skin only** — Amyopathic DM (DM sine myositis): no muscle involvement")
add_bullet("**Skin + Muscle** — Classic Dermatomyositis")
add_bullet("**Muscle only** — Polymyositis (no skin)")

add_h3("Muscle Involvement")
add_bullet("Affects **proximal extensor muscles** (biceps, triceps, quadriceps)")
add_bullet("Bilateral and symmetrical")
add_bullet("Tender on palpation")
add_bullet("Patient cannot comb hair or climb stairs")

add_h3("Pathogenesis")
add_bullet("Environmental trigger + Genetic susceptibility → Immune dysregulation")

add_h4("Environmental Triggers")
add_bullet("**Malignancy** (most important!) — Ovarian cancer, Colon cancer, Breast cancer")
add_bullet("**Infections:** Coronavirus, Echovirus, Coxsackievirus, HIV, Toxoplasma, Borrelia, E. coli")
add_bullet("**Drugs:** NSAIDs, Statins, Hydroxyurea, D-penicillamine, TNF-alpha inhibitors, Cyclophosphamide")

add_h4("Genetic Susceptibility")
add_bullet("HLA-DR3, HLA-DR52, HLA-DR7")
add_bullet("Increased concordance in monozygotic twins")

add_h4("Autoantibodies")
add_bullet("**Anti-Mi-2** (Helicase) — Classic DM")
add_bullet("**Anti-Jo-1** (Histidyl-tRNA synthetase) — Severe form / Antisynthetase syndrome")

add_h4("Malignancy Association")
add_bullet("ONLY with **Adult DM** and **Amyopathic DM**")
add_bullet("NOT in Juvenile DM, NOT in Polymyositis")
add_bullet("Most feared: Ovarian, Colon, Breast cancer")

add_image(3)
add_image(6)

# ─── Topic 2: Dermatomyositis — Cutaneous Features ──────────────────────────
add_h2("Dermatomyositis — Cutaneous Features")

add_bullet("**Poikiloderma** = erythema + telangiectasia + hypo/hyperpigmentation")
add_bullet("Pink-violet color (NOT brown); SEVERELY ITCHY", level=1)
add_bullet("**Heliotrope erythema:** Periorbital violaceous erythema — *pathognomonic*")
add_bullet("**V sign:** Erythema on anterior chest (V-neck distribution)")
add_bullet("**Shawl sign:** Erythema on upper back (like wearing a shawl)")
add_bullet("**Gottron sign:** Erythema/Poikiloderma over joints (elbows, knees) — flat")
add_bullet("**Gottron papules:** Raised violaceous papules over knuckles/MCP joints — *pathognomonic*")
add_bullet("**Nail fold changes:** Ragged cuticle + periungual telangiectasia (seen on dermoscopy)")
add_bullet("**Calcinosis cutis** (dystrophic type — NOT metastatic)")
add_bullet("**Cutaneous small vessel vasculitis**")

add_h3("Differential Diagnosis")
add_bullet("DDx for **malar erythema** — 4 diseases: DM, SLE, Rosacea, Seborrheic dermatitis")
add_bullet("DDx for **Gottron sign**: Psoriasis (silvery white scales on elbows/knees)")

add_image(8)
add_image(10)
add_image(12)

# ─── Topic 3: DM — Systemic Features, Antisynthetase Syndrome & Treatment ───
add_h2("Dermatomyositis — Systemic Features, Antisynthetase Syndrome & Treatment")

add_h3("Smooth Muscle Involvement")
add_bullet("**GI:** GERD, Dysphagia (cricopharyngeal muscle affection)")
add_bullet("**Cardiac:** Arrhythmia")
add_bullet("**Pulmonary:** Diffuse interstitial lung fibrosis")

add_h3("Antisynthetase Syndrome (Anti-Jo-1)")
add_bullet("DM + Interstitial lung fibrosis + Erosive polyarthritis + Mechanic's hands + Fever")
add_bullet("More severe than classic DM (Anti-Mi-2)")

add_h3("Causes of Death")
add_bullet("Respiratory infection")
add_bullet("Cardiac failure")
add_bullet("Malnutrition (dysphagia)")
add_bullet("Carcinoma")

add_h3("Investigations")
add_bullet("Skin biopsy")
add_bullet("ANA → Anti-Mi-2 / Anti-Jo-1")
add_bullet("Muscle enzymes: CPK, Aldolase, LDH, Myoglobin, AST/ALT")
add_bullet("Muscle MRI / Ultrasound")
add_bullet("Electromyography")
add_bullet("Muscle biopsy (from **triceps**): lymphocytic infiltrate, perifascicular/perivascular, muscle atrophy/necrosis")
add_bullet("Urine: Creatinine / Myoglobin")
add_bullet("Malignancy screening")

add_h3("Histopathology (Skin)")
add_bullet("**Vacuolar interface dermatitis** (same category as SLE)")
add_bullet("From Poikiloderma: sparse lymphocytic infiltrate, vasodilation, interstitial mucin deposition")
add_bullet("From Gottron papule: lichenoid infiltrate (more intense), acanthosis, NO epidermal atrophy")

add_h3("Treatment — Skin")
add_bullet("Challenging — not all patients clear")
add_bullet("Photoprotection")
add_bullet("Antimalarials")
add_bullet("Low-dose Methotrexate")
add_bullet("Immunosuppressives")
add_bullet("Topical steroids, Tacrolimus")

add_h3("Treatment — Muscle")
add_bullet("Systemic steroids: **Prednisolone 1 mg/kg/day × 6 months**")
add_bullet("Steroid-sparing drugs (Methotrexate, Azathioprine)")
add_bullet("Rest + Physiotherapy")
add_bullet("*Muscle responds well to treatment; skin does not*")

add_image(14)
add_image(16)
add_image(18)

# ─── Topic 4: Systemic Sclerosis — Pathogenesis ─────────────────────────────
add_h2("Systemic Sclerosis (Scleroderma) — Pathogenesis")

add_bullet("**Definition:** Systemic hardening/fibrosis of skin and internal organs due to excessive collagen production")

add_h3("Pathogenesis Triangle")
add_bullet("Vascular dysfunction → Ischemia → Immune dysregulation (cytokines) → Fibroblast activation → Excess collagen → Fibrosis → Sclerosis")

add_h3("Mechanism")
add_bullet("Starting point: **Blood vessel abnormality**")
add_bullet("Endothelial cell injury → Intimal proliferation → Vessel narrowing → Reduced blood flow → Hypoxia", level=1)
add_bullet("Hypoxia → Cytokines (IL-4, TGF-β, CTGF) → Fibroblast stimulation → Excessive collagen + proteoglycan → Fibrosis/Sclerosis", level=1)

add_h3("Raynaud's Phenomenon")
add_bullet("Abnormal vascular response — vasoconstriction/vasodilation cycle")
add_bullet("Manifests as **white → blue → red** color changes in fingers")

add_h3("Autoantibodies")
add_bullet("**Anti-Topoisomerase I (Scl-70)** = Diffuse Systemic Sclerosis")
add_bullet("**Anti-Centromere** = Limited / CREST syndrome")

add_image(22)
add_image(24)

# ─── Topic 5: Systemic Sclerosis — Classification, Clinical Features & CREST ─
add_h2("Systemic Sclerosis — Classification, Clinical Features & CREST")

add_h3("Two Types")
add_bullet("**Diffuse Systemic Sclerosis** (Scleroderma)")
add_bullet("**Limited Systemic Sclerosis:**")
add_bullet("Acrosclerosis (bilateral symmetric, hands/feet only)", level=1)
add_bullet("CREST syndrome", level=1)

add_h3("CREST Syndrome")
add_bullet("**C** — Calcinosis cutis")
add_bullet("**R** — Raynaud's phenomenon")
add_bullet("**E** — Esophageal dysmotility")
add_bullet("**S** — Sclerodactyly")
add_bullet("**T** — Telangiectasia")

add_h3("Clinical Features of Scleroderma")
add_bullet("Raynaud's phenomenon (white → blue → red)")
add_bullet("Sclerodactyly (fingers contracted from skin tightening)")
add_bullet("Digital ulcers")
add_bullet("**Salt-and-pepper leukoderma** (perifollicular pigment preservation + surrounding depigmentation)")
add_bullet("**Matted/squared-off telangiectasia** (clustered, like mosquito bites)")
add_bullet("Fish mouth (microstomia)")
add_bullet("Calcinosis cutis (dystrophic)")

add_h3("Internal Organ Involvement")
add_bullet("GIT: GERD, Dysphagia")
add_bullet("Lungs: Interstitial lung disease, Pulmonary hypertension")
add_bullet("Renal")
add_bullet("Cardiac")

add_h3("Diagnostic Criteria")
add_bullet("**1 Major:** Symmetric cutaneous sclerosis proximal to MCP/MTP joints")
add_bullet("**OR 2+ Minor:** Sclerodactyly, Digital pitted scars, Loss of finger pad fat, Bibasilar pulmonary fibrosis")

add_h3("Serology")
add_bullet("ANA → **Anti-Topoisomerase (Scl-70)** for diffuse")
add_bullet("ANA → **Anti-Centromere** for CREST")
add_bullet("Both start bilateral/symmetric from hands/feet/face; Diffuse extends to entire body; Limited stops at extremities")

add_image(26)
add_image(28)
add_image(30)
add_image(32)

# ─── Topic 6: Systemic Sclerosis — Treatment ────────────────────────────────
add_h2("Systemic Sclerosis — Treatment")

add_bullet("**NO effective cure** for the sclerosis itself")

add_h3("Cutaneous Sclerosis")
add_bullet("D-penicillamine")
add_bullet("Minocycline")
add_bullet("Methotrexate")
add_bullet("PUVA / UVA1")

add_h3("Raynaud's Phenomenon")
add_bullet("Avoid cold, Stop smoking")
add_bullet("**Calcium channel blockers (Nifedipine)** — first-line vasodilator")
add_bullet("**PDE-5 inhibitors (Sildenafil/Viagra)** — third-line")

add_h3("Digital Ulcers")
add_bullet("Avoid excessive debridement")
add_bullet("Moist non-adherent dressings")
add_bullet("Low-dose Aspirin")

add_h3("Calcinosis Cutis")
add_bullet("Sodium thiosulfate")
add_bullet("Calcium channel blockers")
add_bullet("Surgical excision")
add_bullet("Low-dose Warfarin")

add_h3("DDx for Scleroderma-like Conditions")
add_bullet("Chronic GVHD")
add_bullet("**Scleredema** (adultorum of Buschke) — mucin deposition on back, associated with diabetes")
add_bullet("**Scleromyxedema** — systemic mucinosis with paraproteinemia")
add_bullet("Eosinophilic fasciitis")
add_bullet("Generalized morphea")
add_bullet("POEMS syndrome")
add_bullet("Carcinoid syndrome")

add_image(34)
add_image(36)

# ─── Topic 7: Morphea — Definition, Classification & Difference from Scleroderma ─
add_h2("Morphea — Definition, Classification & Difference from Scleroderma")

add_bullet("**Morphea** = Inflammatory skin disease ONLY (NO internal organ involvement); Localized sclerosis")
add_bullet("Same pathogenesis as Scleroderma (blood vessel → ischemia → cytokines → collagen) but LOCALIZED")

add_h3("Key Differences from Systemic Sclerosis")
add_bullet("**Morphea:** Patchy/linear, asymmetric; No Raynaud's; No internal organs; ANA usually negative")
add_bullet("**Scleroderma:** Symmetric, starts hands/feet; Raynaud's present; Internal organs affected; ANA positive (Scl-70 or Centromere)")
add_bullet("Histopathology similar but Morphea has **more intense inflammatory infiltrate**")

add_h3("Types — Mnemonic 'BB GDL'")
add_bullet("**P**laque (B)")
add_bullet("**B**ullous (B)")
add_bullet("**G**eneralized (G)")
add_bullet("**D**eep (D)")
add_bullet("**L**inear (L)")

add_h3("Plaque Morphea")
add_bullet("**Most common** type")
add_bullet("Asymmetric, multiple on trunk")
add_bullet("Variants: Superficial, Nodular (keloid-like), Guttate")
add_bullet("Has **Lilac border** in early inflammatory stage")

add_h3("Linear Morphea")
add_bullet("**Most common in CHILDREN**")
add_bullet("3 variants:")
add_bullet("Linear (trunk/limbs)", level=1)
add_bullet("**En coup de sabre** (forehead — 'sword strike')", level=1)
add_bullet("**Parry-Romberg disease** (unilateral facial fat loss → facial asymmetry)", level=1)

add_h3("Deep Morphea")
add_bullet("Subcutaneous fat involved")
add_bullet("Skin tethered/bound down (cup-shaped depression)")
add_bullet("Can cause SCC")
add_bullet("Can affect muscle/bone")

add_h3("Generalized Morphea — Criteria '2, 3, 4'")
add_bullet("≥**2** anatomical sites")
add_bullet("≥**3** cm each")
add_bullet("≥**4** plaques")

add_h3("Bullous Morphea")
add_bullet("Can complicate any type of Morphea")

add_image(38)
add_image(40)
add_image(42)
add_image(44)

# ─── Topic 8: Morphea — Treatment ───────────────────────────────────────────
add_h2("Morphea — Treatment")

add_h3("Topical (for Localized Plaque Morphea)")
add_bullet("Topical steroids")
add_bullet("Intralesional steroids")
add_bullet("Topical calcineurin inhibitors (Tacrolimus/Pimecrolimus)")
add_bullet("Vitamin D derivatives (Calcipotriol BID)")
add_bullet("Imiquimod (increases IFN-gamma → alters cytokine process)")
add_bullet("Targeted phototherapy (UVA1 or NB-UVB on lesion)")

add_h3("Systemic (for Generalized, Linear, Facial, Joint-involving, Deep Morphea)")
add_bullet("Corticosteroids: Methylprednisolone or Prednisolone 1-2 mg/kg/day")
add_bullet("**Methotrexate:** 5-25 mg (adult), 0.3 mg/kg (children)")
add_bullet("**Mycophenolate mofetil (MMF):** specifically inhibits Type 1 collagen expression")

add_h3("Treatment Algorithm")
add_bullet("**Generalized →** Start Phototherapy → if no response → Methotrexate + Steroids → if no response → MMF")
add_bullet("**Linear →** Start directly with Methotrexate + Steroids (skip phototherapy)")

add_h3("Key Point")
add_bullet("Once sclerosis/scarring occurs, it does **NOT** reverse")
add_bullet("Goal: stop the inflammatory process")
add_bullet("Surgery (fat grafting, fillers) for cosmetic correction of deformities")

add_image(46)

# ─── Topic 9: Lichen Sclerosus et Atrophicus ─────────────────────────────────
add_h2("Lichen Sclerosus et Atrophicus")

add_bullet("Name breakdown: **Lichen** (interface dermatitis) + **Sclerosus** (sclerosis) + **Atrophicus** (atrophy)")
add_bullet("Histopathology pattern: Lichenoid interface dermatitis (band-like lymphocytic infiltrate + vacuolar degeneration of BMZ)")

add_h3("Pathogenesis")
add_bullet("Oxidative stress (H₂O₂/free radicals) on skin → triggers IgG autoantibodies against **Extracellular Matrix Protein 1 (ECM1)** → sclerosis + atrophy")
add_bullet("**Genetic:** HLA-DQ7")
add_bullet("More common in females (10:1)")
add_bullet("Bimodal: prepubertal + postmenopausal")

add_h3("Lipoid Proteinosis Connection")
add_bullet("Genetic DEFECT in ECM1 gene = **Lipoid Proteinosis** (Hyalinosis cutis et mucosae)")
add_bullet("Features: waxy papules, pitted scars on nose/forehead, hoarse voice/crying since birth")

add_h3("Distribution")
add_bullet("**Genital (70-80%)** vs **Extragenital (20-30%)**")

add_h3("Genital — Female (Kraurosis Vulvae)")
add_bullet("Vulva + perianal (**Figure of 8 pattern**)")
add_bullet("Severe pruritus")
add_bullet("Shiny hypopigmented sclerotic atrophic plaque")
add_bullet("Complications: Scarring, Dyspareunia, **Precancerous (SCC risk)**")

add_h3("Genital — Male (Balanitis Xerotica Obliterans)")
add_bullet("Stenosis, Dysuria")
add_bullet("Precancerous")

add_h3("Extragenital")
add_bullet("Sites: Trunk, proximal extremities, neck, shoulders, flexor wrists, pressure sites")
add_bullet("Lesion: Ivory-colored sclerotic scar-like papules/plaques with shiny surface")
add_bullet("May have telangiectasia, follicular plugging (comedone-like)")
add_bullet("Advanced: hemorrhagic bullae (from fragile DEJ)")

add_h3("Histopathology")
add_bullet("Lichenoid interface dermatitis")
add_bullet("**MOST CHARACTERISTIC:** Edema in the upper dermis (pale-staining homogeneous band)")
add_bullet("Late: hyperkeratosis, epidermal atrophy, orthohyperkeratosis (especially follicular openings → plugging)")
add_bullet("Vacuolar degeneration of basal layer, flattened rete ridges")
add_bullet("Key: problems confined to **UPPER dermis** (deep dermis and subcut fat normal — unlike Morphea/Scleroderma)")

add_h3("Differential Diagnosis")
add_bullet("Vitiligo (genital)")
add_bullet("Sexual abuse (in children)")
add_bullet("Morphea (extragenital)")
add_bullet("Lichen planus (genital)")

add_h3("Treatment")
add_bullet("Same as Morphea: Topical steroids, Intralesional steroids, Topical calcineurin inhibitors, Imiquimod")
add_bullet("Systemic: Penicillin, Antimalarials, Corticosteroids, Vitamin E")

add_image(48)
add_image(50)
add_image(52)
add_image(54)

# ─── Topic 10: Bonus — End-of-Lecture SCE Questions ──────────────────────────
add_h2("Bonus — End-of-Lecture SCE Questions")

add_h3("Q1: Fumaric Acid Ester Side Effect")
add_bullet("**Progressive Multifocal Leukoencephalopathy** (demyelinating disease)")
add_bullet("If patient on Fumaric acid develops CNS symptoms → **stop drug + MRI urgently**")

add_h3("Q2: Pemphigus Vulgaris Resistant to Conventional Therapy")
add_bullet("Answer: **Rituximab** (anti-CD20, anti-B cell monoclonal antibody)")
add_bullet("CD markers: CD3 = all T-cells, CD4 = T-helper, CD8 = T-cytotoxic, CD20 = B-cells")

add_h3("Q3: Isotretinoin and Night Vision")
add_bullet("Advise patient to avoid driving at night **only IF** they develop night vision problems")
add_bullet("Not a prophylactic restriction — only if symptoms develop")

add_image(58)

# ══════════════════════════════════════════════════════════════════════════════
#  TRANSCRIPT — Lecture 10
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_h2("Lecture 10 — Full Transcript")


add_normal("**Professor:**  Can you all hear me? Could someone please type 'ok'? Is someone speaking and I can't hear? Alright, could someone just speak so I know if I can hear you or not. Hey everyone, doctors, anyone please respond. It seems like you can hear me but I can't hear you. Okay, I know you can hear me and see the screen now, but someone please speak so I know I can hear you too.")
add_normal("")
add_normal("**Student:**  Doctor, we can hear you.")
add_normal("")
add_normal("**Professor:**  Finally someone responded! How are you, doctors? Are you still fasting or what? Has iftar kicked in already? What's the news? How's the studying going?")
add_normal("")
add_normal("**Student:**  Sorry doctor, we're in clinic so we won't be able to answer you much.")
add_normal("")
add_normal("**Professor:**  Oh of course, sorry about that. I'm the one who switched the schedule today; God willing on Tuesday we'll be back to normal and I won't do this again. Alright, whoever can respond, tell me what's going on. How's the studying? How's everything going?")
add_normal("")
add_normal("**Professor:**  Last time we started a new chapter — Connective Tissue Diseases — and we talked about investigations, the most important being the ANA. We discussed the first and most important disease: Lupus Erythematosus in all its forms — Systemic, Acute, Chronic, and Subacute. Today we'll continue with a few more topics in Connective Tissue Diseases. God willing they'll be easy for us.")
add_normal("")
add_normal("**Professor:**  We'll start with our first disease today: **Dermatomyositis**. What does Dermatomyositis mean? It's a multisystemic autoimmune disease affecting the muscle, the skin, and others. From its name: Dermato = skin, Myositis = muscle inflammation.")
add_normal("")
add_normal("**Professor:**  It's a rare disease worldwide with a bimodal distribution — meaning there's Juvenile Dermatomyositis and Adult Dermatomyositis. Which is more dangerous — juvenile or adult?")
add_normal("")
add_normal("**Student:**  Related to cancer.")
add_normal("")
add_normal("**Professor:**  Bravo! Related to cancer. The adult form is more dangerous because of malignancy association.")
add_normal("")
add_normal("**Professor:**  Dermatomyositis classification: **Skin only** = Amyopathic DM (DM sine myositis), **Skin + Muscle** = Classic DM, **Muscle only** = Polymyositis.")
add_normal("")
add_normal("**Professor:**  What characterizes Dermatomyositis muscle problems? The inflammation involves **proximal muscles** — biceps, triceps, quadriceps. The patient presents saying: 'Doctor, I can't raise my hand to comb my hair and I can't climb stairs.' Why would she come to me as a dermatologist? Because of what appears on the skin:")
add_normal("")
add_normal("**Professor:**  First: **Poikiloderma** — appearing in various forms. We'll find erythema, a macular rash on the face, especially around the eyes — that's **Heliotrope erythema**. Then the **V sign** anteriorly and the **Shawl sign** posteriorly. All of this is Poikiloderma. Poikiloderma means: erythema + telangiectasia + hypo/hyperpigmentation.")
add_normal("")
add_normal("**Professor:**  The **Gottron papules** and **Gottron sign** — Gottron sign is erythema and telangiectasia with mottled hypo/hyperpigmentation (Poikiloderma) present over joints. If elevated above the skin as papules on the knuckles, we call them Gottron papules.")
add_normal("")

add_h3("[10:00] — Pathogenesis & Triggers")

add_normal("**Professor:**  The pathogenesis of DM — same concept: environmental trigger in a genetically susceptible patient → immune dysregulation. The triggers: **first and most important = malignancy**, second = infections, third = drugs.")
add_normal("")
add_normal("**Professor:**  What's the most important malignancy associated with DM? Especially since it's adult-onset and more common in women like any autoimmune disease — we worry most about **Ovarian cancer** and **Colon cancer**.")
add_normal("")
add_normal("**Professor:**  Infections that can trigger DM: Coronavirus, Echovirus, Coxsackievirus, HIV, Toxoplasma, Borrelia, E. coli. Drugs: NSAIDs, Statins, Hydroxyurea, D-penicillamine, TNF-alpha inhibitors, Cyclophosphamide.")
add_normal("")
add_normal("**Professor:**  Genetically susceptible — how do we know? More common in monozygotic twins, can have family history. The HLA associations: **HLA-DR3, HLA-DR52, HLA-DR7**.")
add_normal("")
add_normal("**Professor:**  The immunological trigger stimulates **Anti-Jo-1** and **Anti-Mi-2**. What are these? They're components of the nucleus. Mi-2 is the **Helicase**, and Anti-Jo-1 is the **Synthetase** (histidyl-tRNA synthetase). These are autoantibodies against nuclear components.")
add_normal("")
add_normal("**Professor:**  The skin involvement — Poikiloderma, but it's NOT the brown type. It's **pink to violet erythema**. Around the eyes = Heliotrope erythema. On the chest = V sign. On the back = Shawl sign — like wearing a shawl of erythema. On joints = Gottron sign (elbows, knees). If elevated = Gottron papules on knuckles. All of these signs are **severely itchy**.")
add_normal("")
add_normal("**Professor:**  Other findings: **Nail fold changes** — ragged cuticle and telangiectasia visible on dermoscopy. **Calcinosis cutis** — dystrophic type (NOT metastatic). **Cutaneous small vessel vasculitis**.")
add_normal("")

add_h3("[20:00] — Differential Diagnosis & Muscle Involvement")

add_normal("**Professor:**  Differential diagnosis for **malar erythema** — a woman in her 40s with red cheeks: 1) Dermatomyositis, 2) Systemic Lupus Erythematosus, 3) Rosacea, 4) Seborrheic dermatitis. These four should never be forgotten for homogeneous erythema on the face.")
add_normal("")
add_normal("**Professor:**  The Gottron sign DDx — erythema on elbows/knees with scales: **Psoriasis** (silvery white scales). Don't confuse it.")
add_normal("")
add_normal("**Professor:**  The muscle involvement: affects **proximal extensor muscles**, always bilateral and symmetrical, tender on palpation. Patient can't comb hair or climb stairs.")
add_normal("")
add_normal("**Professor:**  DM doesn't just affect skeletal muscle — it also affects **smooth muscle**. This means it affects most organs: **GIT** (GERD, dysphagia from cricopharyngeal muscle affection — don't forget this muscle!), **Cardiac** (arrhythmia), **Pulmonary** (diffuse interstitial lung fibrosis).")
add_normal("")

add_h3("[30:00] — Antisynthetase Syndrome & Investigations")

add_normal("**Professor:**  When you have diffuse interstitial fibrosis with fever, myositis, Gottron papules/signs, and interstitial lung fibrosis — this is a severe form of DM. The patient also has erosive arthritis and **Mechanic's hands** (rough hands from using hands to support walking). What's this syndrome called?")
add_normal("")
add_normal("**Student:**  Antisynthetase syndrome?")
add_normal("")
add_normal("**Professor:**  Yes! The **Antisynthetase syndrome** — autoantibodies against **Jo-1**, NOT Mi-2. Mi-2 is classic DM. Jo-1 is the dangerous one. It includes: DM + Interstitial lung fibrosis + Erosive polyarthritis + Mechanic's hands + Fever.")
add_normal("")
add_normal("**Professor:**  Let me summarize: Environmental triggers (malignancy most important, then infections, then drugs) in genetically susceptible patients produce autoantibodies against Mi-2 (classic) or Jo-1 (severe with lung fibrosis). **Malignancy association** = ONLY with Adult DM and Amyopathic DM. NOT in Juvenile DM, NOT in Polymyositis. Most feared: Ovarian, Breast, Colon cancer. Juvenile DM is better — no malignancy.")
add_normal("")
add_normal("**Professor:**  **Causes of death**: Respiratory infection, Cardiac failure, Malnutrition (dysphagia), Carcinoma.")
add_normal("")
add_normal("**Professor:**  **Investigations**: First — skin biopsy. Then serum ANA (if positive → Anti-Mi-2/Anti-Jo-1). Muscle enzymes: **CPK, Aldolase, LDH, Myoglobin, AST/ALT**. Muscle MRI/US. Electromyography. Muscle biopsy (from **triceps**) — showing lymphocytic inflammatory infiltrate, perifascicular and perivascular, muscle atrophy and necrosis. Urine: Creatinine/Myoglobin. Screen for malignancy.")
add_normal("")
add_normal("**Professor:**  **Skin biopsy**: DM falls under **vacuolar interface dermatitis** (same as SLE). From Poikiloderma: sparse lymphocytic infiltrate, vasodilation, interstitial mucin deposition. From Gottron papules: **lichenoid infiltrate** (more intense), acanthosis, NO epidermal atrophy. This explains the spectrum: vacuolar (less intense) → lichenoid (band-like when lymphocytes increase significantly).")
add_normal("")
add_normal("**Professor:**  **Treatment**: Skin — challenging, doesn't respond well. Photoprotection (lesions are photosensitive). Antimalarials. Low-dose Methotrexate. Immunosuppressives. Topical steroids/Tacrolimus. **Muscle** — Systemic steroids (Prednisolone 1 mg/kg/day × 6 months), steroid-sparing drugs, rest, physiotherapy. Muscle responds well, skin doesn't.")
add_normal("")

add_h3("[40:00] — Systemic Sclerosis (Scleroderma)")

add_normal("**Professor:**  Now we'll discuss **Systemic Sclerosis** or Scleroderma — meaning systemic hardening. The skin becomes extremely tight, then fibrosis occurs, then sclerosis. Fingers contract, ulcers appear, nail problems develop.")
add_normal("")
add_normal("**Professor:**  What causes this? What makes the skin become fibrotic and hard as wood?")
add_normal("")
add_normal("**Student:**  Collagen.")
add_normal("")
add_normal("**Professor:**  Bravo! Excess **collagen production** is the problem. The starting point in Systemic Sclerosis is the **blood vessels**. The smooth muscle of blood vessels keeps opening and closing — vasoconstriction, vasodilation — which manifests as **Raynaud's phenomenon**.")
add_normal("")
add_normal("**Professor:**  What happens: **Endothelial cell injury** → intimal proliferation → wall thickens → vessel narrows → reduced blood flow → **hypoxia** → cytokines released → **fibroblasts stimulated** → excess collagen production → **fibrosis → sclerosis**. The skin becomes like iron, like wood.")
add_normal("")
add_normal("**Professor:**  The pathogenesis diagram: Endothelial injury → intimal proliferation → lumen narrows → hypoxia → immune stimulation (T-cells, IL-4, TGF-β, CTGF) → increased collagen and proteoglycan. Immune dysregulation produces: **Anti-centromere** (→ CREST syndrome) or **Anti-topoisomerase/Scl-70** (→ Diffuse Systemic Sclerosis).")
add_normal("")
add_normal("**Professor:**  Two types: **Diffuse Systemic Sclerosis** (Scleroderma) vs **Limited Systemic Sclerosis**. Limited can be: Acrosclerosis (bilateral symmetric, hands/feet only) OR **CREST syndrome** (Calcinosis, Raynaud's, Esophageal dysmotility, Sclerodactyly, Telangiectasia).")
add_normal("")
add_normal("**Professor:**  Both types start bilateral and symmetric from hands/feet/face. The diffuse type extends to the entire body. The limited type stops at the extremities. Both have Raynaud's phenomenon (white → blue → red).")
add_normal("")

add_h3("[50:00] — Scleroderma Clinical Features & Diagnosis")

add_normal("**Professor:**  Clinical features: **Digital ulcers**. **Salt-and-pepper leukoderma** — perifollicular pigmentation is preserved while surrounding skin becomes white. **Matted/squared-off telangiectasia** — clustered together like mosquito bites. **Calcinosis cutis** — dystrophic type. **Fish mouth** (microstomia). **Sclerodactyly**.")
add_normal("")
add_normal("**Student:**  Vascular.")
add_normal("")
add_normal("**Professor:**  Bravo! The pathogenesis starts with vascular dysfunction.")
add_normal("")
add_normal("**Professor:**  **Serology**: First do ANA. If positive → Anti-topoisomerase (Scl-70) for diffuse; Anti-centromere for CREST/limited. Do both since CREST can progress.")
add_normal("")
add_normal("**Professor:**  **Diagnostic criteria**: Major = symmetric cutaneous sclerosis proximal to MCP/MTP joints. Minor = sclerodactyly, digital pitted scars, loss of finger pad fat, bibasilar pulmonary fibrosis. Diagnosis requires **1 Major OR 2+ Minor** criteria.")
add_normal("")

add_h3("Scleroderma Treatment (Part 2 of Lecture)")

add_normal("**Professor:**  Treatment of Systemic Sclerosis: First, understand there's NO effective treatment for the sclerosis itself. But we use: **D-penicillamine, Minocycline, Methotrexate, PUVA/UVA1**.")
add_normal("")
add_normal("**Professor:**  **Raynaud's treatment**: Avoid cold, stop smoking. Vasodilators: **Calcium channel blockers (Nifedipine)** = first-line. **PDE-5 inhibitors (Sildenafil/Viagra)** = third-line.")
add_normal("")
add_normal("**Professor:**  **Digital ulcers**: Avoid excessive debridement (tissue has vascular problems — cleaning removes tissue and worsens it). Use moist non-adherent dressings. Low-dose Aspirin.")
add_normal("")
add_normal("**Professor:**  **Calcinosis cutis**: Sodium thiosulfate, Calcium channel blockers, Surgical excision of nodules, Low-dose Warfarin.")
add_normal("")
add_normal("**Professor:**  **Differential diagnosis** for scleroderma-like conditions: **Chronic GVHD** (remember it can present as Scleroderma, Morphea, Lichen planus, Lichen striatus). **Scleredema** (adultorum of Buschke) — mucin deposition on back, associated with diabetes. **Scleromyxedema** — systemic mucinosis with paraproteinemia/lymphoproliferative disorders. Eosinophilic fasciitis, Generalized morphea, POEMS syndrome, Carcinoid syndrome.")
add_normal("")

add_h3("[10:00 Part 2] — Morphea")

add_normal("**Professor:**  **Morphea** — what's the difference from Systemic Sclerosis? Morphea is an inflammatory skin disease ONLY. No systemic affection. No internal organ involvement. Unlike Scleroderma which has internal organ involvement.")
add_normal("")
add_normal("**Professor:**  Same pathogenesis as Scleroderma but LOCALIZED: blood vessel → endothelial injury (NO Raynaud's) → intimal proliferation → ischemia → hypoxia → cytokines → fibroblasts → excess collagen → scar-like sclerosis.")
add_normal("")
add_normal("**Professor:**  **Comparison table (from Dr. Ahmed Kamel's textbook)**: Morphea = patchy/linear, asymmetric; No other cutaneous features; No internal organ involvement; More intense inflammatory infiltrate histologically; ANA usually negative. Systemic Sclerosis = symmetric, starts hands/fingers, extends proximally; Raynaud's, ulceration, sclerodactyly, telangiectasia, calcinosis; GIT and Lungs mainly affected; ANA positive (Topoisomerase/Centromere).")
add_normal("")

add_h3("[20:00 Part 2] — Types of Morphea")

add_normal("**Professor:**  Types of Morphea — mnemonic **'BB GDL'**: Plaque (B), Bullous (B), Generalized (G), Deep (D), Linear (L).")
add_normal("")
add_normal("**Professor:**  **Plaque morphea** — most common. Asymmetric, multiple on trunk. Three variants: Superficial, Nodular (keloid-like), Guttate. Has a **Lilac border** in the early inflammatory stage.")
add_normal("")
add_normal("**Professor:**  **Linear morphea** — most common in children. Three types: Linear (trunk/limbs), **En coup de sabre** (forehead — 'sword strike' — we see this a lot especially in young girls, the forehead literally looks divided in two with atrophy of skin and subcutaneous fat), **Parry-Romberg disease** (unilateral facial fat loss → facial asymmetry).")
add_normal("")
add_normal("**Professor:**  **Deep morphea** — involves subcutaneous fat, skin is tethered/bound down (like a cup-shaped depression — reminds us of Lupus panniculitis). Can cause disabling pansclerotic morphea, may affect muscle/bone, can lead to Squamous Cell Carcinoma.")
add_normal("")
add_normal("**Professor:**  **Generalized morphea criteria — '2, 3, 4'**: ≥2 anatomical sites + ≥3 cm each + ≥4 plaques. That's how you diagnose Generalized morphea.")
add_normal("")
add_normal("**Professor:**  **Bullous morphea** — not really a separate type; any type of morphea can develop bullae.")
add_normal("")

add_h3("[30:00 Part 2] — Morphea Treatment")

add_normal("**Professor:**  **Topical treatment** (for localized plaque morphea): Topical steroids, Intralesional steroids, Topical calcineurin inhibitors (Tacrolimus/Pimecrolimus), Calcipotriol (Vitamin D) twice daily, Imiquimod (increases IFN-gamma production → alters the immunological process → reduces cytokines → reduces collagen production), Targeted phototherapy (UVA1 or NB-UVB on the lesion only).")
add_normal("")
add_normal("**Professor:**  **Systemic treatment** (for Generalized, Linear, facial, joint-involving, Deep morphea): Corticosteroids (Methylprednisolone or Prednisolone 1-2 mg/kg/day). **Methotrexate** (5-25 mg adult, 0.3 mg/kg children — remember this dose, it's the same for pediatric Psoriasis). **Mycophenolate mofetil (MMF)** — specifically inhibits Type 1 collagen expression (Type 1 is the main collagen in skin).")
add_normal("")
add_normal("**Professor:**  **Treatment algorithm**: Generalized morphea → start with Phototherapy (UVA1/NB-UVB) → no response → Methotrexate + Systemic steroids → no response after 8 weeks → change to MMF. **Linear morphea** → start directly with Methotrexate + Systemic steroids (skip phototherapy).")
add_normal("")
add_normal("**Professor:**  **Key point**: Once sclerosis has occurred in the skin, it's a scar — it does NOT reverse. The goal is to stop the ongoing inflammatory process. For cosmetic deformities (like En coup de sabre or Parry-Romberg), we can inject fillers, fat grafting, or refer to plastic surgery. But first stop the inflammation, then address the deformity.")
add_normal("")
add_normal("**Student:**  When do you stop treatment? When there's no active inflammatory lesion?")
add_normal("")
add_normal("**Professor:**  Exactly. Once there's no active inflammation, we stop. The sclerosis that already happened won't return to normal skin. Our goal is only to prevent progression.")
add_normal("")

add_h3("[40:00 Part 2] — Lichen Sclerosus et Atrophicus")

add_normal("**Professor:**  Last disease for today: **Lichen Sclerosus et Atrophicus**. From its name: Lichen = interface dermatitis (band-like lymphocytic infiltrate + vacuolar degeneration of the dermoepidermal junction). Sclerosus = sclerosis (collagen/material deposition). Atrophicus = atrophy.")
add_normal("")
add_normal("**Professor:**  **Pathogenesis**: Oxidative stress on the skin (H₂O₂, free radicals) → triggers IgG autoantibodies against **Extracellular Matrix Protein 1 (ECM1)** → sclerosis + atrophy. In genetically susceptible patients: **HLA-DQ7**. More common in females (10:1). Bimodal peak: prepubertal and postmenopausal.")
add_normal("")
add_normal("**Professor:**  **Lipoid Proteinosis connection**: If there's a genetic DEFECT in the ECM1 gene, you get **Lipoid Proteinosis** (Hyalinosis cutis et mucosae). Features: hyaline material deposited in skin → waxy papules → pitted scars (especially nose/forehead), and hoarseness of voice/crying since birth. Very important for exams!")
add_normal("")
add_normal("**Professor:**  **Clinical picture**: Genital (70-80%) vs Extragenital (20-30%). **Female genital** = Kraurosis vulvae — affects vulva + perianal area (Figure of 8 pattern). Severe pruritus. Shiny, hypopigmented, sclerotic, atrophic plaque. Complications: scarring, dyspareunia, **precancerous (SCC risk)**.")
add_normal("")
add_normal("**Professor:**  **Male genital** = Balanitis xerotica obliterans — causes stenosis, dysuria, precancerous. **DDx**: Vitiligo (genital), Sexual abuse (in children), Morphea (extragenital).")
add_normal("")

add_h3("[50:00 Part 2] — Extragenital LS & Histopathology")

add_normal("**Professor:**  **Extragenital**: Frequently asymptomatic (except dryness and pruritus). Sites: trunk, proximal extremities, neck, shoulders, flexor wrist, pressure sites. Lesion: sclerotic scar-like papules/plaques, ivory-colored with shiny surface. Advanced cases: telangiectasia, **follicular plugging** (comedone-like), and **hemorrhagic bullae** (due to fragile DEJ — because it's vacuolar interface dermatitis).")
add_normal("")
add_normal("**Professor:**  **Histopathology**: Lichenoid interface dermatitis. **Most characteristic** = edema in the upper dermis (pale-staining homogeneous band). Late findings: hyperkeratosis, epidermal thinning/atrophy, orthohyperkeratosis (especially in follicular openings → follicular plugging), vacuolar degeneration of basal cell layer, flattened rete ridges. Key difference: problems confined to **upper dermis only** — deep dermis and subcutaneous fat are normal (unlike Morphea/Scleroderma).")
add_normal("")
add_normal("**Professor:**  **Treatment** of Lichen Sclerosus = same as Morphea: Topical steroids, Intralesional corticosteroids, Topical calcineurin inhibitors, Imiquimod. Systemic: Penicillin, Antimalarials, Corticosteroids, Vitamin E.")
add_normal("")
add_normal("**Professor:**  That's it for today's lecture! I hope it wasn't too heavy for you.")
add_normal("")

add_h3("SCE Questions Discussion")

add_normal("**Professor:**  Let me share a few exam-style questions before we finish:")
add_normal("")
add_normal("**Professor:**  **Question 1** — A 50-year-old man on **Fumaric acid ester** for psoriasis in remission. His wife noticed behavioral changes, new-onset clumsiness, and limb weakness suggestive of CNS involvement. What should you do? **Answer**: Stop medication and arrange for MRI scan urgently. Why? Because Fumaric acid causes **Progressive Multifocal Leukoencephalopathy** (a demyelinating disease). Key takeaway: if a patient on Fumaric acid develops CNS symptoms → stop drug + MRI immediately.")
add_normal("")
add_normal("**Professor:**  **Question 2** — A patient with **Pemphigus vulgaris** not responding to treatment (1.5 mg/kg steroids + MMF + Azathioprine caused hepatotoxicity). What biologic do we use? We need one that targets B-cells (since PV involves autoantibodies against Desmoglein). **Answer: Rituximab** — anti-CD20 (B-cell marker). Remember: CD3 = all T-cells, CD4 = T-helper, CD8 = T-cytotoxic, **CD20 = B-cells**.")
add_normal("")
add_normal("**Professor:**  **Question 3** — Isotretinoin and night driving. A patient wants to restart Isotretinoin. What's the most appropriate advice? **Answer**: He should avoid driving at night **only IF he develops night vision problems**. It's not a blanket restriction — only advise this if symptoms develop. This is about the Vitamin A-related side effect of impaired night vision.")
add_normal("")
add_normal("**Student:**  Thank you very much, Doctor!")
add_normal("")
add_normal("**Professor:**  You're welcome! God bless you all. Any questions?")
add_normal("")
add_normal("**Student:**  If possible, what will be the topic of the next lecture so we can prepare?")
add_normal("")
add_normal("**Professor:**  They're supposed to upload the PDF for you before the lecture. Last time there was a mix-up because I changed the curriculum. I'll coordinate with them to make sure it's uploaded consistently. Don't worry about it.")

# ══════════════════════════════════════════════════════════════════════════════
#  Save
# ══════════════════════════════════════════════════════════════════════════════

print("\n💾 Saving document...")
doc.save(OUT_PATH)
print(f"✅ Done! Saved as:\n   {OUT_PATH}")
print(f"   File size: {os.path.getsize(OUT_PATH) / 1024 / 1024:.2f} MB")
